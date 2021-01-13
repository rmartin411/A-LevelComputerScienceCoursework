import os.path
import time
from tkinter import *
from tkinter import messagebox
from functools import partial
from functools import partialmethod

from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

import sqlite3
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from reportlab import *

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from smtplib import SMTPException

import datetime

now = datetime.datetime.now()
Years = list(range(1920, int(now.year)+1))
CompYears = list(range(int(now.year), int(now.year)+3))
print(CompYears)
CompYears.reverse()

Years.reverse()
print(Years)
Days = range(0, 32)


root = Tk()
global Login
Login = False

# create the CompetitionDraw table
conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
cursor.execute(
    'CREATE TABLE IF NOT EXISTS CompetitionDraw (CompetitionID INTEGER, BandID INTEGER, EntryID INTEGER, CircleNumber INTEGER,'
    'CompetitingTime TEXT)')
conn.commit()

# create the BandsEntered table
conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
# create the BandsEntered table
cursor.execute('CREATE TABLE IF NOT EXISTS BandsEntered '
               '(EntryID INTEGER PRIMARY KEY, BandID INTEGER, CompetitionID INTEGER, DateEntered TEXT, Grade TEXT)')
conn.commit()

conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
# create the Results table
cursor.execute(
    'CREATE TABLE IF NOT EXISTS Results (EntryID INTEGER, BandID INTEGER, CompetitionID INTEGER, '
    'Piping1JudgeID INTEGER, Piping1Score INTEGER, Piping1Comments TEXT, '
    'Piping2JudgeID INTEGER, Piping2Score INTEGER, Piping2Comments TEXT, '
    'TotalPipingScore INTEGER,'
    'DrummingJudgeID INTEGER, DrummingScore INTEGER, DrummingComments TEXT, '
    'EnsembleJudgeID INTEGER, EnsembleScore INTEGER, EnsembleComments TEXT,'
    'TotalScore INTEGER, Position INTEGER, BandSheetsFileName TEXT)')
conn.commit()

conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
# create the BandMembers table
cursor.execute(
    'CREATE TABLE IF NOT EXISTS BandMembers (MemberID INTEGER PRIMARY KEY, BandID INTEGER, '
    'Role TEXT, Title TEXT, FirstName TEXT, SecondName TEXT, AddressLine1 TEXT, AddressLine2 TEXT, County TEXT, '
    'Postcode TEXT, DateOfBirth TEXT, Gender TEXT, Signature TEXT, EmailAddress TEXT)')
conn.commit()

conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
# create the Band table to store Band account details
cursor.execute('CREATE TABLE IF NOT EXISTS BandAccount '
               '(BandID INTEGER PRIMARY KEY, BandName TEXT, BandBranch TEXT, BandGrade TEXT, PMFirstName TEXT, '
               'PMSecondName TEXT, HallName TEXT, BandAddressLine1 TEXT, BandAddressLine2 TEXT, BandCounty TEXT, '
               'BandPostcode TEXT, PracticeTime TEXT, Tartan TEXT, EmailAddress TEXT, Username TEXT, Password TEXT, SecurityQ TEXT, SecurityQAns, FirstLogIn BOOLEAN)')
conn.commit()

conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
# create the judge table to store judge account details
cursor.execute('CREATE TABLE IF NOT EXISTS JudgeAccount '
               '(JudgeID INTEGER PRIMARY KEY, Title TEXT, FirstName TEXT, SecondName TEXT, DateOfBirth TEXT, Gender TEXT, '
               'AddressLine1 TEXT, AddressLine2 TEXT, County TEXT, Postcode TEXT, Email TEXT,'
               'Username TEXT, Password TEXT, SecurityQ TEXT, SecurityQAns TEXT, FirstLogIn TEXT)')
conn.commit()

conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
# create the Admin table to store admin account details
cursor.execute('CREATE TABLE IF NOT EXISTS AdminAccount '
               '(AdminID INTEGER PRIMARY KEY, Title TEXT, FirstName TEXT, SecondName TEXT, DateOfBirth TEXT, Gender TEXT, '
               'AddressLine1 TEXT, AddressLine2 TEXT, County TEXT, Postcode TEXT, '
               'Email TEXT, Username TEXT, Password TEXT, SecurityQ TEXT, SecurityQAns TEXT, FirstLogIn TEXT)')
conn.commit()

with sqlite3.connect("RspbaniDB.db") as db:
    # Creates a cursor to search through database
    cursor = db.cursor()
    # SQL Query to select Username of the inital Admin account
InitalAdminQuery = 'SELECT Username FROM AdminAccount WHERE Username = "WPinkerton345"'
# Execute the query to return Username if there is a match
cursor.execute(InitalAdminQuery)
# The query will produce a 2 dimensional tuple
InitalAdminCheck = cursor.fetchall()

# if there is no results produced ie file is being run for the first time then create the iniital account
# otherwise dont create a duplicate of the account.

if InitalAdminCheck == []:
    conn = sqlite3.connect('RspbaniDB.db')
    with conn:
        cursor = conn.cursor()
    #Initally creating a user so that they can create additional accounts and competitions etc.
    cursor.execute('INSERT INTO AdminAccount '
                   '(AdminID, Title, FirstName, SecondName, DateOfBirth, Gender, AddressLine1, AddressLine2, '
                   'County, Postcode, Email, Username, Password, FirstLogIn, SecurityQ, SecurityQAns) '
                   'VALUES(NULL,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                   ("Mr", "Winston", "Pinkerton", "12/10/1978", "Male", "50 Dungonnell Road", "Crumlin",
                    "County Antrim", "BT295GH", "Wintonpink000@gmail.com", "WPinkerton345", "tjGH56H6Uj", "True", "Unknown", "Unknown"))
    conn.commit()

conn = sqlite3.connect('RspbaniDB.db')
with conn:
    cursor = conn.cursor()
# create the Competition table to store competition details
cursor.execute('CREATE TABLE IF NOT EXISTS Competitions '
               '(CompetitionID INTEGER PRIMARY KEY, CompetitionName TEXT, CompetitionLocation TEXT, CompetitionDate TEXT, EntryDeadLine TEXT, Drawn TEXT, '
               'NumG1Ent INTEGER, NumG2Ent INTEGER, '
               'NumG3AEnt INTEGER, NumG3BEnt INTEGER, NumG4AEnt INTEGER, NumG4BEnt INTEGER, TotalNumBands INTEGER, G1BestBass TEXT, '
               'G2BestBass TEXT, G3ABestBass TEXT, G3BBestBass TEXT, G4ABestBass TEXT, G4BBestBass TEXT, '
               'G1BestMD TEXT, G2BestMD TEXT, G3ABestMD TEXT, G3BBestMD TEXT, G4ABestMD TEXT, G4BBestMD TEXT, Collated TEXT, '
               'ResultFileName TEXT, Latest TEXT, SentToBands TEXT)')
conn.commit()

global GListOfGrades
GListOfGrades = ("1", "2", "3A", "3B", "4A", "4B")

# funcion to email user which takes the email of the recipient, the subject and the body. This does not allow for attachments.
def EmailUser(Email, Subject, BodyText):
    # email address of the sender
    email_user = 'rspbani.info@gmail.com'
    # password of the sender
    email_password = 'P@55w0rd123'
    # email address of recipient
    email_send = Email
    # Creating the Emails subject
    subject = str(Subject)

    msg = MIMEMultipart()
    msg['From'] = email_user
    msg['To'] = email_send
    msg['Subject'] = subject

    msg.attach(MIMEText(BodyText, 'plain'))
    text = msg.as_string()

    try:
        # connet to email server and send email
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(email_user, email_password)
        server.sendmail(email_user, email_send, text)

    # error handling if email cant be sent - display a message stating this
    except(smtplib.SMTPException, ConnectionRefusedError, OSError):
        messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
    finally:
        # stop connection with email server
        server.quit()

# Validation function which checks to ensure variable passed only contains letters or numbers.
def CheckOnlyNumORLetters(item):
    # removes spaces from string if any exist.
    item = ''.join(item.split())

    print(item)

    for i in item:
        print(i)
        print(i.isdigit())
        print(i.isalpha())
        if i.isdigit() == False and i.isalpha() == False:
            return False
        else:
            return True

# Validation function which checks to ensure variable passed must contains letters and numbers.
def CheckOnlyNumAndLetters(item):
    # removes spaces from string if any exist.

    item = ''.join(item.split())
    number = False
    letter = False

    for i in item:
        if i.isalpha()==True:
            letter = True
        elif i.isdigit() == True:
            number = True

    if number == True and letter == True:
        return True
    else:
        return False

# Validation function which checks to ensure variable passed only contains letters.
def CheckAllAlpha(item):
    # removes spaces from string if any exist.

    item = ''.join(item.split())

    if item.isalpha() == True:
        return True
    else:
        return False


# Validation function which checks to ensure variable passed only contains numbers.
def CheckAllDigits(item):
    if item.isdigit() == True:
        return True
    else:
        return False

# Validation function which checks to ensure password variable passed is a valid password.
# Must be greater or equal to 8 in length
# Contain a captical letter
# contain a lower case letter
# Contains a digit

def CheckValidPassword(password):
    length = False
    capital = False
    lower = False
    digit = False

    if len(password) >= 8:
        length = True
    for i in password:
        if i.isupper() == True:
            capital = True
        if i.islower() == True:
            lower = True
        if i.isdigit() == True:
            digit = True


    if length == True and capital == True and lower == True and digit == True:
        return True
    else:
        return False

# Validation function which checks to ensure username pass is valid and hasnt already been used.
# Must not be empty
# must be unique (cant be found in database already)
# # Must be greater or equal to 8 characters in length
# mode is edit if the user has decided to edit band details.
# If mode is edit their username is removed from the list so that if they dont change it it wont throw an error to say username has already been taken.

def CheckUsername(username, mode):
    presence = False
    unique = False
    length = False

    # Get a list of all usernames from all three tables
    ListOfUsernames = []

    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through database
        cursor = db.cursor()
        # SQL Query to select all Usernames from the Band Account table
    listOfUsernamesQuery = 'SELECT Username FROM BandAccount'
    # Execute the query to get a list of all the Usernames
    cursor.execute(listOfUsernamesQuery)
    # The query will produce a 2 dimensional tuple
    ListOfUsernamesBandAccount = cursor.fetchall()
    # For each index the tuple append to the new list
    for u in ListOfUsernamesBandAccount:
        for user in u:
            ListOfUsernames.append(user)

    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through database
        cursor = db.cursor()
        # SQL Query to select all Username from the Admin Account table
    listOfUsernamesQuery = 'SELECT Username FROM AdminAccount'
    # Execute the query to get a list of all the Usernames
    cursor.execute(listOfUsernamesQuery)
    # The query will produce a 2 dimensional tuple
    ListOfUsernamesAdminAccount = cursor.fetchall()
    # For each index the tuple append to the new list
    for u in ListOfUsernamesAdminAccount:
        for user in u:
            ListOfUsernames.append(user)

    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through database
        cursor = db.cursor()
        # SQL Query to select all Username from the Admin Account table
    listOfUsernamesQuery = 'SELECT Username FROM JudgeAccount'
    # Execute the query to get a list of all the Usernames
    cursor.execute(listOfUsernamesQuery)
    # The query will produce a 2 dimensional tuple
    ListOfUsernamesJudgeAccount = cursor.fetchall()
    # For each index the tuple append to the new list

    print(ListOfUsernamesJudgeAccount)
    for u in ListOfUsernamesJudgeAccount:
        for name in u:
            ListOfUsernames.append(name)

    print(ListOfUsernames)

    # removes username from list of mode is edit
    if mode == "Edit":
        ListOfUsernames.remove(username)

    match = 0
    # checks list of usernames to see if the wanted username has already been taken
    for user in ListOfUsernames:
        if user == username:
            match += 1
            print(match)

    if match == 0:
        unique = True

    print("match")
    print(match)
    print("unique")
    print(unique)

    if username != "":
        presence = True
    print(presence)

    if len(username) >= 8:
        length = True
    # if all validation has been met then return true
    if presence == True and unique == True and length == True:
        return True
    else:
        return False

# Validation function which checks to ensure postcode is in the format LLNNNLL or LLNNLL.
def CheckPostcode(postcode):
    state1 = False
    state2 = False
    state3 = False

    postcodeCharList = []

    for char in range(len(postcode)):
        postcodeCharList.append(postcode[char])

    postcodeCharList = [char for char in postcodeCharList if char != " "]

    print(postcodeCharList)
    # Checks postcode to see if its in LLNNNLL
    if len(postcodeCharList) == 7:
        for char in range(0, 2):
            print(char)
            print(postcodeCharList[char])
            state1 = postcodeCharList[char].isalpha()
            if state1 == False:
                break
            print(state1)

        print(postcodeCharList)

        if state1 == True:
            for char in range(2, 5):
                print(char)
                print(postcodeCharList[char])
                state2 = postcodeCharList[char].isdigit()
                if state2 == False:
                    break
                print(state2)
        else:
            print("")

        if state2 == True:
            for char in range(5, 7):
                print(char)
                print(postcodeCharList[char])
                state3 = postcodeCharList[char].isalpha()
                if state3 == False:
                    break
                print(state3)
        else:
            print("")

        if state1 == False or state2 == False or state3 == False or len(postcodeCharList) != 7:
            return False
        else:
            return True
    # Checks postcode to see if its in LLNNLL
    elif len(postcodeCharList) == 6:
        for char in range(0, 2):
            print(char)
            print(postcodeCharList[char])
            state1 = postcodeCharList[char].isalpha()
            if state1 == False:
                break
            print(state1)

        print(postcodeCharList)

        if state1 == True:
            for char in range(2, 4):
                print(char)
                print(postcodeCharList[char])
                state2 = postcodeCharList[char].isdigit()
                if state2 == False:
                    break
                print(state2)
        else:
            print("")

        if state2 == True:
            for char in range(4, 6):
                print(char)
                print(postcodeCharList[char])
                state3 = postcodeCharList[char].isalpha()
                if state3 == False:
                    break
                print(state3)
        else:
            print("")

        if state1 == False or state2 == False or state3 == False or len(postcodeCharList) != 6:
            return False
        else:
            return True

# Validation function which checks to ensure email passed is valid.
# Must at least contain "@"
# must contain "." after the "@"

def CheckEmail(email):
    at = 0
    dot = 0

    for i in email:
        if i == "@":
            at += 1
        elif at >= 1:
            if i == ".":
                dot += 1

    if at >= 1 and dot >= 1:
        return True
    else:
        return False

# Creates a list of bands - used to populate list boxes
def getListOfBands():
    global listOfBands

    # Connect to the Database
    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through the data
        cursorListOfBands = db.cursor()
        # SQL Query to select all band names from the Band Account table
    listOfBandsQuery = 'SELECT BandName FROM BandAccount'
    # Execute the query to get a list of all the bands
    cursorListOfBands.execute(listOfBandsQuery)
    # The query will produce a 2 dimensional tuple
    ListOfBands = cursorListOfBands.fetchall()
    listOfBands = []
    # For each index the tuple append to the new list
    for u in ListOfBands:
        for k in u:
            listOfBands.append(k)

    print("listOfBands", listOfBands)

    # getNameOfBandLoggedIn()

    # return listOfBands


getListOfBands()

# gets name of band logged in.
def getNameOfBandLoggedIn():
    with sqlite3.connect("RspbaniDB.db") as db:
        cursorBandname = db.cursor()
    findBandName = 'SELECT bandName FROM BandAccount WHERE Username = ?'
    cursorBandname.execute(findBandName, [BandIDofLoggedIn])
    Bandnameresults = cursorBandname.fetchall()

    if Bandnameresults:

        for i in Bandnameresults:
            i = str(i)
            i = i[2:-3]
            BandnameofLoggedIn = i
            print(i)

# Displays home frame.
def HomeFrame():
    # Creates an instance of the ContentFrame and displays the Home Page to the User once the Button has been clicked
    Home = Frame(contentFrame, width=965, height=600)
    Home.grid(row=0, column=0, sticky="nsew")
    Home.grid_propagate(False)

    lblspacer = Label(Home, text=" ", width=20)
    lblspacer.grid()
    lblTitle = Label(Home, text="", font=("Arvo", 25))
    lblTitle.grid(row=1, column=0, columnspan=4, sticky="n")

    lblspacer1 = Label(Home, text=" ", width=20)
    lblspacer1.grid(row=2)

    lblAboutUs1 = Label(Home, text="The Royal Scottish Pipe Band Association Northern Ireland Branch (RSPBANI) is a "
                                   "governing body to oversee Pipe band competition ",
                        foreground="black", font=("Arvo"))
    lblAboutUs1.grid(row=3, column=0, columnspan=9, sticky="w")

    lblAboutUs2 = Label(Home,
                        text="in Northern Ireland. We represent approximately seventy pipe bands in full membership of the "
                             "association in Northern Ireland and ",
                        foreground="black", font=("Arvo"))
    lblAboutUs2.grid(row=4, column=0, columnspan=9, sticky="w")

    lblAboutUs3 = Label(Home,
                        text="over three thousand individual members.",
                        foreground="black", font=("Arvo"))
    lblAboutUs3.grid(row=5, column=0, columnspan=9, sticky="w")

    # create the canvas, size in pixels
    canvas = Canvas(Home, width=600, height=400)

    # pack the canvas into a frame/form
    canvas.grid(row=6, rowspan=4, column=0, sticky=EW)

    for i in range(2, 12):
        print(i)
        photo = os.path.join('Photos', str(3) + '.png')
        # load the .png image file
        global photoimage
        photoimage = PhotoImage(file=photo)
        # put png image on canvas
        canvas.create_image(0, 0, image=photoimage, anchor=NW)

    lblLatestResults = Label(Home, text="Latest Results", font=("Arvo", 32), fg="#009e0f", width=12)
    lblLatestResults.grid(row=6, column=1, sticky=N)
    LatestR = Text(Home, width=12, height=23, bg="lightgrey")
    LatestR.grid(row=7, rowspan=3, column=1, sticky=NSEW)

    # tries to get the most recent competition result to display it on the home screen
    try:
        with sqlite3.connect("RspbaniDB.db") as db:
            cursor = db.cursor()
        LatestResultQuery = 'SELECT CompetitionID, CompetitionName, CompetitionDate FROM Competitions WHERE Latest = ?'
        cursor.execute(LatestResultQuery, ["True"])
        LatestResult = cursor.fetchall()

        for LR in LatestResult:
            CompID = LR[0]
            CompName = LR[1]
            CompDate = LR[2]

        print(CompID)
        print(CompName)

    except:
        CompID = "No Latest Competition"
        CompName = "No Latest Competition"
        CompDate = "No Latest Competition"

    # Ranked by Position
    AllBandsInCompRankPos = []
    G1BandsInCompRankedPos = []
    G2BandsInCompRankedPos = []
    G3ABandsInCompRankedPos = []
    G3BBandsInCompRankedPos = []
    G4ABandsInCompRankedPos = []
    G4BBandsInCompRankedPos = []

    # Get all the IDs of the bands entered into the latest competition if exists and order by position.
    try:
        with sqlite3.connect("RspbaniDB.db") as db:
            cursor = db.cursor()
        LatestResultQuery = 'SELECT BandID FROM Results WHERE CompetitionID = ? ORDER BY Position'
        cursor.execute(LatestResultQuery, [CompID])
        AllBestBandResults = cursor.fetchall()

        print("AllBestBandResults")
        print(AllBestBandResults)

        if AllBestBandResults:
            for band in AllBestBandResults:
                Temp = []
                Temp.append(band[0])
                print("Temp")
                print(Temp)

                AllBandsInCompRankPos.append(Temp)
                print("AllBandsInCompRankPos")
                print(AllBandsInCompRankPos)
            print("AllBandsInCompRankPos")
            print(AllBandsInCompRankPos)

        # for all bands get their grade
        for I in AllBandsInCompRankPos:
            print("I[0]")
            print(I[0])
            with sqlite3.connect("RspbaniDB.db") as db:
                cursor = db.cursor()
            findGrade = 'SELECT BandGrade, BandName FROM BandAccount WHERE BandID = ?'
            cursor.execute(findGrade, [I[0]])
            BandGrade = cursor.fetchall()

            # sort the bands by grade. Bands will still remain in the correct ordering (by position)
            if BandGrade:
                for G in BandGrade:
                    if G[0] == "1":
                        I.append(G[1])
                        G1BandsInCompRankedPos.append(I)
                    elif G[0] == "2":
                        I.append(G[1])
                        G2BandsInCompRankedPos.append(I)
                    elif G[0] == "3A":
                        I.append(G[1])
                        G3ABandsInCompRankedPos.append(I)
                    elif G[0] == "3B":
                        I.append(G[1])
                        G3BBandsInCompRankedPos.append(I)
                    elif G[0] == "4A":
                        I.append(G[1])
                        G4ABandsInCompRankedPos.append(I)
                    elif G[0] == "4B":
                        I.append(G[1])
                        G4BBandsInCompRankedPos.append(I)

        # for each grade get the winner of each grade
        # if no band entered then display no band entered instead
        try:
            G1Winner = G1BandsInCompRankedPos[0][1]
        except:
            G1Winner = "No Band Entered"
        print("Winners")
        print(G1Winner)

        try:
            G2Winner = G2BandsInCompRankedPos[0][1]
        except:
            G2Winner = "No Band Entered"
        print(G2Winner)

        try:
            G3AWinner = G3ABandsInCompRankedPos[0][1]
        except:
            G3AWinner = "No Band Entered"
        print(G3AWinner)

        try:
            G3BWinner = G3BBandsInCompRankedPos[0][1]
        except:
            G3BWinner = "No Band Entered"
        print(G3BWinner)

        try:
            G4AWinner = G4ABandsInCompRankedPos[0][1]
        except:
            G4AWinner = "No Band Entered"
        print(G4AWinner)

        try:
            G4BWinner = G4BBandsInCompRankedPos[0][1]
        except:
            G4BWinner = "No Band Entered"
        print(G4BWinner)

    except:
        # if no competition is latest then set winners to no latest competition
        G1Winner = "No Latest Competition"
        G2Winner = "No Latest Competition"
        G3AWinner = "No Latest Competition"
        G3BWinner = "No Latest Competition"
        G4AWinner = "No Latest Competition"
        G4BWinner = "No Latest Competition"

    # populate text box with the winning bands
    LatestR.config(font=("Arvo", 12))
    LatestR.insert(END, "Competition Name: \n")
    LatestR.insert(END, CompName + "\n\n")
    LatestR.insert(END, "Competition Date: \n" + CompDate + "\n\n")
    LatestR.insert(END, "Grade 1: \n" + " 1st- " + str(G1Winner) + "\n\n")
    LatestR.insert(END, "Grade 2: \n" + " 1st- " + str(G2Winner) + "\n\n")
    LatestR.insert(END, "Grade 3A: \n" + " 1st- " + str(G3AWinner) + "\n\n")
    LatestR.insert(END, "Grade 3B: \n" + " 1st- " + str(G3BWinner) + "\n\n")
    LatestR.insert(END, "Grade 4A: \n" + " 1st- " + str(G4AWinner) + "\n\n")
    LatestR.insert(END, "Grade 4B: \n" + " 1st- " + str(G4BWinner))

    LatestR.configure(state='disabled')

# Display bands frame to search for bands
def BandsFrame(photoSearchByGrade, photoSearchByBand):
    def DisplayByName(BandNameSearch):

        NameOfBand = BandNameSearch.get()
        print(NameOfBand)

        NameOfBand = NameOfBand.title()
        print(NameOfBand)

        # Creates an instance of the ContentFrame and displays the Bands Page to the User once the Button has been clicked
        Bands = Frame(contentFrame, width=965, height=600)
        Bands.grid(row=0, column=0, sticky="nsew")

        SearchGrade = StringVar()
        # Create Title for the Bands Page
        lblSpacer = Label(Bands, text="Bands ", height=2, font=("Arvo", 32))
        lblSpacer.grid(row=0, column=1, columnspan=6, sticky="ew")

        lblSpacer = Label(Bands, text="", width=5)
        lblSpacer.grid(row=0, column=0, sticky="e")

        lblSpacer = Label(Bands, text="", height=1)
        lblSpacer.grid(row=5, column=0, sticky="e")

        # Create all the filtering options to the User if they want to search by Grade and place them on the Frame
        Radiobutton(Bands, text="All Grades", padx=5, variable=SearchGrade, value="All", font=("Arvo")).grid(row=4,
                                                                                                             column=1,
                                                                                                             sticky="w",
                                                                                                             padx=20)
        Radiobutton(Bands, text="Grade 1", padx=5, variable=SearchGrade, value="1", font=("Arvo")).grid(row=4, column=2,
                                                                                                        sticky="w",
                                                                                                        padx=20)
        Radiobutton(Bands, text="Grade 2", padx=5, variable=SearchGrade, value="2", font=("Arvo")).grid(row=4, column=3,
                                                                                                        sticky="w",
                                                                                                        padx=20)
        Radiobutton(Bands, text="Grade 3A", padx=5, variable=SearchGrade, value="3A", font=("Arvo")).grid(row=4,
                                                                                                          column=4,
                                                                                                          sticky="w",
                                                                                                          padx=20)
        Radiobutton(Bands, text="Grade 3B", padx=5, variable=SearchGrade, value="3B", font=("Arvo")).grid(row=6,
                                                                                                          column=1,
                                                                                                          sticky="w",
                                                                                                          padx=20)
        Radiobutton(Bands, text="Grade 4A", padx=5, variable=SearchGrade, value="4A", font=("Arvo")).grid(row=6,
                                                                                                          column=2,
                                                                                                          sticky="w",
                                                                                                          padx=20)
        Radiobutton(Bands, text="Grade 4B", padx=5, variable=SearchGrade, value="4B", font=("Arvo")).grid(row=6,
                                                                                                          column=3,
                                                                                                          sticky="w",
                                                                                                          padx=20)

        lblSearchByBand = Label(Bands, text="Band Name: ", font=("Arvo"))
        lblSearchByBand.grid(row=4, column=5, sticky="e")

        BandNameSearch = StringVar()
        EtySearchByBand = Entry(Bands, textvar=BandNameSearch)
        EtySearchByBand.grid(row=4, column=6, sticky="w", padx=2)

        BandGradeBP = SearchGrade.get()
        print(BandGradeBP)

        SearchGradeBtn = Button(Bands, image=photoSearchByGrade, command=lambda: SearchByGrade(SearchGrade))
        SearchGradeBtn.grid(row=6, column=4, padx=0)

        SearchBandBtn = Button(Bands, image=photoSearchByBand, command=lambda: DisplayByName(BandNameSearch))
        SearchBandBtn.grid(row=6, column=6)

        lblSpacer1 = Label(Bands, height=2)
        lblSpacer1.grid(row=7)

        Band = []

        with sqlite3.connect("RspbaniDB.db") as db:
            cursorBandPage = db.cursor()
        BandPageInfoGetQuery = 'SELECT BandName, BandBranch, BandGrade, HallName, BandAddressLine1, BandAddressLine2, ' \
                               'BandCounty, BandPostcode, PracticeTime, Tartan FROM BandAccount WHERE BandName = ? ORDER BY BandName'
        cursorBandPage.execute(BandPageInfoGetQuery, [NameOfBand, ])
        BandPageInfoGet = cursorBandPage.fetchall()

        print(BandPageInfoGet)

        if BandPageInfoGet == []:
            messagebox.showinfo(message="No bands found matching your results")
            BandsFrame(photoSearchByGrade, photoSearchByBand)
        else:


            for bandRecord in BandPageInfoGet:

                for eachField in bandRecord:
                    Band.append(eachField)
            print(Band)


            def getNumberOfRecords():
                global numRecords

                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorNumRecords = db.cursor()
                NumRecordsQuery = 'SELECT COUNT(*) FROM BandAccount WHERE BandName = ?'
                cursorNumRecords.execute(NumRecordsQuery, [NameOfBand, ])
                NumRecords2D = cursorNumRecords.fetchall()
                NumRecords1D = NumRecords2D[0]
                numRecords = NumRecords1D[0]

            getNumberOfRecords()

            def getAllBandFields():
                global numRecordsBands, BandTableBandName, ListBandName, BandTableGrade, ListGrade, BandTableSection, \
                    ListSection, BandTableAdress, ListAddress, BandTablePracticeTime, ListPracticeTime, BandTableTartan, \
                    ListTartan

                numRecordsBands = numRecords
                # All BandNames
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfBands = db.cursor()
                listOfBandsQuery = 'SELECT BandName FROM BandAccount WHERE BandName = ? ORDER BY BandName'
                cursorListOfBands.execute(listOfBandsQuery, [NameOfBand, ])
                listOfBands2D = cursorListOfBands.fetchall()

                ListBandName = []

                for numRecordsBands in range(len(listOfBands2D)):
                    var = listOfBands2D[numRecordsBands]

                    ListBandName.append(var[0])

                    numRecordsBands = + 1

                ListBandName.reverse()
                # All BandGrades
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfGrades = db.cursor()
                listOfGradesQuery = 'SELECT BandGrade FROM BandAccount WHERE BandName = ? ORDER BY BandName'
                cursorListOfGrades.execute(listOfGradesQuery, [NameOfBand, ])
                listOfGrades2D = cursorListOfGrades.fetchall()

                ListGrade = []

                for numRecordsGrades in range(len(listOfGrades2D)):
                    var = listOfGrades2D[numRecordsGrades]

                    ListGrade.append(var[0])

                    numRecordsGrade = + 1
                ListGrade.reverse()

                # All BandSections
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfSection = db.cursor()
                listOfSectionQuery = 'SELECT BandBranch FROM BandAccount WHERE BandName = ? ORDER BY BandName'
                cursorListOfSection.execute(listOfSectionQuery, [NameOfBand, ])
                listOfSection2D = cursorListOfSection.fetchall()

                ListSection = []

                for numRecordsSection in range(len(listOfSection2D)):
                    var = listOfSection2D[numRecordsSection]

                    ListSection.append(var[0])

                    numRecordsSection = + 1
                ListSection.reverse()

                # All BandAddress
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfAddress = db.cursor()
                listOfAddressQuery = 'SELECT HallName, BandAddressLine1, BandPostcode FROM BandAccount WHERE BandName = ? ORDER BY BandName'
                cursorListOfAddress.execute(listOfAddressQuery, [NameOfBand, ])
                listOfAddress2D = cursorListOfAddress.fetchall()
                print(listOfAddress2D)

                ListAddress = []

                for numRecordsBands in range(len(listOfAddress2D)):
                    var = listOfAddress2D[numRecordsBands]
                    print(var)

                    ListAddress.append(var[0] + ", " + var[1] + ", " + var[2])

                    numRecordsBands = + 1
                ListAddress.reverse()

                # All BandPracticeTime
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfPracticeTimes = db.cursor()
                listOfPracticeTimesQuery = 'SELECT PracticeTime FROM BandAccount WHERE BandName = ? ORDER BY BandName'
                cursorListOfPracticeTimes.execute(listOfPracticeTimesQuery, [NameOfBand, ])
                listOfPracticeTimes2D = cursorListOfPracticeTimes.fetchall()

                ListPracticeTime = []

                for numRecordsBands in range(len(listOfPracticeTimes2D)):
                    var = listOfPracticeTimes2D[numRecordsBands]

                    ListPracticeTime.append(var[0])

                    numRecordsBands = + 1
                ListPracticeTime.reverse()

                # All BandPracticeTime
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfTartan = db.cursor()
                listOfTartanQuery = 'SELECT Tartan FROM BandAccount WHERE BandName = ? ORDER BY BandName'
                cursorListOfTartan.execute(listOfTartanQuery, [NameOfBand, ])
                listOfTartan2D = cursorListOfTartan.fetchall()

                ListTartan = []

                for numRecordsBands in range(len(listOfTartan2D)):
                    var = listOfTartan2D[numRecordsBands]

                    ListTartan.append(var[0])

                    numRecordsBands = + 1
                ListTartan.reverse()

            getAllBandFields()

        TableFrame = Frame(Bands, height=400, width=600)
        TableFrame.grid(row=8, column=1, columnspan=7, sticky="ew")

        lblbandName = Label(TableFrame, text="Band", font=("Arvo"))
        lblbandName.grid(row=0, column=0, sticky="ew")

        BandListBox = Listbox(TableFrame, height=20, width=18)
        BandListBox.grid(row=1, column=0)

        lblGrade = Label(TableFrame, text="Grade", font=("Arvo"))
        lblGrade.grid(row=0, column=1, sticky="ew")

        GradeListBox = Listbox(TableFrame, height=20, width=7)
        GradeListBox.grid(row=1, column=1)

        lblBranch = Label(TableFrame, text="Branch", font=("Arvo"))
        lblBranch.grid(row=0, column=2, sticky="ew")

        SectionListBox = Listbox(TableFrame, height=20, width=10)
        SectionListBox.grid(row=1, column=2)

        lblAddress = Label(TableFrame, text="Address", font=("Arvo"))
        lblAddress.grid(row=0, column=3, sticky="ew")

        AddressListBox = Listbox(TableFrame, height=20, width=30)
        AddressListBox.grid(row=1, column=3)

        lblPracticeTimes = Label(TableFrame, text="Practice Time(s)", font=("Arvo"))
        lblPracticeTimes.grid(row=0, column=4, sticky="ew")

        PracticeTimeListBox = Listbox(TableFrame, height=20, width=20)
        PracticeTimeListBox.grid(row=1, column=4)

        lblTartan = Label(TableFrame, text="Tartan", font=("Arvo"))
        lblTartan.grid(row=0, column=5, sticky="ew")

        TartanListBox = Listbox(TableFrame, height=20, width=10)
        TartanListBox.grid(row=1, column=5)

        def yview(*args):
            BandListBox.yview(*args)
            GradeListBox.yview(*args)
            SectionListBox.yview(*args)
            AddressListBox.yview(*args)
            PracticeTimeListBox.yview(*args)
            TartanListBox.yview(*args)

        Scrollbary = Scrollbar(TableFrame, orient=VERTICAL, command=yview, )
        Scrollbary.grid(row=1, column=6)

        BandListBox.config(yscrollcommand=Scrollbary.set)
        GradeListBox.config(yscrollcommand=Scrollbary.set)
        SectionListBox.config(yscrollcommand=Scrollbary.set)
        AddressListBox.config(yscrollcommand=Scrollbary.set)
        PracticeTimeListBox.config(yscrollcommand=Scrollbary.set)
        TartanListBox.config(yscrollcommand=Scrollbary.set)

        # Populating Band name listbox
        for numRecordsBands in range(len(ListBandName)):
            BandListBox.insert(0, ListBandName[numRecordsBands])
            numRecordsBands = + 1

        # Populating Grade listbox
        for numRecordBands in range(len(ListGrade)):
            GradeListBox.insert(0, ListGrade[numRecordBands])
            numRecordsBands = + 1

        # Populating Section listbox
        for numRecordBands in range(len(ListSection)):
            SectionListBox.insert(0, ListSection[numRecordBands])
            numRecordsBands = + 1

        # Populating Address listbox
        for numRecordBands in range(len(ListAddress)):
            AddressListBox.insert(0, ListAddress[numRecordBands])
            numRecordsBands = + 1

        # Populating PracticeTime listbox
        for numRecordBands in range(len(ListPracticeTime)):
            PracticeTimeListBox.insert(0, ListPracticeTime[numRecordBands])
            numRecordsBands = + 1

        # Populating Tartan Name listbox
        for numRecordBands in range(len(ListTartan)):
            TartanListBox.insert(0, ListTartan[numRecordBands])
            numRecordsBands = + 1

    def DisplayGrade(Grade):
        # Creates an instance of the ContentFrame and displays the Bands Page to the User once the Button has been clicked
        Bands = Frame(contentFrame, width=965, height=600)
        Bands.grid(row=0, column=0, sticky="nsew")

        SearchGrade = StringVar()
        # Create Title for the Bands Page
        lblSpacer = Label(Bands, text="Bands ", height=2, font=("Arvo", 32))
        lblSpacer.grid(row=0, column=1, columnspan=6, sticky="ew")

        lblSpacer = Label(Bands, text="", width=5)
        lblSpacer.grid(row=0, column=0, sticky="e")

        lblSpacer = Label(Bands, text="", height=1)
        lblSpacer.grid(row=5, column=0, sticky="e")

        # Create all the filtering options to the User if they want to search by Grade and place them on the Frame
        Radiobutton(Bands, text="All Grades", padx=5, variable=SearchGrade, value="All", font=("Arvo")).grid(row=4,
                                                                                                             column=1,
                                                                                                             sticky="w",
                                                                                                             padx=20)
        Radiobutton(Bands, text="Grade 1", padx=5, variable=SearchGrade, value="1", font=("Arvo")).grid(row=4, column=2,
                                                                                                        sticky="w",
                                                                                                        padx=20)
        Radiobutton(Bands, text="Grade 2", padx=5, variable=SearchGrade, value="2", font=("Arvo")).grid(row=4, column=3,
                                                                                                        sticky="w",
                                                                                                        padx=20)
        Radiobutton(Bands, text="Grade 3A", padx=5, variable=SearchGrade, value="3A", font=("Arvo")).grid(row=4,
                                                                                                          column=4,
                                                                                                          sticky="w",
                                                                                                          padx=20)
        Radiobutton(Bands, text="Grade 3B", padx=5, variable=SearchGrade, value="3B", font=("Arvo")).grid(row=6,
                                                                                                          column=1,
                                                                                                          sticky="w",
                                                                                                          padx=20)
        Radiobutton(Bands, text="Grade 4A", padx=5, variable=SearchGrade, value="4A", font=("Arvo")).grid(row=6,
                                                                                                          column=2,
                                                                                                          sticky="w",
                                                                                                          padx=20)
        Radiobutton(Bands, text="Grade 4B", padx=5, variable=SearchGrade, value="4B", font=("Arvo")).grid(row=6,
                                                                                                          column=3,
                                                                                                          sticky="w",
                                                                                                          padx=20)

        lblSearchByBand = Label(Bands, text="Band Name: ", font=("Arvo"))
        lblSearchByBand.grid(row=4, column=5, sticky="e")

        EtySearchByBand = Entry(Bands)
        ''', textvar=BandNameSearch'''
        EtySearchByBand.grid(row=4, column=6, sticky="w", padx=2)

        BandGradeBP = SearchGrade.get()
        print(BandGradeBP)

        SearchGradeBtn = Button(Bands, image=photoSearchByGrade, command=lambda: SearchByGrade(SearchGrade))
        ''',command=BandPageGSearch'''
        SearchGradeBtn.grid(row=6, column=4, padx=0)

        SearchBandBtn = Button(Bands, image=photoSearchByBand, command=lambda: DisplayByName(BandNameSearch))
        SearchBandBtn.grid(row=6, column=6)

        lblSpacer1 = Label(Bands, height=2)
        lblSpacer1.grid(row=7)

        Band = []

        print("Grade")
        print(Grade)

        with sqlite3.connect("RspbaniDB.db") as db:
            cursorBandPage = db.cursor()
        BandPageInfoGetQuery = 'SELECT BandName, BandBranch, BandGrade, HallName, BandAddressLine1, BandAddressLine2, ' \
                               'BandCounty, BandPostcode, PracticeTime, Tartan FROM BandAccount WHERE BandGrade = ? ORDER BY BandName'
        cursorBandPage.execute(BandPageInfoGetQuery, [Grade, ])
        BandPageInfoGet = cursorBandPage.fetchall()

        for bandRecord in BandPageInfoGet:

            for eachField in bandRecord:
                Band.append(eachField)
        print(Band)

        # frame = Frame()
        # frame.grid(row=1, columnspan=2, padx=2, pady=2, sticky="nsew")

        def getNumberOfRecords():
            global numRecords

            with sqlite3.connect("RspbaniDB.db") as db:
                cursorNumRecords = db.cursor()
            NumRecordsQuery = 'SELECT COUNT(*) FROM BandAccount WHERE BandGrade = ?'
            cursorNumRecords.execute(NumRecordsQuery, [Grade, ])
            NumRecords2D = cursorNumRecords.fetchall()
            NumRecords1D = NumRecords2D[0]
            numRecords = NumRecords1D[0]

        getNumberOfRecords()

        def getAllBandFields():
            global numRecordsBands, BandTableBandName, ListBandName, BandTableGrade, ListGrade, BandTableSection, \
                ListSection, BandTableAdress, ListAddress, BandTablePracticeTime, ListPracticeTime, BandTableTartan, \
                ListTartan

            numRecordsBands = numRecords
            # All BandNames
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorListOfBands = db.cursor()
            listOfBandsQuery = 'SELECT BandName FROM BandAccount WHERE BandGrade = ? ORDER BY BandName'
            cursorListOfBands.execute(listOfBandsQuery, [Grade, ])
            listOfBands2D = cursorListOfBands.fetchall()

            ListBandName = []

            if len(listOfBands2D) == 0:
                messagebox.showinfo(message="No bands are registered in this grade")
                BandsFrame(photoSearchByGrade, photoSearchByBand)
            else:

                for numRecordsBands in range(len(listOfBands2D)):
                    var = listOfBands2D[numRecordsBands]

                    ListBandName.append(var[0])

                    numRecordsBands = + 1

                ListBandName.reverse()
                # All BandGrades
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfGrades = db.cursor()
                listOfGradesQuery = 'SELECT BandGrade FROM BandAccount WHERE BandGrade = ? ORDER BY BandName'
                cursorListOfGrades.execute(listOfGradesQuery, [Grade, ])
                listOfGrades2D = cursorListOfGrades.fetchall()

                ListGrade = []

                for numRecordsGrades in range(len(listOfGrades2D)):
                    var = listOfGrades2D[numRecordsGrades]

                    ListGrade.append(var[0])

                    numRecordsGrade = + 1
                ListGrade.reverse()

                # All BandSections
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfSection = db.cursor()
                listOfSectionQuery = 'SELECT BandBranch FROM BandAccount WHERE BandGrade = ? ORDER BY BandName'
                cursorListOfSection.execute(listOfSectionQuery, [Grade, ])
                listOfSection2D = cursorListOfSection.fetchall()

                ListSection = []

                for numRecordsSection in range(len(listOfSection2D)):
                    var = listOfSection2D[numRecordsSection]

                    ListSection.append(var[0])

                    numRecordsSection = + 1
                ListSection.reverse()

                # All BandAddress
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfAddress = db.cursor()
                listOfAddressQuery = 'SELECT HallName, BandAddressLine1, BandPostcode FROM BandAccount WHERE BandGrade = ? ORDER BY BandName'
                cursorListOfAddress.execute(listOfAddressQuery, [Grade, ])
                listOfAddress2D = cursorListOfAddress.fetchall()
                print(listOfAddress2D)

                ListAddress = []

                for numRecordsBands in range(len(listOfAddress2D)):
                    var = listOfAddress2D[numRecordsBands]
                    print(var)

                    ListAddress.append(var[0] + ", " + var[1] + ", " + var[2])

                    numRecordsBands = + 1
                ListAddress.reverse()

                # All BandPracticeTime
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfPracticeTimes = db.cursor()
                listOfPracticeTimesQuery = 'SELECT PracticeTime FROM BandAccount WHERE BandGrade = ? ORDER BY BandName'
                cursorListOfPracticeTimes.execute(listOfPracticeTimesQuery, [Grade, ])
                listOfPracticeTimes2D = cursorListOfPracticeTimes.fetchall()

                ListPracticeTime = []

                for numRecordsBands in range(len(listOfPracticeTimes2D)):
                    var = listOfPracticeTimes2D[numRecordsBands]

                    ListPracticeTime.append(var[0])

                    numRecordsBands = + 1
                ListPracticeTime.reverse()

                # All BandPracticeTime
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorListOfTartan = db.cursor()
                listOfTartanQuery = 'SELECT Tartan FROM BandAccount WHERE BandGrade = ? ORDER BY BandName'
                cursorListOfTartan.execute(listOfTartanQuery, [Grade, ])
                listOfTartan2D = cursorListOfTartan.fetchall()

                ListTartan = []

                for numRecordsBands in range(len(listOfTartan2D)):
                    var = listOfTartan2D[numRecordsBands]

                    ListTartan.append(var[0])

                    numRecordsBands = + 1
                ListTartan.reverse()

        getAllBandFields()

        TableFrame = Frame(Bands, height=400, width=600)
        TableFrame.grid(row=8, column=1, columnspan=7, sticky="ew")

        lblbandName = Label(TableFrame, text="Band", font=("Arvo"))
        lblbandName.grid(row=0, column=0, sticky="ew")

        BandListBox = Listbox(TableFrame, height=20, width=18)
        BandListBox.grid(row=1, column=0)

        lblGrade = Label(TableFrame, text="Grade", font=("Arvo"))
        lblGrade.grid(row=0, column=1, sticky="ew")

        GradeListBox = Listbox(TableFrame, height=20, width=7)
        GradeListBox.grid(row=1, column=1)

        lblBranch = Label(TableFrame, text="Branch", font=("Arvo"))
        lblBranch.grid(row=0, column=2, sticky="ew")

        SectionListBox = Listbox(TableFrame, height=20, width=10)
        SectionListBox.grid(row=1, column=2)

        lblAddress = Label(TableFrame, text="Address", font=("Arvo"))
        lblAddress.grid(row=0, column=3, sticky="ew")

        AddressListBox = Listbox(TableFrame, height=20, width=30)
        AddressListBox.grid(row=1, column=3)

        lblPracticeTimes = Label(TableFrame, text="Practice Time(s)", font=("Arvo"))
        lblPracticeTimes.grid(row=0, column=4, sticky="ew")

        PracticeTimeListBox = Listbox(TableFrame, height=20, width=20)
        PracticeTimeListBox.grid(row=1, column=4)

        lblTartan = Label(TableFrame, text="Tartan", font=("Arvo"))
        lblTartan.grid(row=0, column=5, sticky="ew")

        TartanListBox = Listbox(TableFrame, height=20, width=10)
        TartanListBox.grid(row=1, column=5)

        def yview(*args):
            BandListBox.yview(*args)
            GradeListBox.yview(*args)
            SectionListBox.yview(*args)
            AddressListBox.yview(*args)
            PracticeTimeListBox.yview(*args)
            TartanListBox.yview(*args)

        Scrollbary = Scrollbar(TableFrame, orient=VERTICAL, command=yview, )
        Scrollbary.grid(row=1, column=6)

        BandListBox.config(yscrollcommand=Scrollbary.set)
        GradeListBox.config(yscrollcommand=Scrollbary.set)
        SectionListBox.config(yscrollcommand=Scrollbary.set)
        AddressListBox.config(yscrollcommand=Scrollbary.set)
        PracticeTimeListBox.config(yscrollcommand=Scrollbary.set)
        TartanListBox.config(yscrollcommand=Scrollbary.set)

        # Populating Band name listbox
        for numRecordsBands in range(len(ListBandName)):
            BandListBox.insert(0, ListBandName[numRecordsBands])
            numRecordsBands = + 1

        # Populating Grade listbox
        for numRecordBands in range(len(ListGrade)):
            GradeListBox.insert(0, ListGrade[numRecordBands])
            numRecordsBands = + 1

        # Populating Section listbox
        for numRecordBands in range(len(ListSection)):
            SectionListBox.insert(0, ListSection[numRecordBands])
            numRecordsBands = + 1

        # Populating Address listbox
        for numRecordBands in range(len(ListAddress)):
            AddressListBox.insert(0, ListAddress[numRecordBands])
            numRecordsBands = + 1

        # Populating PracticeTime listbox
        for numRecordBands in range(len(ListPracticeTime)):
            PracticeTimeListBox.insert(0, ListPracticeTime[numRecordBands])
            numRecordsBands = + 1

        # Populating Tartan Name listbox
        for numRecordBands in range(len(ListTartan)):
            TartanListBox.insert(0, ListTartan[numRecordBands])
            numRecordsBands = + 1

    def SearchByGrade(SearchGrade):
        Grade = SearchGrade.get()
        if Grade == "":
            messagebox.showinfo(message="No grade selected to filter by")
        elif Grade == "All":
            BandsFrame(photoSearchByGrade, photoSearchByBand)
        else:
            DisplayGrade(Grade)

    # Creates an instance of the ContentFrame and displays the Bands Page to the User once the Button has been clicked
    Bands = Frame(contentFrame, width=965, height=600)
    Bands.grid(row=0, column=0, sticky="nsew")

    SearchGrade = StringVar()
    # Create Title for the Bands Page
    lblSpacer = Label(Bands, text="Bands ", height=2, font=("Arvo", 32))
    lblSpacer.grid(row=0, column=1, columnspan=6, sticky="ew")

    lblSpacer = Label(Bands, text="", width=5)
    lblSpacer.grid(row=0, column=0, sticky="e")

    lblSpacer = Label(Bands, text="", height=1)
    lblSpacer.grid(row=5, column=0, sticky="e")

    # Create all the filtering options to the User if they want to search by Grade and place them on the Frame
    Radiobutton(Bands, text="All Grades", padx=5, variable=SearchGrade, value="All", font=("Arvo")).grid(row=4,
                                                                                                         column=1,
                                                                                                         sticky="w",
                                                                                                         padx=20)
    Radiobutton(Bands, text="Grade 1", padx=5, variable=SearchGrade, value="1", font=("Arvo")).grid(row=4, column=2,
                                                                                                    sticky="w", padx=20)
    Radiobutton(Bands, text="Grade 2", padx=5, variable=SearchGrade, value="2", font=("Arvo")).grid(row=4, column=3,
                                                                                                    sticky="w", padx=20)
    Radiobutton(Bands, text="Grade 3A", padx=5, variable=SearchGrade, value="3A", font=("Arvo")).grid(row=4, column=4,
                                                                                                      sticky="w",
                                                                                                      padx=20)
    Radiobutton(Bands, text="Grade 3B", padx=5, variable=SearchGrade, value="3B", font=("Arvo")).grid(row=6, column=1,
                                                                                                      sticky="w",
                                                                                                      padx=20)
    Radiobutton(Bands, text="Grade 4A", padx=5, variable=SearchGrade, value="4A", font=("Arvo")).grid(row=6, column=2,
                                                                                                      sticky="w",
                                                                                                      padx=20)
    Radiobutton(Bands, text="Grade 4B", padx=5, variable=SearchGrade, value="4B", font=("Arvo")).grid(row=6, column=3,
                                                                                                      sticky="w",
                                                                                                      padx=20)

    lblSearchByBand = Label(Bands, text="Band Name: ", font=("Arvo"))
    lblSearchByBand.grid(row=4, column=5, sticky="e")

    BandNameSearch = StringVar()
    EtySearchByBand = Entry(Bands, textvar=BandNameSearch)
    EtySearchByBand.grid(row=4, column=6, sticky="w", padx=2)

    BandGradeBP = SearchGrade.get()
    print(BandGradeBP)

    SearchGradeBtn = Button(Bands, image=photoSearchByGrade, command=lambda: SearchByGrade(SearchGrade))
    SearchGradeBtn.grid(row=6, column=4, padx=0)

    SearchBandBtn = Button(Bands, image=photoSearchByBand, command=lambda: DisplayByName(BandNameSearch))
    SearchBandBtn.grid(row=6, column=6)

    lblSpacer1 = Label(Bands, height=2)
    lblSpacer1.grid(row=7)

    Band = []

    with sqlite3.connect("RspbaniDB.db") as db:
        cursorBandPage = db.cursor()
    BandPageInfoGetQuery = 'SELECT BandName, BandBranch, BandGrade, HallName, BandAddressLine1, BandAddressLine2, ' \
                           'BandCounty, BandPostcode, PracticeTime, Tartan FROM BandAccount ORDER BY BandName'
    cursorBandPage.execute(BandPageInfoGetQuery)
    BandPageInfoGet = cursorBandPage.fetchall()

    for bandRecord in BandPageInfoGet:

        for eachField in bandRecord:
            Band.append(eachField)
    print(Band)

    # frame = Frame()
    # frame.grid(row=1, columnspan=2, padx=2, pady=2, sticky="nsew")

    def getNumberOfRecords():
        global numRecords

        with sqlite3.connect("RspbaniDB.db") as db:
            cursorNumRecords = db.cursor()
        NumRecordsQuery = 'SELECT COUNT(*) FROM BandAccount '
        cursorNumRecords.execute(NumRecordsQuery)
        NumRecords2D = cursorNumRecords.fetchall()
        NumRecords1D = NumRecords2D[0]
        numRecords = NumRecords1D[0]

    getNumberOfRecords()

    def getAllBandFields():
        global numRecordsBands, BandTableBandName, ListBandName, BandTableGrade, ListGrade, BandTableSection, \
            ListSection, BandTableAdress, ListAddress, BandTablePracticeTime, ListPracticeTime, BandTableTartan, \
            ListTartan

        numRecordsBands = numRecords
        # All BandNames
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorListOfBands = db.cursor()
        listOfBandsQuery = 'SELECT BandName FROM BandAccount ORDER BY BandName'
        cursorListOfBands.execute(listOfBandsQuery)
        listOfBands2D = cursorListOfBands.fetchall()

        ListBandName = []

        for numRecordsBands in range(len(listOfBands2D)):
            var = listOfBands2D[numRecordsBands]

            ListBandName.append(var[0])

            numRecordsBands = + 1

        ListBandName.reverse()
        # All BandGrades
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorListOfGrades = db.cursor()
        listOfGradesQuery = 'SELECT BandGrade FROM BandAccount ORDER BY BandName'
        cursorListOfGrades.execute(listOfGradesQuery)
        listOfGrades2D = cursorListOfGrades.fetchall()

        ListGrade = []

        for numRecordsGrades in range(len(listOfGrades2D)):
            var = listOfGrades2D[numRecordsGrades]

            ListGrade.append(var[0])

            numRecordsGrade = + 1
        ListGrade.reverse()

        # All BandSections
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorListOfSection = db.cursor()
        listOfSectionQuery = 'SELECT BandBranch FROM BandAccount ORDER BY BandName'
        cursorListOfSection.execute(listOfSectionQuery)
        listOfSection2D = cursorListOfSection.fetchall()

        ListSection = []

        for numRecordsSection in range(len(listOfSection2D)):
            var = listOfSection2D[numRecordsSection]

            ListSection.append(var[0])

            numRecordsSection = + 1
        ListSection.reverse()

        # All BandAddress
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorListOfAddress = db.cursor()
        listOfAddressQuery = 'SELECT HallName, BandAddressLine1, BandPostcode FROM BandAccount ORDER BY BandName'
        cursorListOfAddress.execute(listOfAddressQuery)
        listOfAddress2D = cursorListOfAddress.fetchall()
        print(listOfAddress2D)

        ListAddress = []

        for numRecordsBands in range(len(listOfAddress2D)):
            var = listOfAddress2D[numRecordsBands]
            print(var)

            ListAddress.append(var[0] + ", " + var[1] + ", " + var[2])

            numRecordsBands = + 1
        ListAddress.reverse()

        # All BandPracticeTime
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorListOfPracticeTimes = db.cursor()
        listOfPracticeTimesQuery = 'SELECT PracticeTime FROM BandAccount ORDER BY BandName'
        cursorListOfPracticeTimes.execute(listOfPracticeTimesQuery)
        listOfPracticeTimes2D = cursorListOfPracticeTimes.fetchall()

        ListPracticeTime = []

        for numRecordsBands in range(len(listOfPracticeTimes2D)):
            var = listOfPracticeTimes2D[numRecordsBands]

            ListPracticeTime.append(var[0])

            numRecordsBands = + 1
        ListPracticeTime.reverse()

        # All BandPracticeTime
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorListOfTartan = db.cursor()
        listOfTartanQuery = 'SELECT Tartan FROM BandAccount ORDER BY BandName'
        cursorListOfTartan.execute(listOfTartanQuery)
        listOfTartan2D = cursorListOfTartan.fetchall()

        ListTartan = []

        for numRecordsBands in range(len(listOfTartan2D)):
            var = listOfTartan2D[numRecordsBands]

            ListTartan.append(var[0])

            numRecordsBands = + 1
        ListTartan.reverse()

    getAllBandFields()

    TableFrame = Frame(Bands, height=400, width=600)
    TableFrame.grid(row=8, column=1, columnspan=7, sticky="ew")

    lblbandName = Label(TableFrame, text="Band", font=("Arvo"))
    lblbandName.grid(row=0, column=0, sticky="ew")

    BandListBox = Listbox(TableFrame, height=20, width=18)
    BandListBox.grid(row=1, column=0)

    lblGrade = Label(TableFrame, text="Grade", font=("Arvo"))
    lblGrade.grid(row=0, column=1, sticky="ew")

    GradeListBox = Listbox(TableFrame, height=20, width=7)
    GradeListBox.grid(row=1, column=1)

    lblBranch = Label(TableFrame, text="Branch", font=("Arvo"))
    lblBranch.grid(row=0, column=2, sticky="ew")

    SectionListBox = Listbox(TableFrame, height=20, width=10)
    SectionListBox.grid(row=1, column=2)

    lblAddress = Label(TableFrame, text="Address", font=("Arvo"))
    lblAddress.grid(row=0, column=3, sticky="ew")

    AddressListBox = Listbox(TableFrame, height=20, width=30)
    AddressListBox.grid(row=1, column=3)

    lblPracticeTimes = Label(TableFrame, text="Practice Time(s)", font=("Arvo"))
    lblPracticeTimes.grid(row=0, column=4, sticky="ew")

    PracticeTimeListBox = Listbox(TableFrame, height=20, width=20)
    PracticeTimeListBox.grid(row=1, column=4)

    lblTartan = Label(TableFrame, text="Tartan", font=("Arvo"))
    lblTartan.grid(row=0, column=5, sticky="ew")

    TartanListBox = Listbox(TableFrame, height=20, width=10)
    TartanListBox.grid(row=1, column=5)

    def yview(*args):
        BandListBox.yview(*args)
        GradeListBox.yview(*args)
        SectionListBox.yview(*args)
        AddressListBox.yview(*args)
        PracticeTimeListBox.yview(*args)
        TartanListBox.yview(*args)

    Scrollbary = Scrollbar(TableFrame, orient=VERTICAL, command=yview, )
    Scrollbary.grid(row=1, column=6)

    BandListBox.config(yscrollcommand=Scrollbary.set)
    GradeListBox.config(yscrollcommand=Scrollbary.set)
    SectionListBox.config(yscrollcommand=Scrollbary.set)
    AddressListBox.config(yscrollcommand=Scrollbary.set)
    PracticeTimeListBox.config(yscrollcommand=Scrollbary.set)
    TartanListBox.config(yscrollcommand=Scrollbary.set)

    # Populating Band name listbox
    for numRecordsBands in range(len(ListBandName)):
        BandListBox.insert(0, ListBandName[numRecordsBands])
        numRecordsBands = + 1

    # Populating Grade listbox
    for numRecordBands in range(len(ListGrade)):
        GradeListBox.insert(0, ListGrade[numRecordBands])
        numRecordsBands = + 1

    # Populating Section listbox
    for numRecordBands in range(len(ListSection)):
        SectionListBox.insert(0, ListSection[numRecordBands])
        numRecordsBands = + 1

    # Populating Address listbox
    for numRecordBands in range(len(ListAddress)):
        AddressListBox.insert(0, ListAddress[numRecordBands])
        numRecordsBands = + 1

    # Populating PracticeTime listbox
    for numRecordBands in range(len(ListPracticeTime)):
        PracticeTimeListBox.insert(0, ListPracticeTime[numRecordBands])
        numRecordsBands = + 1

    # Populating Tartan Name listbox
    for numRecordBands in range(len(ListTartan)):
        TartanListBox.insert(0, ListTartan[numRecordBands])
        numRecordsBands = + 1

# On competition button press
def CompetitionFrame():
    # Display Drawn for grade selected if the drawn has been taken place else display to user that the competition has not been drawn yet.
    def DisplayDraw(CompetitionSelectedID, NumEntered, Grade, ViewDraw):
        lblNumBands = Label(ViewDraw, text=" The Number of Bands Entered grade "+ str(Grade) + " is:    " + str(NumEntered))
        lblNumBands.grid(row=2, column=1, columnspan=4)

        lblspacer = Label(ViewDraw)
        lblspacer.grid(row=4)

        #get All bands entered into the competiton grade
        conn = sqlite3.connect('RspbaniDB.db')
        with conn:
            cursor = conn.cursor()
        getComps = 'SELECT EntryID, BandID FROM BandsEntered WHERE CompetitionID = ? AND Grade = ?'
        cursor.execute(getComps, [CompetitionSelectedID, Grade])
        CompetitionResults = cursor.fetchall()

        print("CompetitionResults")
        print(CompetitionResults)

        # for each band get the draw details to display
        if CompetitionResults:
            n = 6
            for c in CompetitionResults:
                EntryID = c[0]
                BandID = c[1]

                with conn:
                    cursor = conn.cursor()
                getComps = 'SELECT BandName FROM BandAccount WHERE BandID = ?'
                cursor.execute(getComps, [BandID])
                CompetitionResultsBandName = cursor.fetchall()

                if CompetitionResultsBandName:
                    for e in CompetitionResultsBandName:
                        BandName = e[0]

                    with conn:
                        cursor = conn.cursor()
                    getComps = 'SELECT CircleNumber, CompetitingTime FROM CompetitionDraw WHERE BandID = ? AND CompetitionID = ?'
                    cursor.execute(getComps, [BandID, CompetitionSelectedID])
                    CompetitionResultsDrawDetails = cursor.fetchall()

                    print("CompetitionResultsDrawDetails")
                    print(CompetitionResultsDrawDetails)

                    if CompetitionResultsDrawDetails:
                        for y in CompetitionResultsDrawDetails:
                            CircleNum = y[0]
                            CompetingTime = y[1]

                            lblBandName = Label(ViewDraw, text=BandName, width=25)
                            lblBandName.grid(row=n, column=1, sticky=W)

                            lblCircleNum = Label(ViewDraw, text=CircleNum, width=15)
                            lblCircleNum.grid(row=n, column=2, sticky=W)

                            lblCompetingTime = Label(ViewDraw, text=CompetingTime, width=15)
                            lblCompetingTime.grid(row=n, column=3, sticky=W)
                    else:

                        NotDrawn = Frame(contentFrame, width=965, height=600)
                        NotDrawn.grid(row=0, column=0, sticky="nsew")

                        lblspacer = Label(NotDrawn, text="", width=35)
                        lblspacer.grid(row=0, column=0)

                        lblSpacer = Label(NotDrawn, text="View Draw", height=2, font=("Arvo", 32))
                        lblSpacer.grid(row=0, column=2, columnspan=4, sticky="ew")

                        lblText = Label(NotDrawn, text="This grade has not been drawn yet! ", height=3,
                                        font=("Arvo", 14))
                        lblText.grid(row=2, column=2, columnspan=6, sticky="ew")

                        ButtonBack = Button(NotDrawn, image=photoBack,
                                            command=lambda: SelectGradeDraw(CompetitionSelected))
                        ButtonBack.grid(row=1, column=0, sticky=W)
                n += 1

            lblspacer = Label(ViewDraw)
            lblspacer.grid(row=n + 1)
            ButtonBack = Button(ViewDraw, image=photoBack, command=lambda: SelectGradeDraw(CompetitionSelected))
            ButtonBack.grid(row=0, column=0)
        else:

            NoDraw = Frame(contentFrame, width=965, height=600)
            NoDraw.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(NoDraw, width=40)
            lblspacer.grid()

            lblSpacer = Label(NoDraw, text="View Draw", height=2, font=("Arvo", 32))
            lblSpacer.grid(row=0, column=2, columnspan=4, sticky="ew")

            lblText = Label(NoDraw, text="No bands have entered this grade! ", height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=2, columnspan=6, sticky="ew")

            ButtonBack = Button(NoDraw, image=photoBack, command=lambda: SelectGradeDraw(CompetitionSelected))
            ButtonBack.grid(row=0, column=0)

    def ViewDraw(CompetitionSelected, CompetitionSelectedID, listOfGradeDrawBox):
        # get selection from list box
        Selection = listOfGradeDrawBox.curselection()

        print("Selection")
        print(Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select a grade to view the draw!")
        else:

            for i in Selection:
                # get value of selection
                DrawGrade = listOfGradeDrawBox.get(i)
                print(str(DrawGrade))

            # get competition details to display
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            getComps = 'SELECT CompetitionName, NumG1Ent, NumG2Ent, NumG3AEnt, NumG3BEnt, NumG4AEnt, NumG4BEnt FROM Competitions WHERE Drawn = ? AND CompetitionID = ?'
            cursor.execute(getComps, ["True", CompetitionSelectedID])
            CompetitionResults = cursor.fetchall()

            if CompetitionResults:
                for q in CompetitionResults:
                    CompName = q[0]
                    NumG1Ent = q[1]
                    NumG2Ent = q[2]
                    NumG3AEnt = q[3]
                    NumG3BEnt = q[4]
                    NumG4AEnt = q[5]
                    NumG4BEnt = q[6]

                ViewDraw = Frame(contentFrame, width=965, height=600)
                ViewDraw.grid(row=0, column=0, sticky="nsew")

                lblspacer = Label(ViewDraw, width=15)
                lblspacer.grid()

                lblSpacer = Label(ViewDraw, text="View Draw", height=2, font=("Arvo", 32))
                lblSpacer.grid(row=0, column=1, columnspan=4, sticky="ew")

                lblCompName = Label(ViewDraw, text="You have selected:   " + str(CompName))
                lblCompName.grid(row=1,column=1, columnspan=6)

                lblBandName = Label(ViewDraw, text="Band:  ", width=25, font=("Arvo", 16))
                lblBandName.grid(row=5, column=1, sticky=W)

                lblCircleNum = Label(ViewDraw, text="Circle Number:", width=15, font=("Arvo", 16))
                lblCircleNum.grid(row=5, column=2, sticky=W)

                lblCompetingTime = Label(ViewDraw, text="Competing Time:", width=15, font=("Arvo", 16))
                lblCompetingTime.grid(row=5, column=3, sticky=W)

                # depending on grade selected run the DisplayDraw function
                if DrawGrade == "1":
                    DisplayDraw(CompetitionSelectedID, NumG1Ent, DrawGrade, ViewDraw)
                elif DrawGrade == "2":
                    DisplayDraw(CompetitionSelectedID, NumG2Ent, DrawGrade, ViewDraw)
                elif DrawGrade == "3A":
                    DisplayDraw(CompetitionSelectedID, NumG3AEnt, DrawGrade, ViewDraw)
                elif DrawGrade == "3B":
                    DisplayDraw(CompetitionSelectedID, NumG3BEnt, DrawGrade, ViewDraw)
                elif DrawGrade == "4A":
                    DisplayDraw(CompetitionSelectedID, NumG4AEnt, DrawGrade, ViewDraw)
                elif DrawGrade == "4B":
                    DisplayDraw(CompetitionSelectedID, NumG4BEnt, DrawGrade, ViewDraw)

            else:

                NoDraw = Frame(contentFrame, width=965, height=600)
                NoDraw.grid(row=0, column=0, sticky="nsew")

                lblspacer = Label(NoDraw, width=4)
                lblspacer.grid()

                lblSpacer = Label(NoDraw, text="View Draw", height=2, font=("Arvo", 32))
                lblSpacer.grid(row=0, column=2, columnspan=4, sticky="ew")

                lblText = Label(NoDraw, text="The Draw for this competition has not taken place yet! ", height=3,
                                font=("Arvo", 14))
                lblText.grid(row=1, column=2, columnspan=6, sticky="ew")

                ButtonBack = Button(NoDraw, image=photoBack, command=lambda: SelectGradeDraw(CompetitionSelected))
                ButtonBack.grid(row=0, column=0)

    # after selecting competition ask user to select grade to view results
    def SelectGradeDraw(CompetitionSelected):
        # A list of all the competitions, the date its held on, the location if it and the number of bands that have entered.
        # Click on the competition and see the bands that have entered the competition for each grade.
        # Once the draw has been made then include the the circle number and the competing time for each band

        CompetitionSelectedID = CompetitionSelected.get()

        print(CompetitionSelectedID)
        print("CompetitionSelectedID")


        if CompetitionSelectedID == "":
            messagebox.showinfo(message="Please select a competition by clicking on the radio button beside the competition that you want to view.")

        conn = sqlite3.connect('RspbaniDB.db')
        with conn:
            cursor = conn.cursor()
        getComps = 'SELECT CompetitionID FROM Competitions WHERE Drawn = ? AND CompetitionID = ?'
        cursor.execute(getComps, ["True", CompetitionSelectedID])
        CompetitionDrawn = cursor.fetchall()

        print("CompetitionDrawn")
        print(CompetitionDrawn)


        if CompetitionDrawn == []:
            messagebox.showinfo(message="The Draw for this competition has not taken place yet! ")
        else:

            CompetitionSelectGrade = Frame(contentFrame, width=965, height=600)
            CompetitionSelectGrade.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(CompetitionSelectGrade, width=35)
            lblspacer.grid()

            lblSpacer = Label(CompetitionSelectGrade, text="Select Grade", height=2, font=("Arvo", 32))
            lblSpacer.grid(row=0, column=2, columnspan=4, sticky="ew")

            lblText = Label(CompetitionSelectGrade, text="Please select the Grade you would like to see the draw for. ",
                            height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=2, columnspan=6, sticky="ew")

            listOfGradeDrawBox = Listbox(CompetitionSelectGrade, selectmode=SINGLE, width=10, height=7)
            listOfGradeDrawBox.grid(row=3, column=4)
            # populate all grades into the list box
            listOfGradeDrawBox.insert(0, "1")
            listOfGradeDrawBox.insert(1, "2")
            listOfGradeDrawBox.insert(2, "3A")
            listOfGradeDrawBox.insert(3, "3B")
            listOfGradeDrawBox.insert(4, "4A")
            listOfGradeDrawBox.insert(5, "4B")

            SelectGradeButton = Button(CompetitionSelectGrade, image=photoSelectGradeButton
                                       , command=lambda: ViewDraw(CompetitionSelected, CompetitionSelectedID,
                                                                  listOfGradeDrawBox))
            SelectGradeButton.grid(row=4, column=3, columnspan=3, pady="20")

            ButtonBack = Button(CompetitionSelectGrade, image=photoBack, command=lambda: CompetitionFrame())
            ButtonBack.grid(row=0, column=0)

    # Creates and displays the competitions page which allows user to select competition

    Competition = Frame(contentFrame, width=965, height=600)
    Competition.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(Competition, width=5)
    lblspacer.grid()

    lblSpacer = Label(Competition, text="Competitions", height=3, font=("Arvo", 32))
    lblSpacer.grid(row=0, column=2, columnspan=4, sticky="n")

    # Create all the filtering options to the User if they want to search by Grade and place them on the Frame
    CompetitionSelected = StringVar()

    lblSelection = Label(Competition, text="Select to\n view draw", font=("Arvo", 16))
    lblSelection.grid(row=2, column=1)

    lblCompetitionName = Label(Competition, text="Competition Name", width=30, font=("Arvo", 16))
    lblCompetitionName.grid(row=2, column=2)

    lblLocation = Label(Competition, text="Location", width=20, font=("Arvo", 16))
    lblLocation.grid(row=2, column=3)

    lblDate = Label(Competition, text="Date", width=15, font=("Arvo", 16))
    lblDate.grid(row=2, column=4)

    lblspacer = Label(Competition, height=2)
    lblspacer.grid(row=3)

    conn = sqlite3.connect('RspbaniDB.db')
    with conn:
        cursor = conn.cursor()
    getComps = 'SELECT CompetitionID, CompetitionName, CompetitionLocation, CompetitionDate FROM Competitions ORDER BY CompetitionName'
    cursor.execute(getComps)
    CompetitionResults = cursor.fetchall()

    print(CompetitionResults)

    c = 4
    for i in CompetitionResults:
        lblName = Label(Competition, text=i[1], width=30)
        lblName.grid(row=c, column=2, sticky="w")

        lblLoc = Label(Competition, text=i[2], width=20)
        lblLoc.grid(row=c, column=3, sticky="w")

        lblD = Label(Competition, text=i[3], width=15)
        lblD.grid(row=c, column=4, sticky="w")

        Radiobutton(Competition, padx=5, variable=CompetitionSelected, value=i[0]).grid(row=c, column=1)
        c += 1

    lblspacer = Label(Competition)
    lblspacer.grid(row=c + 1)

    ButtonViewDraw = Button(Competition, image=photoViewDraw, command=lambda: SelectGradeDraw(CompetitionSelected))
    ButtonViewDraw.grid(row=c + 2, column=1)


def ResultsFrame():
    def ViewResults(CompetitionSelected, CompetitionID, listOfGradeDrawBox, listOfCompsNameBox):
        def getResultsfromGrade(GradeSelected, BandsInComp, CompetitionID):

            GradeResults = []

            if BandsInComp == []:

                lblSpacer = Label(Fullresults, height=5)
                lblSpacer.grid(row=4)

                lblNoBandsEntered = Label(Fullresults, text = "No Bands have entered this grade. Use the back button to select \nanother grade in " + CompetitionSelected + ".")
                lblNoBandsEntered.grid(row=5, column=3, columnspan=5)


            else:

                for bandID in BandsInComp:

                    # for each Band get its Results from the results table
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    getBandID = 'SELECT Piping1Score, Piping2Score, TotalPipingScore, DrummingScore, EnsembleScore, TotalScore, Position From Results WHERE BandID = ? AND CompetitionID = ?'
                    cursor.execute(getBandID, [bandID, CompetitionID])
                    BandResults = cursor.fetchall()

                    print("BandResults")
                    print(BandResults)

                    if BandResults:
                        for R in BandResults:
                            Piping1 = R[0]
                            Piping2 = R[1]
                            TotalPiping = R[2]
                            Drumming = R[3]
                            Ensemble = R[4]
                            TotalScore = R[5]
                            Position = R[6]

                            BandResultsList = []
                            BandResultsList.append(bandID)
                            BandResultsList.append(Piping1)
                            BandResultsList.append(Piping2)
                            BandResultsList.append(TotalPiping)
                            BandResultsList.append(Drumming)
                            BandResultsList.append(Ensemble)
                            BandResultsList.append(TotalScore)
                            BandResultsList.append(Position)

                            print("BandResultsList")
                            print(BandResultsList)

                            GradeResults.append(BandResultsList)

                print("GradeResults")
                print(GradeResults)

                # sort 2D array by Total score (element 5) and then by ensemble Score (element 4) afterwards
                SortedGradeResults = sorted(GradeResults, key=lambda x: (x[6], x[5]))

                print("SortedGradeResults")
                print(SortedGradeResults)

                r = 5

                for BandScore in SortedGradeResults:

                    # Get the BandName from its ID
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    getBandName = 'SELECT BandName From BandAccount WHERE BandID = ?'
                    cursor.execute(getBandName, [BandScore[0]])
                    BandNameResults = cursor.fetchall()

                    if BandNameResults:
                        for n in BandNameResults:
                            BandName = n[0]

                    lblBand = Label(Fullresults, text=BandName, width=25)
                    lblBand.grid(row=r, column=2, sticky="e")

                    lblPiping1 = Label(Fullresults, text=BandScore[1], width=8)
                    lblPiping1.grid(row=r, column=3)

                    lblPiping2 = Label(Fullresults, text=BandScore[2], width=8)
                    lblPiping2.grid(row=r, column=4)

                    lblTotalPiping = Label(Fullresults, text=BandScore[3], width=8)
                    lblTotalPiping.grid(row=r, column=5)

                    lblDrumming = Label(Fullresults, text=BandScore[4], width=8)
                    lblDrumming.grid(row=r, column=6)

                    lblEnsemble = Label(Fullresults, text=BandScore[5], width=8)
                    lblEnsemble.grid(row=r, column=7)

                    lblTotalScore = Label(Fullresults, text=BandScore[6], width=8)
                    lblTotalScore.grid(row=r, column=8)

                    lblPosition = Label(Fullresults, text=BandScore[7], width=8, font=("Arvo", 14))
                    lblPosition.grid(row=r, column=9)
                    r += 1

        AllBandIDInComp = []
        G1BandsInComp = []
        G2BandsInComp = []
        G3ABandsInComp = []
        G3BBandsInComp = []
        G4ABandsInComp = []
        G4BBandsInComp = []

        Selection = listOfGradeDrawBox.curselection()

        if Selection == ():
            messagebox.showinfo(message="Please select a grade.")
        else:

            for i in Selection:
                # get value of selection
                GradeSelected = listOfGradeDrawBox.get(i)
                print(str(GradeSelected))

            Fullresults = Frame(contentFrame, width=965, height=600)
            Fullresults.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(Fullresults, width=1)
            lblspacer.grid()

            lblSpacer = Label(Fullresults, text="Results", height=2, font=("Arvo", 32))
            lblSpacer.grid(row=0, column=3, columnspan=4, sticky="ew")

            lblText = Label(Fullresults, text=str(CompetitionSelected) + " - Grade " + str(GradeSelected),
                            height=3,
                            font=("Arvo", 20))
            lblText.grid(row=1, column=2, columnspan=7, sticky="ew")

            lblSpacer = Label(Fullresults)
            lblSpacer.grid(row=2)

            lblBand = Label(Fullresults, text="        Band", width=25, font=("Arvo", 14))
            lblBand.grid(row=3, column=2, sticky=E)

            lblPiping1 = Label(Fullresults, text="Piping 1", width=8, font=("Arvo", 14))
            lblPiping1.grid(row=3, column=3)

            lblPiping2 = Label(Fullresults, text="Piping 2", width=8, font=("Arvo", 14))
            lblPiping2.grid(row=3, column=4)

            lblTotalPiping = Label(Fullresults, text="Total Piping", width=8, font=("Arvo", 14))
            lblTotalPiping.grid(row=3, column=5)

            lblDrumming = Label(Fullresults, text="Drumming", width=8, font=("Arvo", 14))
            lblDrumming.grid(row=3, column=6)

            lblEnsemble = Label(Fullresults, text="Ensemble", width=8, font=("Arvo", 14))
            lblEnsemble.grid(row=3, column=7)

            lblTotalScore = Label(Fullresults, text="Total Score", width=8, font=("Arvo", 14))
            lblTotalScore.grid(row=3, column=8)

            lblPosition = Label(Fullresults, text="Position", width=8, font=("Arvo", 14))
            lblPosition.grid(row=3, column=9)

            ButtonBack = Button(Fullresults, image=photoBack, command=lambda: ResultsFrame())
            ButtonBack.grid(row=0, column=0)

            # gets the all EntryID's from CompetitionID
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            getEntryID = 'SELECT EntryID From BandsEntered WHERE CompetitionID = ?'
            cursor.execute(getEntryID, [CompetitionID])
            EntryIDResults = cursor.fetchall()

            if EntryIDResults:
                for q in EntryIDResults:
                    EntryID = q[0]

                    # if there is bands entered into the competition get their ID's
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    getBandID = 'SELECT BandID From BandsEntered WHERE EntryID = ? AND CompetitionID = ?'
                    cursor.execute(getBandID, [EntryID, CompetitionID])
                    BandIDResults = cursor.fetchall()

                    if BandIDResults:
                        for d in BandIDResults:
                            BandID = d[0]

                            AllBandIDInComp.append(BandID)

                print(AllBandIDInComp)

                for eachBand in AllBandIDInComp:
                    # for each band in the competition get its grade and then sort them into graded lists from each bands ID
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    getBandGrade = 'SELECT BandGrade From BandAccount WHERE BandID = ?'
                    cursor.execute(getBandGrade, [eachBand])
                    BandGradeResults = cursor.fetchall()

                    if BandGradeResults:
                        for G in BandGradeResults:
                            Grade = G[0]

                        if Grade == "1":
                            G1BandsInComp.append(eachBand)
                        elif Grade == "2":
                            G2BandsInComp.append(eachBand)
                        elif Grade == "3A":
                            G3ABandsInComp.append(eachBand)
                        elif Grade == "3B":
                            G3BBandsInComp.append(eachBand)
                        elif Grade == "4A":
                            G4ABandsInComp.append(eachBand)
                        elif Grade == "4B":
                            G4BBandsInComp.append(eachBand)

                print("G1" + str(G1BandsInComp))
                print("G2" + str(G2BandsInComp))
                print("G3A" + str(G3ABandsInComp))
                print("G3b" + str(G3BBandsInComp))
                print("G4a" + str(G4ABandsInComp))
                print("G4b" + str(G4BBandsInComp))

                if GradeSelected == "1":
                    getResultsfromGrade(GradeSelected, G1BandsInComp, CompetitionID)
                elif GradeSelected == "2":
                    getResultsfromGrade(GradeSelected, G2BandsInComp, CompetitionID)
                elif GradeSelected == "3A":
                    getResultsfromGrade(GradeSelected, G3ABandsInComp, CompetitionID)
                elif GradeSelected == "3B":
                    getResultsfromGrade(GradeSelected, G3BBandsInComp, CompetitionID)
                elif GradeSelected == "4A":
                    getResultsfromGrade(GradeSelected, G4ABandsInComp, CompetitionID)
                elif GradeSelected == "4B":
                    getResultsfromGrade(GradeSelected, G4BBandsInComp, CompetitionID)

    #                         # for each band get its name from its ID
    #                         conn = sqlite3.connect('RspbaniDB.db')
    #                         with conn:
    #                             cursor = conn.cursor()
    #                         getBandID = 'SELECT BandName From BandAccount WHERE BandID = ?'
    #                         cursor.execute(getBandID, [BandID])
    #                         BandNameResults = cursor.fetchall()
    #
    #                         if BandNameResults:
    #                             for n in BandNameResults:
    #                                 BandName = n[0]
    #                             # for each Band get its Results from the results table
    #                             conn = sqlite3.connect('RspbaniDB.db')
    #                             with conn:
    #                                 cursor = conn.cursor()
    #                             getBandID = 'SELECT Piping1Score, Piping2Score, TotalPipingScore, DrummingScore, EnsembleScore, TotalScore, Position From Results WHERE EntryID = ?'
    #                             cursor.execute(getBandID, [EntryID])
    #                             BandResults = cursor.fetchall()
    #
    #                             print("BandResults")
    #
    #                             print(BandResults)
    #
    #                             if BandResults:
    #                                 for R in BandResults:
    #                                     Piping1 = R[0]
    #                                     Piping2 = R[1]
    #                                     TotalPiping = R[2]
    #                                     Drumming = R[3]
    #                                     Ensemble = R[4]
    #                                     TotalScore = R[5]
    #
    #                                 print("BandResults")
    #                                 print(BandResults)
    #                             else:
    #
    #                                 NoResults = Frame(contentFrame, width=965, height=600)
    #                                 NoResults.grid(row=0, column=0, sticky="nsew")
    #
    #                                 lblspacer = Label(NoResults, width=40)
    #                                 lblspacer.grid()
    #
    #                                 lblSpacer = Label(NoResults, text="Results", height=2, font=("Arvo", 32))
    #                                 lblSpacer.grid(row=0, column=2, columnspan=4, sticky="ew")
    #
    #                                 lblText = Label(NoResults,
    #                                                 text="The results have not been published yet for this grade", height=3,
    #                                                 font=("Arvo", 14))
    #                                 lblText.grid(row=1, column=2, columnspan=6, sticky="ew")
    #
    #                                 ButtonBack = Button(NoResults, image=photoBack, command=lambda: ResultsFrame())
    #                                 ButtonBack.grid(row=0, column=0)
    #
    #
    #                 lblBand = Label(Fullresults, text=BandName, width=25)
    #                 lblBand.grid(row=v, column=2, sticky=W)
    #
    #                 lblPiping1 = Label(Fullresults, text=Piping1, width=8)
    #                 lblPiping1.grid(row=v, column=3)
    #
    #                 lblPiping2 = Label(Fullresults, text=Piping2, width=8)
    #                 lblPiping2.grid(row=v, column=4)
    #
    #                 lblTotalPiping = Label(Fullresults, text=TotalPiping, width=8)
    #                 lblTotalPiping.grid(row=v, column=5)
    #
    #                 lblDrumming = Label(Fullresults, text=Drumming, width=8)
    #                 lblDrumming.grid(row=v, column=6)
    #
    #                 lblEnsemble = Label(Fullresults, text=Ensemble, width=8)
    #                 lblEnsemble.grid(row=v, column=7)
    #
    #                 lblTotalScore = Label(Fullresults, text=TotalScore, width=8)
    #                 lblTotalScore.grid(row=v, column=8)
    #
    #                 lblPosition = Label(Fullresults, text=v-4, width=8, font=("Arvo", 14))
    #                 lblPosition.grid(row=v, column=9)
    #                 v += 1
    #         else:
    #
    #             NoResults = Frame(contentFrame, width=965, height=600)
    #             NoResults.grid(row=0, column=0, sticky="nsew")
    #
    #             lblspacer = Label(NoResults, width=40)
    #             lblspacer.grid()
    #
    #             lblSpacer = Label(NoResults, text="Results", height=2, font=("Arvo", 32))
    #             lblSpacer.grid(row=0, column=2, columnspan=4, sticky="ew")
    #
    #             lblText = Label(NoResults, text="No Bands have entered this grade", height=3,
    #                             font=("Arvo", 14))
    #             lblText.grid(row=1, column=2, columnspan=6, sticky="ew")
    #
    #             ButtonBack = Button(NoResults, image=photoBack, command=lambda: ResultsFrame())
    #             ButtonBack.grid(row=0, column=0)
    #
    #
    #
    #

    # get EntryID for all bands in grade from CompID
    # for each entryID get the BandID
    # For each BandID get BandName
    # for Each EntryID get the Bands Results

    def SelectGradeForResults(listOfCompsNameBox):
        Selection = listOfCompsNameBox.curselection()

        if Selection == ():
            messagebox.showinfo(message="Please select a competition to view results for by selecting the competition name in the list box above.")
        else:

            for i in Selection:
                # get value of selection
                CompetitionSelected = listOfCompsNameBox.get(i)
                print(str(CompetitionSelected))

            # gets the Competition ID from its name
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            getComps = 'SELECT CompetitionID From Competitions WHERE CompetitionName = ?'
            cursor.execute(getComps, [CompetitionSelected])
            CompetitionResults = cursor.fetchall()

            if CompetitionResults:
                for q in CompetitionResults:
                    CompetitionID = q[0]

            CompetitionSelectGrade = Frame(contentFrame, width=965, height=600)
            CompetitionSelectGrade.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(CompetitionSelectGrade, width=35)
            lblspacer.grid()

            lblSpacer = Label(CompetitionSelectGrade, text="Results", height=2, font=("Arvo", 32))
            lblSpacer.grid(row=0, column=2, columnspan=4, sticky="ew")

            lblText = Label(CompetitionSelectGrade, text="Please select the Grade you would like to see the results for. ",
                            height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=2, columnspan=6, sticky="ew")

            listOfGradeDrawBox = Listbox(CompetitionSelectGrade, selectmode=SINGLE, width=10, height=7)
            listOfGradeDrawBox.grid(row=3, column=4)
            # populate all grades into the list box
            listOfGradeDrawBox.insert(0, "1")
            listOfGradeDrawBox.insert(1, "2")
            listOfGradeDrawBox.insert(2, "3A")
            listOfGradeDrawBox.insert(3, "3B")
            listOfGradeDrawBox.insert(4, "4A")
            listOfGradeDrawBox.insert(5, "4B")

            # go to next screen
            SelectGradeButton = Button(CompetitionSelectGrade, image=photoSelectGradeButton
                                       , command=lambda: ViewResults(CompetitionSelected, CompetitionID,
                                                                     listOfGradeDrawBox, listOfCompsNameBox))
            SelectGradeButton.grid(row=4, column=3, columnspan=3, pady="20")

            # go back to the list of competitions screen
            ButtonBack = Button(CompetitionSelectGrade, image=photoBack, command=lambda: ResultsFrame())
            ButtonBack.grid(row=0, column=0)

    Results = Frame(contentFrame, width=965, height=600)
    Results.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(Results, width=30)
    lblspacer.grid()

    lblSpacer = Label(Results, text="Results", height=2, font=("Arvo", 32))
    lblSpacer.grid(row=0, column=1, columnspan=6, sticky="ew")

    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through the data
        cursorListOfComps = db.cursor()
        # SQL Query to select all Competitions from the competitions table
    listOfCompsQuery = 'SELECT CompetitionName, CompetitionDate FROM Competitions WHERE Collated = "True" '
    # Execute the query to get a list of all the competitions
    cursorListOfComps.execute(listOfCompsQuery)
    # The query will produce a 2 dimensional tuple
    ListOfComps = cursorListOfComps.fetchall()

    lblText = Label(Results, text="Please select the competition you would like to view the results for ", height=3,
                    font=("Arvo", 14))
    lblText.grid(row=1, column=1, columnspan=8, sticky="ew")

    lblCompetiton = Label(Results, text="Competition Name")
    lblCompetiton.grid(row=3, column=2, columnspan=2, sticky=EW)

    lblDate = Label(Results, text="Competition Date")
    lblDate.grid(row=3, column=4, columnspan=2, sticky=W)

    listOfCompsNameBox = Listbox(Results, selectmode=SINGLE, width=30)
    listOfCompsNameBox.grid(row=4, column=2, columnspan=2, sticky=E)

    for u in ListOfComps:
        n = 0
        listOfCompsNameBox.insert(n, str(u[0]))
        n = + 1
    # listbox to show competition date but it is disabled so it cant be selected
    listOfCompDateBox = Listbox(Results, selectmode=DISABLED, width=20, selectbackground="white")
    listOfCompDateBox.grid(row=4, column=4, columnspan=2, sticky=W)
    for u in ListOfComps:
        n = 0
        listOfCompDateBox.insert(n, str(u[1]))
        n = + 1

    SelectButton = Button(Results, image=photoSelectCompetitionButton,
                          command=lambda: SelectGradeForResults(listOfCompsNameBox))
    SelectButton.grid(row=6, column=2, columnspan=4, pady="10")


def LoginFrame(photoLoginPageButton):
    # initialise variables
    username = StringVar()
    password = StringVar()

    Login = Frame(contentFrame, width=965, height=600)
    Login.grid(row=0, column=0, sticky="nsew")

    Login.grid_columnconfigure(0, weight=1)
    Login.grid_columnconfigure(1, weight=1)

    lblSpacer = Label(Login, text="")
    lblSpacer.grid(row=0, column=0, sticky="e")

    lblSpacer1 = Label(Login, text="")
    lblSpacer1.grid(row=1, column=0, sticky="e")

    lblSpacer2 = Label(Login, text="")
    lblSpacer2.grid(row=3, column=0, sticky="e")

    lblSpacer3 = Label(Login, text="")
    lblSpacer3.grid(row=4, column=0, sticky="e")

    # create the login frame widgets
    lblLogin = Label(Login, text="Login", font=("Arvo", 40))
    lblLogin.grid(row=2, column=1, sticky="w", pady=20, padx=30)

    lblUsername = Label(Login, text="Username: ", font=("Arvo"))
    lblUsername.grid(row=5, column=0, sticky="e")
    EtyUsername = Entry(Login, textvar=username)
    EtyUsername.grid(row=5, column=1, sticky="w")

    lblPassword = Label(Login, text="Password: ", font=("Arvo"))
    lblPassword.grid(row=6, column=0, sticky="e")
    EtyPassword = Entry(Login, textvar=password, show='•')
    EtyPassword.grid(row=6, column=1, sticky="w")

    ButtonPassFrgt = Button(Login, text="Forgot Password?", command=lambda: PassFrgtFrame(photoResetButton),
                            font=("Arvo", 10), fg="blue", underline=True)
    ButtonPassFrgt.grid(row=7, column=1, sticky="wn")

    lblSpacer4 = Label(Login, text="")
    lblSpacer4.grid(row=8, column=0, sticky="e")
    # change 'UserDatabaseCheck' to 'AllFeatures' to bypass not working login to see forms for new band, judge and admin
    ButtonLogin = Button(Login, image=photoLoginPageButton,
                         command=lambda: UserDatabaseCheck(username, password, EtyUsername, EtyPassword))
    ButtonLogin.grid(row=9, column=1, sticky="w", padx="18")

    global Username
    global Password
    Username = username.get()
    Password = password.get()


def PassFrgtFrame(photoResetButton):
    def PassFrgtFrameSQ(Question, Answer, photoEnterButton, attempt):

        print("attempt")
        print(attempt)
        # This code will be run only if the value of attempt is less than 3 meaning they will only be allowed
        # to enter their answer three wrong times
        if attempt < 3:

            SQAnswer = StringVar()

            PassFrgtSQ = Frame(contentFrame, width=965, height=600)
            PassFrgtSQ.grid(row=0, column=0, sticky="nsew")

            # create the Password Reset frame widgets
            lblSpacer = Label(PassFrgtSQ, width=35)
            lblSpacer.grid()

            lblSpacer1 = Label(PassFrgtSQ, text="")
            lblSpacer1.grid(row=1)

            lblTitle = Label(PassFrgtSQ, text="Password Reset", font=("Arvo", 40))
            lblTitle.grid(row=2, column=1, sticky="w", pady=20, padx=30)

            lblSpacer3 = Label(PassFrgtSQ, text="")
            lblSpacer3.grid(row=3)

            lblText = Label(PassFrgtSQ, text="Please answer the Account Security Question:")
            lblText.grid(row=4, column=1, sticky="ew")

            lblSpacer4 = Label(PassFrgtSQ)
            lblSpacer4.grid(row=5)
            # populate a label with their chosen security question
            lblQuestion = Label(PassFrgtSQ, text=Question, font=("Arvo"))
            lblQuestion.grid(row=6, column=1, sticky="sew")

            lblSpacer7 = Label(PassFrgtSQ)
            lblSpacer7.grid(row=7)

            # Ask them for their Answer
            EtyAnswer = Entry(PassFrgtSQ, textvar=SQAnswer, width=30)
            EtyAnswer.grid(row=8, column=1)

            ButtonEnter = Button(PassFrgtSQ, image=photoEnterButton,
                                 command=lambda: CheckSQAns(SQAnswer, Answer, photoEnterButton, Question, attempt))
            ButtonEnter.grid(row=10, column=1)

        else:
            # if they have attempted more than three times then dont allow them to try again
            # need to add validation to not allow them to start the process again - will involve adding another
            # field to table.
            messagebox.showwarning(
                message="Sorry you have entered the wrong answer too many times! Please contact an administrator on 02894396763")
            HomeFrame()

    def PasswordDatabaseCheck(usernamePassReset, emailPassReset):
        # If an account has been found with matching credientials then get the Security Question and answer belonging to the account.
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorBand = db.cursor()
        findBand = 'SELECT SecurityQ, SecurityQAns FROM BandAccount WHERE Username = ? AND EmailAddress = ?'
        cursorBand.execute(findBand, [usernamePassReset.get(), emailPassReset.get()])
        SecurityQA = cursorBand.fetchall()
        if SecurityQA:
            for i in SecurityQA:
                Question = i[0]
                Answer = i[1]
            # Define an incremental variable so that the user can only try 3 times
            attempt = 0
            # Calls the function that asks the user for their security Questions
            PassFrgtFrameSQ(Question, Answer, photoEnterButton, attempt)
        else:
            # If no account found then create a message box and ask them to try again
            if messagebox.askretrycancel(message="No account found with these details") == 1:
                # Call the function where it asks them for their username and email address again
                PassFrgtFrame(photoResetButton)
            else:
                # if they dont want to retry then this function brings them back to the log in screen
                LoginFrame(photoLoginPageButton)

    def UsernameFrgt():
        # Displays a message to user if they have forgotten their username
        messagebox.showwarning("Forgot Username?", message="Please contact the administration team on 02891747865")
        HomeFrame()

    def ChangePassword(NewPassword1, NewPassword2, photoEnterButton, Question, attempt):
        def SaveNewPasswordToDatebase(newPassword, username):
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            cursor.execute('UPDATE BandAccount SET Password = ? WHERE Username = ? ',
                           (newPassword, username))
            conn.commit()

            messagebox.showinfo(message="Password Successfully Changed")
            LoginFrame(photoLoginPageButton)

        username = usernamePassReset.get()
        # if both new passwords match then change update the password stored in the database linked to their account
        if NewPassword1.get() == NewPassword2.get():
            newPassword = NewPassword1.get()
            # if pasword meets validation criteria then save it to the database
            if CheckValidPassword(newPassword) == True:
                SaveNewPasswordToDatebase(newPassword, username)
                # tell the user its not valid and ask them to enter it again
            else:
                messagebox.showerror(message="Password is not valid. Your password must be at least 8 characters in length and contain at least 1 uppercase letter, lowercase letter and number. Please try again ")
                CheckSQAns("True", "True", photoEnterButton, Question, attempt)

        # tell user passwords dont match
        else:
            messagebox.showerror(message="Passwords do not match, please try again ")
            CheckSQAns("True", "True", photoEnterButton, Question, attempt)

    def CheckSQAns(SQAnswer, Answer, photoEnterButton, Question, attempt):

        NewPassword1 = StringVar()
        NewPassword2 = StringVar()

        # if their answer given matches the answer stored then ask them to double enter their new password.
        if SQAnswer.get() == Answer:

            ChangePassWFrame = Frame(contentFrame, width=965, height=600)
            ChangePassWFrame.grid(row=0, column=0, sticky="nsew")

            # create the Password Reset frame widgets
            lblSpacer = Label(ChangePassWFrame, width=35)
            lblSpacer.grid()

            lblSpacer1 = Label(ChangePassWFrame, text="")
            lblSpacer1.grid(row=1)

            lblTitle = Label(ChangePassWFrame, text="Change Password", font=("Arvo", 40))
            lblTitle.grid(row=2, column=1, columnspan=2, sticky="w", pady=20, padx=30)

            lblSpacer3 = Label(ChangePassWFrame, text="")
            lblSpacer3.grid(row=3)

            lblText = Label(ChangePassWFrame, text="Please enter your new password:")
            lblText.grid(row=4, column=1, columnspan=2, sticky="ew")

            lblSpacer4 = Label(ChangePassWFrame)
            lblSpacer4.grid(row=5)

            lblNewPassword1 = Label(ChangePassWFrame, text="New Password:", font=("Arvo"))
            lblNewPassword1.grid(row=6, column=1, sticky="e")

            EtyNewPassword1 = Entry(ChangePassWFrame, textvar=NewPassword1, width=20, show='•')
            EtyNewPassword1.grid(row=6, column=2)

            lblNewPassword2 = Label(ChangePassWFrame, text="Re-enter Password:", font=("Arvo"))
            lblNewPassword2.grid(row=7, column=1, sticky="e")

            EtyNewPassword2 = Entry(ChangePassWFrame, textvar=NewPassword2, width=20, show='•')
            EtyNewPassword2.grid(row=7, column=2)

            ButtonEnter = Button(ChangePassWFrame, image=photoEnterButton,
                                 command=lambda: ChangePassword(NewPassword1, NewPassword2, photoEnterButton, Question,
                                                                attempt))
            ButtonEnter.grid(row=10, column=2)

        else:
            # if their answer doesnt not match then display message, increment attempt and ask them to enter their security Question answer again
            if messagebox.askokcancel(message="Incorrect Answer, please try again") == 1:
                attempt = attempt + 1
                PassFrgtFrameSQ(Question, Answer, photoEnterButton, attempt)

            else:
                # if they dont want to try again they will be directed to the Home page
                HomeFrame()

    global usernamePassReset
    usernamePassReset = StringVar()
    emailPassReset = StringVar()

    # Display a frame that allows user to enter their username and email address
    PassFrgt = Frame(contentFrame, width=965, height=600)
    PassFrgt.grid(row=0, column=0, sticky="nsew")

    PassFrgt.grid_columnconfigure(0, weight=1)
    PassFrgt.grid_columnconfigure(1, weight=1)

    lblSpacer = Label(PassFrgt, text="")
    lblSpacer.grid(row=0, column=0, sticky="e")

    lblSpacer1 = Label(PassFrgt, text="")
    lblSpacer1.grid(row=1, column=0, sticky="e")

    lblSpacer2 = Label(PassFrgt, text="")
    lblSpacer2.grid(row=3, column=0, sticky="e")

    lblSpacer3 = Label(PassFrgt, text="")
    lblSpacer3.grid(row=4, column=0, sticky="e")

    # create the Password Reset frame widgets
    lblLogin = Label(PassFrgt, text="Password Reset ", font=("Arvo", 40))
    lblLogin.grid(row=2, column=0, columnspan=2, sticky="ew", pady=20, padx=30)

    lblUserName = Label(PassFrgt, text="Username: ", font=("Arvo"))
    lblUserName.grid(row=5, column=0, sticky="e")

    EtyUsername = Entry(PassFrgt, textvar=usernamePassReset)
    EtyUsername.grid(row=5, column=1, sticky="w")

    lblEmailAddress = Label(PassFrgt, text="Email Address: ", font=("Arvo"))
    lblEmailAddress.grid(row=6, column=0, sticky="e")
    EtyEmail = Entry(PassFrgt, textvar=emailPassReset)
    EtyEmail.grid(row=6, column=1, sticky="w")
    ButtonPassFrgt = Button(PassFrgt, text="Forgot Username? ", command=UsernameFrgt, font=("Arvo", 10), fg="blue",
                            underline=True)
    ButtonPassFrgt.grid(row=7, column=1, sticky="wn")

    lblSpacer4 = Label(PassFrgt, text="")
    lblSpacer4.grid(row=8, column=0, sticky="e")
    # Pressing button calls function to check their inputted username and password
    ButtonReset = Button(PassFrgt, image=photoResetButton,
                         command=lambda: PasswordDatabaseCheck(usernamePassReset, emailPassReset))
    ButtonReset.grid(row=9, column=1, sticky="w", padx="18")


def AllFeatures(username):
    BandFeatures(photoEditBandDetails, photoEnterCompetitions, photoBookTransport, photoRegisterMember,
                 photoTransferMember, photoEditMemberDetails, username)
    JudgeFeatures()
    AdminFeatures(photoNewBand, photoNewJudge, photoNewAdmin, photoNewCompetition, photoEditCompetition,
                  photoDrawCompetition)


def UserDatabaseCheck(username, password, EtyUsername, EtyPassword):
    def SetAccountDetails(Band, Admin, Judge, Username, Password, Valid):
        def checkSQ(SecurityQuesAns, SecurityQues, Band, Admin, Judge, Username, Password):
            def SaveSQToDatebase(SecurityQuesAns, SecurityQues, Band, Admin, Judge, Username, Password):

                if Band == True:

                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # Insert the fields which have been fetched and insert them into the database
                    # Sets FirstLog in to False so the user is not asked to set these details again.
                    cursor.execute(
                        'UPDATE BandAccount SET SecurityQ = ?, SecurityQAns = ?, FirstLogIn = ? WHERE Username = ? AND Password = ?',
                        (SecurityQues, SecurityQuesAns, "False", Username, Password))
                    conn.commit()
                    # close the window to bring the user to home page
                    PopupWindow.destroy()

                elif Admin == True:
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # Insert the fields which have been fetched and insert them into the database
                    cursor.execute(
                        'UPDATE AdminAccount SET SecurityQ = ?, SecurityQAns = ?, FirstLogIn = ? WHERE Username = ? AND Password = ?',
                        (SecurityQues, SecurityQuesAns, "False", Username, Password))
                    conn.commit()
                    PopupWindow.destroy()


                elif Judge == True:
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # Insert the fields which have been fetched and insert them into the database
                    cursor.execute(
                        'UPDATE JudgeAccount SET SecurityQ = ?, SecurityQAns = ?, FirstLogIn = ? WHERE Username = ? AND Password = ?',
                        (SecurityQues, SecurityQuesAns, "False", Username, Password))
                    conn.commit()
                    PopupWindow.destroy()

            Answer = SecurityQuesAns.get()
            Question = SecurityQues.get()

            AnswerCheck = False
            QuestionCheck = False

            if Question != "Select Question":
                QuestionCheck = True
            if Answer != "":
                AnswerCheck = True

            if QuestionCheck == True:
                if AnswerCheck == True:
                    messagebox.showinfo(message="Account set up complete, thank you!")
                    print(Answer, Question)
                    SaveSQToDatebase(Answer, Question, Band, Admin, Judge, Username, Password)

                else:
                    messagebox.showerror(message="Please enter your answer, to your selected question!")
                    SecurityQFrame(Band, Admin, Judge, Username, Password)

            else:
                messagebox.showerror(message="Please select a security question from the list provided!")
                SecurityQFrame(Band, Admin, Judge, Username, Password)

        def SecurityQFrame(Band, Admin, Judge, Username, Password):

            # create frame to ask user for security question

            SecurityQues = StringVar()
            SecurityQuesAns = StringVar()

            SQ = Frame(MainFrame, width=500, height=250)
            SQ.grid(row=0, sticky=NSEW)

            lblspacer = Label(SQ, height=3)
            lblspacer.grid(row=1)

            lblText = Label(SQ,
                            text="So that we can recover your account if you were to \nforget your password, please choose a security question",
                            font=("Arvo", 12))
            lblText.grid(row=1, column=1, columnspan=7, sticky="ew")

            lblspacer = Label(SQ)
            lblspacer.grid(row=2)

            lblPassword1 = Label(SQ, text="Please select your security question: ", font=("Arvo", 12))
            lblPassword1.grid(row=3, column=1, sticky="e")

            SQList = ["What was your childhood nickname?", "What is the middle name of your oldest child?",
                      "What is your favorite team?", "What is your favorite movie?", "In what town was your first job?",
                      "What school did you attend for sixth grade?", "What was the make of your first car?"]
            SQDropList = OptionMenu(SQ, SecurityQues, *SQList)
            SQDropList.config(width=20, font=("Arvo", 12))
            SecurityQues.set("Select Question")
            SQDropList.grid(row=3, column=2, sticky="w")

            lblSQAns = Label(SQ, text="Enter your Answer: ", font=("Arvo", 12))
            lblSQAns.grid(row=4, column=1, sticky="e")

            EtySecurityQAns = Entry(SQ, width=20, textvar=SecurityQuesAns)
            EtySecurityQAns.grid(row=4, column=2, sticky="w")

            lblspacer = Label(SQ)
            lblspacer.grid(row=5)

            ButtonSQ = Button(SQ, image=photoFinish,
                              command=lambda: checkSQ(SecurityQuesAns, SecurityQues, Band, Admin, Judge, Username,
                                                      Password))
            ButtonSQ.grid(row=6, column=2)

        def CheckPassword(Band, Admin, Judge, Username, Password):
            def SaveNewPassToDatabase(Band, Admin, Judge, Username, Password, NewPassword):
                if Band == True:

                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # Insert the fields which have been fetched and insert them into the database
                    cursor.execute(
                        'UPDATE BandAccount SET Password = ? WHERE Username = ? AND Password = ?',
                        (NewPassword, Username, Password))
                    conn.commit()

                    SecurityQFrame(Band, Admin, Judge, Username, NewPassword)

                elif Admin == True:
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # Insert the fields which have been fetched and insert them into the database
                    cursor.execute(
                        'UPDATE AdminAccount SET Password = ? WHERE Username = ? AND Password = ?',
                        (NewPassword, Username, Password))
                    conn.commit()
                    SecurityQFrame(Band, Admin, Judge, Username, NewPassword)

                elif Judge == True:
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # Insert the fields which have been fetched and insert them into the database
                    cursor.execute(
                        'UPDATE JudgeAccount SET Password = ? WHERE Username = ? AND Password = ?',
                        (NewPassword, Username, Password))
                    conn.commit()
                    SecurityQFrame(Band, Admin, Judge, Username, NewPassword)

            password1 = newPassword1.get()
            password2 = newPassword2.get()

            if password1 == password2:

                NewPassword = password1
                passwordCheck = False

                if CheckValidPassword(NewPassword) == True:
                    passwordCheck = True

                if passwordCheck == True:
                    SaveNewPassToDatabase(Band, Admin, Judge, Username, Password, NewPassword)
                else:
                    PopupWindow.destroy()
                    SetAccountDetails(Band, Admin, Judge, Username, Password, "True")

            else:
                PopupWindow.destroy()
                # recursion - ask the user to enter their passoword again - Valid now true so will populate screen to tell user to enter password again.
                SetAccountDetails(Band, Admin, Judge, Username, Password, "True")

        # create a pop up window for user to change password and set Security Questions
        PopupWindow = Toplevel(root)
        PopupWindow.title('Account Set-up')
        # Set the size of the root to the dimensions of my laptop
        PopupWindow.geometry("500x250+500+275")
        PopupWindow.config(bg="purple")

        PopupWindow.grid_columnconfigure(0, minsize=500)
        PopupWindow.grid_rowconfigure(0, minsize=250)

        newPassword1 = StringVar()
        newPassword2 = StringVar()

        MainFrame = Frame(PopupWindow, bg="Orange")
        MainFrame.grid(row=0, sticky=NSEW)

        MainFrame.grid_columnconfigure(0, minsize=500)
        MainFrame.grid_rowconfigure(0, minsize=250)

        BodyFrame = Frame(MainFrame, bg="White")
        BodyFrame.grid(row=0, sticky=NSEW)

        lblTitle = Label(BodyFrame, text="Welcome to your account", font=("Arvo", 32))
        lblTitle.grid(row=0, column=0, columnspan=3, sticky="ew")

        lblText = Label(BodyFrame, text="We need to finalise a few details from you to make your account secure.",
                        font=("Arvo", 12))
        lblText.grid(row=1, column=0, columnspan=3, sticky="ew")

        if Valid == "False":
            lblspacer = Label(BodyFrame)
            lblspacer.grid(row=2)
        else:
            lblspacer = Label(BodyFrame, text="Invalid password - please try again", font=("Arvo", 12), fg="Red")
            lblspacer.grid(row=2, column=1, columnspan=2)

        lblPassword1 = Label(BodyFrame, text="Please enter your new password: ", font=("Arvo", 12))
        lblPassword1.grid(row=3, column=1, sticky="e")

        EtyPassword1 = Entry(BodyFrame, textvar=newPassword1, show='•')
        EtyPassword1.grid(row=3, column=2, sticky="w")

        lblPassword2 = Label(BodyFrame, text="Please enter your password again: ", font=("Arvo", 12))
        lblPassword2.grid(row=4, column=1, sticky="e")

        EtyPassword2 = Entry(BodyFrame, textvar=newPassword2, show='•')
        EtyPassword2.grid(row=4, column=2, sticky="w")

        lblspacer = Label(BodyFrame)
        lblspacer.grid(row=5)

        passwordButton = Button(BodyFrame, image=photoNextButton,
                                command=lambda: CheckPassword(Band, Admin, Judge, Username, Password))
        passwordButton.grid(row=6, column=2)

    # I am trying to check each table for matching credentials that have been entered and then depending
    # on if the person is a judge band or admin display the correct buttons available to them

    global Admin
    Admin = False

    global BandnameofLoggedIn
    global BandIDofLoggedIn
    successful = False
    Login = False

    # check the BandAccount table from my database to see if the username and password entered are a match with a record
    with sqlite3.connect("RspbaniDB.db") as db:
        cursorBand = db.cursor()
    findBand = 'SELECT * FROM BandAccount WHERE Username = ? AND Password =?'
    cursorBand.execute(findBand, [username.get(), password.get()])
    CurUsername = username.get()
    CurPassword = password.get()
    print(Username)
    Bandresults = cursorBand.fetchall()
    print("B", Bandresults)
    # if results then welcome the user
    if Bandresults:
        messagebox.showinfo("Welcome",
                            message="Hello " + str(Bandresults[0][1]) + ". You have been successfully logged in.")
        Band = True

        print("Bandresults[0][18]")
        print(Bandresults[0][18])
        # if the value of the "FirstLogin" Field is True then as them to change their details
        if Bandresults[0][18] == "True":
            # call function to update their password
            SetAccountDetails(Band, "", "", CurUsername, CurPassword, "False")
        # Populate the application with only the features availbale to bands
        BandFeatures(photoEditBandDetails, photoEnterCompetitions, photoBookTransport, photoRegisterMember,
                     photoTransferMember, photoEditMemberDetails, username)
        # Show the Home Screen
        HomeFrame()
        # Change the top frame to Logged in state
        Loggedin(photoLogin2, photoSmallLogo)
        successful = True
        Login = True

        if Login == True:
            # Once a band is logged in get the name and ID of the band from the database using their unique username
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorBandname = db.cursor()
            findBandName = 'SELECT bandName, BandID FROM BandAccount WHERE Username = ?'
            cursorBandname.execute(findBandName, [username.get()])
            Bandnameresults = cursorBandname.fetchall()

            if Bandnameresults:

                for i in Bandnameresults:
                    print("i", i)
                    BandnameofLoggedIn = i[0]
                    BandIDofLoggedIn = i[1]

                    print("BandnameofLoggedIn", BandnameofLoggedIn)
                    print("BandIDofLoggedIn", BandIDofLoggedIn)

    else:
        # if there are no results to this username and password combination found in the BandAccount table then
        # check in the judge table instead
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorJudge = db.cursor()
        findJudge = 'SELECT * FROM JudgeAccount WHERE Username = ? AND Password = ?'
        cursorJudge.execute(findJudge, [username.get(), password.get()])
        Judgeresults = cursorJudge.fetchall()
        CurUsername = username.get()
        CurPassword = password.get()
        print("J", Judgeresults)

        if Judgeresults:
            messagebox.showinfo("Welcome",
                                message="Hello " + str(Judgeresults[0][2]) + ". You have been successfully logged in.")
            Judge = True

            # if the value of the "FirstLogin" Field is True then as them to change their details
            if Judgeresults[0][12] == "True":
                # call function to update their password
                SetAccountDetails("", "", Judge, CurUsername, CurPassword, "False")
            # Populate the application with only the features available to Judges
            JudgeFeatures()
            # Show home frame
            HomeFrame()
            # Change the top frame to Logged in state
            Loggedin(photoLogin2, photoSmallLogo)
            successful = True
            Login = True

            # Get the ID of the Judge Logged in and make it global to be used throughout the rest of the program
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorJudgeID = db.cursor()
            findJudgeID = 'SELECT JudgeID FROM JudgeAccount WHERE Username = ?'
            cursorJudgeID.execute(findJudgeID, [username.get()])
            Judgeresults = cursorJudgeID.fetchall()

            if Judgeresults:
                global JudgeIDOfLoggedIn
                for i in Judgeresults:
                    for t in i:
                        JudgeIDOfLoggedIn = t

        else:
            # if there are no results to this username and password combination found in the JudgeAccount table then
            # check in the Admin table instead
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorAdmin = db.cursor()
            findAdmin = 'SELECT * FROM AdminAccount WHERE Username = ? AND Password = ?'
            cursorAdmin.execute(findAdmin, [username.get(), password.get()])
            Adminresults = cursorAdmin.fetchall()
            CurUsername = username.get()
            CurPassword = password.get()
            print("A", Adminresults)

            if Adminresults:
                messagebox.showinfo("Welcome", message="Hello " + str(
                    Adminresults[0][2]) + ". You have been successfully logged in.")
                Admin = True
                # if the value of the "FirstLogin" Field is True then as them to change their details
                if Adminresults[0][11] == "True":
                    # call function to update their password
                    SetAccountDetails("", Admin, "", CurUsername, CurPassword, "False")
                # Populate the application with All features
                AllFeatures(username)
                # show home frame
                HomeFrame()
                # Change the top frame to Logged in state
                Loggedin(photoLogin2, photoSmallLogo)
                successful = True
                Login = True

            elif successful == False:
                # If not match is found in any of the tables then show a message box and
                # remove the entered data from the entry boxes
                messagebox.showinfo("-- ERROR --",
                                    "No account has been found on this system with these credentials. Please make sure you have entered your details correctly!",
                                    icon="warning")
                EtyUsername.delete(0, END)
                EtyPassword.delete(0, END)
    print(Admin)


def BandFeatures(photoEditBandDetails, photoEnterCompetitions, photoBookTransport, photoRegisterMember,
                 photoTransferMember, photoEditMemberDetails, username):
    lblSpacer = Label(editFrame, text="", width="22")
    lblSpacer.grid(row=0, column=0)

    # Create the widgets for the additional features that logged on bands can access on the Edit Frame
    ButtonEditBand = Button(editFrame, image=photoEditBandDetails,
                            command=lambda: EditBandDetailsFrame(photoUpdateBandButton, photoSelectBandButton,
                                                                 username))
    ButtonEditBand.grid(row=0, column=1, padx=5)
    ButtonEnterComp = Button(editFrame, image=photoEnterCompetitions,
                             command=lambda: EnterCompFrame(photoSelectCompetitionsButton))
    ButtonEnterComp.grid(row=0, column=2, padx=5)
    ButtonBookTransport = Button(editFrame, image=photoBookTransport,
                                 command=lambda: BookTransportFrame(photoSubmitButton))
    ButtonBookTransport.grid(row=0, column=3, padx=5)
    ButtonRegMember = Button(editFrame, image=photoRegisterMember,
                             command=lambda: RegMemberFrame(photoRegisterMemberFormButton))
    ButtonRegMember.grid(row=0, column=4, padx=5)
    ButtonTransferMember = Button(editFrame, image=photoTransferMember,
                                  command=lambda: TransferMemberFrame(photoTransferMemberFormButton))
    ButtonTransferMember.grid(row=0, column=5, padx=5)
    ButtonEditMember = Button(editFrame, image=photoEditMemberDetails,
                              command=lambda: EditMemberDetailsFrame(photoUpdateMemberButton, photoSelectBandButton,
                                                                     photoSelectMemberButton))
    ButtonEditMember.grid(row=0, column=6, padx=5)


def JudgeFeatures():
    # Create the widgets for the additional features that Judges can access on the JudgeMenu Frame
    ButtonJudgeCompetition = Button(JudgeMenu, image=photoJudgeCompetition,
                                    command=lambda: JudgeCompetitionFrame(photoSelectCompetitionButton))
    ButtonJudgeCompetition.grid(row=0, column=1, padx=200)


def AdminFeatures(photoNewBand, photoNewJudge, photoNewAdmin, photoNewCompetition, photoEditCompetition,
                  photoDrawCompetition):
    # Create the widgets for the additional features that admins can access on the adminMenu Frame
    buttonNewBand = Button(adminMenu, command=lambda: RegBandFrame(photoRegisterBandButton))
    buttonNewBand.config(image=photoNewBand)
    buttonNewBand.grid(row=0)

    buttonNewJudge = Button(adminMenu, command=lambda: RegJudgeFrame(photoRegisterJudgeButton, Years))
    buttonNewJudge.config(image=photoNewJudge)
    buttonNewJudge.grid(row=2)

    buttonNewAdmin = Button(adminMenu, command=lambda: RegAdminFrame(photoRegisterAdminButton, Years))
    buttonNewAdmin.config(image=photoNewAdmin)
    buttonNewAdmin.grid(row=4)

    buttonNewCompetition = Button(adminMenu, command=lambda: AddCompetitionFrame(photoCreateCompetitionButton, CompYears))
    buttonNewCompetition.config(image=photoNewCompetition)
    buttonNewCompetition.grid(row=6)

    buttonEditCompetition = Button(adminMenu, command=lambda: EditCompetitionFrame(photoSelectCompetitionButton, CompYears))
    buttonEditCompetition.config(image=photoEditCompetition)
    buttonEditCompetition.grid(row=8)

    buttonDrawCompetition = Button(adminMenu, command=lambda: DrawCompetitionFrame(photoSelectCompetitionButton, Years))
    buttonDrawCompetition.config(image=photoDrawCompetition)
    buttonDrawCompetition.grid(row=10)

    buttonCollateResults = Button(adminMenu, command=CollateCompResults)
    buttonCollateResults.config(image=photoCollateResults)
    buttonCollateResults.grid(row=12)

    buttonSendResults = Button(adminMenu, command=SendResultstoBands)
    buttonSendResults.config(image=photoSendResults)
    buttonSendResults.grid(row=14)


def EditBandDetailsFrame(photoUpdateBandButton, photoSelectBandButton, username):
    def DisplayBandDetails(BandID):
        def DeleteBand(BandID, BandName):

            # ask them if they want to delete the Band
            Answer = messagebox.askquestion("Delete Band?",
                                            "Are you sure you want to delete " + BandName + " and all of its members?")
            if Answer == "yes":

                # Get all EntryID's of the competition.
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorEntryID = db.cursor()
                findEntryID = 'SELECT EntryID FROM BandsEntered WHERE BandID = ?'
                cursorEntryID.execute(findEntryID, [BandID])
                EntryIDresults = cursorEntryID.fetchall()

                if EntryIDresults:

                    for a in EntryIDresults:
                        Entry = a[0]
                        print("Entry")
                        print(Entry)
                        # Deletes every record that is stored about/ has a relationship with  the competition to be deleted from database
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute("DELETE FROM BandsEntered WHERE EntryID = ?", (Entry,))
                        conn.commit()
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute("DELETE FROM CompetitionDraw WHERE EntryID = ?", (Entry,))
                        conn.commit()
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute("DELETE FROM Results WHERE EntryID = ?", (Entry,))
                        conn.commit()

                # Delete Band from database
                conn = sqlite3.connect('RspbaniDB.db')
                with conn:
                    cursor = conn.cursor()
                cursor.execute("DELETE FROM BandAccount WHERE BandID = ?", (BandID,))
                conn.commit()
                # Delete all members registered to the Band
                with conn:
                    cursor = conn.cursor()
                cursor.execute("DELETE FROM BandMembers WHERE BandID = ?", (BandID,))
                conn.commit()
                getListOfBands()
                EditBandDetailsFrame(photoUpdateBandButton, photoSelectBandButton, username)

        # Define all the textvar variables to store entered data
        UptBandName = StringVar()
        UptBandBranch = StringVar()
        UptBandGrade = StringVar()
        UptPMFirstName = StringVar()
        UptPMSecondName = StringVar()
        UptBandHallName = StringVar()
        UptBandAddressLine1 = StringVar()
        UptBandAddressLine2 = StringVar()
        UptBandCounty = StringVar()
        UptBandPostcode = StringVar()
        UptBandUsername = StringVar()
        UptBandPassword = StringVar()
        UptBandPracticeTime = StringVar()
        UptBandTartan = StringVar()
        UptBandEmail = StringVar()

        EditBandDetails = Frame(contentFrame, width=965, height=600)
        EditBandDetails.grid(row=0, column=0, sticky="nsew")

        lblspacer = Label(EditBandDetails, text="", width=15)
        lblspacer.grid()

        lblSpacer = Label(EditBandDetails, text="Edit Band Details", height=3, font=("Arvo", 32))
        lblSpacer.grid(row=0, column=1, columnspan=5, sticky="ew")

        # get all the data stored about the Band logged in or selected (depends on if admin is logged in or not)
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorCurrentBand = db.cursor()
        FindDetailsCurrentBand = 'SELECT * FROM BandAccount WHERE BandID = ?'
        cursorCurrentBand.execute(FindDetailsCurrentBand, [BandID])
        CurrentBandDetails = cursorCurrentBand.fetchall()

        print(CurrentBandDetails)

        for i in CurrentBandDetails:
            # Assign all the queried data to variables to be populated into the entry boxes
            BandID = i[0]
            CurrBandName = i[1]
            CurrBranch = i[2]
            CurrGrade = i[3]
            CurrPMFirstName = i[4]
            CurrPMSecondName = i[5]
            CurrHallName = i[7]
            CurrAddress1 = i[8]
            CurrAddress2 = i[9]
            CurrCounty = i[10]
            CurrPostcode = i[11]
            CurrPracticeTime = i[12]
            CurrTartan = i[13]
            CurrUsername = i[14]
            CurrPassword = i[15]
            CurrEmail = i[6]



        print("CurrBandName", CurrBandName)
        getListOfBands()

        # populate the frame with all the fields stored about a band and popuate them with the contents
        # of the variables above
        lblBandName = Label(EditBandDetails, text="Band Name: ", font=("Arvo"))
        lblBandName.grid(row=2, column=1, sticky="e")

        EtyUptBandName = Entry(EditBandDetails, textvar=UptBandName)
        EtyUptBandName.insert(0, CurrBandName)
        EtyUptBandName.grid(row=2, column=2, sticky="w")

        lblBranch = Label(EditBandDetails, text="Branch: ", font=("Arvo"))
        lblBranch.grid(row=3, column=1, sticky="e")

        listBranch = ['Antrim', 'Down', 'Fermanagh', 'ISPBA', 'Londonderry']
        droplistUptBranch = OptionMenu(EditBandDetails, UptBandBranch, *listBranch)
        droplistUptBranch.config(width=20)
        UptBandBranch.set(CurrBranch)
        droplistUptBranch.grid(row=3, column=2, sticky="w")

        lblGrade = Label(EditBandDetails, text="Grade", font=("Arvo"))
        lblGrade.grid(row=4, column=1, sticky="e")

        listGrade = ['1', '2', '3A', '3B', '4A', '4B']
        droplistUptGrade = OptionMenu(EditBandDetails, UptBandGrade, *listGrade)
        droplistUptGrade.config(width=20)
        UptBandGrade.set(CurrGrade)
        droplistUptGrade.grid(row=4, column=2, sticky="w")

        lblAddress = Label(EditBandDetails, text="Address:", font=("Arvo", 10))
        lblAddress.grid(row=5, column=1, sticky="e")

        lblspacer = Label(EditBandDetails)
        lblspacer.grid(row=5)

        lblBandHallName = Label(EditBandDetails, text="Hall Name ", font=("Arvo"))
        lblBandHallName.grid(row=6, column=1, sticky="e")

        EtyUptBandHallName = Entry(EditBandDetails, textvar=UptBandHallName)
        EtyUptBandHallName.insert(0, CurrHallName)
        EtyUptBandHallName.grid(row=6, column=2, sticky="w")

        lblAddressLine1 = Label(EditBandDetails, text="Line 1", font=("Arvo"))
        lblAddressLine1.grid(row=7, column=1, sticky="e")

        EtyUptBandAddressLine1 = Entry(EditBandDetails, textvar=UptBandAddressLine1)
        EtyUptBandAddressLine1.insert(0, CurrAddress1)
        EtyUptBandAddressLine1.grid(row=7, column=2, sticky="w")

        lblAddressLine2 = Label(EditBandDetails, text="Line 2", font=("Arvo"))
        lblAddressLine2.grid(row=8, column=1, sticky="e")

        EtyUptBandAddressLine2 = Entry(EditBandDetails, textvar=UptBandAddressLine2)
        EtyUptBandAddressLine2.insert(0, CurrAddress2)
        EtyUptBandAddressLine2.grid(row=8, column=2, sticky="w")

        lblCounty = Label(EditBandDetails, text="County", font=("Arvo"))
        lblCounty.grid(row=9, column=1, sticky="e")
        listCounty = ['County Antrim', 'County Down', 'County Armagh', 'County Fermanagh', 'County Londonderry',
                      'Ireland']
        dropUptlist = OptionMenu(EditBandDetails, UptBandCounty, *listCounty)
        dropUptlist.config(width=20)
        UptBandCounty.set(CurrCounty)
        dropUptlist.grid(row=9, column=2, sticky="w")

        lblBandPostcode = Label(EditBandDetails, text="Postcode", font=("Arvo"))
        lblBandPostcode.grid(row=10, column=1, sticky="e")

        EtyUptBandPostcode = Entry(EditBandDetails, textvar=UptBandPostcode)
        EtyUptBandPostcode.insert(0, CurrPostcode)
        EtyUptBandPostcode.grid(row=10, column=2, sticky="w")

        lblBandPracticeTime = Label(EditBandDetails, text="Practice Time(s): ", font=("Arvo"))
        lblBandPracticeTime.grid(row=11, column=1, sticky="e")

        EtyUptBandPracticeTime = Entry(EditBandDetails, textvar=UptBandPracticeTime)
        EtyUptBandPracticeTime.insert(0, CurrPracticeTime)
        EtyUptBandPracticeTime.grid(row=11, column=2, sticky="w")

        lblBandTartan = Label(EditBandDetails, text="Tartan: ", font=("Arvo"))
        lblBandTartan.grid(row=12, column=1, sticky="e")

        EtyUptBandTartan = Entry(EditBandDetails, textvar=UptBandTartan)
        EtyUptBandTartan.insert(0, CurrTartan)
        EtyUptBandTartan.grid(row=12, column=2, sticky="w")

        lblspacer = Label(EditBandDetails, width=5)
        lblspacer.grid(column=3)

        lblPMFirstName = Label(EditBandDetails, text="PM First Name", font=("Arvo"))
        lblPMFirstName.grid(row=2, column=4, sticky="e")

        EtyUptPMFirstName = Entry(EditBandDetails, textvar=UptPMFirstName)
        EtyUptPMFirstName.insert(0, CurrPMFirstName)
        EtyUptPMFirstName.grid(row=2, column=5, sticky="w")

        lblPMSecondName = Label(EditBandDetails, text="PM Second Name", font=("Arvo"))
        lblPMSecondName.grid(row=3, column=4, sticky="e")

        EtyUptPMSecondName = Entry(EditBandDetails, textvar=UptPMSecondName)
        EtyUptPMSecondName.insert(0, CurrPMSecondName)
        EtyUptPMSecondName.grid(row=3, column=5, sticky="w")

        lblBandEmail = Label(EditBandDetails, text="Email Address:", font=("Arvo"))
        lblBandEmail.grid(row=4, column=4)

        EtyUptBandEmail = Entry(EditBandDetails, textvar=UptBandEmail)
        EtyUptBandEmail.insert(0, CurrEmail)
        EtyUptBandEmail.grid(row=4, column=5, sticky="w")

        lblBandUsername = Label(EditBandDetails, text="Username: ", font=("Arvo"))
        lblBandUsername.grid(row=6, column=4)

        EtyUptBandUsername = Entry(EditBandDetails, textvar=UptBandUsername)
        EtyUptBandUsername.insert(0, CurrUsername)
        EtyUptBandUsername.grid(row=6, column=5, sticky="w")

        lblBandPassword = Label(EditBandDetails, text="Password: ", font=("Arvo"))
        lblBandPassword.grid(row=7, column=4, sticky="e")

        EtyUptBandPassword = Entry(EditBandDetails, textvar=UptBandPassword)
        EtyUptBandPassword.insert(0, CurrPassword)
        EtyUptBandPassword.grid(row=7, column=5, sticky="w")

        # Button to update all the data about the band
        Button(EditBandDetails, image=photoUpdateBandButton,
               command=lambda: UpdateBandDetailsDatabase(UptBandName, UptBandBranch, UptBandGrade, UptPMFirstName,
                                                         UptPMSecondName, UptBandHallName, UptBandAddressLine1,
                                                         UptBandAddressLine2, UptBandCounty, UptBandPostcode,
                                                         UptBandPracticeTime, \
                                                         UptBandTartan, UptBandUsername, UptBandPassword,
                                                         UptBandEmail)).grid(row=20, column=2, columnspan=3, padx="20",
                                                                             pady="10")

        if Admin == True:
            # allow admin to delete the band
            Button(EditBandDetails, image=photoDeleteButton,
                   command=lambda: DeleteBand(BandID, CurrBandName)).grid(row=20, column=1, sticky="w", padx="20",
                                                                          pady="10")

    def getUpdatedBandName():
        # The bands name must then be queried again as it is possible that is has been changed with the
        # username as it can not be changed
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorBandname = db.cursor()
        findBandName = 'SELECT bandName FROM BandAccount WHERE Username = ?'
        cursorBandname.execute(findBandName, [username.get()])
        Bandnameresults = cursorBandname.fetchall()

        if Bandnameresults:

            for i in Bandnameresults:
                i = str(i)
                i = i[2:-3]
                BandnameofLoggedIn = i
                print(i)
        # Now update the list of bands as well
        getListOfBands()

    def UpdateBandDetailsDatabase(UptBandName, UptBandBranch, UptBandGrade, UptPMFirstName, UptPMSecondName,
                                  UptBandHallName, \
                                  UptBandAddressLine1, UptBandAddressLine2, UptBandCounty, UptBandPostcode,
                                  UptBandPracticeTime, \
                                  UptBandTartan, UptBandUsername, UptBandPassword, UptBandEmail):

        def SavetoDatabase(BandID, Band, Branch, Grade, PMFirstname, PMSecondname, HallName, AddressLine1, AddressLine2,
                           County, Postcode, PracticeTime, Tartan, Username, Password, Email):
            def SendEmailToPM():

                # email address of the sender
                email_user = 'rspbani.info@gmail.com'
                # password of the sender
                email_password = 'P@55w0rd123'
                # email address of recipient
                email_send = Email

                # Creating the Emails subject
                subject = 'Band Details Updated'

                msg = MIMEMultipart()

                msg['From'] = email_user
                msg['To'] = email_send
                msg['Subject'] = subject

                # creating the contents of the email
                body = "Dear Pipe Major, \n\n\nWe have successfully updated your Bands details according to the changes you made on the Edit Band Details section. \n\n\nIf this change was not made by you please contact one of the systems administrators so we can resolve this issue. \n\n\n\nThe Royal Scottish Pipe Band Association \nNorthern Ireland Branch "
                msg.attach(MIMEText(body, 'plain'))



                text = msg.as_string()

                try:
                    # connet to email server and send email
                    server = smtplib.SMTP('smtp.gmail.com', 587)
                    server.starttls()
                    server.login(email_user, email_password)

                    server.sendmail(email_user, email_send, text)

                # error handling if email cant be sent - display a message stating this
                except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                    messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
                finally:
                    # stop connection with email server
                    server.quit()

            # update all the fields in the database with the variables above with the ID of the band Selected
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            cursor.execute('UPDATE BandAccount SET BandName = ?, BandBranch = ?, BandGrade = ?, PMFirstName = ?,'
                           'PMSecondName = ?, HallName = ?, BandAddressLine1 = ?, BandAddressLine2 = ?, BandCounty = ?,'
                           'BandPostcode = ?, PracticeTime = ?, Tartan = ?, Username = ?, Password = ?, EmailAddress = ?'
                           'WHERE BandID = ?', (Band, Branch, Grade, PMFirstname, PMSecondname, HallName, AddressLine1,
                                                AddressLine2, County, Postcode, PracticeTime, Tartan, Username,
                                                Password, Email,
                                                BandID))
            conn.commit()
            # get the new name of the band and the new list of bands as the name may have changed
            getUpdatedBandName()
            #update list of bands incase Band name has changed
            getListOfBands()
            #send email to PM
            SendEmailToPM()
            # show a message to the user when all data is saved
            messagebox.showinfo(message="Account updated successfully")



        # get all the data from the fields
        Band = UptBandName.get()
        Band = Band.title()
        Branch = UptBandBranch.get()
        Grade = UptBandGrade.get()
        PMFirstname = UptPMFirstName.get()
        PMSecondname = UptPMSecondName.get()
        HallName = UptBandHallName.get()
        AddressLine1 = UptBandAddressLine1.get()
        AddressLine2 = UptBandAddressLine2.get()
        County = UptBandCounty.get()
        Postcode = UptBandPostcode.get()
        PracticeTime = UptBandPracticeTime.get()
        Tartan = UptBandTartan.get()
        Username = UptBandUsername.get()
        Password = UptBandPassword.get()
        Email = UptBandEmail.get()

        print("Email")
        print(Email)

        BandCheck = False
        BranchCheck = False
        GradeCheck = False
        PMFirstnameCheck = False
        PMSecondnameCheck = False
        HallNameCheck = True
        AddressLine1Check = False
        AddressLine2Check = True
        CountyCheck = False
        PostcodeCheck = False
        UsernameCheck = False
        PasswordCheck = False
        PracticeTimeCheck = True
        TartanCheck = True
        EmailCheck = False

        #checks to ensure each individual field if valid
        if CheckAllAlpha(Band) == True:
            BandCheck = True
        if Branch != "Select Branch":
            BranchCheck = True
        if Grade != "Select Grade":
            GradeCheck = True
        if CheckAllAlpha(PMFirstname) == True:
            PMFirstnameCheck = True
        if CheckAllAlpha(PMSecondname) == True:
            PMSecondnameCheck = True
        if CheckOnlyNumAndLetters(AddressLine1) == True:
            AddressLine1Check = True
        if County != "Select your County":
            CountyCheck = True
        if CheckPostcode(Postcode) == True:
            PostcodeCheck = True
        if CheckUsername(Username, "Edit") == True:
            UsernameCheck = True
        if CheckValidPassword(Password) == True:
            PasswordCheck = True
        if CheckEmail(Email) == True:
            EmailCheck = True

        print(BandCheck, BranchCheck, GradeCheck, PMFirstnameCheck, PMSecondnameCheck, HallNameCheck,
              AddressLine1Check, AddressLine2Check, CountyCheck, PostcodeCheck, UsernameCheck, PasswordCheck,
              PracticeTimeCheck, TartanCheck, EmailCheck)

        # if all fields are valid then save to database - otherwise tell user the first field which is invalid
        if BandCheck == True:
            if BranchCheck == True:
                if GradeCheck == True:
                    if PMFirstnameCheck == True:
                        if PMSecondnameCheck == True:
                            if HallNameCheck == True:
                                if AddressLine1Check == True:
                                    if AddressLine2Check == True:
                                        if CountyCheck == True:
                                            if PostcodeCheck == True:
                                                if UsernameCheck == True:
                                                    if PasswordCheck == True:
                                                        if EmailCheck == True:
                                                            if Admin == True:
                                                                # if admin then call the save to database with the ID of the band the admin choose
                                                                SavetoDatabase(AdminBandIDSelected, Band, Branch,
                                                                               Grade,
                                                                               PMFirstname, PMSecondname, HallName,
                                                                               AddressLine1,
                                                                               AddressLine2, County,
                                                                               Postcode, PracticeTime,
                                                                               Tartan,
                                                                               Username, Password, Email)


                                                                # bring the admin back to the list box of bands to change another bands detials if they wish
                                                                EditBandDetailsFrame(photoUpdateBandButton,
                                                                                     photoSelectBandButton,
                                                                                     username)
                                                            else:
                                                                # call the save to database with the ID of the band logged in
                                                                SavetoDatabase(BandIDofLoggedIn, Band, Branch,
                                                                               Grade,
                                                                               PMFirstname, PMSecondname, HallName,
                                                                               AddressLine1,
                                                                               AddressLine2, County,
                                                                               Postcode, PracticeTime,
                                                                               Tartan,
                                                                               Username, Password, Email)

                                                                HomeFrame()
                                                        else:
                                                            messagebox.showinfo(
                                                                message="Email field is not valid. Please re-enter your email")
                                                    else:
                                                        messagebox.showinfo(
                                                            message="Password field is not valid. Please re-enter your password")
                                                else:
                                                    messagebox.showinfo(
                                                        message="Username field is not valid. Please re-enter your username")
                                            else:
                                                messagebox.showinfo(
                                                    message="Postcode field is not valid. Please re-enter your postcode")
                                        else:
                                            messagebox.showinfo(
                                                message="County field is not valid. Please re-enter your county")
                                    else:
                                        messagebox.showinfo(
                                            message="Address line 2 field is not valid. Please re-enter your address line 2")
                                else:
                                    messagebox.showinfo(
                                        message="Address line 1 field is not valid. Please include your house number and road.")
                            else:
                                messagebox.showinfo(
                                    message="Hall name field is not valid. Please re-enter it")
                        else:
                            messagebox.showinfo(
                                message="Pipe Majors second name field is not valid. Please re-enter it")
                    else:
                        messagebox.showinfo(
                            message="Pipe Majors first name field is not valid. Please re-enter it")
                else:
                    messagebox.showinfo(message="Grade field is not valid. Please re-enter your grade")
            else:
                messagebox.showinfo(message="Branch field is not valid. Please re-enter your branch")
        else:
            messagebox.showinfo(message="Band field is not valid. Please re-enter your band")

    def SpecBandPage():

        global AdminBandIDSelected

        AdminBandIDSelected = StringVar()
        # Get the position of the selection
        Selection = EditMemDListBandBox.curselection()
        print("selection", Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select a band that you wish to make these changes on behalf of.")
        else:

            # get the value of the position selected which is the Band name
            AdminBandSelected = EditMemDListBandBox.get(Selection[0])
            print(AdminBandSelected)
            # with the Band name get the BandID of the Band selected
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorCurrentBand = db.cursor()
            FindDetailsCurrentBand = 'SELECT bandID FROM BandAccount WHERE BandName = ?'
            cursorCurrentBand.execute(FindDetailsCurrentBand, [AdminBandSelected])
            CurrentBandDetails = cursorCurrentBand.fetchall()

            for q in CurrentBandDetails:
                for y in q:
                    AdminBandIDSelected = y

            print(AdminBandIDSelected)

            AdminEditBandDetails = Frame(contentFrame, width=965, height=600)
            AdminEditBandDetails.grid(row=0, column=0, sticky="nsew")

            # run this function and display the details of the band the Admin selected
            DisplayBandDetails(AdminBandIDSelected)

    if Admin == True:
        # If user is admin then ask them to select which band they wish to edit.
        EditBandDetails = Frame(contentFrame, width=965, height=600)
        EditBandDetails.grid(row=0, column=0, sticky="nsew")

        lblspacerCenter = Label(EditBandDetails, text="", width=20)
        lblspacerCenter.grid(row=0, column=0)

        Title = Label(EditBandDetails, text="Edit Band Details", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        lblText = Label(EditBandDetails, text="Please select the Band to edit ", height=3, font=("Arvo", 14))
        lblText.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)
        # create a list box and populate it with all registered bands
        EditMemDListBandBox = Listbox(EditBandDetails, selectmode=EXTENDED, width=20)
        EditMemDListBandBox.grid(row=3, column=4, columnspan=2)

        for k in listOfBands:
            n = 0
            # k = k[2:-3]
            EditMemDListBandBox.insert(n, str(k))
            n = + 1
        # create a button which can be pressed after the admin has selected a band
        SelectButton = Button(EditBandDetails, image=photoSelectBandButton, command=SpecBandPage)
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    else:
        # if the user is not an admin (ie a band) then display the details of the band that is logged in
        DisplayBandDetails(BandIDofLoggedIn)


def EnterCompFrame(photoSelectCompetitionsButton):
    def selectedComps(now, BandIDComps, listOfCompBox):

        # get all the positions of the Competitions selected
        Selection = listOfCompBox.curselection()
        # for each competition selected get the ID of the competiition

        print("Selection")
        print(Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select one or more competitions which you would like to enter")
        else:
            for i in Selection:
                j = listOfCompBox.get(i)

                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCompID = db.cursor()
                findCompID = 'SELECT CompetitionID FROM Competitions WHERE CompetitionName = ?'
                cursorCompID.execute(findCompID, [j])
                CompIDresults = cursorCompID.fetchall()

                if CompIDresults:

                    for a in CompIDresults:
                        a = str(a)
                        CompID = a[1:-2]

                    # get the current date to save to the database
                    CurrentDate = now.strftime('%d/%m/%Y')
                    print(CurrentDate)
                    # get the Grade of the Band entering the competitions
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    findBandGrade = 'SELECT BandGrade, EmailAddress FROM BandAccount WHERE BandID = ? '
                    cursor.execute(findBandGrade, [BandIDComps])
                    BandGradeResults = cursor.fetchall()

                    if BandGradeResults:
                        for i in BandGradeResults:
                            Grade = i[0]
                            Email = i[1]

                    #Table created
                    # Save the BandID, CompetitionID, Date Entered and the Grade to be used in the future
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # create the BandsEntered table if it doesnt already exist to store competition entry details
                    cursor.execute('CREATE TABLE IF NOT EXISTS BandsEntered '
                                   '(EntryID INTEGER PRIMARY KEY, BandID INTEGER, CompetitionID INTEGER, DateEntered TEXT, Grade TEXT)')
                    # Insert the fields which have been fetched and insert them into the database
                    cursor.execute('INSERT INTO BandsEntered '
                                   '(EntryID, BandID, CompetitionID, DateEntered, Grade) '
                                   'VALUES(NULL, ?,?,?,?)',
                                   (BandIDComps, CompID, CurrentDate, Grade))
                    conn.commit()

                    # get the competition details to send an email to the PM of band just entered
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    findBandGrade = 'SELECT CompetitionName, CompetitionLocation, CompetitionDate FROM Competitions WHERE CompetitionID = ? '
                    cursor.execute(findBandGrade, [CompID])
                    CompetitonResults = cursor.fetchall()

                    if CompetitonResults:
                        for i in CompetitonResults:
                            Name = i[0]
                            Location = i[1]
                            Date = i[2]

                    subject = "Competition entry confirmation"
                    body = "Dear Pipe Major, \n\nYou have successfully entered the following competition: " \
                           "\n\nCompetition: " + str(Name) + "\nCompetition Location: " + str(Location) + "\n" \
                                                                                                          "Competition Date: " + str(
                        Date) + "\n\n Thank you for your registration.\nWe look forward to seeing you at " + str(
                        Location) + ".\n\n" \
                                    "The Royal Scottish Pipe Band Association\nNorthern Ireland Branch"
                    EmailUser(Email, subject, body)

            messagebox.showinfo(message="You have successfully entered your selected competitions! You will receive an email as confirmation to these entries.")
            EnterCompFrame(photoSelectCompetitionsButton)

    def ListComps():
        def diff(list1, list2):
            # compare two lists and return a list of elements that dont appear in both.
            c = set(list1).union(set(list2))  # c = set(list1) | set(list2)
            d = set(list1).intersection(set(list2))  # d = set(list1) & set(list2)
            return list(c - d)

        def DisplayComps(BandIDComps):
            def EnteredComps():
                def WithdrawComp(listOfCompBox,listOfCompDateBox, listOfDeadLineDateBox):
                    # get all the positions of the Competitions selected
                    Selection = listOfCompBox.curselection()
                    # for each competition selected get the ID of the competiiton

                    if Selection == ():
                        messagebox.showinfo(message="Please select one competition you would like to withdraw from.")
                    else:
                        for i in Selection:
                            CompName = listOfCompBox.get(i)
                            CompDate = listOfCompDateBox.get(i)
                            CompEntDead = listOfDeadLineDateBox.get(i)

                        if messagebox.askyesno(message="Are you sure you want to withdraw from " + str(CompName)) == True:
                            with sqlite3.connect("RspbaniDB.db") as db:
                                cursor = db.cursor()
                            CompIDQuery = 'SELECT CompetitionID FROM Competitions WHERE CompetitionName = ? AND Drawn = "False"'
                            cursor.execute(CompIDQuery, [CompName])
                            CompID = cursor.fetchall()

                            try:
                                for e in CompID:
                                    for i in e:
                                        CompID = i

                                conn = sqlite3.connect('RspbaniDB.db')
                                with conn:
                                    cursor = conn.cursor()
                                cursor.execute(
                                    "DELETE FROM BandsEntered WHERE BandID = ? AND CompetitionID = ?",
                                    (BandIDComps, CompID,))
                                conn.commit()

                                with sqlite3.connect("RspbaniDB.db") as db:
                                    cursor = db.cursor()
                                EmailQuery = 'SELECT Emailaddress FROM BandAccount WHERE BandID = ?'
                                cursor.execute(EmailQuery, [BandIDComps])
                                EmailQueryResults = cursor.fetchall()

                                Email = EmailQueryResults[0][0]

                                subject = "Withdrawal confirmation"
                                body = "Dear Pipe Major, \n\nThis email is of confirmation of your withdrawal from " + str(CompName) + " " \
                                        "on " + str(CompDate) + ".\n If you wish to re-enter this competition, the entry deadline is " + str(CompEntDead) +". " \
                                        "\n\nThe Royal Scottish Pipe Band Association\nNorthern Ireland Branch"

                                EmailUser(Email, subject, body)
                                messagebox.showinfo(message="You have successfully withdrawn from " + str(CompName))

                            except:
                                messagebox.showinfo(
                                    message="Competition can not be withdrawn from due to the drawn having already taken "
                                            "place! You are still entered into " + str(CompName))

                    EnteredComps()

                EnteredComp = Frame(contentFrame, width=965, height=600)
                EnteredComp.grid(row=0, column=0, sticky="nsew")

                lblspacer = Label(EnteredComp, text="", width=15)
                lblspacer.grid()

                lblSpacer = Label(EnteredComp, text="Entered Competitions & Withdraw", height=3, font=("Arvo", 32))
                lblSpacer.grid(row=0, column=3, columnspan=6, sticky="ew")

                lblCompetitionName = Label(EnteredComp, text="Competition Name")
                lblCompetitionName.grid(row=2, column=4, columnspan=2, sticky=W)

                lblCompetitionDate = Label(EnteredComp, text="Competition Date")
                lblCompetitionDate.grid(row=2, column=6, columnspan=2, sticky=W)

                lblDeadLineDate = Label(EnteredComp, text="Entry Deadline")
                lblDeadLineDate.grid(row=2, column=8, columnspan=2, sticky=W)

                listOfCompBox = Listbox(EnteredComp, selectmode=SINGLE, width=35)
                listOfCompBox.grid(row=3, column=4, columnspan=2)

                listOfCompDateBox = Listbox(EnteredComp, selectmode=DISABLED, width=20, selectbackground="white")
                listOfCompDateBox.grid(row=3, column=6, columnspan=2)

                listOfDeadLineDateBox = Listbox(EnteredComp, selectmode=DISABLED, width=20, selectbackground="white")
                listOfDeadLineDateBox.grid(row=3, column=8, columnspan=2, sticky=W)

                # gets a list of all the competitions that have not yet been drawn and is still possible to enter
                with sqlite3.connect("RspbaniDB.db") as db:
                    # Creates a cursor to search through the data
                    cursorListOfComps = db.cursor()
                    # SQL Query to select all Competitions from the competitions table
                listOfCompsQuery = 'SELECT CompetitionID FROM BandsEntered WHERE BandID = ?'
                # Execute the query to get a list of all the bands
                cursorListOfComps.execute(listOfCompsQuery, [BandIDComps])
                # The query will produce a 2 dimensional tuple
                ListOfComps = cursorListOfComps.fetchall()

                for i in ListOfComps:
                    for j in i:
                        with sqlite3.connect("RspbaniDB.db") as db:
                            cursor = db.cursor()
                        listOfCompsEnteredQuery = 'SELECT CompetitionName, CompetitionDate, EntryDeadLine FROM Competitions WHERE CompetitionID = ?'
                        cursor.execute(listOfCompsEnteredQuery, [j])
                        ListOfCompetitionsEntered = cursor.fetchall()

                        for competition in ListOfCompetitionsEntered:
                            n = 0
                            listOfCompBox.insert(n, competition[0])
                            listOfCompDateBox.insert(n, competition[1])
                            listOfDeadLineDateBox.insert(n, competition[2])
                            n += 1

                ButtonWithdraw = Button(EnteredComp, image=photoWithdraw, command=lambda: WithdrawComp(listOfCompBox,listOfCompDateBox, listOfDeadLineDateBox))
                ButtonWithdraw.grid(row=4, column=4, columnspan=8, pady="10")

                BackButton = Button(EnteredComp, image=photoBack, command=lambda: DisplayComps(BandIDComps))
                BackButton.grid(row=0, column=0)

            EnterComp = Frame(contentFrame, width=965, height=600)
            EnterComp.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(EnterComp, text="", width=15)
            lblspacer.grid()

            lblSpacer = Label(EnterComp, text="Enter Competitions", height=3, font=("Arvo", 32))
            lblSpacer.grid(row=0, column=4, columnspan=6, sticky="ew")

            lblCompetitionName = Label(EnterComp, text="Competition Name")
            lblCompetitionName.grid(row=2, column=4, columnspan=2, sticky=W)

            lblCompetitionDate = Label(EnterComp, text="Competition Date")
            lblCompetitionDate.grid(row=2, column=6, columnspan=2, sticky=W)

            lblDeadLineDate = Label(EnterComp, text="Entry Deadline")
            lblDeadLineDate.grid(row=2, column=8, columnspan=2, sticky=W)

            # gets a list of all the competitions that have not yet been drawn and is still possible to enter
            with sqlite3.connect("RspbaniDB.db") as db:
                # Creates a cursor to search through the data
                cursorListOfComps = db.cursor()
                # SQL Query to select all Competitions from the competitions table
            listOfCompsQuery = 'SELECT CompetitionID FROM Competitions'
            # Execute the query to get a list of all the bands
            cursorListOfComps.execute(listOfCompsQuery)
            # The query will produce a 2 dimensional tuple
            ListOfComps = cursorListOfComps.fetchall()

            listOfComps = []
            print(listOfComps)
            for i in ListOfComps:
                for j in i:
                    listOfComps.append(str(j))
                    print(listOfComps)

            with sqlite3.connect("RspbaniDB.db") as db:
                # Creates a cursor to search through the data
                cursorListOfComps = db.cursor()
                # SQL Query to select all Competitions from the competitions table
            listOfCompsQuery = 'SELECT CompetitionID FROM Competitions WHERE Drawn = ?'
            # Execute the query to get a list of all the bands
            cursorListOfComps.execute(listOfCompsQuery, ["True"])
            # The query will produce a 2 dimensional tuple
            ListOfCompTrue = cursorListOfComps.fetchall()

            ListCompsAlreadyDrawn = []

            for i in ListOfCompTrue:
                for j in i:
                    ListCompsAlreadyDrawn.append(str(j))

            # get a list of the Competitions the Band has already entered
            with sqlite3.connect("RspbaniDB.db") as db:
                cursor = db.cursor()
            listOfCompsEnteredQuery = 'SELECT CompetitionID FROM BandsEntered WHERE BandID = ?'
            cursor.execute(listOfCompsEnteredQuery, [BandIDComps])
            ListOfCompsEntered = cursor.fetchall()

            listOfCompsEntered = []
            for i in ListOfCompsEntered:
                for j in i:
                    listOfCompsEntered.append(str(j))
                    print(listOfCompsEntered)

            # compare these two lists and return the competitions that havent been entered. This list still has competitions that cant be entered.
            CompsNotEntered = diff(listOfCompsEntered, listOfComps)
            # initialise a dictionary to store the competition name and date from all the competitions that can
            # still be entered

            for element in ListCompsAlreadyDrawn:
                try:
                    CompsNotEntered.remove(str(element))
                except:
                    None

            CompNameDateNotEntered = {}
            CompNameEntryDeadLineNotEntered = {}

            # for each competition get the name and date of it and append it to the dictionary where the
            # CompetitionID is the index
            for i in CompsNotEntered:
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursor = db.cursor()
                listOfCompsEnteredQuery = 'SELECT CompetitionName, CompetitionDate, EntryDeadLine FROM Competitions WHERE CompetitionID = ?'
                cursor.execute(listOfCompsEnteredQuery, [i])
                ListOfCompsNOTEntered = cursor.fetchall()

                for i in ListOfCompsNOTEntered:
                    CompNameDateNotEntered[i[0]] = i[1]
                    CompNameEntryDeadLineNotEntered[i[0]] = i[2]

            print("CompNameDateNotEntered")
            print(CompNameDateNotEntered)

            print("CompNameEntryDeadLineNotEntered")
            print(CompNameEntryDeadLineNotEntered)

            # create three list boxes; one to display the competition Name and the competition date and the Entry Deadline
            # Allow the user to select more than one competition to enter
            listOfCompBox = Listbox(EnterComp, selectmode=MULTIPLE, width=35)
            listOfCompBox.grid(row=3, column=4, columnspan=2)

            listOfCompDateBox = Listbox(EnterComp, selectmode=DISABLED, width=20)
            listOfCompDateBox.grid(row=3, column=6, columnspan=2)

            listOfDeadLineDateBox = Listbox(EnterComp, selectmode=DISABLED, width=20, selectbackground="white")
            listOfDeadLineDateBox.grid(row=3, column=8, columnspan=2, sticky=W)

            for k in CompNameDateNotEntered:
                n = 0
                # k = k[2:-3]
                print("k")
                print(k)
                print("K[0]")
                print(CompNameDateNotEntered[k])
                listOfCompBox.insert(n, k)
                listOfCompDateBox.insert(n, CompNameDateNotEntered[k])
                n = + 1

            for k in CompNameEntryDeadLineNotEntered:
                n = 0
                listOfDeadLineDateBox.insert(n, CompNameEntryDeadLineNotEntered[k])
                n = + 1

            listOfComps = []
            # For each index the tuple append to the new list
            for k in ListOfComps:
                listOfComps.append(k[0])

            lblText = Label(EnterComp,
                            text="Please select all the competitions you would like to enter \n Only the Competitions you havent entered will appear ",
                            height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=3, columnspan=8, sticky="ew")

            SelectButton = Button(EnterComp, image=photoSelectCompetitionsButton,
                                  command=lambda: selectedComps(now, BandIDComps, listOfCompBox))
            SelectButton.grid(row=4, column=4, columnspan=2)

            WithdrawButton = Button(EnterComp, image=photoViewEnteredCompsWithdraw, command=EnteredComps)
            WithdrawButton.grid(row=4, column=6, columnspan=3)

        if Admin == True:

            AdminBandIDComps = StringVar()
            # get the position of the selected Band
            Selection = EnteCompListBandBox.curselection()
            print("selection", Selection)
            # get the value of the position  (ie. selected Band)

            if Selection == ():
                messagebox.showinfo(message="Please select a band that you wish to make these changes on behalf of.")
            else:

                AdminBandSelected = EnteCompListBandBox.get(Selection[0])
                print(AdminBandSelected)

                # get the band ID of the selected Band
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCurrentBand = db.cursor()
                FindDetailsCurrentBand = 'SELECT bandID FROM BandAccount WHERE BandName = ?'
                cursorCurrentBand.execute(FindDetailsCurrentBand, [AdminBandSelected])
                CurrentBandDetails = cursorCurrentBand.fetchall()

                for q in CurrentBandDetails:
                    for y in q:
                        AdminBandIDComps = y
                print("AdminBandIDComps")
                print(AdminBandIDComps)
                # Display the list of comps frame with a list of competitions dependant on the Band the Admin selected
                DisplayComps(AdminBandIDComps)

        else:
            # Display the list of comps frame with a list of competitions dependant on the Band Logged in
            DisplayComps(BandIDofLoggedIn)

    if Admin == True:
        # If user is admin then ask them to select which band they wish to edit.
        EnterComp = Frame(contentFrame, width=965, height=600)
        EnterComp.grid(row=0, column=0, sticky="nsew")

        lblspacer = Label(EnterComp, text="", width=25)
        lblspacer.grid()

        lblSpacer = Label(EnterComp, text="Enter Competitions", height=3, font=("Arvo", 32))
        lblSpacer.grid(row=0, column=1, columnspan=8, sticky="ew")

        lblText = Label(EnterComp, text="Please select the Band to edit ", height=3, font=("Arvo", 14))
        lblText.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)
        # create a list box and populate it with all registered bands
        EnteCompListBandBox = Listbox(EnterComp, selectmode=EXTENDED, width=25)
        EnteCompListBandBox.grid(row=3, column=3, columnspan=4)

        for k in listOfBands:
            n = 0
            EnteCompListBandBox.insert(n, str(k))
            n = + 1

        SelectButton = Button(EnterComp, image=photoSelectBandButton, command=ListComps)
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    else:
        # Show a list of the competitions to the user
        ListComps()


def BookTransportFrame(photoSubmitButton):
    global PickDay, pickPoint

    def BookTransportPro(BandName):
        def CheckEntries(trip, PickDay, PickMonth, PickYear, pickPoint, pickTime, destination, RtnDay, RtnMonth,
                         RtnYear, returnTime, partySize, numInst):

            ETrip = trip.get()
            EPickDay = PickDay.get()
            EPickMonth = PickMonth.get()
            EPickYear = PickYear.get()
            EPickPoint = pickPoint.get()
            EpickTime = pickTime.get()
            Edestination = destination.get()
            ERtnDay = RtnDay.get()
            ERtnMonth = RtnMonth.get()
            ERtnYear = RtnYear.get()
            EReturnTime = returnTime.get()
            EpartySize = partySize.get()
            EnumInst = numInst.get()


            ETripCheck = False
            EPickDayCheck = False
            EPickMonthCheck = False
            EPickYearCheck = False
            EPickPointCheck = False
            EpickTimeCheck = False
            EdestinationCheck = False
            ERtnDayCheck = False
            ERtnMonthCheck = False
            ERtnYearCheck = False
            EReturnTimeCheck = False
            EpartySizeCheck = False
            EnumInstCheck = False

            if ETrip == "One Way" or "Return":
                ETripCheck = True
            if ETrip == "":
                ETripCheck = False
            if EPickPoint != "":
                EPickPointCheck = True
            if EPickDay != "Select Day":
                EPickDayCheck = True
            if EPickMonth != "Select Month":
                EPickMonthCheck = True
            if EPickYear != "Select Year":
                EPickYearCheck = True
            if EpickTime != "":
                EpickTimeCheck = True
            if Edestination != "":
                EdestinationCheck = True
            if ERtnDay != "Select Day":
                ERtnDayCheck = True
            if ERtnMonth != "Select Month":
                ERtnMonthCheck = True
            if ERtnYear != "Select Year":
                ERtnYearCheck = True
            if EReturnTime != "":
                EReturnTimeCheck = True
            if EpartySize != "":
                EpartySizeCheck = True
            if EnumInst != "":
                EnumInstCheck = True

            if ETripCheck == True:
                if EPickDayCheck == True:
                    if EPickMonthCheck == True:
                        if EPickYearCheck == True:
                            if EPickPointCheck == True:
                                if EpickTimeCheck == True:
                                    if EdestinationCheck == True:
                                        if ERtnDayCheck == True:
                                            if ERtnMonthCheck == True:
                                                if ERtnYearCheck == True:
                                                    if EReturnTimeCheck == True:
                                                        if EpartySizeCheck == True:
                                                            if EnumInstCheck == True:
                                                                generateDoc(ETrip, EPickDay, EPickMonth, EPickYear,
                                                                            EPickPoint, EpickTime, Edestination,
                                                                            ERtnDay, ERtnMonth,
                                                                            ERtnYear, EReturnTime, EpartySize, EnumInst)
                                                            else:
                                                                messagebox.showinfo(
                                                                    message="Number instruments field is not valid. Please select the number of intruments to transported")
                                                        else:
                                                            messagebox.showinfo(
                                                                message="Party size field is not valid. Please select your party size")
                                                    else:
                                                        messagebox.showinfo(
                                                            message="Return time field is not valid. Please select a time")
                                                else:
                                                    messagebox.showinfo(
                                                        message="Return year field is not valid. Please select a year")
                                            else:
                                                messagebox.showinfo(
                                                    message="Return month field is not valid. Please select a month")
                                        else:
                                            messagebox.showinfo(
                                                message="Return day field is not valid. Please select a day")
                                    else:
                                        messagebox.showinfo(
                                            message=" Destination field is not valid. Please enter your destination")
                                else:
                                    messagebox.showinfo(
                                        message="Pick-up time field is not valid. Please enter a pick-up time ")
                            else:
                                messagebox.showinfo(
                                    message="Pick-up point field is not valid. Please enter a pick-up point")
                        else:
                            messagebox.showinfo(
                                message="Pick-up year field is not valid. Please select a year")
                    else:
                        messagebox.showinfo(
                            message="Pick-up month field is not valid. Please select a month")
                else:
                    messagebox.showinfo(
                        message="Pick-up day field is not valid. Please select a day")
            else:
                messagebox.showinfo(
                    message="Trip Type field is not valid. Please ensure you have checked one of the radio buttons; one way or return")

        def gmail(pdfFileName):

            with sqlite3.connect("RspbaniDB.db") as db:
                cursorCurrentBand = db.cursor()
            FindDetailsCurrentBand = 'SELECT EmailAddress FROM BandAccount WHERE BandName = ?'
            cursorCurrentBand.execute(FindDetailsCurrentBand, [BandName])
            CurrentBandDetails = cursorCurrentBand.fetchall()

            for q in CurrentBandDetails:
                for y in q:
                    EmailAddress = y

            # email address of the sender
            email_user = 'rspbani.booktransport@gmail.com'
            # password of the sender
            email_password = 'P@55w0rd321'
            # email address of recipient
            email_send = EmailAddress

            # Creating the Emails subject
            subject = 'Transport Enquiry - ' + str(BandName)

            msg = MIMEMultipart()

            msg['From'] = email_user
            msg['To'] = email_send
            msg['Subject'] = subject

            # creating the contents of the email
            body = "Dear Pipe Major,\nPlease find attached a copy of your recent enquiry which has been sent to the bus company. " \
                   "They will be directly incontact with you in the near future to confirm this itinerary .\n\nThe Royal Scottish Pipe Band Association" \
                   "\nNorthern Ireland Branch"
            msg.attach(MIMEText(body, 'plain'))

            # adding the PDF to the email as an attachment
            filename = pdfFileName
            attachment = open(filename, 'rb')

            part = MIMEBase('application', 'octet-stream')
            part.set_payload((attachment).read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', "attachment; filename= " + filename)

            msg.attach(part)
            text = msg.as_string()

            try:
                # connet to email server and send email
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                server.login(email_user, email_password)

                server.sendmail(email_user, email_send, text)

                # clear the fields in the book transport frame
                EtyPickPoint.delete(0, END)
                PickDay.set('Select Day')
                PickMonth.set('Select Month')
                PickYear.set('Select Year')
                EtyPickTime.delete(0, END)
                EtyDestiniation.delete(0, END)
                RtnDay.set('Select Day')
                RtnMonth.set('Select Month')
                RtnYear.set('Select Year')
                EtyReturnTime.delete(0, END)
                EtyPartySize.delete(0, END)
                EtyNumInst.delete(0, END)

            # error handling if email cant be sent - display a message stating this
            except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                messagebox.showinfo("Alert ", message="Confirmation not sent to Pipe Major \n\nPlease try again.")
            finally:
                # stop connection with email server
                server.quit()

            # email address of the sender
            email_user = 'rspbani.booktransport@gmail.com'
            # password of the sender
            email_password = 'P@55w0rd321'
            # email address of recipient
            email_send = 'thecompanybustravel@gmail.com'

            # Creating the Emails subject
            subject = 'Transport Enquiry - ' + str(BandName)

            msg = MIMEMultipart()

            msg['From'] = email_user
            msg['To'] = email_send
            msg['Subject'] = subject

            # creating the contents of the email
            body = "Dear Sirs, \n\nPlease see attached enquiry for " + str(BandName) + " who wish to make travel arrangements with you. " \
                   "Please contact the band directly on the email address provided. \n\n The Royal Scottish Pipe Band Association" \
                    "\nNorthern Ireland Branch "
            msg.attach(MIMEText(body, 'plain'))

            # adding the document to the email as an attachment
            filename = pdfFileName
            attachment = open(filename, 'rb')

            part = MIMEBase('application', 'octet-stream')
            part.set_payload((attachment).read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', "attachment; filename= " + filename)

            msg.attach(part)
            text = msg.as_string()

            try:
                # connet to email server and send email
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                server.login(email_user, email_password)

                server.sendmail(email_user, email_send, text)
                messagebox.showinfo("Alert ", message="Enquiry Sent!\n\nThe company will be in contact with you through"
                                                      " the email registered to your account.")
                # clear the fields in the book transport frame
                EtyPickPoint.delete(0, END)
                PickDay.set('Select Day')
                PickMonth.set('Select Month')
                PickYear.set('Select Year')
                EtyPickTime.delete(0, END)
                EtyDestiniation.delete(0, END)
                RtnDay.set('Select Day')
                RtnMonth.set('Select Month')
                RtnYear.set('Select Year')
                EtyReturnTime.delete(0, END)
                EtyPartySize.delete(0, END)
                EtyNumInst.delete(0, END)

            # error handling if email cant be sent - display a message stating this
            except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                messagebox.showinfo("Alert ", message="Message not sent to Bus Company\n\nPlease try again.")
            finally:
                # stop connection with email server
                server.quit()

        def generateDoc(ETrip, PickDay, PickMonth, PickYear, pickPoint, pickTime, destination, RtnDay, RtnMonth,
                        RtnYear, returnTime, partySize, numInst):

            #create document and add logo and title
            document = Document()
            document.add_picture('RSPBANI_Logo.png', width=Inches(2))
            document.add_heading('              RSPBANI Booking Request', 0)

            # add text
            lblCompetition = document.add_paragraph()
            lblCompetition.add_run(str(
                BandName) + ' Would like to book transport with you. Please find their booking requirements below:   ').bold = True
            lblCompetition.add_run()

            #create table
            Draw = document.add_table(rows=1, cols=2)

            # add a row for each field
            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Trip Type:  '
            Content_cells[1].text = str(ETrip)

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Pick-up Location:  '
            Content_cells[1].text = str(pickPoint)

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Pick-up Date:  '
            Content_cells[1].text = str(PickDay) + " / " + str(PickMonth) + " / " + str(PickYear)

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Pick-up Departure Time:  '
            Content_cells[1].text = str(pickTime)

            Content_cells = Draw.add_row().cells

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Destination:  '
            Content_cells[1].text = str(destination)

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Return Date:  '
            Content_cells[1].text = str(RtnDay) + " / " + str(RtnMonth) + " / " + str(RtnYear)

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Return Departure Time:  '
            Content_cells[1].text = str(returnTime)

            Content_cells = Draw.add_row().cells

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Party Size:  '
            Content_cells[1].text = str(partySize)

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = 'Number Instruments:  '
            Content_cells[1].text = str(numInst)

            Content_cells = Draw.add_row().cells

            with sqlite3.connect("RspbaniDB.db") as db:
                cursorCurrentBand = db.cursor()
            FindPMEmailAddressQuery = 'SELECT EmailAddress FROM BandAccount WHERE BandName = ?'
            cursorCurrentBand.execute(FindPMEmailAddressQuery, [BandName])
            EmailAddressResults = cursorCurrentBand.fetchall()

            for q in EmailAddressResults:
                for y in q:
                    PMEmailAddress = y

            Content_cells = Draw.add_row().cells
            Content_cells[0].text = "Pipe Major's Email Address:  "
            Content_cells[1].text = str(PMEmailAddress)

            # save document and name it as:
            document.save('Transport booking enquiry for ' + str(BandName) + '.docx')
            fileName = 'Transport booking enquiry for ' + str(BandName) + '.docx'

            gmail(fileName)


        # initalise all the variables to be used to store the entered data (textvars)
        trip = StringVar()
        PickDay = StringVar()
        PickMonth = StringVar()
        PickYear = StringVar()
        pickPoint = StringVar()
        pickTime = StringVar()
        destination = StringVar()
        RtnDay = StringVar()
        RtnMonth = StringVar()
        RtnYear = StringVar()
        returnTime = StringVar()
        partySize = StringVar()
        numInst = StringVar()

        # populate the Frame will the fields to book transport
        BookTransport = Frame(contentFrame, width=965, height=600)
        BookTransport.grid(row=0, column=0, sticky="nsew")

        lblSpacer = Label(BookTransport, text="Book Transport", height=3, font=("Arvo", 32))
        lblSpacer.grid(row=0, column=1, columnspan=6, sticky="ew")

        Radiobutton(BookTransport, text="One Way", padx=5, variable=trip, value="One Way", font=("Arvo")).grid(row=1,
                                                                                                               column=2,
                                                                                                               sticky="w")
        Radiobutton(BookTransport, text="Return", padx=5, variable=trip, value="Return", font=("Arvo")).grid(row=1,
                                                                                                             column=2,
                                                                                                             sticky="e",
                                                                                                             padx=15)

        lblspacer = Label(BookTransport, text="", width=30)
        lblspacer.grid(column=0)

        lblPickPoint = Label(BookTransport, text="Pick up point: ", font=("Arvo"))
        lblPickPoint.grid(row=2, column=1, sticky="e")

        EtyPickPoint = Entry(BookTransport, textvar=pickPoint)
        EtyPickPoint.grid(row=2, column=2, sticky="w")

        lblPickUpDate = Label(BookTransport, text="Pick up Date: ", font=("Arvo"))
        lblPickUpDate.grid(row=3, column=1, sticky="e")

        listDays = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18',
                    '19', '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
        droplistDay = OptionMenu(BookTransport, PickDay, *listDays)
        droplistDay.config(width=20)
        PickDay.set('Select Day')
        droplistDay.grid(row=3, column=2, sticky="w")

        listMonths = ['April', 'May', 'June', 'July', 'August', 'September']
        droplistMonth = OptionMenu(BookTransport, PickMonth, *listMonths)
        droplistMonth.config(width=20)
        PickMonth.set('Select Month')
        droplistMonth.grid(row=4, column=2, sticky="w")

        listYears = Years
        droplistYear = OptionMenu(BookTransport, PickYear, *listYears)
        droplistYear.config(width=20)
        PickYear.set('Select Year')
        droplistYear.grid(row=5, column=2, sticky="w")

        lblPickTime = Label(BookTransport, text="Pick up time:", font=("Arvo"))
        lblPickTime.grid(row=6, column=1, sticky="e")

        EtyPickTime = Entry(BookTransport, textvar=pickTime)
        EtyPickTime.grid(row=6, column=2, sticky="w")

        lblDestination = Label(BookTransport, text="Destination", font=("Arvo"))
        lblDestination.grid(row=7, column=1, sticky="e")

        EtyDestiniation = Entry(BookTransport, textvar=destination)
        EtyDestiniation.grid(row=7, column=2, sticky="w")

        lblReturnDate = Label(BookTransport, text="Return Date:", font=("Arvo"))
        lblReturnDate.grid(row=8, column=1, sticky="e")

        droplistDayRtn = OptionMenu(BookTransport, RtnDay, *listDays)
        droplistDayRtn.config(width=20)
        RtnDay.set('Select Day')
        droplistDayRtn.grid(row=8, column=2, sticky="w")

        droplistMonthRtn = OptionMenu(BookTransport, RtnMonth, *listMonths)
        droplistMonthRtn.config(width=20)
        RtnMonth.set('Select Month')
        droplistMonthRtn.grid(row=9, column=2, sticky="w")

        droplistYearRtn = OptionMenu(BookTransport, RtnYear, *listYears)
        droplistYearRtn.config(width=20)
        RtnYear.set('Select Year')
        droplistYearRtn.grid(row=10, column=2, sticky="w")

        lblReturnTime = Label(BookTransport, text="Return Time:", font=("Arvo"))
        lblReturnTime.grid(row=11, column=1, sticky="e")

        EtyReturnTime = Entry(BookTransport, textvar=returnTime)
        EtyReturnTime.grid(row=11, column=2, sticky="w")

        lblPartySize = Label(BookTransport, text="Party Size:", font=("Arvo"))
        lblPartySize.grid(row=12, column=1, sticky="e")

        EtyPartySize = Entry(BookTransport, textvar=partySize)
        EtyPartySize.grid(row=12, column=2, sticky="w")

        lblNumInst = Label(BookTransport, text="Number of Instruments:", font=("Arvo"))
        lblNumInst.grid(row=13, column=1, sticky="e")

        EtyNumInst = Entry(BookTransport, textvar=numInst)
        EtyNumInst.grid(row=13, column=2, sticky="w")

        # submit button
        submit_mail = Button(BookTransport, image=photoSubmitButton,
                             command=lambda: CheckEntries(trip, PickDay, PickMonth, PickYear, pickPoint, pickTime,
                                                          destination, RtnDay, RtnMonth, RtnYear, returnTime, partySize,
                                                          numInst))
        submit_mail.grid(row=14, column=2, sticky="w", padx="28", pady="10")

    def getBandNameAdmin():

        # get the position of the selected band
        MemSelection = EditMemDListBandBox.curselection()
        print("selection", MemSelection)
        # get the value of the selected position

        if MemSelection == ():
            messagebox.showinfo(message="Please select a band that you wish to make these changes on behalf of.")
        else:

            BandNameAdminLogged = EditMemDListBandBox.get(MemSelection[0])

            print("BandNameAdminLogged")
            print(BandNameAdminLogged)
            # call the function with the Band the Admin choose
            BookTransportPro(BandNameAdminLogged)

    if Admin == True:
        # If user is admin then ask them to select which band they wish to edit.
        BookTransport = Frame(contentFrame, width=965, height=600)
        BookTransport.grid(row=0, column=0, sticky="nsew")

        lblspacer = Label(BookTransport, text="", width=20)
        lblspacer.grid(row=0, column=0)

        Title = Label(BookTransport, text="Edit Member Details ", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        Title = Label(BookTransport, text="Please select the Band to edit ", height=3, font=("Arvo", 14))
        Title.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)

        # create a list box and populate it with all registered bands
        EditMemDListBandBox = Listbox(BookTransport, selectmode=EXTENDED, width=25)
        EditMemDListBandBox.grid(row=3, column=3, columnspan=4)

        for k in listOfBands:
            print("k")
            print(k)
            n = 0
            EditMemDListBandBox.insert(n, str(k))
            n = + 1

        SelectButton = Button(BookTransport, image=photoSelectBandButton, command=getBandNameAdmin)
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    else:
        # call the function with the Band logged in
        BookTransportPro(BandnameofLoggedIn)


def TransferMemberFrame(photoSelectBandButton):
    def SaveTransfer(MemberID, NewBandName, FirstName, SecondName):
        # get bandID of current band and PM's email address to send him an email
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorNewBand = db.cursor()
        FindDetailsNewBand = 'SELECT BandID, EmailAddress FROM BandAccount WHERE BandName = ?'
        cursorNewBand.execute(FindDetailsNewBand, [NewBandName])
        NewBandDetails = cursorNewBand.fetchall()

        for q in NewBandDetails:
            BandID = q[0]
            PMEmailAddress = q[1]

            # update the BandID in the Members Table to link them with their new band
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            cursor.execute(
                'UPDATE BandMembers SET bandID = ? WHERE MemberID = ? ',
                (BandID, MemberID))
            conn.commit()

            # get email address of member to email them
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorNewBand = db.cursor()
            FindDetailsNewBand = 'SELECT EmailAddress FROM BandMembers WHERE MemberID = ?'
            cursorNewBand.execute(FindDetailsNewBand, [MemberID])
            MemberEmailQuery = cursorNewBand.fetchall()

            for q in MemberEmailQuery:
                MemberEmail = q[0]

            #sending email to Member
            # email address of the sender
            email_user = 'rspbani.info@gmail.com'
            # password of the sender
            email_password = 'P@55w0rd123'
            # email address of recipient
            email_send = MemberEmail
            # Creating the Emails subject
            subject = "Transfer Confirmation"

            msg = MIMEMultipart()
            msg['From'] = email_user
            msg['To'] = email_send
            msg['Subject'] = subject

            BodyText = "Dear " + str(
                FirstName) + ", \nYou are receiving this email as confirmation of your transfer to " + str(
                NewBandName) + "\n\nIf this transfer was not agreed upon, please contact our administrative team on 0829048539 to " \
                               "resolve this issue and revert this transfer. \n\n\nThe Royal Scottish Pipe Band Association \n" \
                               "Northern Ireland Branch"

            msg.attach(MIMEText(BodyText, 'plain'))
            text = msg.as_string()

            try:
                # connet to email server and send email
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                server.login(email_user, email_password)
                server.sendmail(email_user, email_send, text)

            # error handling if email cant be sent - display a message stating this
            except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
            finally:
                # stop connection with email server
                server.quit()

            #sending email to PM
            # email address of the sender
            email_user = 'rspbani.info@gmail.com'
            # password of the sender
            email_password = 'P@55w0rd123'
            # email address of recipient
            email_send = PMEmailAddress
            # Creating the Emails subject
            subject = "Transfer Confirmation"

            msg = MIMEMultipart()
            msg['From'] = email_user
            msg['To'] = email_send
            msg['Subject'] = subject

            BodyText = "Dear Pipe Major, \nYour recent member transfer is complete. " + str(FirstName) + " " + str(
                SecondName) + " has been transfered " \
                              "to " + str(
                NewBandName) + "\n\nIf this change was not made by you please contact our administrative team on 0829048539 to " \
                               "resolve this issue and revert this transfer. \n\n\nThe Royal Scottish Pipe Band Association \n" \
                               "Northern Ireland Branch"

            msg.attach(MIMEText(BodyText, 'plain'))
            text = msg.as_string()

            try:
                # connet to email server and send email
                server = smtplib.SMTP('smtp.gmail.com', 587)
                server.starttls()
                server.login(email_user, email_password)
                server.sendmail(email_user, email_send, text)

            # error handling if email cant be sent - display a message stating this
            except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
            finally:
                # stop connection with email server
                server.quit()
            HomeFrame()

    def TransferMemberCheck(NewBandList, FirstName, SecondName, MemberID):
        NewBandSelected = NewBandList.curselection()

        print("selection", NewBandSelected)

        if NewBandSelected == ():
            messagebox.showinfo(message="Please select the new band for " + str(FirstName))
        else:

            # gets the value of the position selected
            NewBandName = NewBandList.get(NewBandSelected[0])

            CheckTransfer = messagebox.askquestion(title="Transfer Member",
                                                   message="You are about to transfer " + str(FirstName) + " " + str(
                                                       SecondName) + " to " + str(
                                                       NewBandName) + ".\nAre you sure you wish to proceed? You can not revert this once complete.")

            if CheckTransfer == "yes":
                SaveTransfer(MemberID, NewBandName, FirstName, SecondName)

    def ShowTransferSelection(MemDListBandBox, FirstName, SecondName, BandSelected):
        MemberSelection = MemDListBandBox.curselection()

        if MemberSelection == ():
            messagebox.showinfo(message="Please select a member to transfer!")

        else:

            # gets the value of the position selected
            MemberSelected = MemDListBandBox.get(MemberSelection[0])

            MemberSelected = MemberSelected.split(" ")
            MemberSelected = MemberSelected[2]
            MemberIDSelected = MemberSelected[1:-1]

            TransferMember = Frame(contentFrame, width=965, height=600)
            TransferMember.grid(row=0, column=0, sticky="nsew")

            lblspacerCenter = Label(TransferMember, text="", width=15)
            lblspacerCenter.grid(row=0, column=0)

            Title = Label(TransferMember, text="Transfer Member", height=3, font=("Arvo", 32))
            Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

            lblText = Label(TransferMember,
                            text="Please select the band you would like to transfer " + str(FirstName) + " " + str(
                                SecondName) + " to: ", height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=1, columnspan=6, sticky="ew", padx=170)

            ListOfBands = []
            for i in listOfBands:
                ListOfBands.append(i)
            ListOfBands.remove(BandSelected)

            NewBandList = Listbox(TransferMember, selectmode=EXTENDED, width=25)
            NewBandList.grid(row=3, column=2, columnspan=4)

            n = 0
            for k in ListOfBands:
                n = 0
                NewBandList.insert(n, str(k))
                n = +1

            SelectButton = Button(TransferMember, image=photoSelectBandButton,
                                  command=lambda: TransferMemberCheck(NewBandList, FirstName, SecondName, MemberIDSelected))
            SelectButton.grid(row=4, column=3, columnspan=2, pady="10")

    def SelectMemberTransferFromList(ListBandBox):
        BandSelection = ListBandBox.curselection()
        print("selection", BandSelection)

        if BandSelection == ():
            messagebox.showinfo(message="Please select a band that you wish to make these changes on behalf of.")
        else:

            # gets the value of the position selected
            AdminBandSelected = ListBandBox.get(BandSelection[0])

            SelectMemberTransfer(AdminBandSelected)

    def SelectMemberTransfer(BandSelected):

        TransferMember = Frame(contentFrame, width=965, height=600)
        TransferMember.grid(row=0, column=0, sticky="nsew")

        lblspacerCenter = Label(TransferMember, text="", width=20)
        lblspacerCenter.grid(row=0, column=0)

        Title = Label(TransferMember, text="Transfer Member", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        lblText = Label(TransferMember, text="Please select the member to complete the transfer with ", height=3,
                        font=("Arvo", 14))
        lblText.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)

        MemDListBandBox = Listbox(TransferMember, selectmode=EXTENDED, width=25)
        MemDListBandBox.grid(row=3, column=3, columnspan=4)

        with sqlite3.connect("RspbaniDB.db") as db:
            cursor = db.cursor()
        FindBandID = 'SELECT BandID FROM BandAccount WHERE BandName = ?'
        cursor.execute(FindBandID, [BandSelected])
        BandIDSearch = cursor.fetchall()

        for q in BandIDSearch:
            BandID = q[0]

        with sqlite3.connect("RspbaniDB.db") as db:
            cursorCurrentBand = db.cursor()
        FindDetailsCurrentBand = 'SELECT MemberID, firstName, secondName FROM BandMembers WHERE bandID = ?'
        cursorCurrentBand.execute(FindDetailsCurrentBand, [BandID])
        CurrentBandDetails = cursorCurrentBand.fetchall()
        n = 0
        for q in CurrentBandDetails:
            MemberID = q[0]
            FirstName = q[1]
            SecondName = q[2]
            FullName = str(FirstName) + " " + str(SecondName) + " (" + str(MemberID) + ")"
            MemDListBandBox.insert(n, str(FullName))
            n += 1

        SelectButton = Button(TransferMember, image=photoSelectBandButton,
                              command=lambda: ShowTransferSelection(MemDListBandBox, FirstName, SecondName,
                                                                    BandSelected))
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    if Admin == True:
        # If user is admin then ask them to select which band they wish to edit.

        AdminTransferMember = Frame(contentFrame, width=965, height=600)
        AdminTransferMember.grid(row=0, column=0, sticky="nsew")

        lblspacerCenter = Label(AdminTransferMember, text="", width=20)
        lblspacerCenter.grid(row=0, column=0)

        Title = Label(AdminTransferMember, text="Transfer Member", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        lblText = Label(AdminTransferMember, text="Please select the band to complete the transfer with ", height=3,
                        font=("Arvo", 14))
        lblText.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)

        ListBandBox = Listbox(AdminTransferMember, selectmode=EXTENDED, width=25)
        ListBandBox.grid(row=3, column=3, columnspan=4)

        for k in listOfBands:
            n = 0
            ListBandBox.insert(n, str(k))
            n = +1

        SelectButton = Button(AdminTransferMember, image=photoSelectBandButton,
                              command=lambda: SelectMemberTransferFromList(ListBandBox))
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")


    else:

        with sqlite3.connect("RspbaniDB.db") as db:
            cursorCurrentBand = db.cursor()
        FindDetailsCurrentBand = 'SELECT BandName FROM BandAccount WHERE BandID = ?'
        cursorCurrentBand.execute(FindDetailsCurrentBand, [BandIDofLoggedIn])
        CurrentBandDetails = cursorCurrentBand.fetchall()
        n = 0
        for q in CurrentBandDetails:
            BandName = q[0]
        SelectMemberTransfer(BandName)


def EditMemberDetailsFrame(photoUpdateMemberButton, photoSelectBandButton, photoSelectMemberButton):
    def DeleteRecordFromDatabase(ID, name):
        # ask them if they want to delete the user
        Answer = messagebox.askquestion("Delete user?", "Are you sure you want to delete " + name)
        if Answer == "yes":
            # Delete user from database
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            cursor.execute("DELETE FROM BandMembers WHERE MemberID = ?", (ID,))
            conn.commit()
            EditMemberDetailsFrame(photoUpdateMemberButton, photoSelectBandButton, photoSelectMemberButton)

    def AdminDisplayMemberList():

        if Admin == True:
            # Gets the position selected
            MemSelection = EditMemDListBandBox.curselection()
            print("selection", MemSelection)

            if MemSelection == ():
                messagebox.showinfo(message="Please select a band that you wish to make these changes on behalf of.")
            else:

                # gets the value of the position selected
                AdminBandSelected = EditMemDListBandBox.get(MemSelection[0])
                # get band ID
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCurrentBand = db.cursor()
                FindDetailsCurrentBand = 'SELECT bandID FROM BandAccount WHERE BandName = ?'
                cursorCurrentBand.execute(FindDetailsCurrentBand, [AdminBandSelected])
                CurrentBandDetails = cursorCurrentBand.fetchall()

                for q in CurrentBandDetails:
                    for y in q:
                        AdminBandIDEditMem = y

                DisplayMemberList(AdminBandIDEditMem)

    def DisplayMemberList(BandID):

        EditMemberDetails = Frame(contentFrame, width=965, height=600)
        EditMemberDetails.grid(row=0, column=0, sticky="nsew")

        lblspacer = Label(EditMemberDetails, text="", width=20)
        lblspacer.grid(row=0, column=0)

        Title = Label(EditMemberDetails, text="Edit Member Details ", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        Title = Label(EditMemberDetails, text="Please select the member to edit ", height=3, font=("Arvo", 14))
        Title.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)

        global EditMemDBandListBox
        # Creates a list box to display Members Names
        EditMemDBandListBox = Listbox(EditMemberDetails, selectmode=EXTENDED)
        EditMemDBandListBox.grid(row=3, column=4, columnspan=2)

        MembersDetails = {}
        # gets firstName, secondName, MemberID which are saved as a tuple
        with sqlite3.connect("RspbaniDB.db") as db:
            cursorBandPage = db.cursor()
        BandPageInfoGetQuery = 'SELECT firstName, secondName, MemberID FROM BandMembers WHERE bandID = ?'
        cursorBandPage.execute(BandPageInfoGetQuery, [BandID])
        BandPageInfoGet = cursorBandPage.fetchall()

        print("BandPageInfoGet", BandPageInfoGet)

        for bandRecord in BandPageInfoGet:
            n = 0
            # concatenates first and second name together
            MemberName = str(bandRecord[0] + " " + bandRecord[1])
            # Inserts the full name into the list box
            EditMemDBandListBox.insert(n, MemberName)
            # Maps the Members ID to the Members name
            MembersDetails[MemberName] = bandRecord[2]
            n = + 1

        print("MembersDetails")
        print(MembersDetails)

        SelectButton = Button(EditMemberDetails, image=photoSelectMemberButton,
                              command=lambda: SpecMembPage(Years, MembersDetails))
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    def UpdateMembersDetailsDatabase(Title, FirstName, SecondName, AddressLine1, AddressLine2, County, Postcode,
                                     DateOfBirth, Gender, Signature, Email, MemberID):

        # update the details of the user with these new variables
        conn = sqlite3.connect('RspbaniDB.db')
        with conn:
            cursor = conn.cursor()
        cursor.execute(
            'UPDATE BandMembers SET Title = ?, firstName = ?, secondName = ?, addressLine1 = ?, addressLine2 = ?,'
            'county = ?, postcode = ?, dateOfBirth = ?, Gender = ?, signature = ?, EmailAddress = ? WHERE MemberID = ? ',
            (Title, FirstName, SecondName, AddressLine1, AddressLine2, County, Postcode, DateOfBirth, Gender, Signature,
             Email,
             MemberID))
        conn.commit()

        with sqlite3.connect("RspbaniDB.db") as db:
            cursorCurrentMember = db.cursor()
        FindDetailsCurrentMember = 'SELECT bandID FROM BandMembers WHERE MemberID = ?'
        cursorCurrentMember.execute(FindDetailsCurrentMember, [MemberID])
        CurrentMemberDetails = cursorCurrentMember.fetchall()

        for i in CurrentMemberDetails:
            for k in i:
                BandID = k

        with sqlite3.connect("RspbaniDB.db") as db:
            cursorCurrentMember = db.cursor()
        FindDetailsCurrentMember = 'SELECT EmailAddress, BandName FROM BandAccount WHERE BandID = ?'
        cursorCurrentMember.execute(FindDetailsCurrentMember, [BandID])
        CurrentBandDetails = cursorCurrentMember.fetchall()

        for i in CurrentBandDetails:
            PMEmail = i[0]
            BandName = i[1]

        # email address of the sender
        email_user = 'rspbani.info@gmail.com'
        # password of the sender
        email_password = 'P@55w0rd123'
        # email address of recipient
        email_send = PMEmail
        # Creating the Emails subject
        subject = 'Welcome ' + str(FirstName)

        msg = MIMEMultipart()
        msg['From'] = email_user
        msg['To'] = email_send
        msg['Subject'] = subject

        # creating the contents of the email
        body = "Dear Pipe Major, \n\n\nYou are receiving this email to confirm the changes you have made about a member in your band, " \
               "If this change was not made by you please contact one of the systems administrators so we can resolve this issue. \n\n\n\nThe Royal Scottish Pipe " \
               "Band Association \nNorthern Ireland Branch "
        msg.attach(MIMEText(body, 'plain'))
        text = msg.as_string()

        try:
            # connet to email server and send email
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(email_user, email_password)
            server.sendmail(email_user, email_send, text)

        # error handling if email cant be sent - display a message stating this
        except(smtplib.SMTPException, ConnectionRefusedError, OSError):
            messagebox.showinfo("Alert ", message="Message not sent to PM\n\n")
        finally:
            # stop connection with email server
            server.quit()

        # email address of the sender
        email_user = 'rspbani.info@gmail.com'
        # password of the sender
        email_password = 'P@55w0rd123'
        # email address of recipient
        email_send = Email
        # Creating the Emails subject
        subject = "Your Details have been updated"
        msg = MIMEMultipart()
        msg['From'] = email_user
        msg['To'] = email_send
        msg['Subject'] = subject
        # creating the contents of the email
        body = "Dear Member of " + str(
            BandName) + ", \n\n\nYou are receiving this email to confirm the changes that have been made about you. " \
                        "\n\nTitle:   " + str(Title) + "\nFirst Name:   " + str(FirstName) + "\n" \
                                                                                             "Second Name:   " + str(
            SecondName) + "\nAddress Line One:   " + str(AddressLine1) + "\n" \
                                                                         "Address Line Two:   " + str(
            AddressLine2) + "\n" \
                            "County:   " + str(County) + "\n" \
                                                         "Postcode:   " + str(Postcode) + "\n" \
                                                                                          "Date of Birth:   " + str(
            DateOfBirth) + "\n" \
                           "Gender:   " + str(Gender) + "\n" \
                                                        "Please check your details below and if they are not correct please contact your Pipe Major and he will update your details." \
                                                        " \n\n\n\nThe Royal Scottish Pipe Band Association \nNorthern Ireland Branch "
        msg.attach(MIMEText(body, 'plain'))
        text = msg.as_string()
        try:
            # connet to email server and send email
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(email_user, email_password)
            server.sendmail(email_user, email_send, text)
        # error handling if email cant be sent - display a message stating this
        except(smtplib.SMTPException, ConnectionRefusedError, OSError):
            messagebox.showinfo("Alert ", message="Message not sent to member")
        finally:
            # stop connection with email server
            server.quit()

        messagebox.showinfo(message="Band member successfully updated")
        # Brings the user back to the list of Members
        EditMemberDetailsFrame(photoUpdateMemberButton, photoSelectBandButton, photoSelectMemberButton)
        getListOfBands()

    def SpecMembPage(Years, MembersDetails):
        def CheckEditMemberDetails(UptTitle, UptFirstName, UptSecondName, UptAddressLine1, UptAddressLine2, UptCounty,
                                   UptPostcode, UptDateOfBirthDay, UptDateOfBirthMonth, UptDateOfBirthYear, UptGender,
                                   UptEmailAddress, UptSignature, MemberIDSelected):

            # get all the data from the entry boxes
            Title = UptTitle.get()
            FirstName = UptFirstName.get()
            SecondName = UptSecondName.get()
            AddressLine1 = UptAddressLine1.get()
            AddressLine2 = UptAddressLine2.get()
            County = UptCounty.get()
            Postcode = UptPostcode.get()
            DateOfBirthDay = UptDateOfBirthDay.get()
            DateOfBirthMonth = UptDateOfBirthMonth.get()
            DateOfBirthYear = UptDateOfBirthYear.get()

            UptDateOfBirth = str(DateOfBirthDay) + "/" + str(DateOfBirthMonth) + "/" + str(DateOfBirthYear)

            print("DateOfBirthDay")
            print(DateOfBirthDay)
            print(DateOfBirthMonth)
            print(DateOfBirthYear)

            Gender = UptGender.get()
            Email = UptEmailAddress.get()
            Signature = UptSignature.get()

            TitleCheck = False
            FirstNameCheck = False
            SecondNameCheck = False
            AddressLine1Check = False
            AddressLine2Check = True
            CountyCheck = False
            PostcodeCheck = False
            DOBDayCheck = False
            DOBMonthCheck = False
            DOBYearCheck = False
            GenderCheck = False
            EmailCheck = False
            SignatureCheck = False

            if Title != "Select Title":
                TitleCheck = True
            if CheckAllAlpha(FirstName):
                FirstNameCheck = True
            if CheckAllAlpha(SecondName):
                SecondNameCheck = True
            if CheckOnlyNumAndLetters(AddressLine1) == True:
                AddressLine1Check = True
            if County != "Select Country":
                CountyCheck = True
            if CheckPostcode(Postcode):
                PostcodeCheck = True
            if DateOfBirthDay != "Select Day":
                DOBDayCheck = True
            if DateOfBirthMonth != "Select Month":
                DOBMonthCheck = True
            if DateOfBirthYear != "Select Year":
                DOBYearCheck = True
            if Gender != "Select your Gender":
                GenderCheck = True
            if CheckEmail(Email):
                EmailCheck = True
            if CheckAllAlpha(Signature):
                SignatureCheck = True

            if TitleCheck == True:
                if FirstNameCheck == True:
                    if SecondNameCheck == True:
                        if AddressLine1Check == True:
                            if AddressLine2Check == True:
                                if CountyCheck == True:
                                    if PostcodeCheck == True:
                                        if DOBDayCheck == True:
                                            if DOBMonthCheck == True:
                                                if DOBYearCheck == True:
                                                    if GenderCheck == True:
                                                        if EmailCheck == True:
                                                            if SignatureCheck == True:
                                                                UpdateMembersDetailsDatabase(Title, FirstName,
                                                                                             SecondName, AddressLine1,
                                                                                             AddressLine2,
                                                                                             County, Postcode,
                                                                                             UptDateOfBirth, Gender,
                                                                                             Signature, Email,
                                                                                             MemberIDSelected)

                                                            else:
                                                                messagebox.showinfo(
                                                                    message="Members signature field is not valid. Please re-enter it")
                                                        else:
                                                            messagebox.showinfo(
                                                                message="Email Address field is not valid. Please re-enter it")
                                                    else:
                                                        messagebox.showinfo(
                                                            message="Gender field is not valid. Please re-enter it")
                                                else:
                                                    messagebox.showinfo(
                                                        message="Year of birth field is not valid. Please re-enter it")
                                            else:
                                                messagebox.showinfo(
                                                    message="Month of birth field is not valid. Please re-enter it")
                                        else:
                                            messagebox.showinfo(
                                                message="Day of birth field is not valid. Please re-enter it")
                                    else:
                                        messagebox.showinfo(
                                            message="Postcode field is not valid. Please re-enter it")
                                else:
                                    messagebox.showinfo(
                                        message="County field is not valid. Please re-enter it")
                            else:
                                messagebox.showinfo(
                                    message="Address line 2 field is not valid. Please re-enter it")
                        else:
                            messagebox.showinfo(
                                message="Address line 1 field is not valid. Please include your house number and road.")
                    else:
                        messagebox.showinfo(
                            message="Second name field is not valid. Please re-enter it")
                else:
                    messagebox.showinfo(
                        message="First name field is not valid. Please re-enter it")
            else:
                messagebox.showinfo(
                    message="Title field is not valid. Please re-enter it")


        # gets the position of the selection
        Selection = EditMemDBandListBox.curselection()
        print("selection", Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select a member!")
        else:

            EditMember = Frame(contentFrame, width=965, height=600)
            EditMember.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(EditMember, text="", width=15)
            lblspacer.grid()

            # gets the value of the selection
            MemberSelected = EditMemDBandListBox.get(Selection[0])

            # Gets the ID of the Member Selected
            MemberIDSelected = MembersDetails[MemberSelected]

            print(MemberSelected)
            print("MemberIDSelected")
            print(MemberIDSelected)

            # Initalising Variables to be used as textvar variables.
            UptTitle = StringVar()
            UptFirstName = StringVar()
            UptSecondName = StringVar()
            UptRole = StringVar()
            UptAddressLine1 = StringVar()
            UptAddressLine2 = StringVar()
            UptCounty = StringVar()
            UptPostcode = StringVar()
            UptDateOfBirthDay = StringVar()
            UptDateOfBirthMonth = StringVar()
            UptDateOfBirthYear = StringVar()
            UptGender = StringVar()
            UptSignature = StringVar()
            UptEmailAddress = StringVar()

            lblSpacer = Label(EditMember, text="Edit Member Details", height=3, font=("Arvo", 32))
            lblSpacer.grid(row=0, column=1, columnspan=6, sticky="ew")

            # Get all the members details dependant on their ID
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorCurrentMember = db.cursor()
            FindDetailsCurrentMember = 'SELECT * FROM BandMembers WHERE MemberID = ?'
            cursorCurrentMember.execute(FindDetailsCurrentMember, [MemberIDSelected])
            CurrentMemberDetails = cursorCurrentMember.fetchall()
            print("CurrentMemberDetails", CurrentMemberDetails)

            for i in CurrentMemberDetails:
                # Assign each field to a variable to be used to populate each entry box
                # The order of the fields depends on the order of the fields in the database
                CurrRole = i[2]
                CurrTitle = i[3]
                CurrFirstName = i[4]
                CurrSecondName = i[5]
                CurrAddressLine1 = i[6]
                CurrAddressLine2 = i[7]
                CurrCounty = i[8]
                CurrPostcode = i[9]
                CurrDateOfBirth = i[10]
                CurrGender = i[11]
                # CurrSignature = i[15]
                CurrEmailAddress = i[13]

            print("CurrFirstName", CurrFirstName)
            getListOfBands()

            # split the date of birth at evert "/"
            CurrDateOfBirth = CurrDateOfBirth.split('/')
            # The first part is the day
            CurrDateOfBirthDay = CurrDateOfBirth[0]
            # The second part is the Month
            CurrDateOfBirthMonth = CurrDateOfBirth[1]
            # The last part is the year
            CurrDateOfBirthYear = CurrDateOfBirth[2]

            print("CurrDateOfBirthDay", CurrDateOfBirthDay)
            print("CurrDateOfBirthMonth", CurrDateOfBirthMonth)
            print("CurrDateOfBirthYear", CurrDateOfBirthYear)


            lblspacer = Label(EditMember)
            lblspacer.grid(row=8)

            lblTitle = Label(EditMember, text="Title:", font=("Arvo"))
            lblTitle.grid(row=9, column=1, sticky="e")

            listTitles = ['Mr', 'Mrs', 'Miss', 'Ms', 'Dr', 'Rev', 'Master']
            droplistTitle = OptionMenu(EditMember, UptTitle, *listTitles)
            droplistTitle.config(width=20)
            UptTitle.set(CurrTitle)
            droplistTitle.grid(row=9, column=2, sticky="w")

            lblFirstName = Label(EditMember, text="First Name:", font=("Arvo"))
            lblFirstName.grid(row=10, column=1, sticky="e")

            EtyUptFirstName = Entry(EditMember, textvar=UptFirstName)
            # insert the value from the database at position 0 in the entry box
            EtyUptFirstName.insert(0, CurrFirstName)
            EtyUptFirstName.grid(row=10, column=2, sticky="w")

            lblSecondName = Label(EditMember, text="Second Name:", font=("Arvo"))
            lblSecondName.grid(row=11, column=1, sticky="e")

            EtyUptSecondName = Entry(EditMember, textvar=UptSecondName)
            EtyUptSecondName.insert(0, CurrSecondName)
            EtyUptSecondName.grid(row=11, column=2, sticky="w")

            lblGender = Label(EditMember, text="Gender:", font=("Arvo"))
            lblGender.grid(row=12, column=1, sticky="e")

            listGender = ['Male', 'Female', 'Other']
            EtyGender = OptionMenu(EditMember, UptGender, *listGender)
            EtyGender.config(width=20)
            UptGender.set(CurrGender)
            EtyGender.grid(row=12, column=2, sticky="w")

            lblAddress = Label(EditMember, text="Address", font=("Arvo"))
            lblAddress.grid(row=9, column=3, sticky="e")

            lblAddressLine1 = Label(EditMember, text="Line 1:", font=("Arvo"))
            lblAddressLine1.grid(row=10, column=3, sticky="e")

            EtyUptAddressLine1 = Entry(EditMember, textvar=UptAddressLine1)
            EtyUptAddressLine1.insert(0, CurrAddressLine1)
            EtyUptAddressLine1.grid(row=10, column=4, sticky="w")

            lblAddressLine2 = Label(EditMember, text="Line 2:", font=("Arvo"))
            lblAddressLine2.grid(row=11, column=3, sticky="e")

            EtyUptAddressLine2 = Entry(EditMember, textvar=UptAddressLine2)
            EtyUptAddressLine2.insert(0, CurrAddressLine2)
            EtyUptAddressLine2.grid(row=11, column=4, sticky="w")

            lblCounty = Label(EditMember, text="County:", font=("Arvo"))
            lblCounty.grid(row=12, column=3, sticky="e")
            listCounty = ['County Antrim', 'County Down', 'County Armagh', 'County Fermanagh', 'County Londonderry',
                          'Ireland']
            droplist = OptionMenu(EditMember, UptCounty, *listCounty)
            droplist.config(width=20)
            # Set the drop down with the current value stored in the database
            UptCounty.set(CurrCounty)
            droplist.grid(row=12, column=4, sticky="w")

            lblPostcode = Label(EditMember, text="Postcode:", font=("Arvo"))
            lblPostcode.grid(row=13, column=3, sticky="e")

            EtyUptPostcode = Entry(EditMember, textvar=UptPostcode)
            EtyUptPostcode.insert(0, CurrPostcode)
            EtyUptPostcode.grid(row=13, column=4, sticky="w")

            lblDateOfBirth = Label(EditMember, text="Date of Birth:", font=("Arvo"))
            lblDateOfBirth.grid(row=14, column=1, sticky="e")

            Days = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19',
                    '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
            listDays = Days
            droplistDOBDay = OptionMenu(EditMember, UptDateOfBirthDay, *listDays)
            droplistDOBDay.config(width=20)
            UptDateOfBirthDay.set(CurrDateOfBirthDay)
            droplistDOBDay.grid(row=14, column=2, sticky="w")

            Months = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
                      'November', 'December']
            listMonths = Months
            droplistDOBMonth = OptionMenu(EditMember, UptDateOfBirthMonth, *listMonths)
            droplistDOBMonth.config(width=20)
            UptDateOfBirthMonth.set(CurrDateOfBirthMonth)
            droplistDOBMonth.grid(row=15, column=2, sticky="w")

            listYears = Years
            droplistDOBYear = OptionMenu(EditMember, UptDateOfBirthYear, *listYears)
            droplistDOBYear.config(width=20)
            UptDateOfBirthYear.set(CurrDateOfBirthYear)
            droplistDOBYear.grid(row=16, column=2, sticky="w")

            lblEmailAddress = Label(EditMember, text="Email Address:")
            lblEmailAddress.grid(row=14, column=3, sticky="e")

            EtyUptEmailAddress = Entry(EditMember, textvar=UptEmailAddress)
            EtyUptEmailAddress.insert(0, CurrEmailAddress)
            EtyUptEmailAddress.grid(row=14, column=4, sticky="w")

            lblSignature = Label(EditMember, text="Member's Signature:", font=("Arvo"))
            lblSignature.grid(row=16, column=3, sticky="e")

            EtyUptSignature = Entry(EditMember, textvar=UptSignature)
            EtyUptSignature.grid(row=16, column=4, sticky="w")

            lblspacer = Label(EditMember)
            lblspacer.grid(row=17)

            Button(EditMember, image=photoDeleteButton,
                   command=lambda: DeleteRecordFromDatabase(MemberIDSelected, MemberSelected)).grid(row=20, column=1,
                                                                                                    sticky="w", padx="20",
                                                                                                    pady="10")

            Button(EditMember, image=photoUpdateMemberButton,
                   command=lambda: CheckEditMemberDetails(UptTitle, UptFirstName, UptSecondName, UptAddressLine1,
                                                          UptAddressLine2, UptCounty, UptPostcode, UptDateOfBirthDay,
                                                          UptDateOfBirthMonth, UptDateOfBirthYear, UptGender,
                                                          UptEmailAddress, UptSignature, MemberIDSelected)).grid(row=20,
                                                                                                                 column=3,
                                                                                                                 sticky="w",
                                                                                                                 padx="20",
                                                                                                                 pady="10")

    if Admin == True:
        # If user is admin then ask them to select which band they wish to edit.
        EditMemberDetails = Frame(contentFrame, width=965, height=600)
        EditMemberDetails.grid(row=0, column=0, sticky="nsew")

        lblspacer = Label(EditMemberDetails, text="", width=20)
        lblspacer.grid(row=0, column=0)

        Title = Label(EditMemberDetails, text="Edit Member Details ", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        Title = Label(EditMemberDetails, text="Please select the Band to edit ", height=3, font=("Arvo", 14))
        Title.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)

        EditMemDListBandBox = Listbox(EditMemberDetails, selectmode=EXTENDED, width=25)
        EditMemDListBandBox.grid(row=3, column=3, columnspan=4)

        for k in listOfBands:
            print("k")
            print(k)
            n = 0
            EditMemDListBandBox.insert(n, str(k))
            n = + 1

        SelectButton = Button(EditMemberDetails, image=photoSelectBandButton, command=lambda: AdminDisplayMemberList())
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    else:

        DisplayMemberList(BandIDofLoggedIn)


def RegMemberFrame(photoRegisterMemberFormButton):
    def SelectBandPage():

        global AdminBandIDRegMem
        # Initialise variable
        AdminBandIDRegMem = StringVar()
        # get position of band selected
        Selection = EditMemDListBandBox.curselection()
        print("selection", Selection)

        if Selection ==():
            messagebox.showinfo(message="Please select a band that you wish to make these changes on behalf of.")
        else:

            # get the value of the position selected (the Band name)
            AdminBandSelectedRegMem = EditMemDListBandBox.get(Selection[0])
            print(AdminBandSelectedRegMem)

            # gets the Band ID from the Band Name
            with sqlite3.connect("RspbaniDB.db") as db:
                cursorCurrentBand = db.cursor()
            FindDetailsCurrentBand = 'SELECT bandID FROM BandAccount WHERE BandName = ?'
            cursorCurrentBand.execute(FindDetailsCurrentBand, [AdminBandSelectedRegMem])
            CurrentBandDetails = cursorCurrentBand.fetchall()

            for q in CurrentBandDetails:
                for y in q:
                    AdminBandIDRegMem = y

            print(AdminBandIDRegMem)

            RegisterMemberPage(Years)

    def RegisterMemberPage(Years):
        def CheckMemberDetails(EtyFirstName, EtySecondName, EtyAddressLine1, EtyAddressLine2, EtyPostcode, EtySignature,
                               Memrole, MemfirstName,
                               MemsecondName, MemaddressLine1, MemaddressLine2, Memcounty, Mempostcode, MemdateOfBirth,
                               Memsignature, MemDOBDay,
                               MemDOBMonth, MemDOBYear, EtyEmailAddress, MemEmailAddress, Memtitle, Memgender):

            if Admin == True:

                # Assigns the BandID as the BandID which the admin has selected.
                BandID = AdminBandIDRegMem

            else:

                # Assigns the BandID as the ID of the band logged in.
                BandID = BandIDofLoggedIn

            Role = Memrole.get()
            Title = Memtitle.get()
            FirstName = MemfirstName.get()
            SecondName = MemsecondName.get()
            AddressLine1 = MemaddressLine1.get()
            AddressLine2 = MemaddressLine2.get()
            County = Memcounty.get()
            Postcode = Mempostcode.get()
            DOBDay = MemDOBDay.get()
            DOBMonth = MemDOBMonth.get()
            DOBYear = MemDOBYear.get()
            Gender = Memgender.get()
            # Concatenate Day, Month and year to save under one field in the database
            DateOfBirth = str(DOBDay) + "/" + str(DOBMonth) + "/" + str(DOBYear)
            Signature = Memsignature.get()
            Email = MemEmailAddress.get()

            BandIDCheck = True
            RoleCheck = False
            TitleCheck = False
            FirstNameCheck = False
            SecondNameCheck = False
            AddressLine1Check = False
            AddressLine2Check = True
            CountyCheck = False
            PostcodeCheck = False
            DOBDayCheck = False
            DOBMonthCheck = False
            DOBYearCheck = False
            GenderCheck = False
            SignatureCheck = False
            EmailCheck = False

            if Role != "":
                RoleCheck = True
            if Title != "Select Title":
                TitleCheck = True
            if CheckAllAlpha(FirstName):
                FirstNameCheck = True
            if CheckAllAlpha(SecondName):
                SecondNameCheck = True
            if CheckOnlyNumAndLetters(AddressLine1):
                AddressLine1Check = True
            if County != "Select Country":
                CountyCheck = True
            if CheckPostcode(Postcode):
                PostcodeCheck = True
            if DOBDay != "Select Day":
                DOBDayCheck = True
            if DOBMonth != "Select Month":
                DOBMonthCheck = True
            if DOBYear != "Select Year":
                DOBYearCheck = True
            if Gender != "Select your Gender":
                GenderCheck = True
            if CheckAllAlpha(Signature):
                SignatureCheck = True
            if CheckEmail(Email):
                EmailCheck = True

            if BandIDCheck == True:
                if RoleCheck == True:
                    if TitleCheck == True:
                        if FirstNameCheck == True:
                            if SecondNameCheck == True:
                                if AddressLine1Check == True:
                                    if AddressLine2Check == True:
                                        if CountyCheck == True:
                                            if PostcodeCheck == True:
                                                if DOBDayCheck == True:
                                                    if DOBMonthCheck == True:
                                                        if DOBYearCheck == True:
                                                            if GenderCheck == True:
                                                                if EmailCheck == True:
                                                                    if  SignatureCheck == True:
                                                                        addMemberDatabase(EtyFirstName, EtySecondName,
                                                                                          EtyAddressLine1,
                                                                                          EtyAddressLine2,
                                                                                          EtyPostcode, EtySignature,
                                                                                          Role,
                                                                                          FirstName, SecondName,
                                                                                          AddressLine1, AddressLine2,
                                                                                          County, Postcode, DateOfBirth,
                                                                                          Signature, MemDOBDay,
                                                                                          MemDOBMonth, MemDOBYear,
                                                                                          Memcounty,
                                                                                          BandID,
                                                                                          EtyEmailAddress, Email, Title,
                                                                                          Gender)
                                                                    else:
                                                                        messagebox.showinfo(
                                                                            message="Members signature field is not valid. Please re-enter it.")
                                                                else:
                                                                    messagebox.showinfo(message="Email Address field is not valid. Please re-enter it.")
                                                            else:
                                                                messagebox.showinfo(
                                                                    message="Gender field is not valid. Please re-enter it.")
                                                        else:
                                                            messagebox.showinfo(
                                                                message="Please enter your year of birth.")
                                                    else:
                                                        messagebox.showinfo(
                                                            message="Please enter the month that you were born in.")
                                                else:
                                                    messagebox.showinfo(
                                                        message="Please enter the day you were born.")
                                            else:
                                                messagebox.showinfo(
                                                    message="Postcode field is not valid. Please re-enter it. The required format is LLNNNLL or LLNNLL.")
                                        else:
                                            messagebox.showinfo(
                                                message="County field is not valid. Please re-enter it.")
                                    else:
                                        messagebox.showinfo(
                                            message="Address line 2 field is not valid. Please re-enter it.")
                                else:
                                    messagebox.showinfo(
                                        message="Address line 1 field is not valid. Please include your house number and road.")
                            else:
                                messagebox.showinfo(
                                    message="Second name field is not valid. Please re-enter it.")
                        else:
                            messagebox.showinfo(
                                message="First name field is not valid. Please re-enter it.")
                    else:
                        messagebox.showinfo(
                            message="Please select your title.")
                else:
                    messagebox.showinfo(
                        message="Please select your role from the list below!")

        # Initalising Variables to be used as textvar variables.
        Memrole = StringVar()
        Memtitle = StringVar()
        MemfirstName = StringVar()
        MemsecondName = StringVar()
        Memgender = StringVar()
        MemaddressLine1 = StringVar()
        MemaddressLine2 = StringVar()
        Memcounty = StringVar()
        Mempostcode = StringVar()
        MemdateOfBirth = StringVar()
        Memsignature = StringVar()
        MemDOBDay = StringVar()
        MemDOBMonth = StringVar()
        MemDOBYear = StringVar()
        MemEmailAddress = StringVar()

        RegMember = Frame(contentFrame, width=965, height=600)
        RegMember.grid(row=0, column=0, sticky="nsew")

        lblspacer = Label(RegMember, width=10)
        lblspacer.grid()

        lblTitle = Label(RegMember, text="Member Registration", height=2, font=("Arvo", 32))
        lblTitle.grid(row=0, column=1, columnspan=5)

        lblRole = Label(RegMember, text="Please tick as appropriate: ", font=("Arvo", 10))
        lblRole.grid(row=3, column=2, sticky="ew")

        Radiobutton(RegMember, text="Pipe Major", padx=5, variable=Memrole, value="Pipe Major", font=("Arvo")).grid(
            row=5, column=2, sticky="w", padx=20)
        Radiobutton(RegMember, text="Lead Drummer", padx=5, variable=Memrole, value="Lead Drummer", font=("Arvo")).grid(
            row=5, column=3, sticky="w", padx=30)
        Radiobutton(RegMember, text="Pipe Sergeant", padx=5, variable=Memrole, value="Pipe Sergeant",
                    font=("Arvo")).grid(row=5, column=4, sticky="w", padx=30)
        Radiobutton(RegMember, text="Drum Corporal", padx=5, variable=Memrole, value="Drum Corporal",
                    font=("Arvo")).grid(row=6, column=2, sticky="w", padx=20)
        Radiobutton(RegMember, text="Piper", padx=5, variable=Memrole, value="Piper", font=("Arvo")).grid(row=6,
                                                                                                          column=3,
                                                                                                          sticky="w",
                                                                                                          padx=30)
        Radiobutton(RegMember, text="Snare Drummer", padx=5, variable=Memrole, value="Snare Drummer",
                    font=("Arvo")).grid(row=6, column=4, sticky="w", padx=30)
        Radiobutton(RegMember, text="Bass Drummer", padx=5, variable=Memrole, value="Bass Drummer", font=("Arvo")).grid(
            row=7, column=2, sticky="w", padx=20)
        Radiobutton(RegMember, text="Tenor Drummer", padx=5, variable=Memrole, value="Tenor Drummer",
                    font=("Arvo")).grid(row=7, column=3, sticky="w", padx=30)

        lblspacer = Label(RegMember)
        lblspacer.grid(row=8)

        lblTitle = Label(RegMember, text="Title:", font=("Arvo"))
        lblTitle.grid(row=9, column=1, sticky="e")

        listTitles = ['Mr', 'Mrs', 'Miss', 'Ms', 'Dr', 'Rev', 'Master']
        droplistTitle = OptionMenu(RegMember, Memtitle, *listTitles)
        droplistTitle.config(width=20)
        Memtitle.set('Select Title')
        droplistTitle.grid(row=9, column=2, sticky="w")

        lblFirstName = Label(RegMember, text="First Name:", font=("Arvo"))
        lblFirstName.grid(row=10, column=1, sticky="e")

        EtyFirstName = Entry(RegMember, textvar=MemfirstName)
        EtyFirstName.grid(row=10, column=2, sticky="w")

        lblSecondName = Label(RegMember, text="Second Name:", font=("Arvo"))
        lblSecondName.grid(row=11, column=1, sticky="e")

        EtySecondName = Entry(RegMember, textvar=MemsecondName)
        EtySecondName.grid(row=11, column=2, sticky="w")

        lblGender = Label(RegMember, text="Gender:", font=("Arvo"))
        lblGender.grid(row=12, column=1, sticky="e")

        listGender = ['Male', 'Female', 'Other']
        EtyGender = OptionMenu(RegMember, Memgender, *listGender)
        EtyGender.config(width=20)
        Memgender.set('Select your Gender')
        EtyGender.grid(row=12, column=2, sticky="w")

        lblAddress = Label(RegMember, text="Address", font=("Arvo"))
        lblAddress.grid(row=9, column=3, sticky="e")

        lblAddressLine1 = Label(RegMember, text="Line 1:", font=("Arvo"))
        lblAddressLine1.grid(row=10, column=3, sticky="e")

        EtyAddressLine1 = Entry(RegMember, textvar=MemaddressLine1)
        EtyAddressLine1.grid(row=10, column=4, sticky="w")

        lblAddressLine2 = Label(RegMember, text="Line 2:", font=("Arvo"))
        lblAddressLine2.grid(row=11, column=3, sticky="e")

        EtyAddressLine2 = Entry(RegMember, textvar=MemaddressLine2)
        EtyAddressLine2.grid(row=11, column=4, sticky="w")

        lblCounty = Label(RegMember, text="County:", font=("Arvo"))
        lblCounty.grid(row=12, column=3, sticky="e")
        listCounty = ['County Antrim', 'County Down', 'County Armagh', 'County Fermanagh', 'County Londonderry',
                      'Ireland']
        droplist = OptionMenu(RegMember, Memcounty, *listCounty)
        droplist.config(width=20)
        Memcounty.set('Select your County')
        droplist.grid(row=12, column=4, sticky="w")

        lblPostcode = Label(RegMember, text="Postcode:", font=("Arvo"))
        lblPostcode.grid(row=13, column=3, sticky="e")

        EtyPostcode = Entry(RegMember, textvar=Mempostcode)
        EtyPostcode.grid(row=13, column=4, sticky="w")

        lblDateOfBirth = Label(RegMember, text="Date of Birth:", font=("Arvo"))
        lblDateOfBirth.grid(row=14, column=1, sticky="e")

        Days = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19',
                '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
        listDays = Days
        droplistDOBDay = OptionMenu(RegMember, MemDOBDay, *listDays)
        droplistDOBDay.config(width=20)
        MemDOBDay.set('Select Day')
        droplistDOBDay.grid(row=14, column=2, sticky="w")

        Months = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
                  'November', 'December']
        listMonths = Months
        droplistDOBMonth = OptionMenu(RegMember, MemDOBMonth, *listMonths)
        droplistDOBMonth.config(width=20)
        MemDOBMonth.set('Select Month')
        droplistDOBMonth.grid(row=15, column=2, sticky="w")

        listYears = Years
        droplistDOBYear = OptionMenu(RegMember, MemDOBYear, *listYears)
        droplistDOBYear.config(width=20)
        MemDOBYear.set('Select Year')
        droplistDOBYear.grid(row=16, column=2, sticky="w")

        lblEmailAddress = Label(RegMember, text="Email Address:")
        lblEmailAddress.grid(row=14, column=3, sticky="e")

        EtyEmailAddress = Entry(RegMember, textvar=MemEmailAddress)
        EtyEmailAddress.grid(row=14, column=4, sticky="w")

        lblSignature = Label(RegMember, text="Member's Signature:", font=("Arvo"))
        lblSignature.grid(row=16, column=3, sticky="e")

        EtySignature = Entry(RegMember, textvar=Memsignature)
        EtySignature.grid(row=16, column=4, sticky="w")

        lblspacer = Label(RegMember)
        lblspacer.grid(row=17)

        Button(RegMember, image=photoRegisterMemberFormButton, command=lambda: CheckMemberDetails(
            EtyFirstName, EtySecondName,
            EtyAddressLine1, EtyAddressLine2,
            EtyPostcode,
            EtySignature,
            Memrole,
            MemfirstName, MemsecondName,
            MemaddressLine1, MemaddressLine2,
            Memcounty, Mempostcode,
            MemdateOfBirth, Memsignature,
            MemDOBDay, MemDOBMonth,
            MemDOBYear, EtyEmailAddress, MemEmailAddress, Memtitle, Memgender)) \
            .grid(row=20, column=2, columnspan=3, padx="20", pady="10")

    if Admin == True:
        # If user is admin then ask them to select which band they wish to edit.

        EditBandDetails = Frame(contentFrame, width=965, height=600)
        EditBandDetails.grid(row=0, column=0, sticky="nsew")

        lblspacerCenter = Label(EditBandDetails, text="", width=20)
        lblspacerCenter.grid(row=0, column=0)

        Title = Label(EditBandDetails, text="Register New Member", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        lblText = Label(EditBandDetails, text="Please select the band to register a new member ", height=3,
                        font=("Arvo", 14))
        lblText.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)

        EditMemDListBandBox = Listbox(EditBandDetails, selectmode=EXTENDED, width=25)
        EditMemDListBandBox.grid(row=3, column=3, columnspan=4)

        for k in listOfBands:
            n = 0
            EditMemDListBandBox.insert(n, str(k))
            n = + 1

        SelectButton = Button(EditBandDetails, image=photoSelectBandButton, command=SelectBandPage)
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    else:
        # Register the member to the Band Logged in
        RegisterMemberPage(Years)


def RegBandFrame(photoRegisterBandButton):
    def CheckEntries(BandName, BandBranch, BandGrade, PMFirstName, PMSecondName, BandHallName, BandAddressLine1,
                     BandAddressLine2, BandCounty,
                     BandPostcode, BandPracticeTime, BandTartan, BandUsername, BandPassword, BandEmail, EtyBandName,
                     EtyPMFirstName, EtyPMSecondName,
                     EtyBandHallName, EtyBandAddressLine1, EtyBandAddressLine2, EtyBandPostcode, EtyBandPracticeTime,
                     EtyBandTartan, EtyBandUsername,
                     EtyBandPassword, EtyBandEmail):

        # Gets all the user inputs from the Register Band form and saves them to variables
        Band = BandName.get()
        Band = Band.title()
        Branch = BandBranch.get()
        Grade = BandGrade.get()
        PMFirstname = PMFirstName.get()
        PMSecondname = PMSecondName.get()
        HallName = BandHallName.get()
        AddressLine1 = BandAddressLine1.get()
        AddressLine2 = BandAddressLine2.get()
        County = BandCounty.get()
        Postcode = BandPostcode.get()
        PracticeTime = BandPracticeTime.get()
        Tartan = BandTartan.get()
        Username = BandUsername.get()
        Password = BandPassword.get()
        Email = BandEmail.get()

        BandCheck = False
        BranchCheck = False
        GradeCheck = False
        PMFirstnameCheck = False
        PMSecondnameCheck = False
        HallNameCheck = True
        AddressLine1Check = False
        AddressLine2Check = True
        CountyCheck = False
        PostcodeCheck = False
        UsernameCheck = False
        PasswordCheck = False
        PracticeTimeCheck = True
        TartanCheck = True
        EmailCheck = False


        if CheckOnlyNumORLetters(Band) == True:
            BandCheck = True
        if Branch != "Select Branch":
            BranchCheck = True
        if Grade != "Select your Grade":
            GradeCheck = True
        if CheckAllAlpha(PMFirstname) == True:
            PMFirstnameCheck = True
        if CheckAllAlpha(PMSecondname) == True:
            PMSecondnameCheck = True
        if CheckOnlyNumAndLetters(AddressLine1) == True:
            AddressLine1Check = True
        if County != "Select your County":
            CountyCheck = True
        if CheckPostcode(Postcode) == True:
            PostcodeCheck = True
        if CheckUsername(Username, "Reg") == True:
            UsernameCheck = True
        if CheckValidPassword(Password) == True:
            PasswordCheck = True
        if CheckEmail(Email) == True:
            EmailCheck = True

        print(BandCheck, BranchCheck, GradeCheck, PMFirstnameCheck, PMSecondnameCheck, HallNameCheck,
              AddressLine1Check, AddressLine2Check, CountyCheck, PostcodeCheck, UsernameCheck, PasswordCheck,
              PracticeTimeCheck, TartanCheck, EmailCheck)

        if BandCheck == True:
            if BranchCheck == True:
                if GradeCheck == True:
                    if PMFirstnameCheck == True:
                        if PMSecondnameCheck == True:
                            if HallNameCheck == True:
                                if AddressLine1Check == True:
                                    if AddressLine2Check == True:
                                        if CountyCheck == True:
                                            if PostcodeCheck == True:
                                                if UsernameCheck == True:
                                                    if PasswordCheck == True:
                                                        if EmailCheck == True:
                                                            print(Band, Branch, Grade,
                                                                  PMFirstName, PMSecondName, HallName,
                                                                  AddressLine1,
                                                                  AddressLine2, County,
                                                                  Postcode, PracticeTime, Tartan,
                                                                  Username, Password, Email)
                                                            addBandDatabase(Band, Branch, Grade,
                                                                            PMFirstname, PMSecondname, HallName,
                                                                            AddressLine1,
                                                                            AddressLine2, County,
                                                                            Postcode, PracticeTime, Tartan,
                                                                            Username, Password, Email,
                                                                            EtyBandName,
                                                                            EtyPMFirstName, EtyPMSecondName,
                                                                            EtyBandHallName, EtyBandAddressLine1,
                                                                            EtyBandAddressLine2, EtyBandPostcode,
                                                                            EtyBandPracticeTime,
                                                                            EtyBandTartan, EtyBandUsername,
                                                                            EtyBandPassword, EtyBandEmail, BandBranch,
                                                                            BandGrade, BandCounty)

                                                            messagebox.showinfo(message="Band account created!")
                                                        else:
                                                            messagebox.showinfo(
                                                                message="Email field is not valid. Please check your email.")
                                                    else:
                                                        messagebox.showinfo(
                                                            message="Password field is not valid. Please use another password. It must be at least "
                                                                    "8 characters in length and contain at least one upper and lower case "
                                                                    "letter and a number. ")
                                                else:
                                                    messagebox.showinfo(
                                                        message="Username field is not valid or is already in use. Please choose another username.")
                                            else:
                                                messagebox.showinfo(
                                                    message="Postcode field is not valid. Please check your postcode, it must be in the required format of LLNNNLL or LLNNLL.")
                                        else:
                                            messagebox.showinfo(
                                                message="Please select your county from thr drop down list.")
                                    else:
                                        messagebox.showinfo(
                                            message="Address line 2 field is not valid. Please re-enter your address line 2.")
                                else:
                                    messagebox.showinfo(
                                        message="Address line 1 field is not valid. Please include your house number and road.")
                            else:
                                messagebox.showinfo(
                                    message="Hall name field is not valid. Ensure this name only contains letters.")
                        else:
                            messagebox.showinfo(
                                message="Pipe Major's second name field is not valid. Ensure this name only contains letters.")
                    else:
                        messagebox.showinfo(
                            message="Pipe Major's first name field is not valid. Ensure this name only contains letters.")
                else:
                    messagebox.showinfo(message="Please select a grade from the drop down list.")
            else:
                messagebox.showinfo(message="Please select a branch from the drop down list.")
        else:
            messagebox.showinfo(message="Please ensure the name chosen only contains letters and numbers.")

    # Define all the textvar variables to store entered data
    BandName = StringVar()
    BandBranch = StringVar()
    BandGrade = StringVar()
    PMFirstName = StringVar()
    PMSecondName = StringVar()
    BandHallName = StringVar()
    BandAddressLine1 = StringVar()
    BandAddressLine2 = StringVar()
    BandCounty = StringVar()
    BandPostcode = StringVar()
    BandUsername = StringVar()
    BandPassword = StringVar()
    BandPracticeTime = StringVar()
    BandTartan = StringVar()
    BandEmail = StringVar()

    RegBand = Frame(contentFrame, width=965, height=600)
    RegBand.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(RegBand, width=15)
    lblspacer.grid()

    lblTitle = Label(RegBand, text="New Band Registration", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=1, columnspan=5, sticky="ew")

    # Creation of widgets on the RegBand Frame
    lblBandName = Label(RegBand, text="Band Name:", font=("Arvo"))
    lblBandName.grid(row=2, column=1, sticky="e")

    EtyBandName = Entry(RegBand, textvar=BandName)
    EtyBandName.grid(row=2, column=2, sticky="w")

    lblBranch = Label(RegBand, text="Branch:", font=("Arvo"))
    lblBranch.grid(row=3, column=1, sticky="e")

    listBranch = ['Antrim', 'Down', 'Fermanagh', 'ISPBA', 'Londonderry']
    droplistBranch = OptionMenu(RegBand, BandBranch, *listBranch)
    droplistBranch.config(width=20)
    BandBranch.set('Select your Branch')
    droplistBranch.grid(row=3, column=2, sticky="w")

    lblGrade = Label(RegBand, text="Grade:", font=("Arvo"))
    lblGrade.grid(row=4, column=1, sticky="e")

    listGrade = ['1', '2', '3A', '3B', '4A', '4B']
    droplistGrade = OptionMenu(RegBand, BandGrade, *listGrade)
    droplistGrade.config(width=20)
    BandGrade.set('Select your Grade')
    droplistGrade.grid(row=4, column=2, sticky="w")

    lblAddress = Label(RegBand, text="Address", font=("Arvo", 10))
    lblAddress.grid(row=5, column=1, sticky="e")

    lblspacer = Label(RegBand)
    lblspacer.grid(row=5)

    lblBandHallName = Label(RegBand, text="Hall Name:", font=("Arvo"))
    lblBandHallName.grid(row=6, column=1, sticky="e")

    EtyBandHallName = Entry(RegBand, textvar=BandHallName)
    EtyBandHallName.grid(row=6, column=2, sticky="w")

    lblAddressLine1 = Label(RegBand, text="Line 1:", font=("Arvo"))
    lblAddressLine1.grid(row=7, column=1, sticky="e")

    EtyBandAddressLine1 = Entry(RegBand, textvar=BandAddressLine1)
    EtyBandAddressLine1.grid(row=7, column=2, sticky="w")

    lblAddressLine2 = Label(RegBand, text="Line 2:", font=("Arvo"))
    lblAddressLine2.grid(row=8, column=1, sticky="e")

    EtyBandAddressLine2 = Entry(RegBand, textvar=BandAddressLine2)
    EtyBandAddressLine2.grid(row=8, column=2, sticky="w")

    lblCounty = Label(RegBand, text="County:", font=("Arvo"))
    lblCounty.grid(row=9, column=1, sticky="e")
    listCounty = ['County Antrim', 'County Down', 'County Armagh', 'County Fermanagh', 'County Londonderry', 'Ireland']
    droplist = OptionMenu(RegBand, BandCounty, *listCounty)
    droplist.config(width=20)
    BandCounty.set('Select your County')
    droplist.grid(row=9, column=2, sticky="w")

    lblBandPostcode = Label(RegBand, text="Postcode:", font=("Arvo"))
    lblBandPostcode.grid(row=10, column=1, sticky="e")

    EtyBandPostcode = Entry(RegBand, textvar=BandPostcode)
    EtyBandPostcode.grid(row=10, column=2, sticky="w")

    lblBandPracticeTime = Label(RegBand, text="Practice Time(s):", font=("Arvo"))
    lblBandPracticeTime.grid(row=11, column=1, sticky="e")

    EtyBandPracticeTime = Entry(RegBand, textvar=BandPracticeTime)
    EtyBandPracticeTime.grid(row=11, column=2, sticky="w")

    lblBandTartan = Label(RegBand, text="Tartan:", font=("Arvo"))
    lblBandTartan.grid(row=12, column=1, sticky="e")

    EtyBandTartan = Entry(RegBand, textvar=BandTartan)
    EtyBandTartan.grid(row=12, column=2, sticky="w")

    lblspacer = Label(RegBand, width=5)
    lblspacer.grid(column=3)

    lblPMFirstName = Label(RegBand, text="PM First Name:", font=("Arvo"))
    lblPMFirstName.grid(row=2, column=4, sticky="e")

    EtyPMFirstName = Entry(RegBand, textvar=PMFirstName)
    EtyPMFirstName.grid(row=2, column=5, sticky="w")

    lblPMSecondName = Label(RegBand, text="PM Second Name:", font=("Arvo"))
    lblPMSecondName.grid(row=3, column=4, sticky="e")

    EtyPMSecondName = Entry(RegBand, textvar=PMSecondName)
    EtyPMSecondName.grid(row=3, column=5, sticky="w")

    lblBandEmail = Label(RegBand, text="Email Address:", font=("Arvo"))
    lblBandEmail.grid(row=4, column=4, sticky="e")

    EtyBandEmail = Entry(RegBand, textvar=BandEmail)
    EtyBandEmail.grid(row=4, column=5, sticky="w")

    lblBandUsername = Label(RegBand, text="Username:", font=("Arvo"))
    lblBandUsername.grid(row=6, column=4, sticky="e")

    EtyBandUsername = Entry(RegBand, textvar=BandUsername)
    EtyBandUsername.grid(row=6, column=5, sticky="w")

    lblBandPassword = Label(RegBand, text="Password:", font=("Arvo"))
    lblBandPassword.grid(row=7, column=4, sticky="e")

    EtyBandPassword = Entry(RegBand, textvar=BandPassword)
    EtyBandPassword.grid(row=7, column=5, sticky="w")

    Button(RegBand, image=photoRegisterBandButton, command=lambda: CheckEntries(BandName, BandBranch, BandGrade,
                                                                                PMFirstName, PMSecondName,
                                                                                BandHallName, BandAddressLine1,
                                                                                BandAddressLine2, BandCounty,
                                                                                BandPostcode, BandPracticeTime,
                                                                                BandTartan, BandUsername,
                                                                                BandPassword, BandEmail, EtyBandName,
                                                                                EtyPMFirstName, EtyPMSecondName,
                                                                                EtyBandHallName, EtyBandAddressLine1,
                                                                                EtyBandAddressLine2, EtyBandPostcode,
                                                                                EtyBandPracticeTime, EtyBandTartan,
                                                                                EtyBandUsername, EtyBandPassword,
                                                                                EtyBandEmail)).grid(row=20, column=2,
                                                                                                    columnspan=3,
                                                                                                    padx="20",
                                                                                                    pady="10")


def RegJudgeFrame(photoRegisterJudgeButton, Years):
    def CheckEntries(JudgeFirstName, JudgeSecondName, JudgeDOB, JudgeGender, JudgeAddressLine1, JudgeAddressLine2,
                     JudgeCounty, JudgePostcode, JudgeUsername,
                     JudgePassword, JudgeDOBDay, JudgeDOBMonth, JudgeDOBYear, JudgeTitle, EtyJudgeFirstName,
                     EtyJudgeSecondName, EtyJudgeGender,
                     EtyJudgeAddressLine1, EtyJudgeAddressLine2, EtyJudgeCounty, EtyJudgePostcode, EtyJudgeUsername,
                     EtyJudgePassword, JudgeEmail, EtyJudgeEmail):

        # Gets all the user inputs from the Register Judge form and saves them to variables
        Title = JudgeTitle.get()
        FirstName = JudgeFirstName.get()
        SecondName = JudgeSecondName.get()

        DOBDay = JudgeDOBDay.get()
        DOBMonth = JudgeDOBMonth.get()
        DOBYear = JudgeDOBYear.get()

        # Concatenate Day, Month and year to save under one field in the database
        DOB = str(DOBDay) + "/" + str(DOBMonth) + "/" + str(DOBYear)

        Gender = JudgeGender.get()
        AddressLine1 = JudgeAddressLine1.get()
        AddressLine2 = JudgeAddressLine2.get()
        County = JudgeCounty.get()
        Postcode = JudgePostcode.get()
        Email = JudgeEmail.get()
        Username = JudgeUsername.get()
        Password = JudgePassword.get()

        TitleCheck = False
        FirstNameCheck = False
        SecondNameCheck = False
        DOBDayCheck = False
        DOBMonthCheck = False
        DOBYearCheck = False
        GenderCheck = False
        AddressLine1Check = False
        AddressLine2Check = True
        CountyCheck = False
        PostcodeCheck = False
        EmailCheck = False
        UsernameCheck = False
        PasswordCheck = False

        if CheckAllAlpha(FirstName) == True:
            FirstNameCheck = True
        if CheckAllAlpha(SecondName) == True:
            SecondNameCheck = True
        if Title != "Select Title":
            TitleCheck = True
        if DOBDay != "Select Day":
            DOBDayCheck = True
        if DOBMonth != "Select Month":
            DOBMonthCheck = True
        if DOBYear != "Select Year":
            DOBYearCheck = True
        if Gender != "Select your Gender":
            GenderCheck = True
        if CheckOnlyNumAndLetters(AddressLine1) == True:
            AddressLine1Check = True
        if County != "Select your County":
            CountyCheck = True
        if CheckPostcode(Postcode) == True:
            PostcodeCheck = True
        if CheckEmail(Email) == True:
            EmailCheck = True
        if CheckUsername(Username, "Reg") == True:
            UsernameCheck = True
        if CheckValidPassword(Password) == True:
            PasswordCheck = True

        print(TitleCheck, FirstNameCheck, SecondNameCheck, DOBDayCheck, DOBMonthCheck, DOBYearCheck, GenderCheck,
              AddressLine1Check, AddressLine2Check, CountyCheck, PostcodeCheck, UsernameCheck, PasswordCheck)

        if TitleCheck == True:
            if FirstNameCheck == True:
                if SecondNameCheck == True:
                    if DOBDayCheck == True:
                        if DOBMonthCheck == True:
                            if DOBYearCheck == True:
                                if GenderCheck == True:
                                    if AddressLine1Check == True:
                                        if AddressLine2Check == True:
                                            if CountyCheck == True:
                                                if PostcodeCheck == True:
                                                    if EmailCheck == True:
                                                        if UsernameCheck == True:
                                                            if PasswordCheck == True:
                                                                addJudgeDatabase(FirstName, SecondName, DOB, Gender,
                                                                                 AddressLine1, AddressLine2, County,
                                                                                 Postcode, Email, Username, Password,
                                                                                 Title,
                                                                                 EtyJudgeFirstName, EtyJudgeSecondName,
                                                                                 EtyJudgeAddressLine1,
                                                                                 EtyJudgeAddressLine2,
                                                                                 EtyJudgePostcode, EtyJudgeEmail,
                                                                                 EtyJudgeUsername, EtyJudgePassword,
                                                                                 JudgeDOBDay, JudgeDOBMonth,
                                                                                 JudgeDOBYear,
                                                                                 JudgeGender, JudgeCounty, JudgeTitle)

                                                                messagebox.showinfo(message="Judge Created")
                                                            else:
                                                                messagebox.showinfo(
                                                                    message="Password field is not valid. Please re-enter your password")
                                                        else:
                                                            messagebox.showinfo(
                                                                message="Username field is not valid. Please re-enter your username")
                                                    else:
                                                        messagebox.showinfo(
                                                            message="Email field is not valid. Please re-enter your email")
                                                else:
                                                    messagebox.showinfo(
                                                        message="Postcode field is not valid. Please re-enter your postcode")
                                            else:
                                                messagebox.showinfo(
                                                    message="County field is not valid. Please re-enter your county")
                                        else:
                                            messagebox.showinfo(
                                                message="Address line 2 field is not valid. Please re-enter your address line 2")
                                    else:
                                        messagebox.showinfo(
                                            message="Address line 1 field is not valid. Please include your house number and road.")
                                else:
                                    messagebox.showinfo(
                                        message="Gender field is not valid. Please re-enter your gender")
                            else:
                                messagebox.showinfo(
                                    message="Year of birth field is not valid. Please re-enter your tear of birth")
                        else:
                            messagebox.showinfo(
                                message="Month of birth field is not valid. Please re-enter your month of birth")
                    else:
                        messagebox.showinfo(
                            message="Day of birth field is not valid. Please re-enter your day of birth")
                else:
                    messagebox.showinfo(message="Second name field is not valid. Please re-enter your second name")
            else:
                messagebox.showinfo(message="First name field is not valid. Please re-enter your first name")
        else:
            messagebox.showinfo(message="Title field is not valid. Please re-enter your title")

    # Define all the textvar variables to store entered data
    JudgeTitle = StringVar()
    JudgeFirstName = StringVar()
    JudgeSecondName = StringVar()
    JudgeDOB = StringVar()
    JudgeGender = StringVar()
    JudgeAddressLine1 = StringVar()
    JudgeAddressLine2 = StringVar()
    JudgeCounty = StringVar()
    JudgePostcode = StringVar()
    JudgeEmail = StringVar()
    JudgeUsername = StringVar()
    JudgePassword = StringVar()
    JudgeDOBDay = StringVar()
    JudgeDOBMonth = StringVar()
    JudgeDOBYear = StringVar()

    RegJudge = Frame(contentFrame, width=965, height=600)
    RegJudge.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(RegJudge, width=40)
    lblspacer.grid()

    # Creation of widgets on the RegJudge Frame

    lblTitle = Label(RegJudge, text="New Judge Registration", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=1, columnspan=3)

    lblJudgeTitle = Label(RegJudge, text="Title:", font=("Arvo"))
    lblJudgeTitle.grid(row=2, column=1, sticky="e")

    listTitles = ['Mr', 'Mrs', 'Miss', 'Ms', 'Dr', 'Rev', 'Master']
    droplistTitle = OptionMenu(RegJudge, JudgeTitle, *listTitles)
    droplistTitle.config(width=20)
    JudgeTitle.set('Select Title')
    droplistTitle.grid(row=2, column=2, sticky="w")

    # Creates a label to tell the user to enter their First Name into the entry field beside it.
    lblJudgeFirstName = Label(RegJudge, text="First Name:", font=("Arvo"))
    lblJudgeFirstName.grid(row=3, column=1, sticky="e")

    EtyJudgeFirstName = Entry(RegJudge, textvar=JudgeFirstName)
    EtyJudgeFirstName.grid(row=3, column=2, sticky="w")

    # Creates a label to tell the user to enter their Second Name into the entry field beside it.
    lblJudgeSecondName = Label(RegJudge, text="Second Name:", font=("Arvo"))
    lblJudgeSecondName.grid(row=4, column=1, sticky="e")

    EtyJudgeSecondName = Entry(RegJudge, textvar=JudgeSecondName)
    EtyJudgeSecondName.grid(row=4, column=2, sticky="w")

    # Creates a label to tell the user to enter their DOB into the entry field beside it.
    lblJudgeDOB = Label(RegJudge, text="Date of Birth:", font=("Arvo"))
    lblJudgeDOB.grid(row=5, column=1, sticky="e")

    listDays = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19',
                '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
    droplistDOBDay = OptionMenu(RegJudge, JudgeDOBDay, *listDays)
    droplistDOBDay.config(width=20)
    JudgeDOBDay.set('Select Day')
    droplistDOBDay.grid(row=5, column=2, sticky="w")

    listMonths = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
                  'November', 'December']
    droplistDOBMonth = OptionMenu(RegJudge, JudgeDOBMonth, *listMonths)
    droplistDOBMonth.config(width=20)
    JudgeDOBMonth.set('Select Month')
    droplistDOBMonth.grid(row=6, column=2, sticky="w")

    listYears = Years
    droplistDOBYear = OptionMenu(RegJudge, JudgeDOBYear, *listYears)
    droplistDOBYear.config(width=20)
    JudgeDOBYear.set('Select Year')
    droplistDOBYear.grid(row=7, column=2, sticky="w")

    lblJudgeGender = Label(RegJudge, text="Gender:", font=("Arvo"))
    lblJudgeGender.grid(row=8, column=1, sticky="e")

    listGender = ['Male', 'Female', 'Other']
    EtyJudgeGender = OptionMenu(RegJudge, JudgeGender, *listGender)
    EtyJudgeGender.config(width=20)
    JudgeGender.set('Select your Gender')
    EtyJudgeGender.grid(row=8, column=2, sticky="w")

    lblAddress = Label(RegJudge, text="Address", font=("Arvo"))
    lblAddress.grid(row=9, column=1, sticky="e")

    # Creates a label to tell the user to enter the First line of their address into the entry field beside it.
    lblJudgeAddressLine1 = Label(RegJudge, text="Line 1:", font=("Arvo"))
    lblJudgeAddressLine1.grid(row=10, column=1, sticky="e")

    EtyJudgeAddressLine1 = Entry(RegJudge, textvar=JudgeAddressLine1)
    EtyJudgeAddressLine1.grid(row=10, column=2, sticky="w")

    # Creates a label to tell the user to enter the second line of their address into the entry field beside it.
    lblAddressLine2 = Label(RegJudge, text="Line 2:", font=("Arvo"))
    lblAddressLine2.grid(row=11, column=1, sticky="e")

    EtyJudgeAddressLine2 = Entry(RegJudge, textvar=JudgeAddressLine2)
    EtyJudgeAddressLine2.grid(row=11, column=2, sticky="w")

    # Creates a label to tell the user to enter the Country they live in into the entry field beside it.
    lblJudgeCounty = Label(RegJudge, text="County:", font=("Arvo"))
    lblJudgeCounty.grid(row=12, column=1, sticky="e")

    listJudgeCounty = ['County Antrim', 'County Down', 'County Armagh', 'County Fermanagh', 'County Londonderry',
                       'Ireland']
    EtyJudgeCounty = OptionMenu(RegJudge, JudgeCounty, *listJudgeCounty)
    EtyJudgeCounty.config(width=20)
    JudgeCounty.set('Select your county')
    EtyJudgeCounty.grid(row=12, column=2, sticky="w")

    # Creates a label to tell the user to enter their Postcode into the entry field beside it.
    lblJudgePostcode = Label(RegJudge, text="Postcode:", font=("Arvo"))
    lblJudgePostcode.grid(row=13, column=1, sticky="e")

    EtyJudgePostcode = Entry(RegJudge, textvar=JudgePostcode)
    EtyJudgePostcode.grid(row=13, column=2, sticky="w")

    lblEmail = Label(RegJudge, text="Email Address:")
    lblEmail.grid(row=14, column=1, sticky="e")

    EtyJudgeEmail = Entry(RegJudge, textvar=JudgeEmail)
    EtyJudgeEmail.grid(row=14, column=2, sticky="w")

    # Creates a label to tell the user to enter their Username into the entry field beside it.
    lblJudgeUsername = Label(RegJudge, text="Username:", font=("Arvo"))
    lblJudgeUsername.grid(row=15, column=1, sticky="e")

    EtyJudgeUsername = Entry(RegJudge, textvar=JudgeUsername)
    EtyJudgeUsername.grid(row=15, column=2, sticky="w")

    # Creates a label to tell the user to enter their Password into the entry field beside it.
    lblJudgePassword = Label(RegJudge, text="Password:", font=("Arvo"))
    lblJudgePassword.grid(row=16, column=1, sticky="e")

    EtyJudgePassword = Entry(RegJudge, textvar=JudgePassword)
    EtyJudgePassword.grid(row=16, column=2, sticky="w")

    Button(RegJudge, image=photoRegisterJudgeButton, command=lambda: CheckEntries(JudgeFirstName, JudgeSecondName,
                                                                                  JudgeDOB, JudgeGender,
                                                                                  JudgeAddressLine1,
                                                                                  JudgeAddressLine2, JudgeCounty,
                                                                                  JudgePostcode, JudgeUsername,
                                                                                  JudgePassword, JudgeDOBDay,
                                                                                  JudgeDOBMonth, JudgeDOBYear,
                                                                                  JudgeTitle, EtyJudgeFirstName,
                                                                                  EtyJudgeSecondName,
                                                                                  EtyJudgeGender,
                                                                                  EtyJudgeAddressLine1,
                                                                                  EtyJudgeAddressLine2,
                                                                                  EtyJudgeCounty, EtyJudgePostcode,
                                                                                  EtyJudgeUsername,
                                                                                  EtyJudgePassword, JudgeEmail,
                                                                                  EtyJudgeEmail)) \
        .grid(row=19, column=2, sticky="w", padx="18", pady="10")


def RegAdminFrame(photoRegisterAdminButton, Years):
    def CheckEntries(AdminFirstName, AdminSecondName, AdminDOB, AdminDOBDay, AdminDOBMonth, AdminDOBYear, AdminGender,
                     AdminAddressLine1, AdminAddressLine2,
                     AdminCounty, AdminPostcode, AdminEmail, AdminUsername, AdminPassword, AdminTitle,
                     EtyAdminFirstName, EtyAdminSecondName, EtyAdminGender,
                     EtyAdminAddressLine1, EtyAdminAddressLine2, EtyAdminCounty, EtyAdminPostcode, EtyEmail,
                     EtyAdminUsername, EtyAdminPassword):

        Title = AdminTitle.get()
        FirstName = AdminFirstName.get()
        SecondName = AdminSecondName.get()
        DOBDay = AdminDOBDay.get()
        DOBMonth = AdminDOBMonth.get()
        DOBYear = AdminDOBYear.get()
        Gender = AdminGender.get()

        # Concatenate Day, Month and year to save under one field in the database
        DOB = str(DOBDay) + "/" + str(DOBMonth) + "/" + str(DOBYear)

        AddressLine1 = AdminAddressLine1.get()
        AddressLine2 = AdminAddressLine2.get()
        County = AdminCounty.get()
        Postcode = AdminPostcode.get()
        Email = AdminEmail.get()
        Username = AdminUsername.get()
        Password = AdminPassword.get()

        TitleCheck = False
        FirstNameCheck = False
        SecondNameCheck = False
        DOBDayCheck = False
        DOBMonthCheck = False
        DOBYearCheck = False
        GenderCheck = False
        AddressLine1Check = False
        AddressLine2Check = True
        CountyCheck = False
        PostcodeCheck = False
        EmailCheck = False
        UsernameCheck = False
        PasswordCheck = False

        if CheckAllAlpha(FirstName) == True:
            FirstNameCheck = True
        if CheckAllAlpha(SecondName) == True:
            SecondNameCheck = True
        if Title != "Select Title":
            TitleCheck = True
        if DOBDay != "Select Day":
            DOBDayCheck = True
        if DOBMonth != "Select Month":
            DOBMonthCheck = True
        if DOBYear != "Select Year":
            DOBYearCheck = True
        if Gender != "Select your Gender":
            GenderCheck = True
        if CheckOnlyNumAndLetters(AddressLine1) == True:
            AddressLine1Check = True
        if County != "Select your County":
            CountyCheck = True
        if CheckPostcode(Postcode) == True:
            PostcodeCheck = True
        if CheckEmail(Email) == True:
            EmailCheck = True
        if CheckUsername(Username, "Reg") == True:
            UsernameCheck = True
        if CheckValidPassword(Password) == True:
            PasswordCheck = True

        print(TitleCheck, FirstNameCheck, SecondNameCheck, DOBDayCheck, DOBMonthCheck, DOBYearCheck, GenderCheck,
              AddressLine1Check, AddressLine2Check, CountyCheck, PostcodeCheck, UsernameCheck, PasswordCheck)

        if TitleCheck == True:
            if FirstNameCheck == True:
                if SecondNameCheck == True:
                    if DOBDayCheck == True:
                        if DOBMonthCheck == True:
                            if DOBYearCheck == True:
                                if GenderCheck == True:
                                    if AddressLine1Check == True:
                                        if AddressLine2Check == True:
                                            if CountyCheck == True:
                                                if PostcodeCheck == True:
                                                    if EmailCheck == True:
                                                        if UsernameCheck == True:
                                                            if PasswordCheck == True:
                                                                addAdminDatabase(FirstName, SecondName,
                                                                                 DOB, Gender, AddressLine1,
                                                                                 AddressLine2, County,
                                                                                 Postcode, Email, Username,
                                                                                 Password, Title,
                                                                                 EtyAdminFirstName,
                                                                                 EtyAdminSecondName,
                                                                                 EtyAdminGender,
                                                                                 EtyAdminAddressLine1,
                                                                                 EtyAdminAddressLine2,
                                                                                 EtyAdminCounty, EtyAdminPostcode,
                                                                                 EtyEmail, EtyAdminUsername,
                                                                                 EtyAdminPassword, AdminDOBDay,
                                                                                 AdminDOBMonth,
                                                                                 AdminDOBYear, AdminGender, AdminCounty,
                                                                                 AdminTitle)

                                                                messagebox.showinfo(message="Admin Created")
                                                            else:
                                                                messagebox.showinfo(
                                                                    message="Password field is not valid. Please re-enter your password")
                                                        else:
                                                            messagebox.showinfo(
                                                                message="Username field is not valid. Please re-enter your username")
                                                    else:
                                                        messagebox.showinfo(
                                                            message="Email field is not valid. Please re-enter your email")
                                                else:
                                                    messagebox.showinfo(
                                                        message="Postcode field is not valid. Please re-enter your postcode")
                                            else:
                                                messagebox.showinfo(
                                                    message="County field is not valid. Please re-enter your county")
                                        else:
                                            messagebox.showinfo(
                                                message="Address line 2 field is not valid. Please re-enter your address line 2")
                                    else:
                                        messagebox.showinfo(
                                            message="Address line 1 field is not valid. Please include your house number and road.")
                                else:
                                    messagebox.showinfo(
                                        message="Gender field is not valid. Please re-enter your gender")
                            else:
                                messagebox.showinfo(
                                    message="Year of birth field is not valid. Please re-enter your tear of birth")
                        else:
                            messagebox.showinfo(
                                message="Month of birth field is not valid. Please re-enter your month of birth")
                    else:
                        messagebox.showinfo(
                            message="Day of birth field is not valid. Please re-enter your day of birth")
                else:
                    messagebox.showinfo(message="Second name field is not valid. Please re-enter your second name")
            else:
                messagebox.showinfo(message="First name field is not valid. Please re-enter your first name")
        else:
            messagebox.showinfo(message="Title field is not valid. Please re-enter your title")

    # Define all the textvar variables to store entered data
    AdminFirstName = StringVar()
    AdminSecondName = StringVar()
    AdminDOB = StringVar()
    AdminDOBDay = StringVar()
    AdminDOBMonth = StringVar()
    AdminDOBYear = StringVar()
    AdminGender = StringVar()
    AdminAddressLine1 = StringVar()
    AdminAddressLine2 = StringVar()
    AdminCounty = StringVar()
    AdminPostcode = StringVar()
    AdminEmail = StringVar()
    AdminUsername = StringVar()
    AdminPassword = StringVar()
    AdminTitle = StringVar()

    RegAdmin = Frame(contentFrame, width=965, height=600)
    RegAdmin.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(RegAdmin, width=40)
    lblspacer.grid()

    # Creation of widgets on the RegAdmin Frame
    lblTitle = Label(RegAdmin, text="New Admin Registration", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=1, columnspan=3)

    lblAdminTitle = Label(RegAdmin, text="Title:", font=("Arvo"))
    lblAdminTitle.grid(row=2, column=1, sticky="e")

    listTitles = ['Mr', 'Mrs', 'Miss', 'Ms', 'Dr', 'Rev', 'Master']
    droplistTitle = OptionMenu(RegAdmin, AdminTitle, *listTitles)
    droplistTitle.config(width=20)
    AdminTitle.set('Select Title')
    droplistTitle.grid(row=2, column=2, sticky="w")

    # Creates a label to tell the user to enter their First Name into the entry field beside it.
    lblAdminFirstName = Label(RegAdmin, text="First Name:", font=("Arvo"))
    lblAdminFirstName.grid(row=3, column=1, sticky="e")

    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminFirstName = Entry(RegAdmin, textvar=AdminFirstName)
    EtyAdminFirstName.grid(row=3, column=2, sticky="w")

    # Creates a label to tell the user to enter their First Second into the entry field beside it.
    lblAdminSecondName = Label(RegAdmin, text="Second Name:", font=("Arvo"))
    lblAdminSecondName.grid(row=4, column=1, sticky="e")

    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminSecondName = Entry(RegAdmin, textvar=AdminSecondName)
    EtyAdminSecondName.grid(row=4, column=2, sticky="w")

    # Creates a label to tell the user to enter their Date of birth into the entry field beside it.
    lblAdminDOB = Label(RegAdmin, text="Date of Birth:", font=("Arvo"))
    lblAdminDOB.grid(row=6, column=1, sticky="e")

    listDays = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19',
                '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
    droplistDOBDay = OptionMenu(RegAdmin, AdminDOBDay, *listDays)
    droplistDOBDay.config(width=20)
    AdminDOBDay.set('Select Day')
    droplistDOBDay.grid(row=6, column=2, sticky="w")

    listMonths = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
                  'November', 'December']
    droplistDOBMonth = OptionMenu(RegAdmin, AdminDOBMonth, *listMonths)
    droplistDOBMonth.config(width=20)
    AdminDOBMonth.set('Select Month')
    droplistDOBMonth.grid(row=7, column=2, sticky="w")

    listYears = Years
    droplistDOBYear = OptionMenu(RegAdmin, AdminDOBYear, *listYears)
    droplistDOBYear.config(width=20)
    AdminDOBYear.set('Select Year')
    droplistDOBYear.grid(row=8, column=2, sticky="w")

    # Creates a label to tell the user to enter their Gender into the entry field beside it.
    lblAdminGender = Label(RegAdmin, text="Gender:", font=("Arvo"))
    lblAdminGender.grid(row=9, column=1, sticky="e")

    listGender = ['Male', 'Female', 'Other']
    EtyAdminGender = OptionMenu(RegAdmin, AdminGender, *listGender)
    EtyAdminGender.config(width=20)
    AdminGender.set('Select your Gender')
    EtyAdminGender.grid(row=9, column=2, sticky="w")

    lblAddress = Label(RegAdmin, text="Address", font=("Arvo"))
    lblAddress.grid(row=10, column=1, sticky="e")

    lblAdminAddressLine1 = Label(RegAdmin, text="Line 1:", font=("Arvo"))
    lblAdminAddressLine1.grid(row=11, column=1, sticky="e")

    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminAddressLine1 = Entry(RegAdmin, textvar=AdminAddressLine1)
    EtyAdminAddressLine1.grid(row=11, column=2, sticky="w")

    lblAdminAddressLine2 = Label(RegAdmin, text="Line 2:", font=("Arvo"))
    lblAdminAddressLine2.grid(row=12, column=1, sticky="e")

    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminAddressLine2 = Entry(RegAdmin, textvar=AdminAddressLine2)
    EtyAdminAddressLine2.grid(row=12, column=2, sticky="w")

    lblAdminCounty = Label(RegAdmin, text="County:", font=("Arvo"))
    lblAdminCounty.grid(row=13, column=1, sticky="e")
    listAdminCounty = ['County Antrim', 'County Down', 'County Armagh', 'County Fermanagh', 'County Londonderry',
                       'Ireland']
    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminCounty = OptionMenu(RegAdmin, AdminCounty, *listAdminCounty)
    EtyAdminCounty.config(width=20)
    AdminCounty.set('Select your County')
    EtyAdminCounty.grid(row=13, column=2, sticky="w")

    lblAdminPostcode = Label(RegAdmin, text="Postcode:", font=("Arvo"))
    lblAdminPostcode.grid(row=14, column=1, sticky="e")

    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminPostcode = Entry(RegAdmin, textvar=AdminPostcode)
    EtyAdminPostcode.grid(row=14, column=2, sticky="w")

    lblEmailAddress = Label(RegAdmin, text="Email Address:")
    lblEmailAddress.grid(row=15, column=1, sticky="e")

    EtyEmail = Entry(RegAdmin, textvar=AdminEmail)
    EtyEmail.grid(row=15, column=2, sticky="w")

    # Creates a label to tell the user to enter their Username into the entry field beside it.
    lblAdminUsername = Label(RegAdmin, text="Username:", font=("Arvo"))
    lblAdminUsername.grid(row=16, column=1, sticky="e")

    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminUsername = Entry(RegAdmin, textvar=AdminUsername)
    EtyAdminUsername.grid(row=16, column=2, sticky="w")

    # Creates a label to tell the user to enter their Password into the entry field beside it.
    lblAdminPassword = Label(RegAdmin, text="Password:", font=("Arvo"))
    lblAdminPassword.grid(row=17, column=1, sticky="e")

    # Creates and places an entry box to the left of the Label for the user to enter their information
    EtyAdminPassword = Entry(RegAdmin, textvar=AdminPassword)
    EtyAdminPassword.grid(row=17, column=2, sticky="w")

    # This button once pressed saves all the data to the database by running the function required to do this.
    Button(RegAdmin, image=photoRegisterAdminButton, command=lambda: CheckEntries(AdminFirstName, AdminSecondName,
                                                                                  AdminDOB, AdminDOBDay,
                                                                                  AdminDOBMonth, AdminDOBYear,
                                                                                  AdminGender, AdminAddressLine1,
                                                                                  AdminAddressLine2, AdminCounty,
                                                                                  AdminPostcode, AdminEmail,
                                                                                  AdminUsername,
                                                                                  AdminPassword, AdminTitle,
                                                                                  EtyAdminFirstName,
                                                                                  EtyAdminSecondName,
                                                                                  EtyAdminGender,
                                                                                  EtyAdminAddressLine1,
                                                                                  EtyAdminAddressLine2,
                                                                                  EtyAdminCounty, EtyAdminPostcode,
                                                                                  EtyEmail, EtyAdminUsername,
                                                                                  EtyAdminPassword)) \
        .grid(row=18, column=2, sticky="w", padx="18", pady="10")


def AddCompetitionFrame(photoCreateCompetitionButton, CompYears):
    def CheckValidComp(CompetitionName, CompetitionLocation, CompDay, CompMonth, CompYear, DeadLineDay, DeadLineMonth,
                       DeadLineYear, EtyCompLocation, EtyCompName):
        CompName = CompetitionName.get()
        CompLocation = CompetitionLocation.get()
        Day = CompDay.get()
        Month = CompMonth.get()
        Year = CompYear.get()
        EDeadDay = DeadLineDay.get()
        EDeadMonth = DeadLineMonth.get()
        EDeadYear = DeadLineYear.get()

        CompDate = str(Day) + "/" + str(Month) + "/" + str(Year)

        EntryDate = str(EDeadDay) + "/" + str(EDeadMonth) + "/" + str(EDeadYear)

        CompNameCheck = False
        CompLocationCheck = False
        DayCheck = False
        MonthCheck = False
        YearCheck = False
        EDeadDayCheck = False
        EDeadMonthCheck = False
        EDeadYearCheck = False

        if CheckAllAlpha(CompName) == True:
            CompNameCheck = True
        if CompLocation != "":
            CompLocationCheck = True
        if Day != "Select Day":
            DayCheck = True
        if Month != "Select Month":
            MonthCheck = True
        if Year != "Select Year":
            YearCheck = True
        if EDeadDay != "Select Day":
            EDeadDayCheck = True
        if EDeadMonth != "Select Month":
            EDeadMonthCheck = True
        if EDeadYear != "Select Year":
            EDeadYearCheck = True

        if CompNameCheck == True:
            if CompLocationCheck == True:
                if DayCheck == True:
                    if MonthCheck == True:
                        if YearCheck == True:
                            if EDeadDayCheck == True:
                                if EDeadMonthCheck == True:
                                    if EDeadYearCheck == True:
                                        addCompDatabase(CompName, CompLocation, CompDate, EntryDate, EtyCompLocation,
                                                        EtyCompName, CompDay, CompMonth, CompYear, DeadLineDay,
                                                        DeadLineMonth, DeadLineYear)
                                    else:
                                        messagebox.showinfo(
                                            message="Entry deadline year field is not valid. Please re-enter it")
                                else:
                                    messagebox.showinfo(
                                        message="Entry deadline month field is not valid. Please re-enter it")
                            else:
                                messagebox.showinfo(
                                    message="Entry deadline day field is not valid. Please re-enter it")
                        else:
                            messagebox.showinfo(
                                message="Competition year field is not valid. Please re-enter it")
                    else:
                        messagebox.showinfo(
                            message="Competition month field is not valid. Please re-enter it")
                else:
                    messagebox.showinfo(
                        message="Competition day field is not valid. Please re-enter it")
            else:
                messagebox.showinfo(
                    message="Competition location field is not valid. Please re-enter it")
        else:
            messagebox.showinfo(
                message="Competition name field is not valid. Please re-enter it")

    # Define all the textvar variables to store entered data
    CompetitionName = StringVar()
    CompetitionLocation = StringVar()
    CompDay = StringVar()
    CompMonth = StringVar()
    CompYear = StringVar()
    DeadLineDay = StringVar()
    DeadLineMonth = StringVar()
    DeadLineYear = StringVar()

    AddComp = Frame(contentFrame, width=965, height=600)
    AddComp.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(AddComp, width=30)
    lblspacer.grid()

    # Creation of widgets on the RegAdmin Frame
    lblTitle = Label(AddComp, text="Create Competition", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=1, columnspan=3)

    lblCompName = Label(AddComp, text="Competition Name: ", font=("Arvo"))
    lblCompName.grid(row=2, column=1, sticky="e")

    EtyCompName = Entry(AddComp, textvar=CompetitionName)
    EtyCompName.grid(row=2, column=2, sticky="w")

    lblspacer = Label(AddComp)
    lblspacer.grid(row=3, column=1, sticky="e")

    lblCompDate = Label(AddComp, text="Competition Date: ", font=("Arvo"))
    lblCompDate.grid(row=4, column=1, sticky="e")

    listDays = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19',
                '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
    droplistCompDay = OptionMenu(AddComp, CompDay, *listDays)
    droplistCompDay.config(width=20)
    CompDay.set('Select Day')
    droplistCompDay.grid(row=4, column=2, sticky="w")

    listMonths = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October',
                  'November', 'December']
    droplistCompMonth = OptionMenu(AddComp, CompMonth, *listMonths)
    droplistCompMonth.config(width=20)
    CompMonth.set('Select Month')
    droplistCompMonth.grid(row=5, column=2, sticky="w")

    droplistCompYear = OptionMenu(AddComp, CompYear, *CompYears)
    droplistCompYear.config(width=20)
    CompYear.set('Select Year')
    droplistCompYear.grid(row=6, column=2, sticky="w")

    lblspacer = Label(AddComp)
    lblspacer.grid(row=7, column=1, sticky="e")

    lblCompLocation = Label(AddComp, text="Competition Location: ", font=("Arvo"))
    lblCompLocation.grid(row=8, column=1, sticky="e")

    EtyCompLocation = Entry(AddComp, textvar=CompetitionLocation)
    EtyCompLocation.grid(row=8, column=2, sticky="w")

    lblspacer = Label(AddComp)
    lblspacer.grid(row=9, column=1, sticky="e")

    lblEntryDeadLine = Label(AddComp, text="Competition Entry Deadline:", font=("Arvo"))
    lblEntryDeadLine.grid(row=10, column=1, sticky="e")

    droplistCompYear = OptionMenu(AddComp, DeadLineDay, *listDays)
    droplistCompYear.config(width=20)
    DeadLineDay.set('Select Day')
    droplistCompYear.grid(row=10, column=2, sticky="w")

    droplistCompYear = OptionMenu(AddComp, DeadLineMonth, *listMonths)
    droplistCompYear.config(width=20)
    DeadLineMonth.set('Select Month')
    droplistCompYear.grid(row=11, column=2, sticky="w")

    droplistCompYear = OptionMenu(AddComp, DeadLineYear, *CompYears)
    droplistCompYear.config(width=20)
    DeadLineYear.set('Select Year')
    droplistCompYear.grid(row=12, column=2, sticky="w")

    # This button once pressed "Creates" the competiiton as all the data is saved to the database by running the function required to do this.
    Button(AddComp, image=photoCreateCompetitionButton, command=lambda: CheckValidComp(CompetitionName,
                                                                                       CompetitionLocation, CompDay,
                                                                                       CompMonth, CompYear, DeadLineDay,
                                                                                       DeadLineMonth, DeadLineYear,
                                                                                       EtyCompLocation, EtyCompName)) \
        .grid(row=13, column=2, sticky="w", pady="10", padx="7")


def EditCompetitionFrame(photoSelectCompetitionButton, Years):
    def CheckEditedCompetition(UptCompetitionName, UptCompetitionLocation, UptCompDay, UptCompMonth, UptCompYear,
                               CompID, UptEntryDay, UptEntryMonth, UptEntryYear):

        # Concatenate the day, month and year format the data
        # Get the data from the entry boxes and store them to variables to be saved to the database
        Day = UptCompDay.get()
        Month = UptCompMonth.get()
        Year = UptCompYear.get()
        EntryDay = UptEntryDay.get()
        EntryMonth = UptEntryMonth.get()
        EntryYear = UptEntryYear.get()
        CompetitionLocation = UptCompetitionLocation.get()
        CompetitionName = UptCompetitionName.get()

        CompetitionDate = str(Day) + "/" + str(Month) + "/" + str(Year)
        EntryDeadLine = str(EntryDay) + "/" + str(EntryMonth) + "/" + str(EntryYear)

        DayCheck = False
        MonthCheck = False
        YearCheck = False
        EntryDayCheck = False
        EntryMonthCheck = False
        EntryYearCheck = False
        CompetitionLocationCheck = False
        CompetitionNameCheck = False

        if CheckAllAlpha(CompetitionName):
            CompetitionNameCheck = True
        if CheckAllAlpha(CompetitionLocation):
            CompetitionLocationCheck = True
        if Day != "Select Day":
            DayCheck = True
        if Month != "Select Month":
            MonthCheck = True
        if Year != "Select Year":
            YearCheck = True
        if EntryDay != "Select Day":
            EntryDayCheck = True
        if EntryMonth != "Select Month":
            EntryMonthCheck = True
        if EntryYear != "Select Year":
            EntryYearCheck = True

        if DayCheck == True:
            if MonthCheck == True:
                if YearCheck == True:
                    if EntryDayCheck == True:
                        if EntryMonthCheck == True:
                            if EntryYearCheck == True:
                                if CompetitionLocationCheck == True:
                                    if CompetitionNameCheck == True:
                                        UptCompDetailsDatabase(CompetitionName, CompetitionLocation, CompetitionDate,
                                                               CompID, EntryDeadLine)
                                    else:
                                        messagebox.showinfo(
                                            message="Competition name field is not valid. Please re-enter it")
                                else:
                                    messagebox.showinfo(
                                        message="Competition location field is not valid. Please re-enter it")
                            else:
                                messagebox.showinfo(
                                    message="Entry deadline year field is not valid. Please re-enter it")
                        else:
                            messagebox.showinfo(
                                message="Entry deadline month field is not valid. Please re-enter it")
                    else:
                        messagebox.showinfo(
                            message="Entry deadline day field is not valid. Please re-enter it")
                else:
                    messagebox.showinfo(
                        message="Competition year field is not valid. Please re-enter it")
            else:
                messagebox.showinfo(
                    message="Competition month field is not valid. Please re-enter it")
        else:
            messagebox.showinfo(
                message="Competition day field is not valid. Please re-enter it")

    def UptCompDetailsDatabase(CompetitionName, CompetitionLocation, CompetitionDate, CompID, EntryDeadLine):

        # update the competition details with the new details
        conn = sqlite3.connect('RspbaniDB.db')
        with conn:
            cursor = conn.cursor()
        cursor.execute(
            'UPDATE Competitions SET CompetitionDate = ?, CompetitionLocation = ?, CompetitionName = ?, EntryDeadLine = ?'
            'WHERE CompetitionID = ?', (CompetitionDate, CompetitionLocation, CompetitionName, EntryDeadLine, CompID))
        conn.commit()

        messagebox.showinfo(message="Competition updated successfully")
        EditCompetitionFrame(photoSelectCompetitionButton, Years)

    def selectedComp(photoUpdateCompetitionButton):
        def DeleteComp(CompID, CompName):
            # ask them if they want to delete the Competition
            Answer = messagebox.askquestion("Delete Competition?", "Are you sure you want to delete " + CompName + ". All of its data will be lost.")
            if Answer == "yes":

                # Get all EntryID's of the competition.
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorEntryID = db.cursor()
                findEntryID = 'SELECT EntryID FROM BandsEntered WHERE CompetitionID = ?'
                cursorEntryID.execute(findEntryID, [CompID])
                EntryIDresults = cursorEntryID.fetchall()

                if EntryIDresults:

                    for a in EntryIDresults:
                        Entry = a[0]
                        print("Entry")
                        print(Entry)
                        # Deletes every record that is stored about/ has a relationship with  the competition to be deleted from database
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute("DELETE FROM BandsEntered WHERE EntryID = ?", (Entry,))
                        conn.commit()
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute("DELETE FROM CompetitionDraw WHERE EntryID = ?", (Entry,))
                        conn.commit()
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute("DELETE FROM Results WHERE EntryID = ?", (Entry,))
                        conn.commit()

                # Delete Competition from database
                conn = sqlite3.connect('RspbaniDB.db')
                with conn:
                    cursor = conn.cursor()
                cursor.execute("DELETE FROM Competitions WHERE CompetitionID = ?", (CompID,))
                conn.commit()
                EditCompetitionFrame(photoSelectCompetitionButton, Years)

        Selection = listOfCompNameBox.curselection()

        print(Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select a competition to edit.")
        else:

            for i in Selection:
                # gets the name of the competition from the selection
                CompName = listOfCompNameBox.get(i)
                # gets the competition name
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCompID = db.cursor()
                findCompID = 'SELECT CompetitionID FROM Competitions WHERE CompetitionName = ?'
                cursorCompID.execute(findCompID, [CompName])
                CompIDresults = cursorCompID.fetchall()

                if CompIDresults:

                    for a in CompIDresults:
                        a = str(a)
                        CompID = a[1:-2]

                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    GetFieldsComp = 'SELECT CompetitionName, CompetitionLocation, CompetitionDate, EntryDeadLine FROM Competitions WHERE CompetitionID = ?'
                    # Selects the fields which the admin has set
                    cursor.execute(GetFieldsComp, [CompID])
                    GetCompFieldresults = cursor.fetchall()
                    conn.commit()

                    for i in GetCompFieldresults:
                        CurrCompetitionName = i[0]
                        CurrCompetitionLocation = i[1]
                        CurrCompetitionDate = i[2]
                        CurrEntryDeadLine = i[3]

                    # split the date at every "/"

                    CurrCompetitionDate = CurrCompetitionDate.split('/')
                    # The first part is the day
                    CurrCompDay = CurrCompetitionDate[0]
                    # The second part is the Month
                    CurrCompMonth = CurrCompetitionDate[1]
                    # The last part is the year
                    CurrCompYear = CurrCompetitionDate[2]

                    CurrEntryDeadLine = CurrEntryDeadLine.split('/')
                    # The first part is the day
                    CurrEntryDay = CurrEntryDeadLine[0]
                    # The second part is the Month
                    CurrEntryMonth = CurrEntryDeadLine[1]
                    # The last part is the year
                    CurrEntryYear = CurrEntryDeadLine[2]

                    print(CurrCompetitionDate)
                    print(CurrCompDay)
                    print(CurrCompMonth)
                    print(CurrCompYear)

                    # initialising all the textvar variables to store the data entered into the entry boxes
                    UptCompetitionName = StringVar()
                    UptCompetitionLocation = StringVar()
                    UptCompDay = StringVar()
                    UptCompMonth = StringVar()
                    UptCompYear = StringVar()
                    UptEntryDay = StringVar()
                    UptEntryMonth = StringVar()
                    UptEntryYear = StringVar()

                    EditCompDetails = Frame(contentFrame, width=965, height=600)
                    EditCompDetails.grid(row=0, column=0, sticky="nsew")

                    lblspacer = Label(EditCompDetails, width=40)
                    lblspacer.grid()

                    lblTitle = Label(EditCompDetails, text="Edit Competition Details", height=3, font=("Arvo", 32))
                    lblTitle.grid(row=0, column=1, columnspan=3)

                    lblCompName = Label(EditCompDetails, text="Competition Name: ", font=("Arvo"))
                    lblCompName.grid(row=2, column=1, sticky="e")

                    EtyCompName = Entry(EditCompDetails, textvar=UptCompetitionName)
                    # populate the entry boxes with the data already associated with Competition name
                    EtyCompName.insert(0, CurrCompetitionName)
                    EtyCompName.grid(row=2, column=2, sticky="w")

                    lblspacer = Label(EditCompDetails)
                    lblspacer.grid(row=3, column=1, sticky="e")

                    lblCompDate = Label(EditCompDetails, text="Competition Date: ", font=("Arvo"))
                    lblCompDate.grid(row=4, column=1, sticky="e")

                    listDays = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17',
                                '18', '19',
                                '20', '21', '22', '23', '24', '25', '26', '27', '28', '29', '30', '31']
                    droplistCompDay = OptionMenu(EditCompDetails, UptCompDay, *listDays)
                    droplistCompDay.config(width=20)
                    # Populate the drop list with the current data
                    UptCompDay.set(CurrCompDay)
                    droplistCompDay.grid(row=4, column=2, sticky="w")

                    listMonths = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September',
                                  'October',
                                  'November', 'December']
                    droplistCompMonth = OptionMenu(EditCompDetails, UptCompMonth, *listMonths)
                    droplistCompMonth.config(width=20)
                    UptCompMonth.set(CurrCompMonth)
                    droplistCompMonth.grid(row=5, column=2, sticky="w")

                    listYears = Years
                    droplistCompYear = OptionMenu(EditCompDetails, UptCompYear, *listYears)
                    droplistCompYear.config(width=20)
                    UptCompYear.set(CurrCompYear)
                    droplistCompYear.grid(row=6, column=2, sticky="w")

                    lblspacer = Label(EditCompDetails)
                    lblspacer.grid(row=7, column=1, sticky="e")

                    lblCompLocation = Label(EditCompDetails, text="Competition Location: ", font=("Arvo"))
                    lblCompLocation.grid(row=8, column=1, sticky="e")

                    EtyCompLocation = Entry(EditCompDetails, textvar=UptCompetitionLocation)
                    EtyCompLocation.insert(0, CurrCompetitionLocation)
                    EtyCompLocation.grid(row=8, column=2, sticky="w")

                    lblspacer = Label(EditCompDetails)
                    lblspacer.grid(row=7, column=1, sticky="e")

                    droplistCompYear = OptionMenu(EditCompDetails, UptEntryDay, *listDays)
                    droplistCompYear.config(width=20)
                    UptEntryDay.set(CurrEntryDay)
                    droplistCompYear.grid(row=10, column=2, sticky="w")

                    droplistCompYear = OptionMenu(EditCompDetails, UptEntryMonth, *listMonths)
                    droplistCompYear.config(width=20)
                    UptEntryMonth.set(CurrEntryMonth)
                    droplistCompYear.grid(row=11, column=2, sticky="w")

                    droplistCompYear = OptionMenu(EditCompDetails, UptEntryYear, *listYears)
                    droplistCompYear.config(width=20)
                    UptEntryYear.set(CurrEntryYear)
                    droplistCompYear.grid(row=12, column=2, sticky="w")

                    # This button once pressed "Creates" the competiiton as all the data is saved to the database by running the function required to do this.
                    Button(EditCompDetails, image=photoUpdateCompetitionButton,
                           command=lambda: CheckEditedCompetition(UptCompetitionName, UptCompetitionLocation, UptCompDay,
                                                                  UptCompMonth, UptCompYear, CompID, UptEntryDay,
                                                                  UptEntryMonth, UptEntryYear)) \
                        .grid(row=13, column=2, sticky="w", padx="5", pady="10")

                    SelectButton = Button(EditCompDetails, image=photoDeleteButton
                                          , command=lambda: DeleteComp(CompID, CompName))
                    SelectButton.grid(row=13, column=1, columnspan=1, sticky=W)

    EditComp = Frame(contentFrame, width=965, height=600)
    EditComp.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(EditComp, width=30)
    lblspacer.grid()

    lblTitle = Label(EditComp, text="Edit Competition", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=3, columnspan=3)

    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through the data
        cursorListOfComps = db.cursor()
        # SQL Query to select all Competitions from the competitions table
    listOfCompsQuery = 'SELECT CompetitionName, CompetitionDate, EntryDeadLine FROM Competitions'
    # Execute the query to get a list of all the competitions
    cursorListOfComps.execute(listOfCompsQuery)
    # The query will produce a 2 dimensional tuple
    ListOfComps = cursorListOfComps.fetchall()

    lblText = Label(EditComp, text="Please select the competition you would like to edit details for ", height=3,
                    font=("Arvo", 14))
    lblText.grid(row=1, column=1, columnspan=8, sticky="ew")

    lblCompetiton = Label(EditComp, text="Competition Name")
    lblCompetiton.grid(row=3, column=2, columnspan=2, sticky=EW)

    lblDate = Label(EditComp, text="Competition Date")
    lblDate.grid(row=3, column=4, columnspan=2, sticky=W)

    lblDeadLineDate = Label(EditComp, text="Entry Deadline")
    lblDeadLineDate.grid(row=3, column=6, columnspan=2, sticky=W)

    listOfCompNameBox = Listbox(EditComp, selectmode=SINGLE, width=25)
    listOfCompNameBox.grid(row=4, column=2, columnspan=2, sticky=E)

    for u in ListOfComps:
        n = 0
        listOfCompNameBox.insert(n, str(u[0]))
        n = + 1
    # listbox to show competition date but it is disabled so it cant be selected
    listOfCompDateBox = Listbox(EditComp, selectmode=DISABLED, width=15, selectbackground="white")
    listOfCompDateBox.grid(row=4, column=4, columnspan=2, sticky=W)
    for u in ListOfComps:
        n = 0
        listOfCompDateBox.insert(n, str(u[1]))
        n = + 1

    listOfDeadLineDateBox = Listbox(EditComp, selectmode=DISABLED, width=15, selectbackground="white")
    listOfDeadLineDateBox.grid(row=4, column=6, columnspan=2, sticky=W)
    for u in ListOfComps:
        n = 0
        listOfDeadLineDateBox.insert(n, str(u[2]))
        n = + 1

    SelectButton = Button(EditComp, image=photoSelectCompetitionButton
                          , command=lambda: selectedComp(photoUpdateCompetitionButton))
    SelectButton.grid(row=6, column=2, columnspan=6, pady="10")


def CollateCompResults():
    def createCommentsSheets(CompID, EntryID, CompetitionName, Location, Date, BandName, Piping1JudgeID, Piping1,
                             Piping1Comments, Piping2JudgeID, Piping2, Piping2Comments, DrummingJudgeID, Drumming,
                             DrummingComments, EnsembleJudgeID, Ensemble, EnsembleComments):
        #creates document
        document = Document()
        document.add_picture('RSPBANI_Logo.png', width=Inches(2))
        document.add_heading('              Official Results Summary', 0)

        print("Piping1JudgeID")
        print(Piping1JudgeID)

        with sqlite3.connect("RspbaniDB.db") as db:
            # Creates a cursor to search through the database
            cursor = db.cursor()
            # SQL Query to select the judges firstname and second name from the Judge Account table
        JudgeNameQuery = 'SELECT FirstName, SecondName FROM JudgeAccount WHERE JudgeID = ?'
        # Execute the query to get a list of all the bands
        cursor.execute(JudgeNameQuery, [Piping1JudgeID])
        # The query will produce a 2 dimensional tuple
        JudgeNameresults = cursor.fetchall()

        print("JudgeNameresults")
        print(JudgeNameresults)

        # gets names and concatenates them
        for u in JudgeNameresults:
            p1FirstName = u[0]
            p1SecondName = u[1]

        p1AdjudicatorName = p1FirstName + " " + p1SecondName

        lblCompetition = document.add_paragraph()
        lblCompetition.add_run('Competition:   ').bold = True
        lblCompetition.add_run(CompetitionName)
        lblLocation = document.add_paragraph()
        lblLocation.add_run('Location:    ').bold = True
        lblLocation.add_run(Location)
        lblDate = document.add_paragraph()
        lblDate.add_run('Date:    ').bold = True
        lblDate.add_run(Date)
        lblBand = document.add_paragraph()
        lblBand.add_run('Band:    ').bold = True
        lblBand.add_run(BandName)
        lblAdjudicator = document.add_paragraph()
        lblAdjudicator.add_run('Adjudicator:    ').bold = True
        lblAdjudicator.add_run(str(p1AdjudicatorName))
        lblSection = document.add_paragraph()
        lblSection.add_run('Section:    ').bold = True
        lblSection.add_run(Piping1)

        document.add_heading('', 0)

        p = document.add_paragraph(Piping1Comments)

        document.add_page_break()

        document.add_picture('RSPBANI_Logo.png', width=Inches(2))
        document.add_heading('              Official Results Summary', 0)

        with sqlite3.connect("RspbaniDB.db") as db:
            # Creates a cursor to search through the database
            cursor = db.cursor()
            # SQL Query to select the judges firstname and second name from the Judge Account table
        JudgeNameQuery = 'SELECT FirstName, SecondName FROM JudgeAccount WHERE JudgeID = ?'
        # Execute the query to get a list of all the bands
        cursor.execute(JudgeNameQuery, [Piping2JudgeID])
        # The query will produce a 2 dimensional tuple
        JudgeNameresults = cursor.fetchall()
        # gets names and concatenates them
        for u in JudgeNameresults:
            p2FirstName = u[0]
            p2SecondName = u[1]

        p2AdjudicatorName = p2FirstName + " " + p2SecondName

        lblCompetition = document.add_paragraph()
        lblCompetition.add_run('Competition:   ').bold = True
        lblCompetition.add_run(CompetitionName)
        lblLocation = document.add_paragraph()
        lblLocation.add_run('Location:    ').bold = True
        lblLocation.add_run(Location)
        lblDate = document.add_paragraph()
        lblDate.add_run('Date:    ').bold = True
        lblDate.add_run(Date)
        lblBand = document.add_paragraph()
        lblBand.add_run('Band:    ').bold = True
        lblBand.add_run(BandName)
        lblAdjudicator = document.add_paragraph()
        lblAdjudicator.add_run('Adjudicator:    ').bold = True
        lblAdjudicator.add_run(str(p2AdjudicatorName))
        lblSection = document.add_paragraph()
        lblSection.add_run('Section:    ').bold = True
        lblSection.add_run(Piping2)

        document.add_heading('', 0)

        p = document.add_paragraph(Piping2Comments)

        document.add_page_break()
        document.add_picture('RSPBANI_Logo.png', width=Inches(2))
        document.add_heading('              Official Results Summary', 0)

        with sqlite3.connect("RspbaniDB.db") as db:
            # Creates a cursor to search through the database
            cursor = db.cursor()
            # SQL Query to select the judges firstname and second name from the Judge Account table
        JudgeNameQuery = 'SELECT FirstName, SecondName FROM JudgeAccount WHERE JudgeID = ?'
        # Execute the query to get a list of all the bands
        cursor.execute(JudgeNameQuery, [DrummingJudgeID])
        # The query will produce a 2 dimensional tuple
        JudgeNameresults = cursor.fetchall()
        # gets names and concatenates them
        for u in JudgeNameresults:
            DFirstName = u[0]
            DSecondName = u[1]

        DAdjudicatorName = DFirstName + " " + DSecondName

        lblCompetition = document.add_paragraph()
        lblCompetition.add_run('Competition:   ').bold = True
        lblCompetition.add_run(CompetitionName)
        lblLocation = document.add_paragraph()
        lblLocation.add_run('Location:    ').bold = True
        lblLocation.add_run(Location)
        lblDate = document.add_paragraph()
        lblDate.add_run('Date:    ').bold = True
        lblDate.add_run(Date)
        lblBand = document.add_paragraph()
        lblBand.add_run('Band:    ').bold = True
        lblBand.add_run(BandName)
        lblAdjudicator = document.add_paragraph()
        lblAdjudicator.add_run('Adjudicator:    ').bold = True
        lblAdjudicator.add_run(str(DAdjudicatorName))
        lblSection = document.add_paragraph()
        lblSection.add_run('Section:    ').bold = True
        lblSection.add_run(Drumming)

        document.add_heading('', 0)
        p = document.add_paragraph(DrummingComments)

        document.add_page_break()

        document.add_picture('RSPBANI_Logo.png', width=Inches(2))
        document.add_heading('              Official Results Summary', 0)

        with sqlite3.connect("RspbaniDB.db") as db:
            # Creates a cursor to search through the database
            cursor = db.cursor()
            # SQL Query to select the judges firstname and second name from the Judge Account table
        JudgeNameQuery = 'SELECT FirstName, SecondName FROM JudgeAccount WHERE JudgeID = ?'
        # Execute the query to get a list of all the bands
        cursor.execute(JudgeNameQuery, [EnsembleJudgeID])
        # The query will produce a 2 dimensional tuple
        JudgeNameresults = cursor.fetchall()
        # gets names and concatenates them
        for u in JudgeNameresults:
            EFirstName = u[0]
            ESecondName = u[1]

        EAdjudicatorName = EFirstName + " " + ESecondName

        lblCompetition = document.add_paragraph()
        lblCompetition.add_run('Competition:   ').bold = True
        lblCompetition.add_run(CompetitionName)
        lblLocation = document.add_paragraph()
        lblLocation.add_run('Location:    ').bold = True
        lblLocation.add_run(Location)
        lblDate = document.add_paragraph()
        lblDate.add_run('Date:    ').bold = True
        lblDate.add_run(Date)
        lblBand = document.add_paragraph()
        lblBand.add_run('Band:    ').bold = True
        lblBand.add_run(BandName)
        lblAdjudicator = document.add_paragraph()
        lblAdjudicator.add_run('Adjudicator:    ').bold = True
        lblAdjudicator.add_run(str(EAdjudicatorName))
        lblSection = document.add_paragraph()
        lblSection.add_run('Section:    ').bold = True
        lblSection.add_run(Ensemble)

        document.add_heading('', 0)

        p = document.add_paragraph(EnsembleComments)

        document.save(str(EntryID) + '.docx')
        BandSheetFileName = str(EntryID) + '.docx'

        conn = sqlite3.connect('RspbaniDB.db')
        with conn:
            cursor = conn.cursor()
        cursor.execute('UPDATE Results SET BandSheetsFileName = ?'
                       'WHERE EntryID = ?', (BandSheetFileName, EntryID))
        conn.commit()

        conn = sqlite3.connect('RspbaniDB.db')
        with conn:
            cursor = conn.cursor()
        cursor.execute('UPDATE Competitions SET SentToBands = ?'
                       'WHERE CompetitionID = ?', ("False", CompID))
        conn.commit()


    def GetCompToBeCollated(listOfCompNameBox):

        def appendDetailsToList(List, BandName, Piping1Score, Piping2Score, TotalPipingScore, DrummingScore,
                                EnsembleScore, TotalScore):

            List[
                TotalScore] = BandName, Piping1Score, Piping2Score, TotalPipingScore, DrummingScore, EnsembleScore, TotalScore
            print(List)

        # gets the position of the Competition Selected
        Selection = listOfCompNameBox.curselection()

        if Selection ==():
            messagebox.showinfo(message="Please select a competition to collate.")
        else:

            print(Selection)

            for i in Selection:
                # gets the name of the competition
                j = listOfCompNameBox.get(i)
                CompetitionName = j

                #Sort Bands into grades
                G1BandScores = {}
                G2BandScores = {}
                G3ABandScores = {}
                G3BBandScores = {}
                G4ABandScores = {}
                G4BBandScores = {}

                # gets the all the required fields to be outputted into the Results sheet of the competition.
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCompID = db.cursor()
                findCompID = 'SELECT CompetitionLocation, CompetitionDate, G1BestBass, G2BestBass, G3ABestBass, G3BBestBass, G4ABestBass, G4BBestBass, G1BestMD, G2BestMD, G3ABestMD, G3BBestMD,G4ABestMD, G4BBestMD, CompetitionID, NumG1Ent, NumG2Ent, NumG3AEnt, NumG3BEnt, NumG4AEnt, NumG4BEnt FROM Competitions WHERE CompetitionName = ?'
                cursorCompID.execute(findCompID, [j])
                CompIDresults = cursorCompID.fetchall()

                if CompIDresults:

                    for a in CompIDresults:
                        Location = a[0]
                        Date = a[1]
                        G1BestBass = a[2]
                        G2BestBass = a[3]
                        G3ABestBass = a[4]
                        G3BBestBass = a[5]
                        G4ABestBass = a[6]
                        G4BBestBass = a[7]
                        G1BestMD = a[8]
                        G2BestMD = a[9]
                        G3ABestMD = a[10]
                        G3BBestMD = a[11]
                        G4ABestMD = a[12]
                        G4BBestMD = a[13]
                        CompID = a[14]
                        NumG1Ent = a[15]
                        NumG2Ent = a[16]
                        NumG3AEnt = a[17]
                        NumG3BEnt = a[18]
                        NumG4AEnt = a[19]
                        NumG4BEnt = a[20]

                        print("Grade3B BestBass")
                        print(G3BBestBass)

                        if G1BestBass == None:
                            G1BestBass = " "
                        if G2BestBass == None:
                            G2BestBass = " "
                        if G3ABestBass == None:
                            G3ABestBass = " "
                        if G3BBestBass == None:
                            G3BBestBass = " "
                        if G4ABestBass == None:
                            G4ABestBass = " "
                        if G4BBestBass == None:
                            G4BBestBass = " "

                        if G1BestMD == None:
                            G1BestMD = " "
                        if G2BestMD == None:
                            G2BestMD = " "
                        if G3ABestMD == None:
                            G3ABestMD = " "
                        if G3BBestMD == None:
                            G3BBestMD = " "
                        if G4ABestMD == None:
                            G4ABestMD = " "
                        if G4BBestMD == None:
                            G4BBestMD = " "


                        print(CompID)

                # get all EntryID of all entered Bands in the Competition in Grade 1
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCompID = db.cursor()
                findCompID = 'SELECT EntryID, BandID FROM BandsEntered WHERE CompetitionID = ?'
                cursorCompID.execute(findCompID, [CompID])
                EntryIDresults = cursorCompID.fetchall()

                if EntryIDresults:
                    print("EntryIDresults")
                    print(EntryIDresults)

                    for e in EntryIDresults:
                        EntryID = e[0]
                        print("EntryID")
                        print(EntryID)
                        BandID = e[1]

                        # calculate total score for each band
                        with sqlite3.connect("RspbaniDB.db") as db:
                            cursor = db.cursor()
                        findScores = 'SELECT Piping1Score, Piping2Score, TotalPipingScore, DrummingScore, EnsembleScore FROM Results WHERE EntryID = ?'
                        cursor.execute(findScores, [EntryID])
                        BandScores = cursor.fetchall()

                        print("BandScores")
                        print(BandScores)

                        if BandScores:

                            for score in BandScores:
                                Piping1Score = score[0]
                                Piping2Score = score[1]
                                TotalPipingScore = score[2]
                                DrummingScore = score[3]
                                EnsembleScore = score[4]

                            TotalScore = Piping1Score + Piping2Score + DrummingScore + EnsembleScore

                            with sqlite3.connect("RspbaniDB.db") as db:
                                cursorCompID = db.cursor()
                            cursorCompID.execute('UPDATE Results SET TotalScore = ? WHERE EntryID = ?',
                                                 (TotalScore, EntryID))
                            db.commit()

                        with sqlite3.connect("RspbaniDB.db") as db:
                            cursorCompID = db.cursor()
                        findCompID = 'SELECT BandGrade, BandName FROM BandAccount WHERE BandID = ?'
                        cursorCompID.execute(findCompID, [BandID])
                        Graderesults = cursorCompID.fetchall()

                        if Graderesults:
                            print("Graderesults")
                            print(Graderesults)

                            for b in Graderesults:
                                Grade = b[0]
                                BandName = b[1]

                        with sqlite3.connect("RspbaniDB.db") as db:
                            cursorCompID = db.cursor()
                        findCompID = 'SELECT Piping1Score, Piping2Score, TotalPipingScore, DrummingScore, EnsembleScore, TotalScore, Piping1JudgeID, Piping1Comments, Piping2JudgeID, Piping2Comments, DrummingJudgeID, DrummingComments, EnsembleJudgeID, EnsembleComments FROM Results WHERE EntryID = ?'
                        cursorCompID.execute(findCompID, [EntryID])
                        BandScoreresults = cursorCompID.fetchall()

                        print("BandScoreresults")
                        print(BandScoreresults)

                        if BandScoreresults:

                            for BS in BandScoreresults:
                                Piping1Score = BS[0]
                                Piping2Score = BS[1]
                                TotalPipingScore = BS[2]
                                DrummingScore = BS[3]
                                EnsembleScore = BS[4]
                                TotalScore = BS[5]
                                Piping1JudgeID = BS[6]
                                Piping1Comments = BS[7]
                                Piping2JudgeID = BS[8]
                                Piping2Comments = BS[9]
                                DrummingJudgeID = BS[10]
                                DrummingComments = BS[11]
                                EnsembleJudgeID = BS[12]
                                EnsembleComments = BS[13]

                        print("Grade" + Grade)
                        if str(Grade) == "1":
                            appendDetailsToList(G1BandScores, BandName, Piping1Score, Piping2Score, TotalPipingScore,
                                                DrummingScore, EnsembleScore, TotalScore)
                            createCommentsSheets(CompID, EntryID, CompetitionName, Location, Date, BandName, Piping1JudgeID,
                                                 "Piping 1", Piping1Comments, Piping2JudgeID, "Piping 2", Piping2Comments,
                                                 DrummingJudgeID, "Drumming", DrummingComments, EnsembleJudgeID, "Ensemble",
                                                 EnsembleComments)
                        elif str(Grade) == "2":
                            appendDetailsToList(G2BandScores, BandName, Piping1Score, Piping2Score, TotalPipingScore,
                                                DrummingScore, EnsembleScore, TotalScore)
                            createCommentsSheets(CompID, EntryID, CompetitionName, Location, Date, BandName, Piping1JudgeID,
                                                 "Piping 1", Piping1Comments, Piping2JudgeID, "Piping 2", Piping2Comments,
                                                 DrummingJudgeID, "Drumming", DrummingComments, EnsembleJudgeID, "Ensemble",
                                                 EnsembleComments)
                        elif str(Grade) == "3A":
                            appendDetailsToList(G3ABandScores, BandName, Piping1Score, Piping2Score, TotalPipingScore,
                                                DrummingScore, EnsembleScore, TotalScore)
                            createCommentsSheets(CompID, EntryID, CompetitionName, Location, Date, BandName, Piping1JudgeID,
                                                 "Piping 1", Piping1Comments, Piping2JudgeID, "Piping 2", Piping2Comments,
                                                 DrummingJudgeID, "Drumming", DrummingComments, EnsembleJudgeID, "Ensemble",
                                                 EnsembleComments)
                        elif str(Grade) == "3B":
                            appendDetailsToList(G3BBandScores, BandName, Piping1Score, Piping2Score, TotalPipingScore,
                                                DrummingScore, EnsembleScore, TotalScore)
                            createCommentsSheets(CompID, EntryID, CompetitionName, Location, Date, BandName, Piping1JudgeID,
                                                 "Piping 1", Piping1Comments, Piping2JudgeID, "Piping 2", Piping2Comments,
                                                 DrummingJudgeID, "Drumming", DrummingComments, EnsembleJudgeID, "Ensemble",
                                                 EnsembleComments)
                        elif str(Grade) == "4A":
                            appendDetailsToList(G4ABandScores, BandName, Piping1Score, Piping2Score, TotalPipingScore,
                                                DrummingScore, EnsembleScore, TotalScore)
                            createCommentsSheets(CompID, EntryID, CompetitionName, Location, Date, BandName, Piping1JudgeID,
                                                 "Piping 1", Piping1Comments, Piping2JudgeID, "Piping 2", Piping2Comments,
                                                 DrummingJudgeID, "Drumming", DrummingComments, EnsembleJudgeID, "Ensemble",
                                                 EnsembleComments)
                        elif str(Grade) == "4B":
                            appendDetailsToList(G4BBandScores, BandName, Piping1Score, Piping2Score, TotalPipingScore,
                                                DrummingScore, EnsembleScore, TotalScore)
                            createCommentsSheets(CompID, EntryID, CompetitionName, Location, Date, BandName, Piping1JudgeID,
                                                 "Piping 1", Piping1Comments, Piping2JudgeID, "Piping 2", Piping2Comments,
                                                 DrummingJudgeID, "Drumming", DrummingComments, EnsembleJudgeID, "Ensemble",
                                                 EnsembleComments)
                    messagebox.showinfo(message="Sheets for " + CompetitionName + " Successfully Created!")

                print("G2BandScores")
                print(G2BandScores)

                # Ranked by Total Score to work out the position of the bands in each grade from their total score
                AllBandsInCompRank = []
                G1BandsInCompRanked = []
                G2BandsInCompRanked = []
                G3ABandsInCompRanked = []
                G3BBandsInCompRanked = []
                G4ABandsInCompRanked = []
                G4BBandsInCompRanked = []

                with sqlite3.connect("RspbaniDB.db") as db:
                    cursor = db.cursor()
                findposition = 'SELECT BandID FROM Results WHERE CompetitionID = ? ORDER BY TotalScore, EnsembleScore'
                cursor.execute(findposition, [CompID])
                bandIDinposition = cursor.fetchall()

                print("bandIDinposition")
                print(bandIDinposition)

                if bandIDinposition:
                    for q in bandIDinposition:
                        for band in q:
                            AllBandsInCompRank.append(band)

                print("AllBandsInCompRank")
                print(AllBandsInCompRank)

                for I in AllBandsInCompRank:
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursor = db.cursor()
                    findGrade = 'SELECT BandGrade FROM BandAccount WHERE BandID = ?'
                    cursor.execute(findGrade, [I])
                    BandGrade = cursor.fetchall()

                    if BandGrade:
                        for g in BandGrade:
                            for G in g:
                                print("G")
                                print(G)

                                if G == "1":
                                    G1BandsInCompRanked.append(I)
                                elif G == "2":
                                    G2BandsInCompRanked.append(I)
                                elif G == "3A":
                                    G3ABandsInCompRanked.append(I)
                                elif G == "3B":
                                    G3BBandsInCompRanked.append(I)
                                elif G == "4A":
                                    G4ABandsInCompRanked.append(I)
                                elif G == "4B":
                                    G4BBandsInCompRanked.append(I)

                print("G1BandsInCompRanked")
                print(G1BandsInCompRanked)
                print("G2BandsInCompRanked")
                print(G2BandsInCompRanked)
                print("G3ABandsInCompRanked")
                print(G3ABandsInCompRanked)
                print("G3BBandsInCompRanked")
                print(G3BBandsInCompRanked)
                print("G4ABandsInCompRanked")
                print(G4ABandsInCompRanked)
                print("G4BBandsInCompRanked")
                print(G4BBandsInCompRanked)

                try:
                    for Y in range(1, len(G1BandsInCompRanked) + 1):
                        print("G1" + str(Y))
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        # Insert the fields which have been fetched and insert them into the database
                        cursor.execute(
                            'UPDATE Results SET Position = ? WHERE BandID = ? AND CompetitionID = ?',
                            (Y, G1BandsInCompRanked[Y - 1], CompID))
                        conn.commit()
                except:
                    None

                try:
                    for Y in range(1, len(G2BandsInCompRanked) + 1):
                        print("G2" + str(Y))

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        # Insert the fields which have been fetched and insert them into the database
                        cursor.execute(
                            'UPDATE Results SET Position = ? WHERE BandID = ? AND CompetitionID = ?',
                            (Y, G2BandsInCompRanked[Y - 1], CompID))
                        conn.commit()
                except:
                    None
                try:
                    for Y in range(1, len(G3ABandsInCompRanked) + 1):
                        print("G3a" + str(Y))

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        # Insert the fields which have been fetched and insert them into the database
                        cursor.execute(
                            'UPDATE Results SET Position = ? WHERE BandID = ? AND CompetitionID = ?',
                            (Y, G3ABandsInCompRanked[Y - 1], CompID))
                        conn.commit()
                except:
                    None
                try:
                    for Y in range(1, len(G3BBandsInCompRanked) + 1):
                        print("G3b" + str(Y))

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        # Insert the fields which have been fetched and insert them into the database
                        cursor.execute(
                            'UPDATE Results SET Position = ? WHERE BandID = ? AND CompetitionID = ?',
                            (Y, G3BBandsInCompRanked[Y - 1], CompID))
                        conn.commit()
                except:
                    None

                try:
                    for Y in range(1, len(G4ABandsInCompRanked) + 1):
                        print("G4a" + str(Y))

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        # Insert the fields which have been fetched and insert them into the database
                        cursor.execute(
                            'UPDATE Results SET Position = ? WHERE BandID = ? AND CompetitionID = ?',
                            (Y, G4ABandsInCompRanked[Y - 1], CompID))
                        conn.commit()
                except:
                    None

                try:
                    for Y in range(1, len(G4BBandsInCompRanked) + 1):
                        print("G4b" + str(Y))

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        # Insert the fields which have been fetched and insert them into the database
                        cursor.execute(
                            'UPDATE Results SET Position = ? WHERE BandID = ? AND CompetitionID = ?',
                            (Y, G4BBandsInCompRanked[Y - 1], CompID))
                        conn.commit()
                except:
                    None

                # Ranked by Total PipingScore
                AllBandsInCompRank = []
                G1BandsInCompRanked = []
                G2BandsInCompRanked = []
                G3ABandsInCompRanked = []
                G3BBandsInCompRanked = []
                G4ABandsInCompRanked = []
                G4BBandsInCompRanked = []

                with sqlite3.connect("RspbaniDB.db") as db:
                    cursor = db.cursor()
                findposition = 'SELECT BandID FROM Results WHERE CompetitionID = ? ORDER BY TotalPipingScore'
                cursor.execute(findposition, [CompID])
                bandIDinposition = cursor.fetchall()

                print("bandIDinposition")
                print(bandIDinposition)

                if bandIDinposition:
                    for q in bandIDinposition:
                        for band in q:
                            AllBandsInCompRank.append(band)

                print("AllBandsInCompRank")
                print(AllBandsInCompRank)

                for I in AllBandsInCompRank:
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursor = db.cursor()
                    findGrade = 'SELECT BandGrade, BandName FROM BandAccount WHERE BandID = ?'
                    cursor.execute(findGrade, [I])
                    BandGrade = cursor.fetchall()

                    if BandGrade:
                        for G in BandGrade:

                            if G[0] == "1":
                                G1BandsInCompRanked.append(G[1])
                            elif G[0] == "2":
                                G2BandsInCompRanked.append(G[1])
                            elif G[0] == "3A":
                                G3ABandsInCompRanked.append(G[1])
                            elif G[0] == "3B":
                                G3BBandsInCompRanked.append(G[1])
                            elif G[0] == "4A":
                                G4ABandsInCompRanked.append(G[1])
                            elif G[0] == "4B":
                                G4BBandsInCompRanked.append(G[1])

                try:
                    G1BestPipes = G1BandsInCompRanked[0]
                except:
                    G1BestPipes = " "

                try:
                    G2BestPipes = G2BandsInCompRanked[0]

                except:
                    G2BestPipes = " "

                try:
                    G3ABestPipes = G3ABandsInCompRanked[0]
                except:
                    G3ABestPipes = " "

                try:
                    G3BBestPipes = G3BBandsInCompRanked[0]
                except:
                    G3BBestPipes = " "

                try:
                    G4ABestPipes = G4ABandsInCompRanked[0]
                except:
                    G4ABestPipes = " "

                try:
                    G4BBestPipes = G4BBandsInCompRanked[0]
                except:
                    G4BBestPipes = " "

                # Ranked by Total PipingScore
                AllBandsInCompRank = []
                G1BandsInCompRanked = []
                G2BandsInCompRanked = []
                G3ABandsInCompRanked = []
                G3BBandsInCompRanked = []
                G4ABandsInCompRanked = []
                G4BBandsInCompRanked = []

                with sqlite3.connect("RspbaniDB.db") as db:
                    cursor = db.cursor()
                findposition = 'SELECT BandID FROM Results WHERE CompetitionID = ? ORDER BY DrummingScore'
                cursor.execute(findposition, [CompID])
                bandIDinposition = cursor.fetchall()

                print("bandIDinposition")
                print(bandIDinposition)

                if bandIDinposition:
                    for q in bandIDinposition:
                        for band in q:
                            AllBandsInCompRank.append(band)

                print("AllBandsInCompRank")
                print(AllBandsInCompRank)

                for I in AllBandsInCompRank:
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursor = db.cursor()
                    findGrade = 'SELECT BandGrade, BandName FROM BandAccount WHERE BandID = ?'
                    cursor.execute(findGrade, [I])
                    BandGrade = cursor.fetchall()

                    if BandGrade:
                        for G in BandGrade:
                            print("G")
                            print(G)

                            if G[0] == "1":
                                G1BandsInCompRanked.append(G[1])
                            elif G[0] == "2":
                                G2BandsInCompRanked.append(G[1])
                            elif G[0] == "3A":
                                G3ABandsInCompRanked.append(G[1])
                            elif G[0] == "3B":
                                G3BBandsInCompRanked.append(G[1])
                            elif G[0] == "4A":
                                G4ABandsInCompRanked.append(G[1])
                            elif G[0] == "4B":
                                G4BBandsInCompRanked.append(G[1])
                try:
                    G1BestDrum = G1BandsInCompRanked[0]

                except:
                    G1BestDrum = " "
                try:
                    G2BestDrum = G2BandsInCompRanked[0]

                except:
                    G2BestDrum = " "

                try:
                    G3ABestDrum = G3ABandsInCompRanked[0]
                except:
                    G3ABestDrum = " "

                try:
                    G3BBestDrum = G3BBandsInCompRanked[0]
                except:
                    G3BBestDrum = " "

                try:
                    G4ABestDrum = G4ABandsInCompRanked[0]
                except:
                    G4ABestDrum = " "

                try:
                    G4BBestDrum = G4BBandsInCompRanked[0]
                except:
                    G4BBestDrum = " "

                # Ranked by Position
                AllBandsInCompRankPos = []
                G1BandsInCompRankedPos = []
                G2BandsInCompRankedPos = []
                G3ABandsInCompRankedPos = []
                G3BBandsInCompRankedPos = []
                G4ABandsInCompRankedPos = []
                G4BBandsInCompRankedPos = []

                with sqlite3.connect("RspbaniDB.db") as db:
                    cursor = db.cursor()
                findposition = 'SELECT BandID, Piping1Score, Piping2Score, TotalPipingScore, DrummingScore, EnsembleScore, TotalScore, Position FROM Results WHERE CompetitionID = ? ORDER BY Position'
                cursor.execute(findposition, [CompID])
                bandIDRankPosition = cursor.fetchall()

                print("bandIDRankPosition")
                print(bandIDRankPosition)

                if bandIDRankPosition:
                    for band in bandIDRankPosition:
                        Temp = []
                        Temp.append(band[0])
                        Temp.append(band[1])
                        Temp.append(band[2])
                        Temp.append(band[3])
                        Temp.append(band[4])
                        Temp.append(band[5])
                        Temp.append(band[6])
                        Temp.append(band[7])
                        print("Temp")
                        print(Temp)

                        AllBandsInCompRankPos.append(Temp)

                print("AllBandsInCompRankPos")
                print(AllBandsInCompRankPos)

                for I in AllBandsInCompRankPos:
                    print("I[0]")
                    print(I[0])
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursor = db.cursor()
                    findGrade = 'SELECT BandGrade, BandName FROM BandAccount WHERE BandID = ?'
                    cursor.execute(findGrade, [I[0]])
                    BandGrade = cursor.fetchall()

                    if BandGrade:
                        for G in BandGrade:
                            if G[0] == "1":
                                I.append(G[1])
                                G1BandsInCompRankedPos.append(I)
                            elif G[0] == "2":
                                I.append(G[1])
                                G2BandsInCompRankedPos.append(I)
                            elif G[0] == "3A":
                                I.append(G[1])
                                G3ABandsInCompRankedPos.append(I)
                            elif G[0] == "3B":
                                I.append(G[1])
                                G3BBandsInCompRankedPos.append(I)
                            elif G[0] == "4A":
                                I.append(G[1])
                                G4ABandsInCompRankedPos.append(I)
                            elif G[0] == "4B":
                                I.append(G[1])
                                G4BBandsInCompRankedPos.append(I)

                print("G1BandsInCompRankedPos")
                print(G1BandsInCompRankedPos)
                print("G2BandsInCompRankedPos")
                print(G2BandsInCompRankedPos)
                print("G3ABandsInCompRankedPos")
                print(G3ABandsInCompRankedPos)
                print("G3BBandsInCompRankedPos")
                print(G3BBandsInCompRankedPos)
                print("G4ABandsInCompRankedPos")
                print(G4ABandsInCompRankedPos)
                print("G4BBandsInCompRankedPos")
                print(G4BBandsInCompRankedPos)

                document = Document()
                document.add_picture('RSPBANI_Logo.png', width=Inches(2))
                document.add_heading('              Official Results Summary', 0)

                lblCompetition = document.add_paragraph()
                lblCompetition.add_run('Competition:   ').bold = True
                lblCompetition.add_run(CompetitionName)
                lblLocation = document.add_paragraph()
                lblLocation.add_run('Location:    ').bold = True
                lblLocation.add_run(Location)
                lblDate = document.add_paragraph()
                lblDate.add_run('Date:    ').bold = True
                lblDate.add_run(Date)

                document.add_heading('Grade 1 Results', level=1)

                Grade1Results = document.add_table(rows=2, cols=9)
                Grade1Results.alignment = WD_TABLE_ALIGNMENT.CENTER

                hdr_cells = Grade1Results.rows[0].cells
                hdr_cells[1].text = 'Band'
                hdr_cells[2].text = 'Piping 1'
                hdr_cells[3].text = 'Piping 2'
                hdr_cells[4].text = 'Total Piping'
                hdr_cells[5].text = 'Drumming'
                hdr_cells[6].text = 'Ensemble'
                hdr_cells[7].text = 'Total'
                hdr_cells[8].text = 'Overall Position'

                print("G1BandsInCompRankedPos")
                print(G1BandsInCompRankedPos)
                for each in G1BandsInCompRankedPos:
                    x = 1
                    Content_cells = Grade1Results.add_row().cells
                    Content_cells[1].text = str(each[8])
                    Content_cells[2].text = str(each[1])
                    Content_cells[3].text = str(each[2])
                    Content_cells[4].text = str(each[3])
                    Content_cells[5].text = str(each[4])
                    Content_cells[6].text = str(each[5])
                    Content_cells[7].text = str(each[6])
                    Content_cells[8].text = str(each[7])

                    x += 1

                Content_cells = Grade1Results.add_row().cells
                Content_cells = Grade1Results.add_row().cells

                G1Awards = document.add_table(rows=5, cols=3)
                hdr_cells = G1Awards.rows[1].cells
                hdr_cells[0].text = 'Best Pipes:  '
                hdr_cells[1].text = str(G1BestPipes)

                hdr_cells = G1Awards.rows[2].cells
                hdr_cells[0].text = 'Best Drum Corp:  '
                hdr_cells[1].text = str(G1BestDrum)

                hdr_cells = G1Awards.rows[3].cells
                hdr_cells[0].text = 'Best Bass Section:  '
                hdr_cells[1].text = str(G1BestBass)

                hdr_cells = G1Awards.rows[4].cells
                hdr_cells[0].text = 'Best M&D:  '
                hdr_cells[1].text = str(G1BestMD)
                document.add_page_break()

                document.add_heading('Grade 2 Results', level=1)

                Grade2Results = document.add_table(rows=2, cols=9)
                hdr_cells = Grade2Results.rows[0].cells
                hdr_cells[1].text = 'Band'
                hdr_cells[2].text = 'Piping 1'
                hdr_cells[3].text = 'Piping 2'
                hdr_cells[4].text = 'Total Piping'
                hdr_cells[5].text = 'Drumming'
                hdr_cells[6].text = 'Ensemble'
                hdr_cells[7].text = 'Total'
                hdr_cells[8].text = 'Overall Position'

                for each in G2BandsInCompRankedPos:
                    x = 1
                    Content_cells = Grade2Results.add_row().cells
                    Content_cells[1].text = str(each[8])
                    Content_cells[2].text = str(each[1])
                    Content_cells[3].text = str(each[2])
                    Content_cells[4].text = str(each[3])
                    Content_cells[5].text = str(each[4])
                    Content_cells[6].text = str(each[5])
                    Content_cells[7].text = str(each[6])
                    Content_cells[8].text = str(each[7])
                    x += 1

                Content_cells = Grade2Results.add_row().cells
                Content_cells = Grade2Results.add_row().cells

                G1Awards = document.add_table(rows=5, cols=3)

                hdr_cells = G1Awards.rows[1].cells
                hdr_cells[0].text = 'Best Pipes:  '
                hdr_cells[1].text = str(G2BestPipes)

                hdr_cells = G1Awards.rows[2].cells
                hdr_cells[0].text = 'Best Drum Corp:  '
                hdr_cells[1].text = str(G2BestDrum)

                hdr_cells = G1Awards.rows[3].cells
                hdr_cells[0].text = 'Best Bass Section:  '
                hdr_cells[1].text = str(G2BestBass)

                hdr_cells = G1Awards.rows[4].cells
                hdr_cells[0].text = 'Best M&D:  '
                hdr_cells[1].text = str(G2BestMD)
                document.add_page_break()

                document.add_heading('Grade 3A Results', level=1)

                Grade3AResults = document.add_table(rows=2, cols=9)
                hdr_cells = Grade3AResults.rows[0].cells
                hdr_cells[1].text = 'Band'
                hdr_cells[2].text = 'Piping 1'
                hdr_cells[3].text = 'Piping 2'
                hdr_cells[4].text = 'Total Piping'
                hdr_cells[5].text = 'Drumming'
                hdr_cells[6].text = 'Ensemble'
                hdr_cells[7].text = 'Total'
                hdr_cells[8].text = 'Overall Position'

                for each in G3ABandsInCompRankedPos:
                    x = 1
                    Content_cells = Grade3AResults.add_row().cells
                    Content_cells[1].text = str(each[8])
                    Content_cells[2].text = str(each[1])
                    Content_cells[3].text = str(each[2])
                    Content_cells[4].text = str(each[3])
                    Content_cells[5].text = str(each[4])
                    Content_cells[6].text = str(each[5])
                    Content_cells[7].text = str(each[6])
                    Content_cells[8].text = str(each[7])
                    x += 1

                Content_cells = Grade3AResults.add_row().cells
                Content_cells = Grade3AResults.add_row().cells

                G1Awards = document.add_table(rows=5, cols=3)
                hdr_cells = G1Awards.rows[1].cells
                hdr_cells[0].text = 'Best Pipes:  '
                hdr_cells[1].text = str(G3ABestPipes)

                hdr_cells = G1Awards.rows[2].cells
                hdr_cells[0].text = 'Best Drum Corp:  '
                hdr_cells[1].text = str(G3ABestDrum)

                hdr_cells = G1Awards.rows[3].cells
                hdr_cells[0].text = 'Best Bass Section:  '
                hdr_cells[1].text = str(G3ABestBass)

                hdr_cells = G1Awards.rows[4].cells
                hdr_cells[0].text = 'Best M&D:  '
                hdr_cells[1].text = str(G3ABestMD)
                document.add_page_break()

                document.add_heading('Grade 3B Results', level=1)

                Grade3BResults = document.add_table(rows=2, cols=9)
                hdr_cells = Grade3BResults.rows[0].cells
                hdr_cells[1].text = 'Band'
                hdr_cells[2].text = 'Piping 1'
                hdr_cells[3].text = 'Piping 2'
                hdr_cells[4].text = 'Total Piping'
                hdr_cells[5].text = 'Drumming'
                hdr_cells[6].text = 'Ensemble'
                hdr_cells[7].text = 'Total'
                hdr_cells[8].text = 'Overall Position'

                for each in G3BBandsInCompRankedPos:
                    x = 1
                    Content_cells = Grade3BResults.add_row().cells
                    Content_cells[1].text = str(each[8])
                    Content_cells[2].text = str(each[1])
                    Content_cells[3].text = str(each[2])
                    Content_cells[4].text = str(each[3])
                    Content_cells[5].text = str(each[4])
                    Content_cells[6].text = str(each[5])
                    Content_cells[7].text = str(each[6])
                    Content_cells[8].text = str(each[7])
                    x += 1

                Content_cells = Grade3BResults.add_row().cells
                Content_cells = Grade3BResults.add_row().cells

                G1Awards = document.add_table(rows=5, cols=3)
                hdr_cells = G1Awards.rows[1].cells
                hdr_cells[0].text = 'Best Pipes:  '
                hdr_cells[1].text = str(G3BBestPipes)

                hdr_cells = G1Awards.rows[2].cells
                hdr_cells[0].text = 'Best Drum Corp:  '
                hdr_cells[1].text = str(G3BBestDrum)

                hdr_cells = G1Awards.rows[3].cells
                hdr_cells[0].text = 'Best Bass Section:  '
                hdr_cells[1].text = str(G3BBestBass)

                hdr_cells = G1Awards.rows[4].cells
                hdr_cells[0].text = 'Best M&D:  '
                hdr_cells[1].text = str(G3BBestMD)
                document.add_page_break()

                document.add_heading('Grade 4A Results', level=1)

                Grade4AResults = document.add_table(rows=2, cols=9)
                hdr_cells = Grade4AResults.rows[0].cells
                hdr_cells[1].text = 'Band'
                hdr_cells[2].text = 'Piping 1'
                hdr_cells[3].text = 'Piping 2'
                hdr_cells[4].text = 'Total Piping'
                hdr_cells[5].text = 'Drumming'
                hdr_cells[6].text = 'Ensemble'
                hdr_cells[7].text = 'Total'
                hdr_cells[8].text = 'Overall Position'

                for each in G4ABandsInCompRankedPos:
                    x = 1
                    Content_cells = Grade4AResults.add_row().cells
                    Content_cells[1].text = str(each[8])
                    Content_cells[2].text = str(each[1])
                    Content_cells[3].text = str(each[2])
                    Content_cells[4].text = str(each[3])
                    Content_cells[5].text = str(each[4])
                    Content_cells[6].text = str(each[5])
                    Content_cells[7].text = str(each[6])
                    Content_cells[8].text = str(each[7])
                    x += 1

                Content_cells = Grade4AResults.add_row().cells
                Content_cells = Grade4AResults.add_row().cells

                G1Awards = document.add_table(rows=5, cols=3)
                hdr_cells = G1Awards.rows[1].cells
                hdr_cells[0].text = 'Best Pipes:  '
                hdr_cells[1].text = str(G4ABestPipes)

                hdr_cells = G1Awards.rows[2].cells
                hdr_cells[0].text = 'Best Drum Corp:  '
                hdr_cells[1].text = str(G4ABestDrum)

                hdr_cells = G1Awards.rows[3].cells
                hdr_cells[0].text = 'Best Bass Section:  '
                hdr_cells[1].text = str(G4ABestBass)

                hdr_cells = G1Awards.rows[4].cells
                hdr_cells[0].text = 'Best M&D:  '
                hdr_cells[1].text = str(G4ABestMD)

                document.add_page_break()

                document.add_heading('Grade 4B Results', level=1)

                Grade4BResults = document.add_table(rows=2, cols=9)
                hdr_cells = Grade4BResults.rows[0].cells
                hdr_cells[1].text = 'Band'
                hdr_cells[2].text = 'Piping 1'
                hdr_cells[3].text = 'Piping 2'
                hdr_cells[4].text = 'Total Piping'
                hdr_cells[5].text = 'Drumming'
                hdr_cells[6].text = 'Ensemble'
                hdr_cells[7].text = 'Total'
                hdr_cells[8].text = 'Overall Position'

                for each in G4BBandsInCompRankedPos:
                    x = 1
                    Content_cells = Grade4BResults.add_row().cells
                    Content_cells[1].text = str(each[8])
                    Content_cells[2].text = str(each[1])
                    Content_cells[3].text = str(each[2])
                    Content_cells[4].text = str(each[3])
                    Content_cells[5].text = str(each[4])
                    Content_cells[6].text = str(each[5])
                    Content_cells[7].text = str(each[6])
                    Content_cells[8].text = str(each[7])
                    x += 1

                Content_cells = Grade4BResults.add_row().cells
                Content_cells = Grade4BResults.add_row().cells

                G1Awards = document.add_table(rows=5, cols=3)
                hdr_cells = G1Awards.rows[1].cells
                hdr_cells[0].text = 'Best Pipes:  '
                hdr_cells[1].text = str(G4BBestPipes)

                hdr_cells = G1Awards.rows[2].cells
                hdr_cells[0].text = 'Best Drum Corp:  '
                hdr_cells[1].text = str(G4BBestDrum)

                hdr_cells = G1Awards.rows[3].cells
                hdr_cells[0].text = 'Best Bass Section:  '
                hdr_cells[1].text = str(G4BBestBass)

                hdr_cells = G1Awards.rows[4].cells
                hdr_cells[0].text = 'Best M&D:  '
                hdr_cells[1].text = str(G4BBestMD)

                document.save(str(CompID) + ' - ' + CompetitionName + '.docx')

                documentName = str(CompID) + ' - ' + CompetitionName + '.docx'

                # update the competition details with the new details from collating results
                conn = sqlite3.connect('RspbaniDB.db')
                with conn:
                    cursor = conn.cursor()
                cursor.execute('UPDATE Competitions SET Collated = ?, ResultFileName = ?'
                               'WHERE CompetitionID = ?', ("True", documentName, CompID))
                conn.commit()

                messagebox.showinfo(message="Full Result Summary for " + CompetitionName + " has successfully been Created!")
                CollateCompResults()

    CollateComp = Frame(contentFrame, width=965, height=600)
    CollateComp.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(CollateComp, width=25)
    lblspacer.grid()

    lblTitle = Label(CollateComp, text="Collate Competition Results", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=1, columnspan=7)

    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through the data
        cursorListOfComps = db.cursor()
        # SQL Query to select all Competitions from the competitions table
    listOfCompsQuery = 'SELECT CompetitionName, CompetitionDate FROM Competitions WHERE Collated = "False"'
    # Execute the query to get a list of all the competitions
    cursorListOfComps.execute(listOfCompsQuery)
    # The query will produce a 2 dimensional tuple
    ListOfComps = cursorListOfComps.fetchall()

    lblText = Label(CollateComp, text="Please select the competition you would like to produce the overall Results ",
                    height=3,
                    font=("Arvo", 14))
    lblText.grid(row=1, column=1, columnspan=8, sticky="ew")

    lblCompetiton = Label(CollateComp, text="Competition Name")
    lblCompetiton.grid(row=3, column=2, columnspan=3, sticky=EW)

    lblDate = Label(CollateComp, text="Competition Date")
    lblDate.grid(row=3, column=5, columnspan=2, sticky=W)

    listOfCompNameBox = Listbox(CollateComp, selectmode=SINGLE, width=30)
    listOfCompNameBox.grid(row=4, column=2, columnspan=3, sticky=E)

    for u in ListOfComps:
        n = 0
        listOfCompNameBox.insert(n, str(u[0]))
        n = + 1

    # listbox to show competition date but it is disabled so it cant be selected
    listOfCompDateBox = Listbox(CollateComp, selectmode=DISABLED, width=20, selectbackground="white")
    listOfCompDateBox.grid(row=4, column=5, columnspan=2, sticky=W)
    for u in ListOfComps:
        n = 0
        listOfCompDateBox.insert(n, str(u[1]))
        n = + 1

    SelectButton = Button(CollateComp, image=photoSelectCompetitionButton
                          , command=lambda: GetCompToBeCollated(listOfCompNameBox))
    SelectButton.grid(row=6, column=3, columnspan=3, pady="10")

    # select Comp
    # get Comp details
    # Get Results
    # save name of doc to Competitions table


def SendResultstoBands():
    def GetResultsToBeSent(listOfCompNameBox):

        # gets the position of the Competition Selected
        Selection = listOfCompNameBox.curselection()

        print(Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select the competition you would like to send the results of to all bands.")
        else:

            for i in Selection:
                # gets the name of the competition
                j = listOfCompNameBox.get(i)
                CompetitionName = j

            with sqlite3.connect("RspbaniDB.db") as db:
                # Creates a cursor to search through the data
                cursor = db.cursor()
                # SQL Query to select all Competitions from the competitions table
            resultsQuery = 'SELECT CompetitionID, CompetitionDate, ResultFileName FROM Competitions WHERE CompetitionName = ?'
            # Execute the query to get a list of all the competitions
            cursor.execute(resultsQuery, [CompetitionName])
            # The query will produce a 2 dimensional tuple
            CollatedResultComp = cursor.fetchall()

            if CollatedResultComp:

                for t in CollatedResultComp:
                    CompID = t[0]
                    CompDate = [1]
                    AllResultFileName = t[2]

            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            # Set all competitions to false for that the competition selected will be the only comp that latest is true
            cursor.execute('UPDATE Competitions SET Latest = ?', ("False",))
            conn.commit()

            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()

            cursor.execute(
                'UPDATE Competitions SET Latest = ? WHERE CompetitionID = ?',
                ("True", CompID))
            conn.commit()

            # get a List of All Bands Email addresses
            ListOfEmail = ['ross.martin1201@gmail.com']
            with sqlite3.connect("RspbaniDB.db") as db:
                # Creates a cursor to search through the data
                cursor = db.cursor()
                # SQL Query to select all Competitions from the competitions table
            resultsQuery = 'SELECT EmailAddress FROM BandAccount'
            # Execute the query to get a list of all the competitions
            cursor.execute(resultsQuery)
            # The query will produce a 2 dimensional tuple
            CollatedResultComp = cursor.fetchall()

            print(CollatedResultComp)

            if CollatedResultComp:

                for e in CollatedResultComp:
                    ListOfEmail.append(e[0])

            print("List of email")
            print(ListOfEmail)

            for w in ListOfEmail:
                # email address of the sender
                email_user = 'rspbani.results@gmail.com'
                # password of the sender
                email_password = 'P@55w0rd123'
                # email address of recipient
                email_send = w

                # Creating the Emails subject
                subject = str(CompetitionName) + "'s Results"

                msg = MIMEMultipart()

                msg['From'] = email_user
                msg['To'] = email_send
                msg['Subject'] = subject

                # creating the contents of the email
                body = "Dear Pipe Major \n Please find the full results sheet attached from " + str(
                    CompetitionName) + "\n\n\n We  must thank all the competing Bands " \
                                       "and the supporters they bring for their continued " \
                                       "support throughout the season \n\n\n The Royal " \
                                       "Scottish Pipe Band Association"
                msg.attach(MIMEText(body, 'plain'))

                # adding the PDF to the email as an attachment
                filename = AllResultFileName
                attachment = open(filename, 'rb')

                part = MIMEBase('application', 'octet-stream')
                part.set_payload((attachment).read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', "attachment; filename= " + filename)

                msg.attach(part)
                text = msg.as_string()

                try:
                    # connet to email server and send email
                    server = smtplib.SMTP('smtp.gmail.com', 587)
                    server.starttls()
                    server.login(email_user, email_password)

                    server.sendmail(email_user, email_send, text)
                    # clear the fields in the book transport frame

                # error handling if email cant be sent - display a message stating this
                except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                    messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
                finally:
                    # stop connection with email server
                    server.quit()
            messagebox.showinfo("Alert ", message="Full Competition Results Sent!")

            ListOfEntryID = []
            with sqlite3.connect("RspbaniDB.db") as db:
                # Creates a cursor to search through the data
                cursor = db.cursor()
                # SQL Query to select all Competitions from the competitions table
            resultsQuery = 'SELECT EntryID, BandID FROM BandsEntered WHERE CompetitionID = ?'
            # Execute the query to get a list of all the competitions
            cursor.execute(resultsQuery, [CompID])
            # The query will produce a 2 dimensional tuple
            BandResultComp = cursor.fetchall()

            for Band in BandResultComp:
                EntryID = Band[0]
                BandID = Band[1]

                with sqlite3.connect("RspbaniDB.db") as db:
                    # Creates a cursor to search through the data
                    cursor = db.cursor()
                    # SQL Query to select all Competitions from the competitions table
                resultsQuery = 'SELECT EmailAddress, BandName FROM BandAccount WHERE BandID = ?'
                # Execute the query to get a list of all the competitions
                cursor.execute(resultsQuery, [BandID])
                # The query will produce a 2 dimensional tuple
                BandEmail = cursor.fetchall()

                for Band in BandEmail:
                    EmailAddress = Band[0]
                    BandName = Band[1]

                    with sqlite3.connect("RspbaniDB.db") as db:
                        # Creates a cursor to search through the data
                        cursor = db.cursor()
                        # SQL Query to select all Competitions from the competitions table
                    resultsQuery = 'SELECT BandSheetsFileName FROM Results WHERE EntryID = ?'
                    # Execute the query to get a list of all the competitions
                    cursor.execute(resultsQuery, [EntryID])
                    # The query will produce a 2 dimensional tuple
                    FileName = cursor.fetchall()

                    for File in FileName:
                        NameFile = File[0]

                        # email address of the sender
                        email_user = 'rspbani.results@gmail.com'
                        # password of the sender
                        email_password = 'P@55w0rd123'
                        # email address of recipient
                        email_send = EmailAddress

                        # Creating the Emails subject
                        subject = str(BandName) + "'s Results from " + str(CompetitionName)

                        msg = MIMEMultipart()

                        msg['From'] = email_user
                        msg['To'] = email_send
                        msg['Subject'] = subject

                        # creating the contents of the email
                        body = "Dear Pipe Major, \n Please find the comment sheets attached from the " + str(
                            CompetitionName) + "\n\n\n Congratulations on your recent result. We hope to see you are the next competition. \n\n\n The Royal " \
                                               "Scottish Pipe Band Association"
                        msg.attach(MIMEText(body, 'plain'))

                        # adding the PDF to the email as an attachment
                        filename = NameFile
                        attachment = open(filename, 'rb')

                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload((attachment).read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', "attachment; filename= " + filename)

                        msg.attach(part)
                        text = msg.as_string()

                        try:
                            # connet to email server and send email
                            server = smtplib.SMTP('smtp.gmail.com', 587)
                            server.starttls()
                            server.login(email_user, email_password)

                            server.sendmail(email_user, email_send, text)
                            # clear the fields in the book transport frame

                        # error handling if email cant be sent - display a message stating this
                        except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                            messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
                        finally:
                            # stop connection with email server
                            server.quit()
            messagebox.showinfo("Alert ", message="Bands' Comment Sheets Sent!")

            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            # Set Drawn to true so that it can not be drawn again by accident
            cursor.execute(
                'UPDATE Competitions SET SentToBands = ? WHERE CompetitionID = ?',
                ("True", CompID))
            conn.commit()

    SendResults = Frame(contentFrame, width=965, height=600)
    SendResults.grid(row=0, column=0, sticky="nsew")

    lblspacer = Label(SendResults, width=25)
    lblspacer.grid()

    lblTitle = Label(SendResults, text="Send Competition Results to Bands", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=1, columnspan=5)

    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through the data
        cursor = db.cursor()
        # SQL Query to select all Competitions from the competitions table
    resultsQuery = 'SELECT CompetitionName, CompetitionDate FROM Competitions WHERE Collated = "True" AND SentToBands = "False"'
    # Execute the query to get a list of all the competitions
    cursor.execute(resultsQuery)
    # The query will produce a 2 dimensional tuple
    CollatedResultComp = cursor.fetchall()

    lblText = Label(SendResults, text="Please select the competition you would like publish to the bands", height=3,
                    font=("Arvo", 14))
    lblText.grid(row=1, column=1, columnspan=8, sticky="ew", )

    lblCompetiton = Label(SendResults, text="Competition Name")
    lblCompetiton.grid(row=3, column=2, columnspan=2, sticky=EW)

    lblDate = Label(SendResults, text="Competition Date")
    lblDate.grid(row=3, column=4, columnspan=2, sticky=W)

    listOfCompNameBox = Listbox(SendResults, selectmode=SINGLE, width=25)
    listOfCompNameBox.grid(row=4, column=2, columnspan=2, sticky=E)

    for u in CollatedResultComp:
        n = 0
        listOfCompNameBox.insert(n, str(u[0]))
        n = + 1
    # listbox to show competiion date but it is disabled so it cant be selected
    listOfCompDateBox = Listbox(SendResults, selectmode=DISABLED, width=15, selectbackground="white")
    listOfCompDateBox.grid(row=4, column=4, columnspan=2, sticky=W)
    for u in CollatedResultComp:
        n = 0
        listOfCompDateBox.insert(n, str(u[1]))
        n = + 1

    SelectButton = Button(SendResults, image=photoSelectCompetitionButton
                          , command=lambda: GetResultsToBeSent(listOfCompNameBox))
    SelectButton.grid(row=6, column=3, columnspan=2, pady="10")


def DrawCompetitionFrame(photoSelectCompetitionButton, Years):

    global Grade1Drawn, Grade2Drawn, Grade3ADrawn, Grade3BDrawn, Grade4ADrawn, Grade4BDrawn

    Grade1Drawn = False
    Grade2Drawn = False
    Grade3ADrawn = False
    Grade3BDrawn = False
    Grade4ADrawn = False
    Grade4BDrawn = False

    def CountNumBandsInGrade(Grade):
        NumBands = len(Grade)

        return NumBands

    def selectedDrawComp(photoNextButton):
        def AssignTimeToBands(GradeNum, ListBands, NumBandsInGrade, CompID):
            def SaveDraw(CircleNum, CompTimeHrs, CompTimeMin, ListBands, CompID, GradeNum):
                def SendDrawToBands(GradeNum, CompID):
                    def SendMail(EmailAddresses, CompName, Grade, fileName):

                        print("EmailList" + str(EmailAddresses))

                        for email in EmailAddresses:
                            print(email)
                            # email address of the sender
                            email_user = 'rspbani.draw@gmail.com'
                            # password of the sender
                            email_password = 'P@55w0rd123'
                            # email address of recipient
                            email_send = email

                            # Creating the Emails subject
                            subject = "Draw for " + str(CompName) + " - Grade " + str(Grade)

                            msg = MIMEMultipart()

                            msg['From'] = email_user
                            msg['To'] = email_send
                            msg['Subject'] = subject

                            # creating the contents of the email
                            body = "Please see attached Draw for " + str(CompName) + " - Grade " + str(Grade)
                            msg.attach(MIMEText(body, 'plain'))

                            # adding the PDF to the email as an attachment
                            filename = fileName
                            attachment = open(filename, 'rb')

                            part = MIMEBase('application', 'octet-stream')
                            part.set_payload((attachment).read())
                            encoders.encode_base64(part)
                            part.add_header('Content-Disposition', "attachment; filename= " + filename)

                            msg.attach(part)
                            text = msg.as_string()

                            try:
                                # connet to email server and send email
                                server = smtplib.SMTP('smtp.gmail.com', 587)
                                server.starttls()
                                server.login(email_user, email_password)

                                server.sendmail(email_user, email_send, text)


                            # error handling if email cant be sent - display a message stating this
                            except(smtplib.SMTPException, ConnectionRefusedError, OSError):
                                messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
                            finally:
                                # stop connection with email server
                                server.quit()

                        if len(EmailAddresses) >= 1:
                            messagebox.showinfo("Alert ",
                                                message="Draw Sent for Grade " + Grade + "!")

                    Grade = GradeNum.split(" ")
                    Grade = Grade[1]

                    print("Grade")
                    print(Grade)

                    with sqlite3.connect("RspbaniDB.db") as db:
                        # Creates a cursor to search through the data
                        cursor = db.cursor()
                        # SQL Query to select all Competitions from the competitions table
                    CompDetailsQuery = 'SELECT CompetitionName, CompetitionLocation, CompetitionDate FROM Competitions WHERE CompetitionID = ?'
                    # Execute the query to get a list of all the competitions
                    cursor.execute(CompDetailsQuery, [CompID])
                    # The query will produce a 2 dimensional tuple
                    CompDetailsQueryResults = cursor.fetchall()

                    if CompDetailsQueryResults:
                        for i in CompDetailsQueryResults:
                            CompName = i[0]
                            CompLocation = i[1]
                            CompDate = i[2]

                    document = Document()
                    document.add_picture('RSPBANI_Logo.png', width=Inches(2))
                    document.add_heading('              Official Draw for Grade ' + str(Grade), 0)

                    lblCompetition = document.add_paragraph()
                    lblCompetition.add_run('Competition:   ').bold = True
                    lblCompetition.add_run(CompName)
                    lblLocation = document.add_paragraph()
                    lblLocation.add_run('Location:    ').bold = True
                    lblLocation.add_run(CompLocation)
                    lblDate = document.add_paragraph()
                    lblDate.add_run('Date:    ').bold = True
                    lblDate.add_run(CompDate)

                    document.add_heading(GradeNum, level=1)

                    Draw = document.add_table(rows=1, cols=5)
                    hdr_cells = Draw.rows[0].cells
                    hdr_cells[0].text = 'Band'
                    hdr_cells[1].text = 'Circle Number'
                    hdr_cells[2].text = 'Competing Time'

                    with sqlite3.connect("RspbaniDB.db") as db:

                        # Creates a cursor to search through the data
                        cursor = db.cursor()
                        # SQL Query to select all Competitions from the competitions table
                    BandsQuery = 'SELECT EntryID, BandID  FROM BandsEntered WHERE CompetitionID = ? AND Grade = ?'
                    # Execute the query to get a list of all the competitions
                    cursor.execute(BandsQuery, [CompID, Grade])
                    # The query will produce a 2 dimensional tuple
                    BandsEnteredQuery = cursor.fetchall()

                    EmailAddresses = []
                    if BandsEnteredQuery:
                        for i in BandsEnteredQuery:
                            EntryID = i[0]
                            BandID = i[1]

                            with sqlite3.connect("RspbaniDB.db") as db:
                                # Creates a cursor to search through the data
                                cursor = db.cursor()
                                # SQL Query to select all Competitions from the competitions table
                            BandEmailQuery = 'SELECT BandName, EmailAddress FROM BandAccount WHERE BandID = ?'
                            # Execute the query to get a list of all the competitions
                            cursor.execute(BandEmailQuery, [BandID])
                            # The query will produce a 2 dimensional tuple
                            BandsEmailResults = cursor.fetchall()

                            if BandsEmailResults:
                                for f in BandsEmailResults:
                                    BandName = f[0]
                                    Email = f[1]

                                    EmailAddresses.append(Email)

                                    with sqlite3.connect("RspbaniDB.db") as db:
                                        # Creates a cursor to search through the data
                                        cursor = db.cursor()
                                        # SQL Query to select all Competitions from the competitions table
                                    CompetingQuery = 'SELECT CircleNumber, CompetitingTime FROM CompetitionDraw WHERE EntryID = ?'
                                    # Execute the query to get a list of all the competitions
                                    cursor.execute(CompetingQuery, [EntryID])
                                    # The query will produce a 2 dimensional tuple
                                    CircleTimeResults = cursor.fetchall()

                                    if CircleTimeResults:
                                        for c in CircleTimeResults:
                                            CircleNum = c[0]
                                            CompetingTime = c[1]

                                            Content_cells = Draw.add_row().cells
                                            Content_cells[0].text = str(BandName)
                                            Content_cells[1].text = str(CircleNum)
                                            Content_cells[2].text = str(CompetingTime)

                                            document.save(str(CompName) + " Draw for Grade " + str(Grade) + '.docx')
                                            fileName = str(CompName) + " Draw for Grade " + str(Grade) + '.docx'

                        print(EmailAddresses)
                        SendMail(EmailAddresses, CompName, Grade, fileName)


                def SaveValidDraw(GradeCompTimeHrs, GradeCompTimeMin):

                    e = 0
                    for every in GradeCircleNum:
                        with sqlite3.connect("RspbaniDB.db") as db:
                            # Creates a cursor to search through the data
                            cursorEntryID = db.cursor()
                        # SQL Query to select the EntryID from the BandsEntered table
                        EntryIDQuery = 'SELECT EntryID FROM BandsEntered WHERE BandID = ? AND CompetitionID = ?'
                        # Execute the query to get a tuple of the EntryID
                        cursorEntryID.execute(EntryIDQuery, [every, CompID])
                        # The query will produce a 2 dimensional tuple
                        EntryIDResults = cursorEntryID.fetchall()
                        for i in EntryIDResults:
                            for k in i:
                                EntryID = k

                        print("CompID")
                        print(CompID)
                        print("ListBands[e]")
                        print(ListBands[e])
                        print("EntryID")
                        print(EntryID)
                        print("GradeCircleNum[every]")
                        print(GradeCircleNum[every])
                        print("GradeCompTime[every]")
                        # print(GradeCompTime[every])

                        # Insert these drawn times and circle numbers into the Competition Draw Table
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute(
                            'INSERT INTO CompetitionDraw (CompetitionID, BandID, EntryID, CircleNumber, CompetitingTime) VALUES(?,?,?,?,?)',
                            (CompID, ListBands[e], EntryID, GradeCircleNum[every],
                             str(GradeCompTimeHrs[every]) + ":" + str(GradeCompTimeMin[every])))
                        conn.commit()
                        # CompID, BandID, CircleNum of the band, Competing Time
                        # Set all the Judges score to "" in the database
                        blank = "  "
                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute(
                            'INSERT INTO Results (EntryID, BandID, CompetitionID, '
                            'Piping1JudgeID, Piping1Score , Piping1Comments, '
                            'Piping2JudgeID, Piping2Score, Piping2Comments, '
                            'TotalPipingScore,'
                            'DrummingJudgeID, DrummingScore, DrummingComments, '
                            'EnsembleJudgeID, EnsembleScore, EnsembleComments) '
                            'VALUES(?, ?, ?, ?,?,?,?,?,?,?,?,?,?,?,?,?)',
                            (EntryID, ListBands[e], CompID, blank, blank, blank, blank, blank, blank, blank, blank,
                             blank, blank, blank, blank, blank))
                        conn.commit()
                        # increment the band
                        e += 1

                        print("GradeNum")
                        print(GradeNum)
                        global Grade1Drawn, Grade2Drawn, Grade3ADrawn, Grade3BDrawn, Grade4ADrawn, Grade4BDrawn

                        if GradeNum == "Grade 4B":
                            Grade4BDrawn = True
                            DrawGradeFrame()
                        if GradeNum == "Grade 4A":
                            Grade4ADrawn = True
                            DrawGradeFrame()
                        if GradeNum == "Grade 3B":
                            Grade3BDrawn = True
                            DrawGradeFrame()
                        if GradeNum == "Grade 3A":
                            Grade3ADrawn = True
                            DrawGradeFrame()
                        if GradeNum == "Grade 2":
                            Grade2Drawn = True
                            DrawGradeFrame()
                        if GradeNum == "Grade 1":
                            Grade1Drawn = True
                            DrawGradeFrame()
                    SendDrawToBands(GradeNum, CompID)

                def getCompTimeMin(GradeCompTimeHrs):
                    GradeCompTimeMin = {}
                    gctMin = 0

                    validCount = 0

                    for entries in CompTimeMin:
                        validMin = False
                        # get every input for the competing time
                        CompTimeDrawn = entries.get()
                        if CompTimeDrawn != "":
                            if CompTimeDrawn.isdigit() == True:
                                if int(CompTimeDrawn) >= 0 and int(CompTimeDrawn) <= 59:
                                    validMin = True
                                    # map the bandID to the competing time drawn
                                    GradeCompTimeMin[ListBands[gctMin]] = CompTimeDrawn
                                    # increment the BandID list
                                    gctMin += 1
                                else:
                                    messagebox.showerror(
                                        message="One or more fields are invalid. Make sure all times are between 09:00 & 18:00 (mins)")
                                    validCount += 1

                            else:
                                messagebox.showerror(
                                    message="One or more fields are invalid. Make sure all times are of the correct time format")
                                validCount += 1
                        else:
                            messagebox.showerror(
                                message="One or more fields are invalid. Please make sure to enter a competing time for all bands")
                            validCount += 1

                    print("Valid" + str(valid))
                    if validCount == 0:
                        print("validCount" + str(validCount))
                        SaveValidDraw(GradeCompTimeHrs, GradeCompTimeMin)

                def getCompTimeHrs():

                    GradeCompTimeHrs = {}
                    gctHrs = 0

                    validCount = 0
                    for entries in CompTimeHrs:
                        valid = False
                        # get every input for the competing time
                        CompTimeDrawn = entries.get()
                        if CompTimeDrawn != "":
                            if CompTimeDrawn.isdigit() == True:
                                if int(CompTimeDrawn) > 8 and int(CompTimeDrawn) < 18:
                                    valid = True
                                    # map the bandID to the competing time drawn
                                    GradeCompTimeHrs[ListBands[gctHrs]] = CompTimeDrawn
                                    # increment the BandID list
                                    gctHrs += 1
                                else:
                                    messagebox.showerror(
                                        message="One or more fields are invalid. Make sure all times are between 09:00 & 18:00 (Hrs)")
                                    validCount += 1

                            else:
                                messagebox.showerror(
                                    message="One or more fields are invalid. Make sure all times are of the correct time format")
                                validCount += 1

                        else:
                            messagebox.showerror(
                                message="One or more fields are invalid. Please make sure to enter a competing time for all bands")
                            validCount += 1

                    print("Valid" + str(valid))
                    if validCount == 0:
                        getCompTimeMin(GradeCompTimeHrs)

                GradeCircleNum = {}
                gcn = 0
                validCount = 0
                for entry in CircleNum:
                    valid = False

                    # get every input for the circle num
                    CricleNumberDrawn = entry.get()
                    if CricleNumberDrawn != "":
                        if CricleNumberDrawn.isdigit() == True:
                            if int(CricleNumberDrawn) < 7 and int(CricleNumberDrawn) > 0:

                                valid = True
                                # map the bandID to the Circle number drawn
                                GradeCircleNum[ListBands[gcn]] = CricleNumberDrawn
                                # increment the BandID list
                                gcn += 1
                            else:
                                messagebox.showerror(
                                    message="One or more fields are invalid. Make sure circle number is between one and six.")
                                validCount += 1
                        else:
                            messagebox.showerror(
                                message="One or more fields are invalid, please enter only numbers for the circle number")
                            validCount += 1
                    else:
                        messagebox.showerror(message="One or more fields are invalid, please enter the circle number")
                        validCount += 1
                        # AssignTimeToBands(GradeNum, ListBands, NumBandsInGrade, CompID)

                print("Valid" + str(valid))
                if validCount == 0:
                    getCompTimeHrs()

            SelDrawCompetition = Frame(contentFrame, width=965, height=600)
            SelDrawCompetition.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(SelDrawCompetition, width=25)
            lblspacer.grid()

            # Creation of widgets on the RegAdmin Frame
            lblTitle = Label(SelDrawCompetition, text="Draw " + GradeNum, height=3, font=("Arvo", 32))
            lblTitle.grid(row=0, column=2, columnspan=3)

            lblText = Label(SelDrawCompetition,
                            text="Please enter the competing time and circle number for each of the following Bands. ",
                            font=("Arvo", 14))
            lblText.grid(row=1, column=1, columnspan=7, sticky="ew")

            lblBandID = Label(SelDrawCompetition, text="Band ID: ", font=("Arvo", 14))
            lblBandID.grid(row=2, column=1)

            lblCircleNum = Label(SelDrawCompetition, text="Circle Number: ", font=("Arvo", 14))
            lblCircleNum.grid(row=2, column=2)

            lblTime = Label(SelDrawCompetition, text="Competing Time: ", font=("Arvo", 14))
            lblTime.grid(row=2, column=3, columnspan=3)

            CircleNum = []
            CompTimeHrs = []
            CompTimeMin = []
            r = 5
            print(NumBandsInGrade)

            # iterate through the number of bands entered and for each ask the user to input the circle
            # number and competing time for the band

            for i in range(NumBandsInGrade):
                r += 1

                lblBandID = Label(SelDrawCompetition, text=ListBands[i])
                lblBandID.grid(row=i + 4, column=1)

                EtyCircleNum = Entry(SelDrawCompetition, width=5)
                EtyCircleNum.grid(row=i + 4, column=2)
                # append the value of each entry box to a list
                CircleNum.append(EtyCircleNum)

                EtyCompTimeHrs = Entry(SelDrawCompetition, width=5)
                EtyCompTimeHrs.grid(row=i + 4, column=3, sticky="e")
                # append the value of each entry box to a list
                CompTimeHrs.append(EtyCompTimeHrs)

                lblColon = Label(SelDrawCompetition, text=" : ")
                lblColon.grid(row=i + 4, column=4)

                EtyCompTimeMin = Entry(SelDrawCompetition, width=5)
                EtyCompTimeMin.grid(row=i + 4, column=5, sticky="w")
                # append the value of each entry box to a list
                CompTimeMin.append(EtyCompTimeMin)


            ButtonSelect = Button(SelDrawCompetition, image=photoNextButton,
                                  command=lambda: SaveDraw(CircleNum, CompTimeHrs, CompTimeMin, ListBands, CompID,
                                                           GradeNum))
            ButtonSelect.grid(row=r, column=2, columnspan=2, pady="18")

        def DrawSelectedGrade(listOfGradeDrawBox, Grade1, Grade2, Grade3A, Grade3B, Grade4A, Grade4B, NumGrade1,
                              NumGrade2, NumGrade3A, NumGrade3B, NumGrade4A, NumGrade4B, CompID):
            # get position of selected grade
            Selection = listOfGradeDrawBox.curselection()

            if Selection ==():
                messagebox.showinfo(message="Please select a grade to draw.")
            else:

                for i in Selection:
                    # get value of selection
                    DrawGrade = listOfGradeDrawBox.get(i)
                    print(str(DrawGrade))

                if DrawGrade == "1" and NumGrade1 != 0:
                    # text, List of Band sentered in the grade, Number of bands, competitionID
                    AssignTimeToBands("Grade 1", Grade1, NumGrade1, CompID)
                elif DrawGrade == "2" and NumGrade2 != 0:
                    AssignTimeToBands("Grade 2", Grade2, NumGrade2, CompID)
                elif DrawGrade == "3A" and NumGrade3A != 0:
                    AssignTimeToBands("Grade 3A", Grade3A, NumGrade3A, CompID)
                elif DrawGrade == "3B" and NumGrade3B != 0:
                    AssignTimeToBands("Grade 3B", Grade3B, NumGrade3B, CompID)
                elif DrawGrade == "4A" and NumGrade4A != 0:
                    AssignTimeToBands("Grade 4A", Grade4A, NumGrade4A, CompID)
                elif DrawGrade == "4B" and NumGrade4B != 0:
                    AssignTimeToBands("Grade 4B", Grade4B, NumGrade4B, CompID)

                else:
                    messagebox.showwarning(message="There are no bands entered in this grade")
                    global Grade1Drawn, Grade2Drawn, Grade3ADrawn, Grade3BDrawn, Grade4ADrawn, Grade4BDrawn

                    if DrawGrade == "1":
                        Grade1Drawn = True
                    elif DrawGrade == "2":
                        Grade2Drawn = True
                    elif DrawGrade == "3A":
                        Grade3ADrawn = True
                    elif DrawGrade == "3B":
                        Grade3BDrawn = True
                    elif DrawGrade == "4A":
                        Grade4ADrawn = True
                    elif DrawGrade == "4B":
                        Grade4BDrawn = True

                    DrawGradeFrame()
        # gets the position of the Competition Selected
        Selection = listOfCompNameBox.curselection()

        print(Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select a competition to draw!")
        else:

            for i in Selection:
                # gets the name of the competition
                j = listOfCompNameBox.get(i)

                # gets the competitionID of the competition
                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCompID = db.cursor()
                findCompID = 'SELECT CompetitionID FROM Competitions WHERE CompetitionName = ?'
                cursorCompID.execute(findCompID, [j])
                CompIDresults = cursorCompID.fetchall()

                if CompIDresults:

                    for a in CompIDresults:
                        a = str(a)
                        CompID = a[1:-2]

            # get a list off all the bands that have applied for the competition and count how many there are.
            with sqlite3.connect("RspbaniDB.db") as db:
                # Creates a cursor to search through the data
                cursorListOfBandsEntered = db.cursor()
                # SQL Query to select all bandIDs from the entered bands for the specific competition
            listOfBandsEnteredQuery = 'SELECT BandID FROM BandsEntered WHERE CompetitionID = ?'
            # Execute the query to get a list of all the bands
            cursorListOfBandsEntered.execute(listOfBandsEnteredQuery, [CompID])
            # The query will produce a 2 dimensional tuple
            ListOfBandsEntered = cursorListOfBandsEntered.fetchall()

            print("ListOfBandsEntered")
            print(ListOfBandsEntered)
            # list of each grade to store the ID of the Band entered
            Grade1 = []
            Grade2 = []
            Grade3A = []
            Grade3B = []
            Grade4A = []
            Grade4B = []

            for o in ListOfBandsEntered:
                for i in o:
                    # i is bandID

                    with sqlite3.connect("RspbaniDB.db") as db:
                        # Creates a cursor to search through the data
                        cursorGradeOfBandEntered = db.cursor()
                        # SQL Query to select all Competitions from the competitions table
                    gradeOfBandsEnteredQuery = 'SELECT BandID, BandGrade FROM BandAccount WHERE BandID = ?'
                    # Execute the query to get a list of all the bands
                    cursorGradeOfBandEntered.execute(gradeOfBandsEnteredQuery, [i])
                    # The query will produce a 2 dimensional tuple
                    GradeOfBandsEntered = cursorGradeOfBandEntered.fetchall()

                    print("GradeOfBandsEntered")
                    print(GradeOfBandsEntered)

                    for y in GradeOfBandsEntered:
                        # y[0] is BandID
                        # y[1] is BandGrade

                        # Append the bandID to the GradeList if the Grade of the Band matches the Grade of the list
                        if y[1] == "1":
                            Grade1.append(y[0])
                        elif y[1] == "2":
                            Grade2.append(y[0])
                        elif y[1] == "3A":
                            Grade3A.append(y[0])
                        elif y[1] == "3B":
                            Grade3B.append(y[0])
                        elif y[1] == "4A":
                            Grade4A.append(y[0])
                        elif y[1] == "4B":
                            Grade4B.append(y[0])

            # Find out how many Bands have entered each grade
            NumGrade1 = CountNumBandsInGrade(Grade1)
            NumGrade2 = CountNumBandsInGrade(Grade2)
            NumGrade3A = CountNumBandsInGrade(Grade3A)
            NumGrade3B = CountNumBandsInGrade(Grade3B)
            NumGrade4A = CountNumBandsInGrade(Grade4A)
            NumGrade4B = CountNumBandsInGrade(Grade4B)

            # Total amount of Bands Entered
            TotalBandsEntered = NumGrade1 + NumGrade2 + NumGrade3A + NumGrade3B + NumGrade4A + NumGrade4B

            print("NumGrade2")
            print(NumGrade4B)

            print("TotalBandsEntered")
            print(TotalBandsEntered)

            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            # Insert the Num of Bands in each grade & total number of bands Entered into the Competition Table
            cursor.execute(
                'UPDATE Competitions SET NumG1Ent =?, NumG2Ent = ?, NumG3AEnt = ?, NumG3BEnt = ?, NumG4AEnt = ?, NumG4BEnt = ?, TotalNumBands = ? WHERE CompetitionID = ?',
                (NumGrade1, NumGrade2, NumGrade3A, NumGrade3B, NumGrade4A, NumGrade4B, TotalBandsEntered, CompID))
            conn.commit()


            #created
            conn = sqlite3.connect('RspbaniDB.db')
            with conn:
                cursor = conn.cursor()
            # create the CompetitionDraw table if it doesnt already exist to store DrawDetails details once made.
            cursor.execute(
                'CREATE TABLE IF NOT EXISTS CompetitionDraw (CompetitionID INTEGER, BandID INTEGER, EntryID INTEGER, CircleNumber INTEGER,'
                'CompetitingTime TEXT)')
            conn.commit()
            List = {}

            def DrawGradeFrame():

                if Grade1Drawn == True and Grade2Drawn == True and Grade3ADrawn == True and Grade3BDrawn == True and Grade4ADrawn == True and Grade4BDrawn == True:
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    # Set Drawn to true so that it can not be drawn again by accident
                    cursor.execute(
                        'UPDATE Competitions SET Drawn = ? WHERE CompetitionID = ?',
                        ("True", CompID))
                    conn.commit()
                    messagebox.showinfo(message="All grades have been drawn!")

                    DrawCompetitionFrame(photoSelectCompetitionButton, Years)

                else:

                    # Creation of widgets on the DrawGrade Frame
                    DrawGrade = Frame(contentFrame, width=965, height=600)
                    DrawGrade.grid(row=0, column=0, sticky="nsew")

                    lblspacer = Label(DrawGrade, width=40)
                    lblspacer.grid()

                    lblTitle = Label(DrawGrade, text="Select Grade", height=3, font=("Arvo", 32))
                    lblTitle.grid(row=0, column=2, columnspan=6)

                    lblText = Label(DrawGrade,
                                    text="Please select the Grade you would like to Draw. \n NOTE Please select Grade4B last!!",
                                    height=3,
                                    font=("Arvo", 14))
                    lblText.grid(row=1, column=2, columnspan=6, sticky="ew")

                    listOfGradeDrawBox = Listbox(DrawGrade, selectmode=SINGLE, width=10, height=7)
                    listOfGradeDrawBox.grid(row=3, column=4)

                    print("G1")
                    print(Grade1Drawn)
                    print("G2")
                    print(Grade2Drawn)
                    print("G3a")
                    print(Grade3ADrawn)
                    print("G3b")
                    print(Grade3BDrawn)
                    print("G4a")
                    print(Grade4ADrawn)
                    print("G4b")
                    print(Grade4BDrawn)

                    i = 0

                    # populate all grades into the list box
                    if Grade1Drawn == False:
                        listOfGradeDrawBox.insert(i, "1")
                        i+=1
                    if Grade2Drawn == False:
                        listOfGradeDrawBox.insert(i, "2")
                        i+=1
                    if Grade3ADrawn == False:
                        listOfGradeDrawBox.insert(i, "3A")
                        i+=1
                    if Grade3BDrawn == False:
                        listOfGradeDrawBox.insert(i, "3B")
                        i+=1
                    if Grade4ADrawn == False:
                        listOfGradeDrawBox.insert(i, "4A")
                        i+=1
                    if Grade4BDrawn == False:
                        listOfGradeDrawBox.insert(i, "4B")
                        i+=1


                    SelectGradeButton = Button(DrawGrade, image=photoSelectGradeButton
                                               , command=lambda: DrawSelectedGrade(listOfGradeDrawBox, Grade1, Grade2, Grade3A,
                                                                                   Grade3B, Grade4A, Grade4B, NumGrade1,
                                                                                   NumGrade2, NumGrade3A, NumGrade3B,
                                                                                   NumGrade4A, NumGrade4B, CompID))
                    SelectGradeButton.grid(row=4, column=3, columnspan=3, pady="20")

            DrawGradeFrame()

    DrawCompetition = Frame(contentFrame, width=965, height=600)
    DrawCompetition.grid(row=0, column=0, sticky="nsew")
    # Creation of widgets on the DrawCompetition Frame to allow an admin to draw the grade

    lblspacer = Label(DrawCompetition, width=30)
    lblspacer.grid()

    lblTitle = Label(DrawCompetition, text="Draw Competition", height=3, font=("Arvo", 32))
    lblTitle.grid(row=0, column=1, columnspan=8)

    # get all the competitions names and dates that haven't been drawn yet
    with sqlite3.connect("RspbaniDB.db") as db:
        # Creates a cursor to search through the data
        cursorListOfComps = db.cursor()
        # SQL Query to select all Competition names and corresponding from the competitions table
    listOfCompsQuery = 'SELECT CompetitionName, CompetitionDate FROM Competitions WHERE Drawn = ?'
    # Execute the query to get a list of all the bands
    cursorListOfComps.execute(listOfCompsQuery, ["False"])
    # The query will produce a 2 dimensional tuple
    ListOfComps = cursorListOfComps.fetchall()

    lblText = Label(DrawCompetition, text="Please select the competition you would like to Draw.", height=3,
                    font=("Arvo", 14))
    lblText.grid(row=1, column=1, columnspan=8, sticky="ew")

    lblCompetiton = Label(DrawCompetition, text="Competition Name")
    lblCompetiton.grid(row=3, column=1, columnspan=2, sticky=EW)

    lblDeadline = Label(DrawCompetition, text="Entry Deadline")
    lblDeadline.grid(row=3, column=4, columnspan=1, sticky=EW)

    lblDate = Label(DrawCompetition, text="Competition Date")
    lblDate.grid(row=3, column=3, columnspan=1, sticky=W)

    # populate a list box with the list of bands names & dates
    listOfCompNameBox = Listbox(DrawCompetition, selectmode=SINGLE, width=25)
    listOfCompNameBox.grid(row=4, column=1, columnspan=2)

    for u in ListOfComps:
        n = 0
        listOfCompNameBox.insert(n, str(u[0]))
        n = + 1
    # date disabled so it cant be selected
    listOfCompDateBox = Listbox(DrawCompetition, selectmode=DISABLED, width=15)
    listOfCompDateBox.grid(row=4, column=3)

    for u in ListOfComps:
        n = 0
        listOfCompDateBox.insert(n, str(u[1]))
        n = + 1

    listOfCompDeadlineBox = Listbox(DrawCompetition, selectmode=DISABLED, width=15)
    listOfCompDeadlineBox.grid(row=4, column=4)

    for u in ListOfComps:
        n = 0
        listOfCompDeadlineBox.insert(n, str(u[1]))
        n = + 1

    SelectButton = Button(DrawCompetition, image=photoSelectCompetitionButton
                          , command=lambda: selectedDrawComp(photoNextButton))
    SelectButton.grid(row=6, column=1, columnspan=4, pady="10")


def JudgeCompetitionFrame(photoSelectCompetitionButton):
    def AdminSelectedJudge(JudgeListBandBox):
        # get position of the selected value
        Selection = JudgeListBandBox.curselection()

        print(Selection)

        if Selection == ():
            messagebox.showinfo(message="Admin, please select the judge you wish to partake as.")
        else:

            for i in Selection:
                # get the value of the selected position
                Name = JudgeListBandBox.get(i)

                Name = Name.split(" ")
                FirstName = Name[0]
                SecondName = Name[1]

                print(FirstName)
                print(SecondName)

                # get all the competitions names and dates that haven't been drawn yet
                with sqlite3.connect("RspbaniDB.db") as db:
                    # Creates a cursor to search through the data
                    cursor = db.cursor()
                    # SQL Query to select all Competition names and corresponding from the competitions table
                JudgeIDQuery = 'SELECT JudgeID FROM JudgeAccount WHERE FirstName = ? AND SecondName = ?'
                # Execute the query to get a list of all the bands
                cursor.execute(JudgeIDQuery, [FirstName, SecondName])
                # The query will produce a 2 dimensional tuple
                JudgeIDResults = cursor.fetchall()

                if JudgeIDResults:

                    for J in JudgeIDResults:
                        AdminSelectedJudgeID = J[0]

                    DisplayCompetitionList(AdminSelectedJudgeID)

    def DisplayCompetitionList(JudgeID):

        #Table Created
        blank = " "
        conn = sqlite3.connect('RspbaniDB.db')
        with conn:
            cursor = conn.cursor()
        # create the Results table if it doesnt already exist to store each Bands Results details
        cursor.execute(
            'CREATE TABLE IF NOT EXISTS Results (EntryID INTEGER, BandID INTEGER, CompetitionID INTEGER, '
            'Piping1JudgeID INTEGER, Piping1Score INTEGER, Piping1Comments TEXT, '
            'Piping2JudgeID INTEGER, Piping2Score INTEGER, Piping2Comments TEXT, '
            'TotalPipingScore INTEGER,'
            'DrummingJudgeID INTEGER, DrummingScore INTEGER, DrummingComments TEXT, '
            'EnsembleJudgeID INTEGER, EnsembleScore INTEGER, EnsembleComments TEXT,'
            'TotalScore INTEGER, Position INTEGER, BandSheetsFileName TEXT)')
        conn.commit()

        JudgeComp = Frame(contentFrame, width=965, height=600)
        JudgeComp.grid(row=0, column=0, sticky="nsew")

        lblspacer = Label(JudgeComp, width=35)
        lblspacer.grid()

        # Creation of widgets on the JudgeComp Frame
        lblTitle = Label(JudgeComp, text="Select Competition", height=3, font=("Arvo", 32))
        lblTitle.grid(row=0, column=2, columnspan=6)

        with sqlite3.connect("RspbaniDB.db") as db:
            # Creates a cursor to search through the data
            cursorListOfComps = db.cursor()
            # SQL Query to select all Competitions from the competitions table
        listOfCompsQuery = 'SELECT CompetitionName, CompetitionDate FROM Competitions WHERE Drawn = ?'
        # Execute the query to get a list of all the bands
        cursorListOfComps.execute(listOfCompsQuery, ["True"])
        # The query will produce a 2 dimensional tuple
        ListOfComps = cursorListOfComps.fetchall()

        lblText = Label(JudgeComp, text="Please select the competition you would like to Judge. ", height=3,
                        font=("Arvo", 14))
        lblText.grid(row=1, column=2, columnspan=6, sticky="ew", )

        # create two lists boxes to display the competition name and date respectively
        listOfCompNameBox = Listbox(JudgeComp, selectmode=SINGLE, width=30)
        listOfCompNameBox.grid(row=3, column=2, columnspan=3)

        listOfCompDateBox = Listbox(JudgeComp, selectmode=DISABLED, width=15, selectbackground="white")
        listOfCompDateBox.grid(row=3, column=5, columnspan=2)

        for u in ListOfComps:
            n = 0
            listOfCompNameBox.insert(n, str(u[0]))
            listOfCompDateBox.insert(n, str(u[1]))
            n =+ 1

        SelectButton = Button(JudgeComp, image=photoSelectCompetitionButton
                              , command=lambda: selectedComp(photoUpdateCompetitionButton, listOfCompNameBox, JudgeID))
        SelectButton.grid(row=4, column=3, columnspan=3, pady="10")

        if Admin == True:
            ButtonBack = Button(JudgeComp, image=photoBack,
                                command=lambda: JudgeCompetitionFrame(photoSelectCompetitionButton))
            ButtonBack.grid(row=0, column=0)

    def selectedGrade(photoUpdateCompetitionButton, listOfGradeBox, BandsNameInGrade1, BandsNameInGrade2,
                      BandsNameInGrade3A, BandsNameInGrade3B, BandsNameInGrade4A, BandsNameInGrade4B, CompID, JudgeID,
                      listOfCompNameBox):
        def SaveSheetDatabase(judgeType, comment, CompID, ListBands, BandIDofSelected, JudgeID):

            # get the comments made by the judge
            CommentSheetData = comment.get("1.0", END)
            print(CommentSheetData)
            print(CompID)

            print("JudgeID")
            print(JudgeID)

            # Save the comments about the band to the piping 1 section that the judge was marking
            if judgeType == "Piping 1":

                conn = sqlite3.connect('RspbaniDB.db')
                with conn:
                    cursor = conn.cursor()

                # can i not just use EntryID as a composite key instead?
                cursor.execute(
                    'UPDATE Results SET Piping1JudgeID = ?, Piping1Comments = ? WHERE CompetitionID = ? AND BandID = ?',
                    (JudgeID, CommentSheetData, CompID, BandIDofSelected))
                conn.commit()
                messagebox.showinfo(message="Comment Sheet Updated")

            # Save the comments about the band to the piping 2 section that the judge was marking
            elif judgeType == "Piping 2":

                conn = sqlite3.connect('RspbaniDB.db')
                with conn:
                    cursor = conn.cursor()
                cursor.execute(
                    'UPDATE Results SET Piping2JudgeID = ?, Piping2Comments = ? WHERE CompetitionID = ? AND BandID = ? ',
                    (JudgeID, CommentSheetData, CompID, BandIDofSelected))
                conn.commit()
                messagebox.showinfo(message="Comment Sheet Updated")

            # Save the comments about the band to the Drumming section that the judge was marking
            elif judgeType == "Drumming":

                conn = sqlite3.connect('RspbaniDB.db')
                with conn:
                    cursor = conn.cursor()
                cursor.execute(
                    'UPDATE Results SET DrummingJudgeID = ?, DrummingComments = ? WHERE CompetitionID = ? AND BandID = ?',
                    (JudgeID, CommentSheetData, CompID, BandIDofSelected))
                conn.commit()
                messagebox.showinfo(message="Comment Sheet Updated")

            # Save the comments about the band to the Ensemble section that the judge was marking
            elif judgeType == "Ensemble":

                conn = sqlite3.connect('RspbaniDB.db')
                with conn:
                    cursor = conn.cursor()
                cursor.execute(
                    'UPDATE Results SET EnsembleJudgeID = ?, EnsembleComments = ? WHERE CompetitionID = ? AND BandID = ?',
                    (JudgeID, CommentSheetData, CompID, BandIDofSelected))
                conn.commit()
                messagebox.showinfo(message="Comment Sheet Updated")

            CreateListBoxOfBands(ListBands, grade, JudgeID, listOfCompNameBox)

        def JudgeSheet(photoUpdateCompetitionButton, listOfJudgeBox, save, ListBands, BandIDofSelected, JudgeID):
            Selection = listOfJudgeBox.curselection()

            if Selection == ():
                messagebox.showinfo(message="Please select a section to judge!")
            else:

                for i in Selection:
                    judgeType = listOfJudgeBox.get(i)
                print("Judge Type")
                print(judgeType)

                JudgeSheet = Frame(contentFrame, width=965, height=600)
                JudgeSheet.grid(row=0, column=0, sticky="nsew")

                lblspacer = Label(JudgeSheet, width=30)
                lblspacer.grid()

                # Creation of widgets on the JudgeSheet Frame
                lblTitle = Label(JudgeSheet, text="Comments Sheet", height=3, font=("Arvo", 32))
                lblTitle.grid(row=0, column=2, columnspan=6)

                lblText = Label(JudgeSheet, text="Please write your comment about the bands performance today.", height=3,
                                font=("Arvo", 14))
                lblText.grid(row=1, column=2, columnspan=6)

                # Depending on the section; get the comments that have already been made - if any
                if judgeType == "Piping 1":
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursorBandCommment = db.cursor()
                    findBandComment = 'SELECT Piping1Comments FROM Results WHERE CompetitionID = ? AND BandID = ?'
                    cursorBandCommment.execute(findBandComment, [CompID, BandIDofSelected])
                    Commentresults = cursorBandCommment.fetchall()

                    for i in Commentresults:
                        for j in i:
                            text = j

                elif judgeType == "Piping 2":

                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursorBandCommment = db.cursor()
                    findBandComment = 'SELECT Piping2Comments FROM Results WHERE CompetitionID = ? AND BandID = ?'
                    cursorBandCommment.execute(findBandComment, [CompID, BandIDofSelected])
                    Commentresults = cursorBandCommment.fetchall()

                    for i in Commentresults:
                        for j in i:
                            text = j

                elif judgeType == "Drumming":
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursorBandCommment = db.cursor()
                    findBandComment = 'SELECT DrummingComments FROM Results WHERE CompetitionID = ? AND BandID = ?'
                    cursorBandCommment.execute(findBandComment, [CompID, BandIDofSelected])
                    Commentresults = cursorBandCommment.fetchall()

                    for i in Commentresults:
                        for j in i:
                            text = j

                elif judgeType == "Ensemble":
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursorBandCommment = db.cursor()
                    findBandComment = 'SELECT EnsembleComments FROM Results WHERE CompetitionID = ? AND BandID = ?'
                    cursorBandCommment.execute(findBandComment, [CompID, BandIDofSelected])
                    Commentresults = cursorBandCommment.fetchall()

                    for i in Commentresults:
                        for j in i:
                            text = j

                # create a textbox where judges can write their comments
                lblCommentsSheet = Text(JudgeSheet, height=20, width=20, bg="lightgrey")
                # populate the textbox with comments if they already exist
                lblCommentsSheet.insert("1.0", text)
                # allow the textbox to scroll if they write over the size of the textbox placed
                bar = Scrollbar(JudgeSheet)
                lblCommentsSheet.config(yscrollcommand=bar.set)
                bar.config(command=lblCommentsSheet.yview)
                lblCommentsSheet.grid(row=3, column=2, columnspan=7, sticky="ew")
                bar.grid(row=3, column=9, sticky="nsw")

                SelectButton = Button(JudgeSheet, image=save
                                      , command=lambda: SaveSheetDatabase(judgeType, lblCommentsSheet, CompID, ListBands,
                                                                          BandIDofSelected, JudgeID))
                SelectButton.grid(row=4, column=3, columnspan=4, pady="10")

        def selectJudge(photoUpdateCompetitionButton, ListBands, listOfBandNameBox, JudgeID):

            SelectionBand = listOfBandNameBox.curselection()

            if SelectionBand == ():
                messagebox.showinfo(message="Please select a band to judge!")
            else:

                for i in SelectionBand:
                    SelectedBand = listOfBandNameBox.get(i)
                    print("SelectedBand")
                    print(SelectedBand)

                    with sqlite3.connect("RspbaniDB.db") as db:
                        # Creates a cursor to search through the data
                        cursorBandID = db.cursor()
                        # SQL Query to select BandID from band name
                    BandIDQuery = 'SELECT BandID FROM BandAccount WHERE BandName = ?'
                    # Execute the query to get a list of all the bandID's
                    cursorBandID.execute(BandIDQuery, [SelectedBand])
                    # The query will produce a 2 dimensional tuple
                    BandIDResults = cursorBandID.fetchall()

                    for h in BandIDResults:
                        BandIDofSelected = h[0]


                JudgeCompJudge = Frame(contentFrame, width=965, height=600)
                JudgeCompJudge.grid(row=0, column=0, sticky="nsew")

                lblspacer = Label(JudgeCompJudge, width=40)
                lblspacer.grid()

                # Creation of widgets on the JudgeCompJudge Frame
                lblTitle = Label(JudgeCompJudge, text="Select Section", height=3, font=("Arvo", 32))
                lblTitle.grid(row=0, column=2, columnspan=6)

                lblText = Label(JudgeCompJudge, text="Please select the Section you wish to Judge. ", height=3,
                                font=("Arvo", 14))
                lblText.grid(row=1, column=2, columnspan=6, sticky="ew", )

                # populate a list of the sections to be judged
                listOfJudgeBox = Listbox(JudgeCompJudge, selectmode=SINGLE, width=25, height=5)
                listOfJudgeBox.grid(row=3, column=4, columnspan=2)
                listOfJudgeBox.insert(0, "Piping 1")
                listOfJudgeBox.insert(1, "Piping 2")
                listOfJudgeBox.insert(2, "Drumming")
                listOfJudgeBox.insert(3, "Ensemble")

                SelectButton = Button(JudgeCompJudge, image=photoSelectCompetitionButton
                                      , command=lambda: JudgeSheet(photoUpdateCompetitionButton, listOfJudgeBox,
                                                                   photoSaveButton, ListBands, BandIDofSelected, JudgeID))
                SelectButton.grid(row=4, column=3, columnspan=4, pady="20")

        def ScoreSheetCheck(Score, ListBands, Section, grade, ValidPossScoreList, listOfJudgeScoreBox):
            def SaveSheet(listOfJudgeScoreBox):
                for each in SectionScore:

                    # each is bandName
                    with sqlite3.connect("RspbaniDB.db") as db:
                        # Creates a cursor to search through the data
                        cursorBandName = db.cursor()
                        # SQL Query to select BandID with Name
                    BandIDQuery = 'SELECT BandID FROM BandAccount WHERE BandName = ?'
                    # Execute the query to get BandID
                    cursorBandName.execute(BandIDQuery, [each])
                    # The query will produce a 2 dimensional tuple
                    BandIDResults = cursorBandName.fetchall()
                    for i in BandIDResults:
                        for k in i:
                            BandID = k

                    print(BandID)

                    # for each possible section; if selected, Update the sore for each section with
                    # the score given by the judge
                    if Section == "Piping 1":

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute(
                            'UPDATE Results SET Piping1Score = ? WHERE CompetitionID = ? AND BandID = ?',
                            (SectionScore[each], CompID, BandID))
                        conn.commit()

                    elif Section == "Piping 2":

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute(
                            'UPDATE Results SET Piping2Score = ? WHERE CompetitionID = ? AND BandID = ?',
                            (SectionScore[each], CompID, BandID))
                        conn.commit()

                    elif Section == "Drumming":

                        BestBassSection = BandBestBass.get()

                        if grade == "1":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G1BestBass = ? WHERE CompetitionID = ?',
                                (BestBassSection, CompID))
                            conn.commit()

                        elif grade == "2":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G2BestBass = ? WHERE CompetitionID = ?',
                                (BestBassSection, CompID))
                            conn.commit()

                        elif grade == "3A":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G3ABestBass = ? WHERE CompetitionID = ?',
                                (BestBassSection, CompID))
                            conn.commit()

                        elif grade == "3B":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G3BBestBass = ? WHERE CompetitionID = ?',
                                (BestBassSection, CompID))
                            conn.commit()

                        elif grade == "4A":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G4ABestBass = ? WHERE CompetitionID = ?',
                                (BestBassSection, CompID))
                            conn.commit()

                        elif grade == "4B":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G4BBestBass = ? WHERE CompetitionID = ?',
                                (BestBassSection, CompID))
                            conn.commit()

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute(
                            'UPDATE Results SET DrummingScore = ? WHERE CompetitionID = ? AND BandID = ?',
                            (SectionScore[each], CompID, BandID))
                        conn.commit()

                    elif Section == "Ensemble":

                        BestMD = BandBestMD.get()

                        if grade == "1":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G1BestMD = ? WHERE CompetitionID = ?',
                                (BestMD, CompID))
                            conn.commit()

                        elif grade == "2":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G2BestMD = ? WHERE CompetitionID = ?',
                                (BestMD, CompID))
                            conn.commit()

                        elif grade == "3A":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G3ABestMD = ? WHERE CompetitionID = ?',
                                (BestMD, CompID))
                            conn.commit()

                        elif grade == "3B":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G3BBestMD = ? WHERE CompetitionID = ?',
                                (BestMD, CompID))
                            conn.commit()

                        elif grade == "4A":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G4ABestMD = ? WHERE CompetitionID = ?',
                                (BestMD, CompID))
                            conn.commit()

                        elif grade == "4B":
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Competitions SET G4BBestMD = ? WHERE CompetitionID = ?',
                                (BestMD, CompID))
                            conn.commit()

                        conn = sqlite3.connect('RspbaniDB.db')
                        with conn:
                            cursor = conn.cursor()
                        cursor.execute(
                            'UPDATE Results SET EnsembleScore = ? WHERE CompetitionID = ? AND BandID = ?',
                            (SectionScore[each], CompID, BandID))
                        conn.commit()

                    if Section == "Piping 1" or "Piping 2":
                        try:
                            # get the scores of both Piping jusges and add them together
                            with sqlite3.connect("RspbaniDB.db") as db:
                                # Creates a cursor to search through the data
                                cursor = db.cursor()
                                # SQL Query to select both piping scores from the results table
                            TotalPipingQuery = 'SELECT Piping1Score, Piping2Score FROM Results WHERE CompetitionID = ? AND BandID = ? '
                            # Execute the query to get these scores
                            cursor.execute(TotalPipingQuery, [CompID, BandID])
                            # The query will produce a 2 dimensional tuple
                            TotalPipingResults = cursor.fetchall()
                            for i in TotalPipingResults:
                                PipingScore1 = i[0]
                                PipingScore2 = i[1]
                            totalScore = int(PipingScore1) + int(PipingScore2)

                            print("Total Piping Score = " + str(totalScore))

                            # save the combined score to the database also
                            conn = sqlite3.connect('RspbaniDB.db')
                            with conn:
                                cursor = conn.cursor()
                            cursor.execute(
                                'UPDATE Results SET TotalPipingScore = ? WHERE CompetitionID = ? AND BandID = ?',
                                (totalScore, CompID, BandID))
                            conn.commit()

                        except:
                            None

                messagebox.showinfo(message="Score Saved")
                JudgeCompetitionFrame(photoSelectCompetitionButton)

            SectionScore = {}
            n = 0
            possibleScore = []
            for i in ValidPossScoreList:
                possibleScore.append(i)

            print("PossibleScore" + str(possibleScore))
            ValidCount = 0
            for entry in Score:
                BandScore = entry.get()

                if BandScore == "":
                    # a number that will never be a possible score
                    # because if the field is left blank then it cant convert "" to a integer
                    BandScore = "0"
                elif BandScore.isdigit() == False:
                    # if input is not a number then it is invalid so i am setting the BandScore to 0 as it is valid
                    # however can never be a possible score so will be caught further on in the validation routine
                    BandScore = "0"

                if int(BandScore) in possibleScore:
                    SectionScore[ListBands[n]] = BandScore
                    possibleScore.remove(int(BandScore))
                    print("PossibleScore" + str(possibleScore))
                    print("List" + str(ValidPossScoreList))
                    n += 1
                else:
                    ValidCount += 1
            if ValidCount == 0:
                print("ValidCount " + str(ValidCount))

                if (Section == "Drumming" and BandBestBass.get() != "Select Best Bass Section") or (
                        Section == "Ensemble" and BandBestMD.get() != "Select Best M&D"):
                    SaveSheet(listOfJudgeScoreBox)
                else:
                    if Section == "Drumming":
                        messagebox.showerror(message="Please select the band with the best Bass Section")
                    elif Section == "Ensemble":
                        messagebox.showerror(message="Please select the band with the best M&D")
                    elif Section == "Piping 1" or "Piping 2":
                        SaveSheet(listOfJudgeScoreBox)

            else:
                messagebox.showwarning(
                    message="You have not entered Valid scores, please make sure enter a score greater than 0 and less "
                            "than " + str(len(ListBands) + 1) + ". You can also not give two bands the same score! ")
                possibleScore = []
                for i in ValidPossScoreList:
                    possibleScore.append(i)

        def ScoreSheet(grade, ListBands, listOfJudgeScoreBox):

            Selection = listOfJudgeScoreBox.curselection()

            if Selection == ():
                messagebox.showinfo(message="Please select the section to enter your scores.")

            for i in Selection:
                # gets the section chosen by the judge
                Section = listOfJudgeScoreBox.get(i)
                print(str(Section))

            ScoreBands = Frame(contentFrame, width=965, height=600)
            ScoreBands.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(ScoreBands, width=30)
            lblspacer.grid()

            # Creation of widgets on the ScoreBands Frame
            lblTitle = Label(ScoreBands, text="Enter Scores", height=3, font=("Arvo", 32))
            lblTitle.grid(row=0, column=2, columnspan=6)

            lblText = Label(ScoreBands,
                            text="Please select the score you wish to enter for each band in Grade " + grade + " - " + Section,
                            height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=2, columnspan=6, sticky="ew", )

            PossScoreList = list(range(1, len(ListBands) + 1))

            print("PossScoreList")
            print(PossScoreList)
            Score = []
            n = 5

            # implement drop down box with possible scores left by removing the choices off the list when they have been selected for Post prototype refinement
            for i in range(len(ListBands)):
                lblBands = Label(ScoreBands, text="Score for " + ListBands[i])
                lblBands.grid(row=i + 3, column=2, sticky="e")

                EtyBandScore = Entry(ScoreBands)
                EtyBandScore.grid(row=i + 3, column=3, sticky="w")
                Score.append(EtyBandScore)
                n += 1

            if Section == "Drumming":
                global BandBestBass
                BandBestBass = StringVar()
                lblBestBass = Label(ScoreBands, text="Best Bass Section ")
                lblBestBass.grid(row=n + 3, column=2, sticky="e")

                BandDropList = OptionMenu(ScoreBands, BandBestBass, *ListBands)
                BandDropList.config(width=20, font=("Arvo", 12))
                BandBestBass.set("Select Best Bass Section")
                BandDropList.grid(row=n + 3, column=3, sticky="w")

            if Section == "Ensemble":
                global BandBestMD
                BandBestMD = StringVar()
                lblBestBass = Label(ScoreBands, text="Best M&D ")
                lblBestBass.grid(row=n + 3, column=2, sticky="e")

                BandDropList = OptionMenu(ScoreBands, BandBestMD, *ListBands)
                BandDropList.config(width=20, font=("Arvo", 12))
                BandBestMD.set("Select Best M&D")
                BandDropList.grid(row=n + 3, column=3, sticky="w")

            lblspacer = Label(ScoreBands)
            lblspacer.grid(row=n+4)

            ButtonSave = Button(ScoreBands, image=photoSaveScores,
                                command=lambda: ScoreSheetCheck(Score, ListBands, Section, grade, PossScoreList,
                                                                listOfJudgeScoreBox))
            ButtonSave.grid(row=n + 5, column=2, columnspan=2)

        def ScoreSection(grade, ListBands):
            JudgeCompJudge = Frame(contentFrame, width=965, height=600)
            JudgeCompJudge.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(JudgeCompJudge, width=30)
            lblspacer.grid()

            # Creation of widgets on the JudgeCompJudge Frame
            lblTitle = Label(JudgeCompJudge, text="Select Section", height=3, font=("Arvo", 32))
            lblTitle.grid(row=0, column=2, columnspan=6)

            lblText = Label(JudgeCompJudge,
                            text="Please select the Section you wish to Enter your Score for in Grade " + grade,
                            height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=2, columnspan=6, sticky="ew", )

            listOfJudgeScoreBox = Listbox(JudgeCompJudge, selectmode=SINGLE, width=25, height=5)
            listOfJudgeScoreBox.grid(row=3, column=4, columnspan=2)
            listOfJudgeScoreBox.insert(0, "Piping 1")
            listOfJudgeScoreBox.insert(1, "Piping 2")
            listOfJudgeScoreBox.insert(2, "Drumming")
            listOfJudgeScoreBox.insert(3, "Ensemble")

            SelectButton = Button(JudgeCompJudge, image=photoSelectCompetitionButton
                                  , command=lambda: ScoreSheet(grade, ListBands, listOfJudgeScoreBox))
            SelectButton.grid(row=4, column=3, columnspan=4, pady="20")

        def CreateListBoxOfBands(ListBands, grade, JudgeID, listOfCompNameBox):
            print("grade")
            print(grade)

            JudgeCompBands = Frame(contentFrame, width=965, height=600)
            JudgeCompBands.grid(row=0, column=0, sticky="nsew")

            lblspacer = Label(JudgeCompBands, width=40)
            lblspacer.grid()

            # Creation of widgets on the JudgeCompBands Frame to select the band to be judged
            lblTitle = Label(JudgeCompBands, text="Select Band", height=3, font=("Arvo", 32))
            lblTitle.grid(row=0, column=2, columnspan=6)

            lblText = Label(JudgeCompBands, text="Please select the Band you wish to Judge. ", height=3,
                            font=("Arvo", 14))
            lblText.grid(row=1, column=2, columnspan=6, sticky="ew", )

            listOfBandNameBox = Listbox(JudgeCompBands, selectmode=SINGLE, width=25)
            listOfBandNameBox.grid(row=3, column=4, columnspan=2)

            for u in ListBands:
                n = 0
                listOfBandNameBox.insert(n, str(u))
                n = + 1

            SelectBandButton = Button(JudgeCompBands, image=photoSelectBandButton
                                      , command=lambda: selectJudge(photoUpdateCompetitionButton, ListBands,
                                                                    listOfBandNameBox, JudgeID))
            SelectBandButton.grid(row=4, column=4, columnspan=2, pady="20")

            if str(grade) == "4B":
                EnterBandsScoreButton = Button(JudgeCompBands, image=photoScoringGrade4B,
                                               command=lambda: ScoreSection(grade, ListBands))
                EnterBandsScoreButton.grid(row=5, column=9, padx="10")
            if str(grade) == "4A":
                EnterBandsScoreButton = Button(JudgeCompBands, image=photoScoringGrade4A,
                                               command=lambda: ScoreSection(grade, ListBands))
                EnterBandsScoreButton.grid(row=5, column=9, padx="10")
            if str(grade) == "3B":
                EnterBandsScoreButton = Button(JudgeCompBands, image=photoScoringGrade3B,
                                               command=lambda: ScoreSection(grade, ListBands))
                EnterBandsScoreButton.grid(row=5, column=9, padx="10")
            if str(grade) == "3A":
                EnterBandsScoreButton = Button(JudgeCompBands, image=photoScoringGrade3A,
                                               command=lambda: ScoreSection(grade, ListBands))
                EnterBandsScoreButton.grid(row=5, column=9, padx="10")
            if str(grade) == "2":
                EnterBandsScoreButton = Button(JudgeCompBands, image=photoScoringGrade2,
                                               command=lambda: ScoreSection(grade, ListBands))
                EnterBandsScoreButton.grid(row=5, column=9, padx="10")
            if str(grade) == "1":
                EnterBandsScoreButton = Button(JudgeCompBands, image=photoScoringGrade1,
                                               command=lambda: ScoreSection(grade, ListBands))
                EnterBandsScoreButton.grid(row=5, column=9, padx="10")



        Selection = listOfGradeBox.curselection()

        if Selection == ():
            messagebox.showinfo(message="Please select a grade to judge!")
        else:

            for i in Selection:
                # get the grade selected in the list box
                grade = listOfGradeBox.get(i)
                print(grade)


            # call the function dependant on the grade selected but dont ley judge select band if there are no bands entered into it
            if grade == "1" and len(BandsNameInGrade1) != 0:
                CreateListBoxOfBands(BandsNameInGrade1, grade, JudgeID, listOfCompNameBox)
            elif grade == "2" and len(BandsNameInGrade2) != 0:
                CreateListBoxOfBands(BandsNameInGrade2, grade, JudgeID, listOfCompNameBox)
            elif grade == "3A" and len(BandsNameInGrade3A) != 0:
                CreateListBoxOfBands(BandsNameInGrade3A, grade, JudgeID, listOfCompNameBox)
            elif grade == "3B" and len(BandsNameInGrade3B) != 0:
                CreateListBoxOfBands(BandsNameInGrade3B, grade, JudgeID, listOfCompNameBox)
            elif grade == "4A" and len(BandsNameInGrade4A) != 0:
                CreateListBoxOfBands(BandsNameInGrade4A, grade, JudgeID, listOfCompNameBox)
            elif grade == "4B" and len(BandsNameInGrade4B) != 0:
                CreateListBoxOfBands(BandsNameInGrade4B, grade, JudgeID, listOfCompNameBox)
            else:
                messagebox.showerror(message="No Bands have entered this grade!")


    def selectedComp(photoUpdateCompetitionButton, listOfCompNameBox, JudgeID):
        def getListOfBandNameIneachgrade(CompID, Grade):

            with sqlite3.connect("RspbaniDB.db") as db:
                # Creates a cursor to search through the data
                cursorListOfComps = db.cursor()
                # SQL Query to select the BandID from BandsEntered where they have entered the competition and grade selected
            listOfCompsQuery = 'SELECT BandID FROM BandsEntered WHERE CompetitionID = ? AND Grade = ?'
            # Execute the query to get a list of all the Bands
            cursorListOfComps.execute(listOfCompsQuery, [CompID, Grade])
            # The query will produce a 2 dimensional tuple
            ListOfBands = cursorListOfComps.fetchall()
            listOfBandNameEntered = []
            for i in ListOfBands:
                for k in i:
                    print("BandID = " + str(k))
                    # gets the Band names from their ID
                    with sqlite3.connect("RspbaniDB.db") as db:
                        cursorBandName = db.cursor()
                    findBandName = 'SELECT BandName FROM BandAccount WHERE BandID = ?'
                    cursorBandName.execute(findBandName, [k])
                    BandNameResults = cursorBandName.fetchall()

                    if BandNameResults:
                        for e in BandNameResults:
                            for u in e:
                                print(u)
                                # append all bandID's to a list
                                listOfBandNameEntered.append(u)
            # returns the list
            return listOfBandNameEntered

        # get position of the selected value
        Selection = listOfCompNameBox.curselection()

        print("Selection")
        print(Selection)

        if Selection == ():
            messagebox.showinfo(message="Please select a competition to judge!")
        else:

            for i in Selection:
                # get the value of the selected position
                j = listOfCompNameBox.get(i)

                with sqlite3.connect("RspbaniDB.db") as db:
                    cursorCompID = db.cursor()
                findCompID = 'SELECT CompetitionID FROM Competitions WHERE CompetitionName = ?'
                cursorCompID.execute(findCompID, [j])
                CompIDresults = cursorCompID.fetchall()

                if CompIDresults:

                    for a in CompIDresults:
                        a = str(a)
                        CompID = a[1:-2]

                    # Sets the Collated Field in the comps table to False
                    conn = sqlite3.connect('RspbaniDB.db')
                    with conn:
                        cursor = conn.cursor()
                    cursor.execute('UPDATE Competitions SET Collated = ? WHERE CompetitionID = ? ',
                                   ("False", CompID))
                    conn.commit()

                    # returns the list of bands in each grade of the selected competition
                    BandsNameInGrade1 = getListOfBandNameIneachgrade(CompID, "1")
                    BandsNameInGrade2 = getListOfBandNameIneachgrade(CompID, "2")
                    BandsNameInGrade3A = getListOfBandNameIneachgrade(CompID, "3A")
                    BandsNameInGrade3B = getListOfBandNameIneachgrade(CompID, "3B")
                    BandsNameInGrade4A = getListOfBandNameIneachgrade(CompID, "4A")
                    BandsNameInGrade4B = getListOfBandNameIneachgrade(CompID, "4B")

                    print("BandsNameInGrade1 " + str(BandsNameInGrade1))
                    print("BandsNameInGrade2 " + str(BandsNameInGrade2))
                    print("BandsNameInGrade3A " + str(BandsNameInGrade3A))
                    print("BandsNameInGrade3B " + str(BandsNameInGrade3B))
                    print("BandsNameInGrade4A " + str(BandsNameInGrade4A))
                    print("BandsNameInGrade4B " + str(BandsNameInGrade4B))

                    JudgeCompGrade = Frame(contentFrame, width=965, height=600)
                    JudgeCompGrade.grid(row=0, column=0, sticky="nsew")

                    lblspacer = Label(JudgeCompGrade, width=40)
                    lblspacer.grid()

                    # Creation of widgets on the JudgeCompGrade Frame
                    lblTitle = Label(JudgeCompGrade, text="Select Grade", height=3, font=("Arvo", 32))
                    lblTitle.grid(row=0, column=2, columnspan=6)

                    lblText = Label(JudgeCompGrade, text="Please select the Grade you would like to Judge. ", height=3,
                                    font=("Arvo", 14))
                    lblText.grid(row=1, column=2, columnspan=6, sticky="ew", )

                    listOfGradeBox = Listbox(JudgeCompGrade, selectmode=SINGLE, width=10, height=7)
                    listOfGradeBox.grid(row=3, column=4)
                    listOfGradeBox.insert(0, "1")
                    listOfGradeBox.insert(1, "2")
                    listOfGradeBox.insert(2, "3A")
                    listOfGradeBox.insert(3, "3B")
                    listOfGradeBox.insert(4, "4A")
                    listOfGradeBox.insert(5, "4B")

                    SelectGradeButton = Button(JudgeCompGrade, image=photoSelectGradeButton
                                               , command=lambda: selectedGrade(photoUpdateCompetitionButton, listOfGradeBox,
                                                                               BandsNameInGrade1, BandsNameInGrade2,
                                                                               BandsNameInGrade3A, BandsNameInGrade3B,
                                                                               BandsNameInGrade4A, BandsNameInGrade4B,
                                                                               CompID, JudgeID, listOfCompNameBox))
                    SelectGradeButton.grid(row=4, column=3, columnspan=3, pady="20")

                    ButtonBack = Button(JudgeCompGrade, image=photoBack, command=lambda: DisplayCompetitionList(JudgeID))
                    ButtonBack.grid(row=0, column=0)

    if Admin == True:
        # If user is admin then ask them to select which band they wish to edit.

        SelectJudge = Frame(contentFrame, width=965, height=600)
        SelectJudge.grid(row=0, column=0, sticky="nsew")

        lblspacerCenter = Label(SelectJudge, text="", width=20)
        lblspacerCenter.grid(row=0, column=0)

        Title = Label(SelectJudge, text="Select Judge", height=3, font=("Arvo", 32))
        Title.grid(row=0, column=1, columnspan=8, sticky="ew", padx=180)

        lblText = Label(SelectJudge, text="Please select the judge you are entering results on behalf of: ", height=3,
                        font=("Arvo", 14))
        lblText.grid(row=1, column=1, columnspan=8, sticky="ew", padx=170)

        JudgeListBandBox = Listbox(SelectJudge, selectmode=EXTENDED, width=25)
        JudgeListBandBox.grid(row=3, column=3, columnspan=4)

        with sqlite3.connect("RspbaniDB.db") as db:
            cursor = db.cursor()
        findJudge = 'SELECT FirstName, SecondName FROM JudgeAccount'
        cursor.execute(findJudge)
        JudgeResults = cursor.fetchall()

        if JudgeResults:
            for k in JudgeResults:
                n = 0
                FirstName = k[0]
                SecondName = k[1]
                Name = FirstName + " " + SecondName
                JudgeListBandBox.insert(n, Name)
                n = + 1

        SelectButton = Button(SelectJudge, image=photoSelectJudge, command=lambda: AdminSelectedJudge(JudgeListBandBox))
        SelectButton.grid(row=4, column=4, columnspan=2, pady="10")

    else:
        DisplayCompetitionList(JudgeIDOfLoggedIn)


def addMemberDatabase(EtyFirstName, EtySecondName, EtyAddressLine1, EtyAddressLine2, EtyPostcode, EtySignature, Role,
                      FirstName, SecondName,
                      AddressLine1, AddressLine2, County, Postcode, DateOfBirth, Signature, MemDOBDay, MemDOBMonth,
                      MemDOBYear, Memcounty,
                    BandID, EtyEmailAddress, Email, Title, Gender):
    def SendEmailToPM(PMEmail):

        # email address of the sender
        email_user = 'rspbani.info@gmail.com'
        # password of the sender
        email_password = 'P@55w0rd123'
        # email address of recipient
        email_send = PMEmail

        # Creating the Emails subject
        subject = 'New Member Registered'

        msg = MIMEMultipart()

        msg['From'] = email_user
        msg['To'] = email_send
        msg['Subject'] = subject

        # creating the contents of the email
        body = "Dear Pipe Major, \n\n\nYou are receiving this email to confirm the registration of your new member " + str(
            FirstName) + " " + str(
            SecondName) + ". \n\n\nThe newly registered member will also receive an email for conformation, please check with them to make sure their details are correct - if not please update them through the Edit Member Details Section. \n\n\n\nThe Royal Scottish Pipe Band Association \nNorthern Ireland Branch "
        msg.attach(MIMEText(body, 'plain'))
        text = msg.as_string()

        try:
            # connet to email server and send email
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(email_user, email_password)

            server.sendmail(email_user, email_send, text)

        # error handling if email cant be sent - display a message stating this
        except(smtplib.SMTPException, ConnectionRefusedError, OSError):
            messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
        finally:
            # stop connection with email server
            server.quit()

    def SendEmailToMember(BandName):
        # email address of the sender
        email_user = 'rspbani.info@gmail.com'
        # password of the sender
        email_password = 'P@55w0rd123'
        # email address of recipient
        email_send = Email

        # Creating the Emails subject
        subject = 'Welcome ' + str(FirstName)

        msg = MIMEMultipart()

        msg['From'] = email_user
        msg['To'] = email_send
        msg['Subject'] = subject

        # creating the contents of the email
        body = "Dear " + str(FirstName) + " " + str(
            SecondName) + ", \n\n\nYou are receiving this email to confirm your registration " \
                          "with " + str(BandName) + ". \n\n\nPlease check to make sure your " \
                                                    "details are correct - if not please let your Pipe Major know. " \
                                                    "\n\nTitle:   " + str(Title) + "\nFirst Name:   " + str(
            FirstName) + "\n" \
                         "Second Name:   " + str(SecondName) + "\nAddress Line One:   " + str(AddressLine1) + "\n" \
                                                                                                              "Address Line Two:   " + str(
            AddressLine2) + "\n" \
                            "County:   " + str(County) + "\n" \
                                                         "Postcode:   " + str(Postcode) + "\n" \
                                                                                          "Date of Birth:   " + str(
            DateOfBirth) + "\n" \
                           "Gender:   " + str(Gender) + "\n" \
                                                        "\n\n\n\n\nThe Royal Scottish Pipe Band Association \nNorthern Ireland Branch "
        msg.attach(MIMEText(body, 'plain'))

        text = msg.as_string()

        try:
            # connet to email server and send email
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(email_user, email_password)

            server.sendmail(email_user, email_send, text)

        # error handling if email cant be sent - display a message stating this
        except(smtplib.SMTPException, ConnectionRefusedError, OSError):
            messagebox.showinfo("Alert ", message="Message not sent \n\nPlease try again.")
        finally:
            # stop connection with email server
            server.quit()

    with sqlite3.connect("RspbaniDB.db") as db:
        cursor = db.cursor()
    findJudge = 'SELECT EmailAddress, BandName FROM BandAccount WHERE BandID = ?'
    cursor.execute(findJudge, [BandID])
    EmailResults = cursor.fetchall()

    if EmailResults:
        for k in EmailResults:
            PMEmail = k[0]
            BandName = k[1]
    #Table Created
    conn = sqlite3.connect('RspbaniDB.db')
    with conn:
        cursor = conn.cursor()
    # create the BandMembers table
    cursor.execute(
        'CREATE TABLE IF NOT EXISTS BandMembers (MemberID INTEGER PRIMARY KEY, BandID INTEGER, '
        'Role TEXT, Title TEXT, FirstName TEXT, SecondName TEXT, AddressLine1 TEXT, AddressLine2 TEXT, County TEXT, '
        'Postcode TEXT, DateOfBirth TEXT, Gender TEXT, Signature TEXT, EmailAddress TEXT)')
    # Insert the fields which have been fetched and insert them into the database
    cursor.execute(
        'INSERT INTO BandMembers (MemberID, BandID, Role, Title, FirstName, SecondName, AddressLine1,'
        ' AddressLine2, County, Postcode, DateOfBirth, Gender, Signature, EmailAddress) VALUES(NULL,?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (BandID, Role, Title, FirstName, SecondName, AddressLine1, AddressLine2, County,
         Postcode, DateOfBirth, Gender, Signature, Email))
    conn.commit()

    SendEmailToPM(PMEmail)
    SendEmailToMember(BandName)

    # Once the data has successfully been saved to the database then remove the entered data and
    # reset the drop down menu
    Memcounty.set('Select your County')
    EtyFirstName.delete(0, END)
    EtySecondName.delete(0, END)
    EtyAddressLine1.delete(0, END)
    EtyAddressLine2.delete(0, END)
    EtyEmailAddress.delete(0, END)
    EtyPostcode.delete(0, END)
    MemDOBMonth.set('Select Month')
    MemDOBDay.set('Select Day')
    MemDOBYear.set('Select Year')
    EtySignature.delete(0, END)

    messagebox.showinfo(message="Member Successfully Registered")
    HomeFrame()


def addBandDatabase(Band, Branch, Grade, PMFirstName, PMSecondName, BandHallName, BandAddressLine1, BandAddressLine2,
                    County, BandPostcode, BandPracticeTime, BandTartan, BandUsername, BandPassword, BandEmail,
                    EtyBandName, EtyPMFirstName, EtyPMSecondName, EtyBandHallName, EtyBandAddressLine1,
                    EtyBandAddressLine2, EtyBandPostcode, EtyBandPracticeTime, EtyBandTartan, EtyBandUsername,
                    EtyBandPassword, EtyBandEmail, BandBranch, BandGrade, BandCounty):
    print(Band, Branch, Grade, PMFirstName, PMSecondName, BandHallName, BandAddressLine1, BandAddressLine2, BandCounty,
          BandPostcode, BandPracticeTime, BandTartan, BandEmail, BandUsername, BandPassword)
    #Table Created
    conn = sqlite3.connect('RspbaniDB.db')
    with conn:
        cursor = conn.cursor()
    # create the Band table if it doesnt already exist to store Band account details
    cursor.execute('CREATE TABLE IF NOT EXISTS BandAccount '
                   '(BandID INTEGER PRIMARY KEY, BandName TEXT, BandBranch TEXT, BandGrade TEXT, PMFirstName TEXT, '
                   'PMSecondName TEXT, HallName TEXT, BandAddressLine1 TEXT, BandAddressLine2 TEXT, BandCounty TEXT, '
                   'BandPostcode TEXT, PracticeTime TEXT, Tartan TEXT, EmailAddress TEXT, Username TEXT, Password TEXT, SecurityQ TEXT, SecurityQAns, FirstLogIn BOOLEAN)')
    # Insert the fields which have been fetched and insert them into the database
    cursor.execute('INSERT INTO BandAccount '
                   '(BandID, BandName, BandBranch, BandGrade, PMFirstName, PMSecondName, HallName, BandAddressLine1, '
                   'BandAddressLine2, BandCounty, BandPostcode, PracticeTime, Tartan, EmailAddress ,Username, Password, FirstLogIn,SecurityQ, SecurityQAns) '
                   'VALUES(NULL,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                   (Band, Branch, Grade, PMFirstName, PMSecondName, BandHallName, BandAddressLine1, BandAddressLine2,
                    County, BandPostcode, BandPracticeTime, BandTartan, BandEmail, BandUsername, BandPassword, "True",
                    "Unknown", "Unknown"))
    conn.commit()

    # Once the data has successfully been saved to the database then remove the entered data and reset the drop down menu
    EtyBandName.delete(0, END)
    BandBranch.set('Select your Branch')
    BandGrade.set('Select your Grade')
    BandCounty.set('Select your County')
    EtyPMFirstName.delete(0, END)
    EtyPMSecondName.delete(0, END)
    EtyBandAddressLine1.delete(0, END)
    EtyBandAddressLine2.delete(0, END)
    EtyBandPostcode.delete(0, END)
    EtyBandHallName.delete(0, END)
    EtyBandPracticeTime.delete(0, END)
    EtyBandTartan.delete(0, END)
    EtyBandUsername.delete(0, END)
    EtyBandPassword.delete(0, END)
    EtyBandEmail.delete(0, END)

    body = "Dear Pipe Major, \n\n\nYou are receiving this email to confirm the registration of this band, " + str(
        Band) + ", under your leadership.\n\n\n Please " \
                "find your initial log in details. Once you log in for the first time, you will be asked to set your own password along with a security question " \
                "which you will need to reset your password through the application. \nAlso through this application you are able to do the following:\n\n* View " \
                "Competitions \n* View Competition Draw \n* View Competition Results \n* View all Bands and their details  \n* Edit " \
                "Band Details \n* Enter Competitions \n* Book Transport \n* Register New Members \n* Edit Member Details \n\n\nUsername: " + str(
        BandUsername) + "\nPassword: " + str(BandPassword) + "\n\n\n\nThe Royal Scottish Pipe " \
                                                             "Band Association \nNorthern Ireland Branch "

    EmailUser(BandEmail, Band, body)

    listOfBands.append(Band)
    return listOfBands


def addJudgeDatabase(FirstName, SecondName, DOB, Gender, AddressLine1, AddressLine2, County, Postcode, Email, Username,
                     Password, Title,
                     EtyJudgeFirstName, EtyJudgeSecondName, EtyJudgeAddressLine1, EtyJudgeAddressLine2,
                     EtyJudgePostcode, EtyJudgeEmail,
                     EtyJudgeUsername, EtyJudgePassword, JudgeDOBDay, JudgeDOBMonth, JudgeDOBYear, JudgeGender,
                     JudgeCounty, JudgeTitle):
    # Table Created
    conn = sqlite3.connect('RspbaniDB.db')
    with conn:
        cursor = conn.cursor()
    # create the judge table if it doesnt already exist to store judge account details
    cursor.execute('CREATE TABLE IF NOT EXISTS JudgeAccount '
                   '(JudgeID INTEGER PRIMARY KEY, Title TEXT, FirstName TEXT, SecondName TEXT, DateOfBirth TEXT, Gender TEXT, '
                   'AddressLine1 TEXT, AddressLine2 TEXT, County TEXT, Postcode TEXT, Email TEXT,'
                   'Username TEXT, Password TEXT, SecurityQ TEXT, SecurityQAns TEXT, FirstLogIn TEXT)')
    # Insert the fields which have been fetched and insert them into the database
    cursor.execute('INSERT INTO JudgeAccount '
                   '(JudgeID, Title, FirstName, SecondName, DateOfBirth, Gender, AddressLine1, AddressLine2, '
                   'County, Postcode, Username, Password, Email, FirstLogIn, SecurityQ, SecurityQAns) '
                   'VALUES(NULL,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                   (Title, FirstName, SecondName, DOB, Gender, AddressLine1, AddressLine2,
                    County, Postcode, Username, Password, Email, "True", "Unknown", "Unknown"))
    conn.commit()
    # Once the data has successfully been saved to the database then remove the entered data and reset the drop down menu
    JudgeTitle.set("Select Title")
    EtyJudgeFirstName.delete(0, END)
    EtyJudgeSecondName.delete(0, END)
    JudgeGender.set('Select your Gender')
    EtyJudgeAddressLine1.delete(0, END)
    EtyJudgeAddressLine2.delete(0, END)
    JudgeCounty.set('Select your county')
    EtyJudgePostcode.delete(0, END)
    EtyJudgeUsername.delete(0, END)
    EtyJudgeEmail.delete(0, END)
    EtyJudgePassword.delete(0, END)
    JudgeDOBYear.set('Select Year')
    JudgeDOBMonth.set('Select Month')
    JudgeDOBDay.set('Select Day')

    body = "Dear " + str(
        FirstName) + ", \n\n\nYou are receiving this email to confirm the registration of this account.\n\n\n Please " \
                     "find your initial log in details. Once you log in for the first time, you will be asked to set your own password along with a security question " \
                     "which you will need to reset your password through the application. \nAlso through this application you are able to do judge " \
                     "competitions through this which replace the current paper based system. \n\n\nUsername: " + str(
        Username) + "\nPassword: " + str(Password) + "\n\n\n\nThe Royal Scottish Pipe " \
                                                     "Band Association \nNorthern Ireland Branch "

    EmailUser(Email, FirstName, body)


def addAdminDatabase(FirstName, SecondName, DOB, Gender,
                     AddressLine1, AddressLine2, County, Postcode, Email, Username, Password,
                     Title, EtyAdminFirstName, EtyAdminSecondName, EtyAdminGender,
                     EtyAdminAddressLine1, EtyAdminAddressLine2, EtyAdminCounty, EtyAdminPostcode, EtyEmail,
                     EtyAdminUsername,
                     EtyAdminPassword, AdminDOBDay, AdminDOBMonth, AdminDOBYear, AdminGender, AdminCounty, AdminTitle):


    #Table Created
    conn = sqlite3.connect('RspbaniDB.db')
    with conn:
        cursor = conn.cursor()
    # create the Admin table if it doesnt already exist to store admin account details
    cursor.execute('CREATE TABLE IF NOT EXISTS AdminAccount '
                   '(AdminID INTEGER PRIMARY KEY, Title TEXT, FirstName TEXT, SecondName TEXT, DateOfBirth TEXT, Gender TEXT, '
                   'AddressLine1 TEXT, AddressLine2 TEXT, County TEXT, Postcode TEXT, '
                   'Email TEXT, Username TEXT, Password TEXT, SecurityQ TEXT, SecurityQAns TEXT, FirstLogIn TEXT)')
    # Insert the fields which have been fetched and insert them into the database
    cursor.execute('INSERT INTO AdminAccount '
                   '(AdminID, Title, FirstName, SecondName, DateOfBirth, Gender, AddressLine1, AddressLine2, '
                   'County, Postcode, Email, Username, Password, FirstLogIn, SecurityQ, SecurityQAns) '
                   'VALUES(NULL,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                   (Title, FirstName, SecondName, DOB, Gender, AddressLine1, AddressLine2,
                    County, Postcode, Email, Username, Password, "True", "Unknown", "Unknown"))
    conn.commit()
    # Once the data has successfully been saved to the database then remove the entered data and reset the drop down menu
    AdminTitle.set('Select Title')
    EtyAdminFirstName.delete(0, END)
    EtyAdminSecondName.delete(0, END)
    AdminGender.set('Select your Gender')
    AdminDOBDay.set('Select Day')
    AdminDOBMonth.set('Select Month')
    AdminDOBYear.set('Select Year')
    EtyAdminAddressLine1.delete(0, END)
    EtyAdminAddressLine2.delete(0, END)
    AdminCounty.set('Select your County')
    EtyAdminPostcode.delete(0, END)
    EtyAdminUsername.delete(0, END)
    EtyAdminPassword.delete(0, END)
    EtyEmail.delete(0, END)

    body = "Dear " + str(
        FirstName) + ", \n\n\nYou are receiving this email to confirm the registration of this administrative account." \
                     "\n\n\n Please find your initial log in details below. Once you log in for the first time, you will be asked to set your own password " \
                     "along with a security question which you will need to reset your password through the application. \nAlso through this application you" \
                     " are able to do the following:\n\n* View Competitions \n* View Competition Draw \n* View Competition Results \n* View all Bands and " \
                     "their details  \n* Edit Band Details on behalf of bands \n* Enter Competitions on behalf of bands \n* Book Transport on behalf of bands" \
                     "\n* Register New Members on behalf of bands\n* Edit Member Details on behalf of bands\n* Register new Bands\n* Register new Judges" \
                     "\n* Register new Admins\n* Create and edit competitions\n* Register new bands\n* Draw Competitions\n* Collate Competition results" \
                     "\n* Send results to Bands\n* Judge Competitons on behalf of a judge" \
                     "\n\n\nUsername: " + str(Username) + "\nPassword: " + str(
        Password) + "\n\n\n\nThe Royal Scottish Pipe " \
                    "Band Association \nNorthern Ireland Branch "

    EmailUser(Email, FirstName, body)


def addCompDatabase(CompetitionName, CompetitionLocation, CompDate, EntryDate, EtyCompLocation, EtyCompName, Day, Month,
                    Year, EDeadDay, EDeadMonth, EDeadYear):
    # Gets all the user inputs from the Create Competition form and saves them to variables
    #Table Created
    conn = sqlite3.connect('RspbaniDB.db')
    with conn:
        cursor = conn.cursor()
    # create the Competition table if it doesnt already exist to store competition  details
    cursor.execute('CREATE TABLE IF NOT EXISTS Competitions '
                   '(CompetitionID INTEGER PRIMARY KEY, CompetitionName TEXT, CompetitionLocation TEXT, CompetitionDate TEXT, EntryDeadLine TEXT, Drawn TEXT, '
                   'NumG1Ent INTEGER, NumG2Ent INTEGER, '
                   'NumG3AEnt INTEGER, NumG3BEnt INTEGER, NumG4AEnt INTEGER, NumG4BEnt INTEGER, TotalNumBands INTEGER, G1BestBass TEXT, '
                   'G2BestBass TEXT, G3ABestBass TEXT, G3BBestBass TEXT, G4ABestBass TEXT, G4BBestBass TEXT, '
                   'G1BestMD TEXT, G2BestMD TEXT, G3ABestMD TEXT, G3BBestMD TEXT, G4ABestMD TEXT, G4BBestMD TEXT, Collated TEXT, '
                   'ResultFileName TEXT, Latest TEXT, SentToBands TEXT)')
    # Insert the fields which have been fetched and insert them into the database
    cursor.execute('INSERT INTO Competitions '
                   '(CompetitionID, CompetitionName, CompetitionLocation, CompetitionDate, Drawn, EntryDeadLine) '
                   'VALUES(NULL, ?,?,?,?,?)',
                   (CompetitionName, CompetitionLocation, CompDate, "False", EntryDate))
    conn.commit()
    # Once the data has successfully been saved to the database then remove the entered data and reset the drop down menus
    EtyCompName.delete(0, END)
    EtyCompLocation.delete(0, END)
    Day.set('Select Day')
    Month.set('Select Month')
    Year.set('Select Year')

    EDeadDay.set('Select Day')
    EDeadMonth.set('Select Month')
    EDeadYear.set('Select Year')


# create all of the main containers
headerFrame = Frame(root, height=130)
editFrame = Frame(root, height=20)
leftMenu = Frame(root, width=200, height=600)
contentFrame = Frame(root, width=965, height=600)
adminMenu = Frame(root, width=115, height=650)
JudgeMenu = Frame(root, width=1170, height=50)

# Place all containers on root
headerFrame.grid(row=0, column=0, columnspan=10, sticky="new")
editFrame.grid(row=1, column=0, columnspan=10, sticky="nsew")
leftMenu.grid(row=2, column=0, columnspan=2, sticky="nsw")
contentFrame.grid(row=2, column=2, columnspan=7, sticky="nsew")
adminMenu.grid(row=2, rowspan=2, column=9, sticky="nsew")
JudgeMenu.grid(row=3, column=0, columnspan=9, sticky="sew")

leftMenu.grid_columnconfigure(0, minsize=200)
adminMenu.grid_columnconfigure(9, minsize=110)

JudgeMenu.grid_rowconfigure(3, minsize=50)
adminMenu.grid_columnconfigure(2, minsize=650)

# HomeFrame()

# create the widgets for the left menu
photoHomePageButton = PhotoImage(file="HomePageButton.png")
photoBandsPageButton = PhotoImage(file="BandsPageButton.png")
photoCompetitionsPageButton = PhotoImage(file="CompetitionsPageButton.png")
photoResultsPageButton = PhotoImage(file="ResultsPageButton.png")

lblMenuSpacer = Label(leftMenu, text="", height=10)

# Create all the Buttons for the left Menu Frame which will load the respected frames when clicked
btnHomepage = Button(leftMenu, image=photoHomePageButton, command=HomeFrame, cursor="hand1")
btnBands = Button(leftMenu, image=photoBandsPageButton,
                  command=lambda: BandsFrame(photoSearchByGrade, photoSearchByBand))
btnCompetition = Button(leftMenu, image=photoCompetitionsPageButton, command=CompetitionFrame)
btnResults = Button(leftMenu, image=photoResultsPageButton, command=ResultsFrame)

# Place widgets on left menu Frame
lblMenuSpacer.grid(row=1, column=0, pady="5")
btnHomepage.grid(row=2, column=0, pady="5")
btnBands.grid(row=3, column=0, pady="5")
btnCompetition.grid(row=4, column=0, pady="5")
btnResults.grid(row=5, column=0, pady="5")

# layout all of the main containers
# root.rowconfigure(0, weight=0)
# root.columnconfigure(0, weight=0)

# create and places the widgets for the top frame
lblSignIn = Label(headerFrame, text="Sign in")
lblSignIn.grid(row=1, column=5, sticky="ns", padx="30")
# Creates a button that once pressed displays the Log in page so a user can log in.
buttonSignIn = Button(headerFrame, justify=RIGHT)
photoLogin = PhotoImage(file="LoginButton.png")
buttonSignIn.config(image=photoLogin, width="64", height="64", command=lambda: LoginFrame(photoLoginPageButton))
buttonSignIn.grid(row=0, column=5, sticky="se", pady="5", padx="30")

# Text and logo on header Frame
Title = Label(headerFrame, text="The Royal Scottish Pipe Band Association NI", font=("Arvo", 40))
Title.grid(row=0, column=2, rowspan=2, columnspan=3, sticky="w", padx="50")

photoSmallLogo = PhotoImage(file="RspbaniLogoSmallResized.png")
BtnLogo = Button(headerFrame, image=photoSmallLogo, command=HomeFrame)
BtnLogo.grid(row=0, rowspan=2, column=0, columnspan=2)


def Loggedin(photoLogin2, photoSmallLogo):
    # Creates an instance of the ContentFrame and displays the Home Page to the User once the Button has been clicked
    # This changes the top frame once a user has logged in
    Loggedin = Frame(root, height=100)
    Loggedin.grid(row=0, column=0, columnspan=10, sticky="nsew")

    lblSignOut = Label(Loggedin, text="Sign out")
    lblSignOut.grid(row=1, column=5, sticky="ns", padx="30")
    # Creates a button that once pressed quits the pogram to log out
    buttonSignOut = Button(Loggedin, justify=RIGHT)
    buttonSignOut.config(image=photoLogin2, width="64", height="64", command=quit)
    buttonSignOut.grid(row=0, column=5, sticky="se", pady="5", padx="30")

    # Displays the logo which if pressed will return the user to the Home Page
    BtnLogo = Button(Loggedin, image=photoSmallLogo, command=HomeFrame)
    BtnLogo.grid(row=0, rowspan=2, column=0, columnspan=2)

    Title = Label(Loggedin, text="The Royal Scottish Pipe Band Association NI", font=("Arvo", 40))
    Title.grid(row=0, column=2, rowspan=2, columnspan=3, sticky="w", padx="50")


# Creating an instance of photoImage for each image and assigning to a variable to be used in my code when needed
photoLogin2 = PhotoImage(file="Login.png")
photoNewBand = PhotoImage(file="NewBandButton.png")
photoNewJudge = PhotoImage(file="NewJudgeButton.png")
photoNewAdmin = PhotoImage(file="NewAdminButton.png")
photoNewCompetition = PhotoImage(file="NewCompetitionButton.png")
photoEditCompetition = PhotoImage(file="EditCompetitionButton.png")
photoDrawCompetition = PhotoImage(file="DrawCompetitionButton.png")
photoCollateResults = PhotoImage(file="CollateResultsButton.png")
photoSendResults = PhotoImage(file="SendResultsButton.png")

photoEditBandDetails = PhotoImage(file="EditBandDetailsButton.png")
photoEnterCompetitions = PhotoImage(file="EnterCompetitionsButton.png")
photoBookTransport = PhotoImage(file="BookTransportButton.png")
photoRegisterMember = PhotoImage(file="RegisterMemberButton.png")
photoTransferMember = PhotoImage(file="TransferMemberButton.png")
photoEditMemberDetails = PhotoImage(file="EditMemberDetailsButton.png")

photoSearchByGrade = PhotoImage(file="SearchByGradeButton.png")
photoSearch = PhotoImage(file="SearchButton.png")
photoSearchByBand = PhotoImage(file="SearchByBandButton.png")
photoLoginPageButton = PhotoImage(file="LoginPageButton.png")
photoRegisterAdminButton = PhotoImage(file="RegisterAdminButton.png")
photoRegisterBandButton = PhotoImage(file="RegisterBandButton.png")
photoRegisterJudgeButton = PhotoImage(file="RegisterJudgeButton.png")

photoRegisterMemberFormButton = PhotoImage(file="RegisterMemberFormButton.png")
photoSelectBandButton = PhotoImage(file="SelectBandButton.png")
photoSelectMemberButton = PhotoImage(file="SelectMemberButton.png")
photoSubmitButton = PhotoImage(file="SubmitButton.png")
photoTransferMemberFormButton = PhotoImage(file="TransferMemberFormButton.png")
photoUpdateBandButton = PhotoImage(file="UpdateBandButton.png")
photoUpdateMemberButton = PhotoImage(file="UpdateMemberButton.png")

photoSelectCompetitionButton = PhotoImage(file="SelectCompetitionButton.png")
photoSelectCompetitionsButton = PhotoImage(file="SelectCompetitionsButton.png")
photoViewCompetitionButton = PhotoImage(file="ViewCompetitionButton.png")
photoCreateCompetitionButton = PhotoImage(file="CreateCompetitionButton.png")
photoUpdateCompetitionButton = PhotoImage(file="UpdateCompetitionButton.png")
photoJudgeCompetition = PhotoImage(file="JudgeCompetitionsButton.png")
photoSelectGradeButton = PhotoImage(file="SelectGradeButton.png")
photoSaveButton = PhotoImage(file="SaveButton.png")
photoNextButton = PhotoImage(file="NextButton.png")
photoDeleteButton = PhotoImage(file="DeleteButton.png")
photoCancelButton = PhotoImage(file="CancelButton.png")
photoFinish = PhotoImage(file="FinishButton.png")
photoBack = PhotoImage(file="BackButton.png")
photoSelectJudge = PhotoImage(file="SelectJudgeButton.png")
photoViewDraw = PhotoImage(file="ViewDrawButton.png")
photoViewResults = PhotoImage(file="ViewResultsButton.png")
photoWithdraw = PhotoImage(file="photoWithdrawButton.png")
photoViewEnteredCompsWithdraw = PhotoImage(file="photoViewEnteredCompsButton.png")

photoScoringGrade1 = PhotoImage(file="ScoringGrade1.png")
photoScoringGrade2 = PhotoImage(file="ScoringGrade2.png")
photoScoringGrade3A = PhotoImage(file="ScoringGrade3A.png")
photoScoringGrade3B = PhotoImage(file="ScoringGrade3B.png")
photoScoringGrade4A = PhotoImage(file="ScoringGrade4A.png")
photoScoringGrade4B = PhotoImage(file="ScoringGrade4B.png")

photoSaveScores = PhotoImage(file="SaveScores.png")

photoResetButton = PhotoImage(file="ResetButton.png")
photoEnterButton = PhotoImage(file="EnterButton.png")

# Set the title of the root to the name of the organisation
root.title('The Royal Scottish Pipe Band Association')
# Set the size of the root to the dimensions of my laptop
root.geometry("1280x800+0+0")
root.config(bg="white")

# run the root
HomeFrame()

root.mainloop()
